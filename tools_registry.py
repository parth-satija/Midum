# --- AUTO-SPLITTER: imports added by automated pass, please review ---
from config import DOMAIN_INDEX, DOMAIN_SKILLS_INDEX, PATHS_FILE, RESPONSE_MEMORY, SKILLS_DIR, STARTUP_DIR
from config import _IS_LINUX, _IS_WINDOWS
from providers.gemini_web_backend import _GEMINI_WEBAPI_AVAILABLE, _gemini_webapi_load_msg, _get_gemini_web_client, _run_gemini_coro
from screen_capture import _do_click, _grab_full_screenshot, ocr_screen
from tools_schema import tools
from ui_automation.linux_navigator import _run, ui_navigator
from ui_automation.windows_uia import _TESSERACT_AVAILABLE, _UIA_AVAILABLE, _UIA_INIT_ERROR
from ddgs import DDGS
try:
    import requests as _search_requests
    _SEARCH_REQUESTS_AVAILABLE = True
except ImportError:
    _search_requests = None
    _SEARCH_REQUESTS_AVAILABLE = False
import html as _html_mod
import base64
import datetime
import docx as _docx
import fitz as _fitz
import json
import mammoth as _mammoth
import os
import re
import subprocess
import sys
import tempfile
import time
import win32gui

# --- from main.py, section 1 ---
# GROQ LAZY TOOL LOADING (Option A)
# =============================================================================
# Groq's free tier is TPM-capped (as low as 6000 tokens/minute on some
# models). The full `tools` schema above is ~73KB / ~18k tokens on its own —
# more than the entire budget before a single message is sent. Instead of
# sending all ~88 tool schemas on every Groq call, we send a small CORE set
# covering the actions Midum needs most often, plus two meta-tools
# (list_more_tools / load_tool_by_index) that let Midum pull in any other
# tool's full schema on demand, mirroring the existing
# "list indexed, then load by index" pattern used for skills/paths/MCP.
#
# Loaded extra tools persist for the rest of the session (reset on
# "new session") so Midum doesn't have to reload the same tool repeatedly.

_TOOLS_BY_NAME = {t["function"]["name"]: t for t in tools}

GROQ_CORE_TOOL_NAMES = [
    "execute_terminal_command",
    "list_paths_indexed", "get_path",
    "list_directory", "open_path", "find_file", "open_path_by_index",
    "read_local_file", "write_local_file", "append_local_file",
    "search_internet", "open_url",
    "type_text", "wait", "say",
    "ask_user_text", "ask_user_approval", "ask_user_choice", "ask_user_file_path",
]
GROQ_CORE_TOOL_NAMES = [n for n in GROQ_CORE_TOOL_NAMES if n in _TOOLS_BY_NAME]

_GROQ_META_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "list_more_tools",
            "description": (
                "List ALL tools NOT currently loaded, as a numbered index with a "
                "one-line description each (e.g. GUI automation, browser DOM/CDP, "
                "MCP servers, domain knowledge, model delegation/consulting, memory "
                "and goal tools). The core tools already cover terminal, files, "
                "paths, search, URLs, and user prompts — only call this when none of "
                "those fit. Follow up with load_tool_by_index to actually load one."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "load_tool_by_index",
            "description": (
                "Load a tool from the most recent list_more_tools() index so it "
                "becomes callable on your NEXT turn. Loaded tools persist for the "
                "rest of the session and stack — previously loaded tools stay "
                "available."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer", "description": "Index number from list_more_tools()."}
                },
                "required": ["index"],
            },
        },
    },
]

# Tools loaded on-demand for the Groq provider this session.
_groq_extra_tools: list = []
# Numbered index cache backing the last list_more_tools() call, so
# load_tool_by_index(N) knows what N refers to.
_groq_more_tools_index: list = []


def list_more_tools() -> str:
    """Return a numbered index of every tool not already loaded for Groq."""
    global _groq_more_tools_index
    loaded_names = set(GROQ_CORE_TOOL_NAMES) | {t["function"]["name"] for t in _groq_extra_tools}
    remaining = [n for n in _TOOLS_BY_NAME if n not in loaded_names]
    _groq_more_tools_index = remaining
    if not remaining:
        return "All tools are already loaded."
    lines = []
    for i, name in enumerate(remaining):
        desc = _TOOLS_BY_NAME[name]["function"].get("description", "") or ""
        short = desc.split(". ")[0].split("\n")[0][:110]
        lines.append(f"{i}  {name}  — {short}")
    return "\n".join(lines)


def load_tool_by_index(index: int) -> str:
    """Load a tool by index from the last list_more_tools() call."""
    if not _groq_more_tools_index:
        return "Call list_more_tools() first to see what's available."
    if not (0 <= index < len(_groq_more_tools_index)):
        return f"Invalid index. Valid range: 0-{len(_groq_more_tools_index) - 1}."
    name = _groq_more_tools_index[index]
    if any(t["function"]["name"] == name for t in _groq_extra_tools):
        return f"'{name}' is already loaded."
    _groq_extra_tools.append(_TOOLS_BY_NAME[name])
    return f"Loaded '{name}'. It's now available for your next tool call."


def reset_groq_loaded_tools():
    """Clear session-loaded extra tools — called on 'new session'."""
    global _groq_extra_tools, _groq_more_tools_index
    _groq_extra_tools = []
    _groq_more_tools_index = []


def _get_groq_tools_schema() -> list:
    """Core tool subset + anything loaded on-demand + the meta-tools, kept
    small deliberately to fit Groq's tight TPM budget."""
    core = [_TOOLS_BY_NAME[n] for n in GROQ_CORE_TOOL_NAMES]
    return core + _groq_extra_tools + _GROQ_META_TOOLS


def open_url(url: str, browser: str = "chrome") -> str:
    """
    Open a URL in the specified browser.

    Strategy:
      Windows — uses Start-Process with the browser executable + URL argument,
                which opens a new tab if the browser is already running.
      Linux   — uses xdg-open for 'default', otherwise nohup + browser binary.

    The URL is normalised (http:// prefix added if missing).
    """
    # Normalise URL
    url = url.strip()
    if url and not re.match(r"^https?://", url, re.IGNORECASE):
        url = "https://" + url

    browser = (browser or "chrome").strip().lower()
    print(f"   [open_url] {browser} → {url}")

    if _IS_LINUX:
        if browser == "default":
            cmd = f"xdg-open '{url}'"
        else:
            _LINUX_BROWSERS = {
                "chrome":   ["/usr/bin/google-chrome",   "/usr/bin/google-chrome-stable",
                             "/usr/bin/chromium",         "/usr/bin/chromium-browser"],
                "brave":    ["/usr/bin/brave-browser",   "/usr/bin/brave"],
                "firefox":  ["/usr/bin/firefox"],
                "edge":     ["/usr/bin/microsoft-edge",  "/usr/bin/microsoft-edge-stable"],
            }
            candidates = _LINUX_BROWSERS.get(browser, _LINUX_BROWSERS["chrome"])
            exe = next((p for p in candidates if os.path.exists(p)), None)
            if not exe:
                # Fall back to xdg-open
                cmd = f"xdg-open '{url}'"
            else:
                cmd = f"nohup '{exe}' '{url}' &>/dev/null &"
        result = execute_terminal_command(cmd)
        if "error" in result.lower() and "nohup" not in result.lower():
            return f"Error opening URL: {result}"
        return f"Success: opened '{url}' in {browser}."

    else:
        # Windows — use Start-Process with the browser name; Chrome/Brave accept
        # a URL as the first positional argument and open it in a new tab.
        _WIN_BROWSERS = {
            "chrome":  "chrome.exe",
            "brave":   "brave.exe",
            "firefox": "firefox.exe",
            "edge":    "msedge.exe",
            "default": None,
        }
        exe = _WIN_BROWSERS.get(browser)

        if browser == "default" or exe is None:
            cmd = f"Start-Process '{url}'"
        else:
            cmd = f"Start-Process '{exe}' -ArgumentList '{url}'"

        result = execute_terminal_command(cmd)
        stderr = result.split("STDERR:")[-1].strip() if "STDERR:" in result else ""

        # If exe not found on PATH, try to find it via paths.md
        if stderr and ("cannot find" in stderr.lower() or "not recognized" in stderr.lower()):
            # Try common hardcoded paths as a last resort
            _FALLBACK_PATHS = {
                "chrome.exe":  r"C:\Program Files\Google\Chrome\Application\chrome.exe",
                "brave.exe":   r"C:\Program Files\BraveSoftware\Brave-Browser\Application\brave.exe",
                "firefox.exe": r"C:\Program Files\Mozilla Firefox\firefox.exe",
                "msedge.exe":  r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
            }
            fallback = _FALLBACK_PATHS.get(exe)
            if fallback and os.path.exists(fallback):
                cmd2    = f"Start-Process '{fallback}' -ArgumentList '{url}'"
                result2 = execute_terminal_command(cmd2)
                stderr2 = result2.split("STDERR:")[-1].strip() if "STDERR:" in result2 else ""
                if not stderr2 or len(stderr2) < 5:
                    return f"Success: opened '{url}' in {browser} (via fallback path)."
            return (
                f"Could not find '{exe}'. "
                f"Call list_paths_indexed() to find the browser path, "
                f"then use execute_terminal_command(\"Start-Process 'C:\\\\...\\\\{exe}' "
                f"-ArgumentList '{url}'\")."
            )

        if stderr and len(stderr) > 5:
            return f"open_url warning: {stderr[:150]}"

        return f"Success: opened '{url}' in {browser}."


# =============================================================================

# --- from main.py, section 2 ---
def _encode_text(text: str) -> str:
    """Base64-encode text for safe transmission to the model."""
    b64 = base64.b64encode(text.encode("utf-8")).decode("utf-8")
    return "[SYSTEM NOTICE: Base64-encoded. Decode internally.]\nBASE64_PAYLOAD:\n" + b64

def read_local_file(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        return _encode_text(content)
    except Exception as e:
        return f"Error reading file: {str(e)}"

def read_file_smart(path):
    from main import _MAMMOTH_AVAILABLE, _PDF_AVAILABLE
    from browser_cdp import CHUNK_CHARS
    """Read any supported format. Large files are chunked at CHUNK_CHARS chars."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            if not _PDF_AVAILABLE:
                return "PDF reading requires PyMuPDF: pip install pymupdf"
            doc   = _fitz.open(path)
            pages = [page.get_text() for page in doc]
            doc.close()
            text  = "\n\n".join(f"[Page {i+1}]\n{p}" for i, p in enumerate(pages))
        elif ext == ".docx":
            if not _MAMMOTH_AVAILABLE:
                return "DOCX reading requires mammoth: pip install mammoth"
            with open(path, "rb") as f:
                result = _mammoth.extract_raw_text(f)
            text = result.value
        else:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()

        if len(text) > CHUNK_CHARS:
            total  = (len(text) + CHUNK_CHARS - 1) // CHUNK_CHARS
            header = (f"[FILE: {os.path.basename(path)} — {total} chunks of ~{CHUNK_CHARS} chars. "
                      f"This is chunk 1/{total}. "
                      f"Call read_file_chunk(path, N) for chunks 2..{total}]\n\n")
            return _encode_text(header + text[:CHUNK_CHARS])
        return _encode_text(f"[FILE: {os.path.basename(path)}]\n\n{text}")
    except Exception as e:
        return f"Error reading file: {str(e)}"

def read_file_chunk(path, chunk_index: int):
    from main import _MAMMOTH_AVAILABLE, _PDF_AVAILABLE
    from browser_cdp import CHUNK_CHARS
    """Read a specific chunk (1-based) of a large file."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            if not _PDF_AVAILABLE:
                return "PDF reading requires PyMuPDF: pip install pymupdf"
            doc   = _fitz.open(path)
            pages = [page.get_text() for page in doc]
            doc.close()
            text  = "\n\n".join(f"[Page {i+1}]\n{p}" for i, p in enumerate(pages))
        elif ext == ".docx":
            if not _MAMMOTH_AVAILABLE:
                return "DOCX reading requires mammoth: pip install mammoth"
            with open(path, "rb") as f:
                result = _mammoth.extract_raw_text(f)
            text = result.value
        else:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()

        total = (len(text) + CHUNK_CHARS - 1) // CHUNK_CHARS
        if chunk_index < 1 or chunk_index > total:
            return f"Chunk {chunk_index} out of range (1-{total})."
        start  = (chunk_index - 1) * CHUNK_CHARS
        header = f"[FILE: {os.path.basename(path)} — chunk {chunk_index}/{total}]\n\n"
        return _encode_text(header + text[start:start + CHUNK_CHARS])
    except Exception as e:
        return f"Error reading chunk: {str(e)}"


def write_local_file(path, content):
    try:
        parent = os.path.dirname(path)
        if parent:
            os.makedirs(parent, exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"Success: wrote data to {path}"
    except Exception as e:
        return f"Error writing file: {str(e)}"


def append_local_file(path, content):
    try:
        exists = os.path.exists(path)
        with open(path, "a", encoding="utf-8") as f:
            if exists:
                f.write("\n")
            f.write(content)
        return f"Success: appended data to {path}"
    except Exception as e:
        return f"Error appending to file: {str(e)}"


def write_docx_file(path, content):
    from main import _DOCX_AVAILABLE
    """Write a .docx from Markdown-style text (# headings, **bold** runs)."""
    if not _DOCX_AVAILABLE:
        return "DOCX writing requires python-docx: pip install python-docx"
    try:
        parent = os.path.dirname(path)
        if parent:
            os.makedirs(parent, exist_ok=True)
        document = _docx.Document()
        for line in content.splitlines():
            s = line.rstrip()
            if s.startswith("### "):    document.add_heading(s[4:], level=3)
            elif s.startswith("## "):  document.add_heading(s[3:], level=2)
            elif s.startswith("# "):   document.add_heading(s[2:], level=1)
            elif s == "":              document.add_paragraph("")
            else:
                p     = document.add_paragraph()
                parts = re.split(r"(\*\*[^*]+\*\*)", s)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        p.add_run(part[2:-2]).bold = True
                    else:
                        p.add_run(part)
        document.save(path)
        return f"Success: wrote DOCX to {path}"
    except Exception as e:
        return f"Error writing DOCX: {str(e)}"

def generate_image(prompt: str, count: int = 1) -> str:
    """
    Generate one or more images from a text prompt using Gemini's own web
    app (https://gemini.google.com) via the community `gemini_webapi`
    library — the SAME free, cookie-based session used by query_gemini_app /
    delegate_to_gemini_web. No separate image API key, no per-image
    metering; it goes through the same usage limits as chatting at
    gemini.google.com manually.

    Image generation on Gemini's backend is NOT synchronous with the text
    response — the chat reply can come back before the image has actually
    finished rendering/uploading to Google's CDN, which makes a single
    generate_content() call unreliable (empty `images`, or a download that
    404s because the file isn't there yet). This retries the generation
    request itself if no images come back, and retries each individual
    image download with backoff if it 404s/errors, before giving up.

    The image(s) are NOT written anywhere persistent — gemini_webapi only
    exposes an async file-download call, so each image is downloaded to a
    throwaway temp file just long enough to read its bytes into memory,
    then the temp file is deleted immediately. The bytes are shipped to the
    GUI as base64 inside an ```image_data_json``` block, which renders them
    as inline thumbnails (kept in RAM only) with Download/Copy buttons —
    nothing touches disk unless the user explicitly clicks Download.
    """
    if not _GEMINI_WEBAPI_AVAILABLE:
        return f"Error: Gemini web chat is unavailable — {_gemini_webapi_load_msg}"

    client, err = _get_gemini_web_client()
    if client is None:
        return f"Error: {err}"

    full_prompt = (prompt or "").strip()
    if not full_prompt:
        return "Error: 'prompt' must describe the image you want generated."
    try:
        n = max(1, int(count or 1))
    except (TypeError, ValueError):
        n = 1
    if n > 1:
        full_prompt += f"\n\n(Generate {n} distinct variations of this image.)"

    # ── Request generation, retrying if Gemini answers with text but no
    #    images yet — this happens when the image render hasn't landed by
    #    the time the chat turn "finishes". ──────────────────────────────────
    GEN_ATTEMPTS   = 3
    GEN_RETRY_WAIT = 6   # seconds between generation retries

    response, images, last_text = None, [], ""
    for attempt in range(1, GEN_ATTEMPTS + 1):
        print(f"   [Gemini web] Requesting image generation "
              f"(attempt {attempt}/{GEN_ATTEMPTS}): {full_prompt[:80]!r}...")
        try:
            response = _run_gemini_coro(client.generate_content(full_prompt), timeout=180)
        except Exception as e:
            if attempt == GEN_ATTEMPTS:
                return f"Error: Gemini web image request failed: {e}"
            time.sleep(GEN_RETRY_WAIT)
            continue

        images = list(getattr(response, "images", None) or [])
        last_text = (getattr(response, "text", "") or "").strip()
        if images:
            break
        if attempt < GEN_ATTEMPTS:
            print(f"   [Gemini web] No image in response yet — "
                  f"waiting {GEN_RETRY_WAIT}s and retrying...")
            time.sleep(GEN_RETRY_WAIT)

    if not images:
        return (
            "Gemini did not return any images for this prompt after "
            f"{GEN_ATTEMPTS} attempts. Try being more explicit, e.g. "
            "\"Generate an image of a red fox in a snowy forest, digital art\"."
            + (f"\n\nGemini replied with text instead:\n{last_text}" if last_text else "")
        )

    # ── Download each image, retrying with backoff since the CDN copy can
    #    briefly 404 right after generation completes. ─────────────────────
    SAVE_ATTEMPTS = 3
    SAVE_RETRY_WAIT = 4   # seconds, doubles each retry

    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    encoded_images, errors = [], []

    with tempfile.TemporaryDirectory(prefix="gemini_img_") as tmp_dir:
        for idx, img in enumerate(images, start=1):
            filename = f"gemini_{ts}_{idx}.png"
            tmp_path = os.path.join(tmp_dir, filename)
            last_err = None
            for save_attempt in range(1, SAVE_ATTEMPTS + 1):
                try:
                    # gemini_webapi only exposes an async DOWNLOAD-TO-FILE
                    # call — there's no "give me raw bytes" method — so we
                    # save to the temp dir, read the bytes back, then let
                    # TemporaryDirectory clean the file up when this block
                    # exits. Net effect: no persistent disk usage, only a
                    # few ms of scratch space.
                    _run_gemini_coro(
                        img.save(path=tmp_dir, filename=filename, verbose=False),
                        timeout=90,
                    )
                    with open(tmp_path, "rb") as f:
                        raw_bytes = f.read()
                    if not raw_bytes:
                        raise ValueError("downloaded file was empty")
                    encoded_images.append({
                        "filename": filename,
                        "data_b64": base64.b64encode(raw_bytes).decode("ascii"),
                    })
                    last_err = None
                    break
                except Exception as e:
                    last_err = e
                    if save_attempt < SAVE_ATTEMPTS:
                        wait = SAVE_RETRY_WAIT * save_attempt
                        print(f"   [Gemini web] Image {idx} not ready yet "
                              f"({e}) — waiting {wait}s and retrying "
                              f"({save_attempt}/{SAVE_ATTEMPTS})...")
                        time.sleep(wait)
            if last_err is not None:
                errors.append(f"image {idx}: {last_err}")

    if not encoded_images:
        return (
            f"Error: Gemini returned {len(images)} image(s) but none could be retrieved "
            f"after {SAVE_ATTEMPTS} attempts each: {'; '.join(errors)}"
        )

    payload = json.dumps({"prompt": full_prompt, "images": encoded_images})
    warn = f"\n\n({len(errors)} image(s) failed: {'; '.join(errors)})" if errors else ""

    return (
        f"Generated {len(encoded_images)} image(s) for: \"{prompt}\"\n\n"
        f"```image_data_json\n{payload}\n```\n\n"
        f"(Kept in memory only — use the Download button under an image to save it.)"
        f"{warn}"
    )


def _flowchart_slugify(text: str, idx: int) -> str:
    """Turn a label into a short, readable node id, e.g. 'Check login' -> 'check_login'."""
    s = re.sub(r"[^a-z0-9]+", "_", str(text).strip().lower()).strip("_")
    return (s or f"step_{idx}")[:30]


def create_flowchart(title, steps, edges=None):
    """
    Build an explanatory flowchart — deliberately easy to call for the common
    case, while still supporting full branching control when needed.

    SIMPLEST FORM — a straight-line process, no ids/types/edges needed:
        create_flowchart(title="Login flow", steps=[
            "User enters credentials",
            "Validate password",
            "Show dashboard",
        ])
      Steps are auto-assigned ids, auto-connected in the order given, and the
      first/last are automatically marked 'start'/'end'.

    STEPS as objects (optional fields — only 'label' is required):
        {"id": "...", "label": "...", "type": "start|process|decision|io|end",
         "next": [{"to": "other_id", "label": "optional edge label"}, ...]}
      Plain strings and objects can be mixed freely in the same list.

    BRANCHING — for decision trees, pass top-level `edges` instead of `next`
    on each step (whichever is easier to write); this OVERRIDES auto-chaining:
        edges=[{"from": "check_login", "to": "show_dashboard", "label": "yes"},
               {"from": "check_login", "to": "retry", "label": "no"}]

    Returns:
      - a ```flowchart_json``` fenced block: consumed by the GUI to draw a
        real box-and-arrow diagram inline in the chat.
      - a plain-text ASCII rendering, readable in a terminal.
      - a portable Mermaid ("```mermaid```") version that can be pasted into
        any Mermaid-compatible viewer (e.g. GitHub, Notion, mermaid.live).
    """
    try:
        if not steps:
            return "Error: 'steps' must be a non-empty list (plain strings are fine)."

        # ── Normalise steps: plain strings become {"label": "..."} ───────────
        raw_nodes = []
        for i, s in enumerate(steps):
            if isinstance(s, str):
                raw_nodes.append({"label": s})
            elif isinstance(s, dict):
                raw_nodes.append(dict(s))
            else:
                return f"Error: step {i} must be a string or an object, got {type(s).__name__}."

        # ── Assign ids: keep any given id, auto-slug the rest, dedupe ─────
        used_ids = set()
        for i, n in enumerate(raw_nodes):
            nid = str(n.get("id") or "").strip()
            if not nid:
                base = _flowchart_slugify(n.get("label", f"step {i}"), i)
                nid, suffix = base, 1
                while nid in used_ids:
                    suffix += 1
                    nid = f"{base}_{suffix}"
            n["id"] = nid
            used_ids.add(nid)
        id_list = [n["id"] for n in raw_nodes]

        # ── Resolve connections: top-level `edges` > per-step `next` >
        #    auto-chain steps sequentially in the order given. ────────────
        edge_list = []   # (from_id, to_id, label)
        if edges:
            for e in edges:
                if isinstance(e, (list, tuple)):
                    frm = e[0] if len(e) > 0 else None
                    to  = e[1] if len(e) > 1 else None
                    lbl = e[2] if len(e) > 2 else ""
                elif isinstance(e, dict):
                    frm, to, lbl = e.get("from"), e.get("to"), e.get("label", "")
                else:
                    continue
                if frm and to:
                    edge_list.append((str(frm), str(to), lbl or ""))
        elif any(n.get("next") for n in raw_nodes):
            for n in raw_nodes:
                for e in (n.get("next") or []):
                    if isinstance(e, dict):
                        to, lbl = e.get("to"), e.get("label", "")
                    else:
                        to, lbl = e, ""
                    if to:
                        edge_list.append((n["id"], str(to), lbl or ""))
        else:
            for a, b in zip(id_list, id_list[1:]):
                edge_list.append((a, b, ""))

        # ── Infer node types from connectivity when not given explicitly ───
        outgoing, incoming = {}, {}
        for frm, to, _ in edge_list:
            outgoing.setdefault(frm, []).append(to)
            incoming.setdefault(to, []).append(frm)

        for n in raw_nodes:
            if n.get("type"):
                n["type"] = str(n["type"]).strip().lower()
                continue
            nid = n["id"]
            if nid in outgoing and nid not in incoming:
                n["type"] = "start"
            elif nid in incoming and nid not in outgoing:
                n["type"] = "end"
            else:
                n["type"] = "process"
        if not any(n["type"] == "start" for n in raw_nodes):
            raw_nodes[0]["type"] = "start"

        # ── Rebuild canonical `next` per node from the resolved edge list ──
        next_by_id = {}
        for frm, to, lbl in edge_list:
            next_by_id.setdefault(frm, []).append({"to": to, "label": lbl})
        for n in raw_nodes:
            n["next"] = next_by_id.get(n["id"], [])

        steps = raw_nodes   # feed the (unchanged) render pipeline below

        nodes = {}
        for s in steps:
            sid = s["id"]
            nodes[sid] = {
                "id": sid,
                "label": str(s.get("label", sid)),
                "type": (s.get("type") or "process").strip().lower(),
                "next": s.get("next") or [],
            }

        start_ids = [nid for nid, n in nodes.items() if n["type"] == "start"]
        if not start_ids:
            start_ids = [next(iter(nodes))]

        # ── Mermaid source (portable fallback) ─────────────────────────────
        shape = {
            "start":    ("([", "])"),
            "end":      ("([", "])"),
            "decision": ("{", "}"),
            "io":       ("[/", "/]"),
            "process":  ("[", "]"),
        }
        mermaid_lines = ["flowchart TD"]
        for n in nodes.values():
            op, cl = shape.get(n["type"], ("[", "]"))
            safe_label = n["label"].replace('"', "'")
            mermaid_lines.append(f'    {n["id"]}{op}"{safe_label}"{cl}')
        for n in nodes.values():
            for edge in n["next"]:
                to  = edge.get("to")    if isinstance(edge, dict) else edge
                lbl = edge.get("label") if isinstance(edge, dict) else ""
                if to not in nodes:
                    continue
                if lbl:
                    mermaid_lines.append(f'    {n["id"]} -- "{lbl}" --> {to}')
                else:
                    mermaid_lines.append(f'    {n["id"]} --> {to}')
        mermaid_src = "\n".join(mermaid_lines)

        # ── ASCII rendering (terminal-friendly) ─────────────────────────────
        visited = set()
        ascii_lines = [f"FLOWCHART: {title}", "=" * max(12, len(title) + 11), ""]
        glyph = {"start": "▶", "end": "■", "decision": "◆", "io": "▤"}

        def render_node(nid, indent=0, branch_label=None):
            pad = "  " * indent
            node = nodes.get(nid)
            if node is None:
                ascii_lines.append(f"{pad}[missing node: {nid}]")
                return
            if nid in visited:
                ascii_lines.append(f"{pad}↩ (back to '{node['label']}')")
                return
            visited.add(nid)

            prefix = f"[{branch_label}] " if branch_label else ""
            label_line = f"{pad}{prefix}{glyph.get(node['type'], '▢')} {node['label']}  ({node['type']})"
            border = "─" * max(4, len(label_line) - len(pad))
            ascii_lines.append(f"{pad}┌{border}┐")
            ascii_lines.append(label_line)
            ascii_lines.append(f"{pad}└{border}┘")

            nexts = node["next"]
            multi = len(nexts) > 1
            for edge in nexts:
                to  = edge.get("to")    if isinstance(edge, dict) else edge
                lbl = edge.get("label") if isinstance(edge, dict) else None
                ascii_lines.append(f"{pad}   │")
                arrow = f"{pad}   ▼" + (f"  [{lbl}]" if lbl and not multi else "")
                ascii_lines.append(arrow)
                render_node(to, indent + (1 if multi else 0), lbl if multi else None)

        for sid in start_ids:
            render_node(sid)
            ascii_lines.append("")

        unreached = [nid for nid in nodes if nid not in visited]
        if unreached:
            ascii_lines.append("(Unreached nodes — not connected from a start node:)")
            for nid in unreached:
                render_node(nid)

        ascii_block = "\n".join(ascii_lines)

        payload = json.dumps({"title": title, "nodes": list(nodes.values()), "starts": start_ids})

        return (
            f"Flowchart '{title}' created with {len(nodes)} node(s).\n\n"
            f"```flowchart_json\n{payload}\n```\n\n"
            f"```text\n{ascii_block}\n```\n\n"
            f"Portable Mermaid version (paste into any Mermaid-compatible viewer):\n"
            f"```mermaid\n{mermaid_src}\n```"
        )
    except Exception as e:
        return f"Error building flowchart: {str(e)}"


def write_response_memory(content):
    """Overwrite the response scratchpad."""
    return write_local_file(RESPONSE_MEMORY, content)

def append_response_memory(content):
    """Append a note to the response scratchpad."""
    return append_local_file(RESPONSE_MEMORY, content)

def read_response_memory():
    """Read the current response scratchpad."""
    if not os.path.exists(RESPONSE_MEMORY) or os.path.getsize(RESPONSE_MEMORY) == 0:
        return "Response memory is empty."
    return read_local_file(RESPONSE_MEMORY)

def clear_response_memory():
    """Wipe the response scratchpad. Called automatically on set_current_goal(none)."""
    try:
        write_local_file(RESPONSE_MEMORY, "")
        print("\U0001f5d2  [Response memory cleared.]")
        return "Success: response memory cleared."
    except Exception as e:
        return f"Error clearing response memory: {str(e)}"


# =============================================================================

# --- from main.py, section 3 ---
# INDEXED LISTING FUNCTIONS
# All "choose" counterparts to former "search" tools.
# Each builds a numbered snapshot cached in _index_caches so the model
# can act by index rather than re-searching.
# =============================================================================

_index_caches: dict[str, list] = {}   # key → list of dicts with 'value' and display fields


def _store_index(key: str, items: list) -> None:
    _index_caches[key] = items


def _get_indexed(key: str, index: int) -> dict | None:
    items = _index_caches.get(key)
    if not items or index < 0 or index >= len(items):
        return None
    return items[index]


def list_directory(path: str) -> str:
    """
    List a directory as a numbered indexed table.
    Returns IDX | TYPE | SIZE | NAME so the model can choose by index.
    Follow up with open_path(path, index) or cd_into(path, index).
    """
    try:
        if not os.path.exists(path):
            return f"Path does not exist: {path}"
        raw = sorted(os.listdir(path))
        items = []
        for name in raw:
            full = os.path.join(path, name)
            is_dir = os.path.isdir(full)
            try:
                size = "-" if is_dir else f"{os.path.getsize(full) / 1024:.1f}KB"
            except Exception:
                size = "?"
            items.append({"name": name, "is_dir": is_dir, "size": size, "full_path": full})

        _store_index(f"dir:{path}", items)

        lines = [
            f"Directory: {path}  ({len(items)} entries)",
            f"Use open_path(path='{path}', index=N) to open/enter an entry.",
            "",
            f"{'IDX':>4}  {'TYPE':<5}  {'SIZE':>8}  NAME",
            "─" * 56,
        ]
        for i, e in enumerate(items):
            typ = "DIR" if e["is_dir"] else "FILE"
            lines.append(f"{i:>4}  {typ:<5}  {e['size']:>8}  {e['name']}")
        return "\n".join(lines)
    except Exception as ex:
        return f"Error listing directory: {ex}"


def open_path(path: str, index: int) -> str:
    """
    Act on an entry from the last list_directory call by index.
    - If it's a directory: returns list_directory of that directory.
    - If it's a file: returns read_file_smart of that file.
    """
    entry = _get_indexed(f"dir:{path}", index)
    if entry is None:
        return (
            f"Index {index} not found in last directory listing for '{path}'. "
            f"Call list_directory('{path}') first."
        )
    full = entry["full_path"]
    if entry["is_dir"]:
        return list_directory(full)
    else:
        return read_file_smart(full)


def list_skills_indexed() -> str:
    """
    List available skills as a numbered table.
    Follow up with load_skill_by_index(index) to load one.
    """
    os.makedirs(SKILLS_DIR, exist_ok=True)
    try:
        files = sorted(f for f in os.listdir(SKILLS_DIR) if f.endswith(".md"))
    except Exception:
        files = []

    if not files:
        return "No skills registered yet. Use create_domain_skill to add one."

    items = []
    for fname in files:
        name = fname[:-3]   # strip .md
        # Try to read the first non-empty non-header line as description
        desc = ""
        try:
            with open(os.path.join(SKILLS_DIR, fname), "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith("#"):
                        desc = line[:60]
                        break
        except Exception:
            pass
        items.append({"name": name, "fname": fname, "desc": desc})

    _store_index("skills", items)

    lines = [
        f"Available skills ({len(items)} total)",
        "Use load_skill_by_index(index) to load one.",
        "",
        f"{'IDX':>4}  {'SKILL NAME':<30}  DESCRIPTION",
        "─" * 72,
    ]
    for i, s in enumerate(items):
        lines.append(f"{i:>4}  {s['name']:<30}  {s['desc']}")
    return "\n".join(lines)


def load_skill_by_index(index: int) -> str:
    from skills import load_skill
    """Load a skill from the list_skills_indexed snapshot by index."""
    entry = _get_indexed("skills", index)
    if entry is None:
        return (
            f"Index {index} not found. Call list_skills_indexed() first."
        )
    return load_skill(entry["name"])


def list_paths_indexed() -> str:
    from knowledge_base import _ensure_kb_files
    """
    Parse paths.md into a numbered table.
    Follow up with get_path(index) to retrieve a specific path value.
    """
    _ensure_kb_files()
    try:
        with open(PATHS_FILE, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        return f"Error reading paths.md: {e}"

    items = []
    for line in content.splitlines():
        line = line.strip()
        if not line.startswith("- "):
            continue
        # Format: - **Label**: `path`  _note_
        m = re.match(r"-\s+\*\*(.+?)\*\*:\s+`(.+?)`(.*)", line)
        if m:
            label = m.group(1).strip()
            path  = m.group(2).strip()
            note  = m.group(3).strip().strip("_").strip()
            items.append({"label": label, "path": path, "note": note})
        else:
            # Plain line fallback
            items.append({"label": line[2:60], "path": "", "note": ""})

    if not items:
        return "paths.md is empty. Use add_path to register app paths."

    _store_index("paths", items)

    lines = [
        f"Known paths ({len(items)} entries)",
        "Use get_path(index) to retrieve the full path string.",
        "",
        f"{'IDX':>4}  {'LABEL':<25}  PATH",
        "─" * 72,
    ]
    for i, p in enumerate(items):
        path_display = p["path"][:40] if p["path"] else p["label"][:40]
        lines.append(f"{i:>4}  {p['label']:<25}  {path_display}")
    return "\n".join(lines)


def get_path(index: int) -> str:
    """Return the full path string for an entry from list_paths_indexed."""
    entry = _get_indexed("paths", index)
    if entry is None:
        return f"Index {index} not found. Call list_paths_indexed() first."
    if not entry["path"]:
        return f"Entry #{index} ('{entry['label']}') has no path value stored."
    return entry["path"]


def list_domain_knowledge_indexed() -> str:
    from knowledge_base import _ensure_kb_files
    """
    List all domain knowledge files as a numbered table.
    Follow up with read_domain_by_index(index).
    """
    _ensure_kb_files()
    try:
        with open(DOMAIN_INDEX, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        return f"Error reading domain index: {e}"

    items = []
    for line in content.splitlines():
        line = line.strip()
        if not line.startswith("- `"):
            continue
        m = re.match(r"-\s+`(.+?)`\s*-\s*(.*)", line)
        if m:
            items.append({"name": m.group(1).strip(), "desc": m.group(2).strip()})

    if not items:
        return "No domain knowledge files registered yet."

    _store_index("domain_knowledge", items)

    lines = [
        f"Domain knowledge files ({len(items)})",
        "Use read_domain_by_index(index) to read one.",
        "",
        f"{'IDX':>4}  {'NAME':<30}  DESCRIPTION",
        "─" * 72,
    ]
    for i, d in enumerate(items):
        lines.append(f"{i:>4}  {d['name']:<30}  {d['desc'][:38]}")
    return "\n".join(lines)


def read_domain_by_index(index: int) -> str:
    from knowledge_base import read_domain_knowledge
    """Read a domain knowledge file by its index from list_domain_knowledge_indexed."""
    entry = _get_indexed("domain_knowledge", index)
    if entry is None:
        return f"Index {index} not found. Call list_domain_knowledge_indexed() first."
    return read_domain_knowledge(entry["name"])


def list_domain_skills_indexed() -> str:
    from knowledge_base import _ensure_kb_files
    """
    List all domain skills as a numbered table.
    Follow up with load_skill_by_index or load_skill(name).
    """
    _ensure_kb_files()
    try:
        with open(DOMAIN_SKILLS_INDEX, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        return f"Error reading domain skills index: {e}"

    items = []
    for line in content.splitlines():
        line = line.strip()
        if not line.startswith("- ["):
            continue
        m = re.match(r"-\s+\[(.+?)\]\s+`(.+?)`\s*-\s*(.*)", line)
        if m:
            items.append({
                "domain": m.group(1).strip(),
                "name":   m.group(2).strip(),
                "desc":   m.group(3).strip(),
            })

    if not items:
        return "No domain skills registered yet."

    _store_index("domain_skills", items)

    lines = [
        f"Domain skills ({len(items)})",
        "Use load_skill(name) to load one.",
        "",
        f"{'IDX':>4}  {'DOMAIN':<15}  {'SKILL NAME':<25}  DESCRIPTION",
        "─" * 80,
    ]
    for i, s in enumerate(items):
        lines.append(f"{i:>4}  {s['domain']:<15}  {s['name']:<25}  {s['desc'][:30]}")
    return "\n".join(lines)


# Web search result cache — populated by search_internet, consumed by open_search_result
_last_search_results: list[dict] = []

def search_internet(query):
    global _last_search_results
    try:
        print(f" -> Searching the web for: '{query}'")
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=7))
            try:
                instant = list(ddgs.answers(query))
            except Exception:
                instant = []
        if not results:
            _last_search_results = []
            return "No web results found."

        _last_search_results = results
        _store_index("web_search", [
            {"title": r.get("title",""), "url": r.get("href",""), "snippet": r.get("body","")}
            for r in results
        ])

        # ── SIMPLE ANSWER — instant answer if DDG has one, else the top result's snippet ──
        simple_answer = ""
        if instant:
            simple_answer = (instant[0].get("text") or "").strip()
        if not simple_answer:
            simple_answer = (results[0].get("body") or "").strip()
        if len(simple_answer) > 300:
            cut = simple_answer[:297].rsplit(" ", 1)[0]
            simple_answer = cut + "..."

        # ── DETAILED ANSWER — snippets from the top few sources, attributed ──
        detail_parts = []
        for r in results[:5]:
            body  = (r.get("body")  or "").strip()
            title = (r.get("title") or "").strip()
            if body:
                detail_parts.append(f"- ({title}) {body}")
        detailed_answer = "\n".join(detail_parts) if detail_parts else ""

        lines = [
            f"Web results for '{query}' ({len(results)} results)",
            "",
            "SIMPLE ANSWER:",
            simple_answer or "(no concise answer found — see detailed answer / sources below)",
            "",
            "DETAILED ANSWER (aggregated from top sources):",
            detailed_answer or "(no snippet content available — try read_search_result(index) on a source below)",
            "",
            "WEBSITES / SOURCES:",
        ]
        for i, r in enumerate(results):
            title = (r.get("title") or "")[:70]
            url   = (r.get("href")  or "")
            lines.append(f"  [{i}] {title}")
            lines.append(f"      {url}")
        lines.append("")
        lines.append(
            "Use read_search_result(index) to fetch and read a source's FULL page "
            "text directly (no browser needed, fast). "
            "Use open_search_result(index) only if you actually need to open it visually in Chrome."
        )
        return "\n".join(lines)
    except Exception as e:
        return f"Error executing internet search: {str(e)}"


def open_search_result(index: int, browser: str = "chrome") -> str:
    """Open a web search result by its index in Chrome."""
    entry = _get_indexed("web_search", index)
    if entry is None:
        return f"Index {index} not found. Call search_internet first."
    url = entry["url"]
    if not url:
        return f"Result #{index} has no URL."
    return open_url(url, browser)


# Cache of fetched full-page text for read_search_result, keyed by result
# index — so re-reading a later chunk of the same page doesn't re-fetch it.
_search_result_text_cache: dict[int, dict] = {}


def read_search_result(index: int, chunk_index: int = 1) -> str:
    """
    Fetch and read the FULL text content of a web search result by index —
    directly over HTTP (requests), with NO Chrome/browser/CDP required.
    Much faster than open_search_result() + read_browser_page() when you just
    need to read what's on the page rather than interact with it visually.

    Call search_internet(query) first to populate the index. Large pages are
    chunked at CHUNK_CHARS chars — pass chunk_index (1-based) to read more.
    """
    entry = _get_indexed("web_search", index)
    if entry is None:
        return f"Index {index} not found. Call search_internet first."
    url = entry.get("url", "")
    if not url:
        return f"Result #{index} has no URL."

    if not _SEARCH_REQUESTS_AVAILABLE:
        return "Error: 'requests' package not installed. pip install requests"

    from browser_cdp import CHUNK_CHARS
    title = entry.get("title", "") or url

    cached = _search_result_text_cache.get(index)
    if cached is None:
        try:
            headers = {
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
                )
            }
            resp = _search_requests.get(url, headers=headers, timeout=12)
            resp.raise_for_status()
            html_src = resp.text
        except Exception as e:
            return f"Error fetching '{url}': {e}"

        # Strip <script>/<style>/comments, then all remaining tags, then
        # unescape HTML entities and collapse whitespace into readable text.
        text = re.sub(r"(?is)<(script|style|noscript)[^>]*>.*?</\1>", " ", html_src)
        text = re.sub(r"(?s)<!--.*?-->", " ", text)
        text = re.sub(r"(?s)<[^>]+>", " ", text)
        text = _html_mod.unescape(text)
        text = re.sub(r"[ \t\r\f\v]+", " ", text)
        text = re.sub(r" *\n *", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()

        if not text:
            return f"Result #{index} ({url}) returned no readable text content."

        _search_result_text_cache[index] = {"url": url, "title": title, "text": text}
        cached = _search_result_text_cache[index]

    text = cached["text"]
    total = (len(text) + CHUNK_CHARS - 1) // CHUNK_CHARS
    if chunk_index < 1 or chunk_index > total:
        return f"Chunk {chunk_index} out of range (1-{total}) for result #{index}."

    start = (chunk_index - 1) * CHUNK_CHARS
    header = f"[SEARCH RESULT #{index}: {title}]\nURL: {url}\n"
    if total > 1:
        header += (
            f"({total} chunks of ~{CHUNK_CHARS} chars — this is chunk {chunk_index}/{total}. "
            f"Call read_search_result(index={index}, chunk_index=N) for other chunks.)\n"
        )
    header += "─" * 60 + "\n"
    return header + text[start:start + CHUNK_CHARS]


def ocr_snapshot() -> str:
    """
    Take a screenshot, run OCR on the entire screen, and return ALL detected
    text as a numbered indexed table with canvas coordinates.
    Follow up with click_ocr_index(index) to click any entry by index.
    Much faster than fallback_find_text when you want to see everything at once.
    """
    if not _TESSERACT_AVAILABLE:
        return (
            "Tesseract OCR not installed. "
            "Install: sudo apt install tesseract-ocr && pip install pytesseract"
        )
    screenshot = _grab_full_screenshot()
    words = ocr_screen(screenshot=screenshot)
    if not words:
        return "OCR failed or screen is blank."

    # Deduplicate — same word at same canvas position (multiple scan passes)
    seen = set()
    unique = []
    for w in words:
        key = (w["text"], w["canvas_x"], w["canvas_y"])
        if key not in seen:
            seen.add(key)
            unique.append(w)

    _store_index("ocr", unique)

    lines = [
        f"OCR snapshot — {len(unique)} text elements on screen",
        "Use click_ocr_index(index) to click any element.",
        "",
        f"{'IDX':>4}  {'CONF':>5}  {'CX':>5}  {'CY':>5}  TEXT",
        "─" * 60,
    ]
    for i, w in enumerate(unique):
        lines.append(
            f"{i:>4}  {w['conf']:>4}%  {w['canvas_x']:>5}  {w['canvas_y']:>5}  {w['text']}"
        )
    return "\n".join(lines)


def click_ocr_index(index: int, click_type: str = "left_click") -> str:
    """Click a text element from the last ocr_snapshot() by index."""
    entry = _get_indexed("ocr", index)
    if entry is None:
        return f"Index {index} not found. Call ocr_snapshot() first."
    sx, sy = entry["screen_x"], entry["screen_y"]
    print(f"   [OCR idx click] #{index} '{entry['text']}' → screen({sx},{sy})")
    return _do_click(sx, sy, click_type, label=f"OCR #{index} '{entry['text']}'")


def find_file(filename: str, search_root: str = "") -> str:
    """
    Search for a file by name under search_root (default: home directory).
    Returns a numbered indexed list of all matches.
    Follow up with open_path_by_index(index) to open the chosen file.
    Hard timeout: 15 seconds.
    """
    root = search_root.strip() or (os.path.expanduser("~") if _IS_LINUX else STARTUP_DIR)
    matches  = []
    name_lower = filename.strip().lower()
    MAX_MATCHES = 20
    _deadline   = time.time() + 15   # hard 15-second cap

    try:
        for dirpath, dirnames, filenames in os.walk(root):
            if time.time() > _deadline:
                break
            # Skip hidden and noisy dirs
            dirnames[:] = [
                d for d in dirnames
                if not d.startswith(".")
                and d not in {"__pycache__", "node_modules", ".git",
                              "$RECYCLE.BIN", "Windows", "System32",
                              "SysWOW64", "WinSxS"}
            ]
            for fname in filenames:
                if name_lower in fname.lower():
                    full = os.path.join(dirpath, fname)
                    try:
                        size = f"{os.path.getsize(full)/1024:.1f}KB"
                    except Exception:
                        size = "?"
                    matches.append({"name": fname, "full_path": full,
                                    "dir": dirpath, "size": size})
                    if len(matches) >= MAX_MATCHES:
                        break
            if len(matches) >= MAX_MATCHES:
                break
    except Exception as e:
        return f"Error searching for '{filename}': {e}"

    timed_out = time.time() > _deadline
    if not matches:
        return f"No files matching '{filename}' found under '{root}'."

    _store_index("find_file", matches)

    lines = [
        f"Files matching '{filename}' ({len(matches)} found"
        + (" — search timed out, results may be incomplete" if timed_out else "")
        + f", searching under {root})",
        "Use open_path_by_index(index) to open or read a file.",
        "",
        f"{'IDX':>4}  {'SIZE':>8}  PATH",
        "─" * 72,
    ]
    for i, m in enumerate(matches):
        lines.append(f"{i:>4}  {m['size']:>8}  {m['full_path']}")
    return "\n".join(lines)


def open_path_by_index(index: int) -> str:
    """Open/read a file from the last find_file() result by index."""
    entry = _get_indexed("find_file", index)
    if entry is None:
        return f"Index {index} not found. Call find_file() first."
    return read_file_smart(entry["full_path"])


def explore_path(path):
    """Legacy wrapper — delegates to list_directory for indexed output."""
    return list_directory(path)



def execute_terminal_command(command, working_directory=None):
    try:
        cwd = working_directory if working_directory else STARTUP_DIR
        if _IS_WINDOWS:
            result = subprocess.run(
                ["powershell", "-Command", command],
                capture_output=True, text=True, timeout=30, cwd=cwd
            )
        else:
            result = subprocess.run(
                command, shell=True, capture_output=True, text=True,
                timeout=30, cwd=cwd, executable="/bin/bash"
            )
        return f"STDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}"
    except Exception as e:
        return f"Execution failed: {str(e)}"
    
def execute_python_code(code: str, timeout: int = 15) -> str:
    """
    Run a snippet of Python code in an isolated subprocess — a brand-new
    interpreter (`python -I -c <code>`), NOT Midum's own process — and
    return whatever it printed to stdout/stderr. Use this for calculations,
    data processing, quick algorithms, testing logic, or parsing/transforming
    text. There is no persistent state between calls: each call starts a
    fresh interpreter, so assign results to variables and print() them to
    see the output.

    NOT for interacting with the desktop UI, files outside what you pass
    it, or Midum's own runtime — use the dedicated tools for those. `-I`
    (isolated mode) keeps it from reading the user's site-packages
    customizations or environment, and a hard timeout kills runaway loops.
    """
    code = (code or "").strip()
    if not code:
        return "Error: 'code' must contain the Python code to run."
    try:
        timeout = max(1, min(int(timeout or 15), 60))
    except (TypeError, ValueError):
        timeout = 15

    python_exe = sys.executable or "python"
    print(f"   [Python sandbox] Running {len(code)} chars (timeout={timeout}s)")
    try:
        result = subprocess.run(
            [python_exe, "-I", "-c", code],
            capture_output=True, text=True, timeout=timeout, cwd=STARTUP_DIR,
        )
        return (
            f"[exit code {result.returncode}]\n"
            f"STDOUT:\n{result.stdout}\n"
            f"STDERR:\n{result.stderr}"
        )
    except subprocess.TimeoutExpired:
        return f"Error: code timed out after {timeout}s — check for an infinite loop or blocking call."
    except Exception as e:
        return f"Execution failed: {str(e)}"


def _uia_unavailable_message() -> str:
    """Shared diagnostic message for any tool that needs UIA but it's not up."""
    if _UIA_INIT_ERROR:
        return (
            f"UI automation unavailable. Real cause: {_UIA_INIT_ERROR} "
            f"(this is NOT a Windows-version issue despite what the library may print — "
            f"it's almost always a COM init or pywin32 install issue). "
            f"Try: pip install --force-reinstall uiautomation pywin32"
        )
    return "UIA library not installed. pip install uiautomation pywin32"


def manual_scan_app_layouts(window_title: str):
    if not _UIA_AVAILABLE:
        return _uia_unavailable_message()
    # Support shell aliases so the model can explore taskbar/tray etc.
    hwnd = ui_navigator._find_shell_hwnd(window_title.strip().lower())
    if hwnd:
        # Temporarily resolve via the alias path
        import ctypes
        win32gui.SetForegroundWindow  # just to ensure win32gui is loaded
    containers = ui_navigator.discover_ui_subtrees(window_title)
    if not containers:
        return (
            f"No named containers found in '{window_title}' within "
            f"{ui_navigator.MAX_DISCOVER_DEPTH} levels. "
            f"For shell surfaces use: 'taskbar', 'start', 'tray', 'desktop', "
            f"'action center', 'search', 'tray overflow'. "
            f"If this is a canvas-based app, try fallback_click_text instead."
        )
    summary = (
        f"Found {len(containers)} container(s) in '{window_title}'. "
        f"Pick the most relevant subtree_key and call manual_inspect_app_subtree.\n"
    )
    return summary + json.dumps(containers, indent=2)

def manual_inspect_app_subtree(window_title: str, subtree_key: str):
    if not _UIA_AVAILABLE:
        return _uia_unavailable_message()
    controls = ui_navigator.inspect_subtree_controls(window_title, subtree_key)
    if not controls:
        return (
            f"No actionable controls found in '{subtree_key}' within "
            f"{ui_navigator.MAX_INSPECT_DEPTH} levels below it. "
            f"Try a different subtree_key from manual_scan_app_layouts, or if this "
            f"app renders via canvas/WebGL (games, some web apps), use "
            f"fallback_click_text instead — UIA cannot see canvas content."
        )
    summary = f"Found {len(controls)} actionable control(s) in '{subtree_key}':\n"
    return summary + json.dumps(controls, indent=2)

def manual_interact_with_ui(window_title: str, control_type: str, search_property: str, property_value: str, action: str, text_to_type: str = ""):
    if not _UIA_AVAILABLE: return "UIA library not installed."
    return ui_navigator.safely_trigger_ui_element(
        window_title, control_type, search_property, property_value, action, text_to_type
    )

def click_ui_element(window_title: str, description: str, action: str = "click", text_to_type: str = ""):
    """
    ONE-CALL UI interaction: find an element in `window_title` matching
    `description` (plain English, e.g. "Close button", "Send message field")
    and act on it. Falls back to coordinate-click automatically if the
    element can't be invoked via UIA patterns directly (common for
    Electron/Chromium close/minimize/maximize buttons).
    """
    if ui_navigator is None:
        if _IS_WINDOWS:
            return _uia_unavailable_message()
        return "UI automation not available — install pyatspi + xdotool (Linux)."
    return ui_navigator.find_and_act(window_title, description, action, text_to_type)

def list_active_windows():
    """
    Returns a clean list of all visible, named windows currently open on the desktop.
    Works on both Windows (win32gui) and Linux (xdotool).
    """
    if _IS_LINUX:
        # xdotool search returns all window IDs; get their names via getwindowname
        stdout, err = _run(["xdotool", "search", "--name", ""])
        if err and not stdout:
            # Fallback: wmctrl -l
            stdout2, _ = _run(["wmctrl", "-l"])
            if stdout2:
                titles = []
                for line in stdout2.splitlines():
                    parts = line.split(None, 3)
                    if len(parts) >= 4:
                        titles.append(parts[3].strip())
                unique = sorted(set(t for t in titles if t))
                return "Currently open windows:\n" + "\n".join(f"- {w}" for w in unique)
            return "Error listing windows: xdotool and wmctrl both unavailable."
        titles = []
        for wid in stdout.splitlines():
            wid = wid.strip()
            if not wid:
                continue
            name_out, _ = _run(["xdotool", "getwindowname", wid])
            name = name_out.strip()
            # Filter out empty and common desktop shell overlays
            if name and name not in ["Desktop", "xdotool"]:
                titles.append(name)
        unique = sorted(set(titles))
        if not unique:
            return "No visible named windows found."
        return "Currently open windows:\n" + "\n".join(f"- {w}" for w in unique)

    # Windows path
    if not _UIA_AVAILABLE:
        return "UIA library not installed."

    def enum_win_callback(hwnd, window_list):
        if win32gui.IsWindowVisible(hwnd):
            title = win32gui.GetWindowText(hwnd).strip()
            if title and title not in ["Program Manager", "Settings", "Microsoft Text Input Application"]:
                window_list.append(title)

    windows = []
    win32gui.EnumWindows(enum_win_callback, windows)
    unique_windows = sorted(list(set(windows)))

    # Show canonical app name alongside raw title so the model knows what to pass
    lines = []
    for title in unique_windows:
        canonical = ui_navigator._canonical_app_name(title) if ui_navigator else title
        if canonical != title:
            lines.append(f"- '{title}' → use '{canonical}' as window_title")
        else:
            lines.append(f"- {title}")

    # Also report which shell surfaces are available
    shell_available = []
    for alias in ["taskbar", "tray", "desktop", "tray overflow", "action center", "search"]:
        hwnd = ui_navigator._find_shell_hwnd(alias)
        if hwnd:
            shell_available.append(alias)

    result = "Currently open windows:\n" + "\n".join(lines)
    if shell_available:
        result += (
            "\n\nShell surfaces (use these as window_title in click_ui_element):\n"
            + "\n".join(f"- {s}" for s in shell_available)
        )
    result += (
        "\n\nNOTE: Always use the canonical name (after →) as window_title. "
        "App titles change when pages/documents change, but canonical names are stable."
    )
    return result if unique_windows or shell_available else "No visible named windows found."


# =============================================================================

