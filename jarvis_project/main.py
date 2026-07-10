import io
import os
import re
import json
import base64
import tempfile
import datetime
import subprocess
import threading
import asyncio
import queue as _queue
import platform as _platform
import ollama
from ddgs import DDGS
import time

# requests is used by OpenRouter (chat completions) and CDP (Chrome tab
# listing). Imported once at module level; both subsystems degrade
# gracefully (return a clear "not installed" message) if missing.
try:
    import requests
except ImportError:
    requests = None

_IS_LINUX   = _platform.system() == "Linux"
_IS_WINDOWS = _platform.system() == "Windows"

# PIL ImageGrab works on Windows natively. On Linux it requires either
# python3-xlib (X11) or a scrot/gnome-screenshot fallback.
try:
    from PIL import ImageGrab as _ImageGrab
    _IMAGEGRAB_AVAILABLE = True
except ImportError:
    _ImageGrab = None
    _IMAGEGRAB_AVAILABLE = False

# keyboard is used for the Ctrl+Q abort shortcut.
# Install with:  pip install keyboard
# Note: on Windows, keyboard requires no extra drivers. Run as admin if hotkeys
# don't fire (rare — usually works in normal user sessions inside a terminal).
try:
    import keyboard as _keyboard
    _KEYBOARD_AVAILABLE = True
except ImportError:
    _KEYBOARD_AVAILABLE = False

# ── Optional file-format libraries ────────────────────────────────────────────
# PDF reading:   pip install pymupdf
# DOCX reading:  pip install mammoth
# DOCX writing:  pip install python-docx
try:
    import fitz as _fitz
    _PDF_AVAILABLE = True
except ImportError:
    _fitz = None
    _PDF_AVAILABLE = False

try:
    import mammoth as _mammoth
    _MAMMOTH_AVAILABLE = True
except ImportError:
    _mammoth = None
    _MAMMOTH_AVAILABLE = False

try:
    import docx as _docx
    _DOCX_AVAILABLE = True
except ImportError:
    _docx = None
    _DOCX_AVAILABLE = False

# MCP (Model Context Protocol) client SDK — lets Midum connect to external
# MCP servers (stdio subprocesses, or remote HTTP/SSE endpoints) and call
# their tools. Install with:  pip install mcp
try:
    import mcp as _mcp_sdk  # noqa: F401
    _MCP_SDK_AVAILABLE = True
except ImportError:
    _mcp_sdk = None
    _MCP_SDK_AVAILABLE = False


# rich renders Markdown in the terminal (headers, bold, code blocks, tables).
# Pure formatting — zero effect on model logic or performance.
# Install with:  pip install rich
try:
    from rich.console import Console as _Console
    from rich.markdown import Markdown as _Markdown
    _console          = _Console()
    _RICH_AVAILABLE   = True
except ImportError:
    _RICH_AVAILABLE   = False

def _print_reply(label: str, text: str):
    """Print Midum's reply, rendering Markdown if rich is available."""
    # Suppress replies that are pure JSON/punctuation leftovers from legacy parsing
    if not text or re.match(r'^[{}\[\]",:\s]*$', text.strip()):
        return
    print(f"\n{label}")
    if _RICH_AVAILABLE and text.strip():
        _console.print(_Markdown(text))
    else:
        print(text)

# =============================================================================
# CONFIGURATION
# =============================================================================
STARTUP_DIR    = os.getcwd()
MODEL_NAME     = "jarvishehe"

# ── PRIMARY MODEL PROVIDER ────────────────────────────────────────────────────
# "ollama"     — use the local Ollama model (MODEL_NAME above) as the primary
#                execution brain. OpenRouter is available as a frequent
#                secondary consultant (see OPENROUTER_CONSULT_MODE below).
# "openrouter" — use OPENROUTER_MODEL as the PRIMARY execution brain instead
#                of Ollama. Every tool-calling turn is a metered API call.
#                Use this if you have an OpenRouter key and want a stronger
#                model driving Midum directly instead of qwen2.5-coder.
# "gemini_web" — use Gemini, accessed through the `gemini_webapi` library
#                (account cookie sign-in, no API key, no per-token metering)
#                as the PRIMARY execution brain. Reuses the same account
#                session as consult_gemini/query_gemini_app (see the
#                GEMINI WEB APP CLIENT section below). Every tool-calling
#                step is a real web request to gemini.google.com's internal
#                endpoints through a persistent ChatSession, so it is slower
#                per-step than Ollama/OpenRouter — expect real latency on
#                multi-hop tool loops. See GEMINI_WEB_MODEL below to pick
#                the model tier.
# "gemini_api" — use Gemini as the PRIMARY execution brain through the
#                OFFICIAL Google Gemini API (an API key, not a browser/cookie
#                hack), talking to Gemini's OpenAI-compatible /chat/completions
#                endpoint with real structured function-calling. Same request
#                shape as "openrouter", so it gets the exact same reliability
#                (native tool_calls field, no scraping, no session juggling)
#                and is fully wired into tool calling / MCP servers exactly
#                like every other provider. See GEMINI_API_MODEL below.
#
# A future GUI will expose this as a dropdown — for now, edit directly.
# "groq"       — use GROQ_MODEL as the PRIMARY execution brain instead of
#                Ollama. GroqCloud offers a genuinely free tier (no credit
#                card) with fast inference and native structured tool
#                calling on models like llama-3.3-70b-versatile and
#                qwen/qwen3-32b. Same request shape as "openrouter" /
#                "gemini_api" (OpenAI-compatible /chat/completions), so it
#                is fully wired into tool calling / MCP servers exactly
#                like every other provider. See GROQ_MODEL below.
MODEL_PROVIDER = "gemini_web"   # "ollama" | "openrouter" | "gemini_web" | "gemini_api" | "groq"

# ── OpenRouter model selection ────────────────────────────────────────────────
# Used when MODEL_PROVIDER == "openrouter" (primary), and always used for
# secondary consultation regardless of MODEL_PROVIDER (see OPENROUTER_CONSULT_MODE).
#
# OpenRouter model IDs: https://openrouter.ai/models
# Free-tier examples (subject to change, check openrouter.ai/models?max_price=0):
#   "meta-llama/llama-3.1-8b-instruct:free"
#   "meta-llama/llama-3.3-70b-instruct:free"
#   "google/gemini-2.0-flash-exp:free"
#   "deepseek/deepseek-chat-v3.1:free"
#   "qwen/qwen3-235b-a22b:free"
# Paid examples (for MODEL_PROVIDER == "openrouter" as a strong primary):
#   "anthropic/claude-3.7-sonnet"
#   "openai/gpt-4o"
#   "google/gemini-2.5-pro"
OPENROUTER_MODEL          = "meta-llama/llama-3.3-70b-instruct:free"   # primary (if selected) AND consult model
OPENROUTER_API_BASE       = "https://openrouter.ai/api/v1"

# ── Fallback free models ──────────────────────────────────────────────────────
# OpenRouter's ":free" models share a global rate-limit pool across ALL
# OpenRouter users, not just you — so a single free model (like the
# llama-3.3-70b:free default above) will very commonly return 429
# ("Provider returned error") during busy periods, no matter how you
# configure retries. When that happens, _openrouter_chat_with_fallback()
# below automatically tries the next model in this list instead of just
# failing the turn. Order = preference. Keep OPENROUTER_MODEL first (it's
# tried first), the rest are only used if it 429s/502s/503s repeatedly.
OPENROUTER_FALLBACK_MODELS = [
    OPENROUTER_MODEL,
    "meta-llama/llama-3.1-8b-instruct:free",
    "google/gemini-2.0-flash-exp:free",
    "deepseek/deepseek-chat-v3.1:free",
    "qwen/qwen3-235b-a22b:free",
]

# ── Secondary consultation mode ───────────────────────────────────────────────
# Controls how often OpenRouter is consulted as a planning brain in addition
# to (or instead of) Gemini. OpenRouter free models are consulted far more
# aggressively than Gemini since there's no desktop-app UI-automation cost —
# it's a direct API call, so it's cheap to call on every non-trivial turn.
#   "always"    — consult OpenRouter on every non-trivial turn (most reliable,
#                 more API usage — fine for free-tier models)
#   "fallback"  — only consult OpenRouter if Gemini (app + API) both fail
#   "off"       — never consult OpenRouter as a secondary planner
OPENROUTER_CONSULT_MODE   = "always"

# ── Gemini API (official) model selection ─────────────────────────────────────
# Used when MODEL_PROVIDER == "gemini_api". This is the REAL Google Gemini API
# (an API key from https://aistudio.google.com/app/apikey), NOT the web-chat
# scraping used by "gemini_web" — no cookies, no browser session, no auto
# model detection. Structured function calling is native, so tool calling is
# reliable step to step.
GEMINI_API_MODEL = "gemini-3.1-flash-lite"   # exact model ID as listed at https://ai.google.dev/gemini-api/docs/models
GEMINI_API_BASE  = "https://generativelanguage.googleapis.com/v1beta/openai"

# ── GroqCloud model selection ─────────────────────────────────────────────────
# Used when MODEL_PROVIDER == "groq" (primary), and always available for
# on-demand consultation/delegation regardless of MODEL_PROVIDER, exactly
# like OpenRouter/Gemini API. GroqCloud's free tier needs only a key from
# https://console.groq.com/keys — no credit card — and gives very fast
# inference plus reliable native structured tool-calling.
#
# Model IDs: https://console.groq.com/docs/models (subject to change):
#   "llama-3.3-70b-versatile"   — strong general-purpose, good tool calling
#   "llama-3.1-8b-instant"      — fastest, smaller
#   "qwen/qwen3-32b"            — strong reasoning/coding
#   "deepseek-r1-distill-llama-70b" — reasoning-heavy
#   "moonshotai/kimi-k2-instruct"   — large MoE, strong tool use
GROQ_MODEL          = "llama-3.3-70b-versatile"   # primary (if selected) AND consult model
GROQ_API_BASE        = "https://api.groq.com/openai/v1"

# GroqCloud's free tier has generous per-minute/per-day request & token
# limits that vary by model, and can occasionally 429 during bursts. When
# that happens, _groq_chat_with_fallback() automatically tries the next
# model in this list instead of just failing the turn. Order = preference.
GROQ_FALLBACK_MODELS = [
    GROQ_MODEL,
    "llama-3.1-8b-instant",
    "qwen/qwen3-32b",
    "moonshotai/kimi-k2-instruct",
]

# ── Gemini-web (gemini_webapi) primary-execution settings ─────────────────────
# Used only when MODEL_PROVIDER == "gemini_web". Separate from consult_gemini/
# query_gemini_app's ad-hoc single-shot calls — this drives the ENTIRE
# tool-calling loop through one persistent gemini_webapi ChatSession per
# conversation (see "GEMINI WEB APP CLIENT" section for the underlying
# client, and "GEMINI-WEB PRIMARY EXECUTION BACKEND" for the loop wiring).
#
# GEMINI_WEB_MODEL: "" (default) = auto-pick the fastest available model via
# client.list_models() at runtime (never hardcode a model name — the web
# app's exposed lineup changes over time and isn't guaranteed to match the
# name you last saw). Set to an exact model_name/display_name string
# (e.g. "gemini-3-flash") to pin one instead.
GEMINI_WEB_MODEL              = "gemini-3-flash"   # exact model name as listed in gemini_webapi client.list_models()

# NOTE: deliberately NOT using a Gem here. Gems proved problematic (flaky
# create/update/fetch round trips, an extra persistent server-side object
# that can drift out of sync, silent "no-gem mode" degradation). Instead
# persona + tool JSON-output-format instructions are sent as a plain-text
# priming message on the first turn of every fresh ChatSession (see
# _gemini_web_persona_prompt()), and the native tool schema is discovered
# on demand via list_native_tools()/show_native_tool_schema(tool_name)
# instead of being inlined into anything persistent. Same capability, no
# Gem dependency.

# Per-hop and whole-task timeout budgets (seconds). Each tool-loop hop is a
# real round trip through gemini.google.com, so these are generous compared
# to the local-model/OpenRouter equivalents.
GEMINI_WEB_HOP_TIMEOUT        = 120
GEMINI_WEB_TOTAL_TASK_TIMEOUT = 1800

# Marker prefixed to every tool-result message injected into the Gemini
# ChatSession, so Gemini can reliably tell an injected tool result apart
# from a genuine new user message. Keep this exact string in sync anywhere
# else results are formatted for Gemini.
GEMINI_WEB_TOOL_RESULT_MARKER = "[TOOL_RESULT]"

# ── Legacy / weak native-tool-calling models ────────────────────────────────
LEGACY_TOOLCALL_MODELS = (
    "qwen2.5-coder",
    "qwen2.5",
    "codeqwen",
    "deepseek-coder",
    "codellama",
)

def _is_legacy_toolcall_model(model_name: str) -> bool:
    """True if model_name or its underlying base model matches a known weak tool-calling family."""
    low = model_name.lower()
    if any(fam in low for fam in LEGACY_TOOLCALL_MODELS):
        return True
    try:
        info = ollama.show(model_name)
        base = info.get("modelinfo", {}).get("general.basename", "").lower()
        return any(fam in base for fam in LEGACY_TOOLCALL_MODELS)
    except Exception:
        return False

# ── Platform-aware paths ──────────────────────────────────────────────────────
if _IS_LINUX:
    _HOME           = os.path.expanduser("~")
    TARGET_DIR      = os.path.join(_HOME, "Jarvis")
    STORAGE_DIR     = os.path.join(TARGET_DIR, "jarvis_project", "storage")
    SKILLS_INDEX    = os.path.join(TARGET_DIR, "jarvis_project", "skills.md")
    SECRETS_FILE    = os.path.join(_HOME, ".config", "JarvisSecrets", "jarvis_secrets.json")
else:
    TARGET_DIR      = r"D:\Jarvis"
    STORAGE_DIR     = r"D:\Jarvis\jarvis_project\storage"
    SKILLS_INDEX    = r"D:\Jarvis\jarvis_project\skills.md"
    SECRETS_FILE    = os.path.join(
        os.path.expanduser("~"), "AppData", "Local", "JarvisSecrets", "jarvis_secrets.json"
    )

COMMANDS_FILE       = os.path.join(STORAGE_DIR, "commands.md")
INSTRUCTIONS_FILE   = os.path.join(STORAGE_DIR, "instructions.md")
PATHS_FILE          = os.path.join(STORAGE_DIR, "paths.md")
DOMAIN_INDEX        = os.path.join(STORAGE_DIR, "domain_index.md")
SKILLS_DIR          = os.path.join(STORAGE_DIR, "skills")
DOMAIN_SKILLS_INDEX = os.path.join(STORAGE_DIR, "domain_skills_index.md")
MASTER_MEMORY       = os.path.join(STORAGE_DIR, "master_memory.md")
SESSION_MEMORY      = os.path.join(STORAGE_DIR, "session_memory.md")
RESPONSE_MEMORY     = os.path.join(STORAGE_DIR, "response_memory.md")
MCP_SERVERS_FILE    = os.path.join(STORAGE_DIR, "mcp_servers.json")
LOG_FILE            = os.path.join(TARGET_DIR, "chat_log.md")

GOAL_SECTION_HEADER = "## Current Goal"
GOAL_SECTION_END    = "## Goal History"

# ── Screen resolution — detected at runtime on Linux, hardcoded on Windows ───
def _detect_screen_resolution() -> tuple[int, int]:
    """Detect screen resolution. Falls back to 1920x1080 if detection fails."""
    if _IS_LINUX:
        try:
            out, _ = subprocess.run(
                ["xdpyinfo"], capture_output=True, text=True, timeout=5
            ).stdout, None
            m = re.search(r"dimensions:\s*(\d+)x(\d+)", out)
            if m:
                return int(m.group(1)), int(m.group(2))
        except Exception:
            pass
        try:
            out, _ = subprocess.run(
                ["xrandr", "--current"], capture_output=True, text=True, timeout=5
            ).stdout, None
            m = re.search(r"current (\d+) x (\d+)", out)
            if m:
                return int(m.group(1)), int(m.group(2))
        except Exception:
            pass
        return 1920, 1080   # safe fallback
    else:
        return 2560, 1600   # Windows default — update if your resolution differs

SCREEN_W, SCREEN_H = _detect_screen_resolution()

# The model's internal canvas — Ollama vision models downscale to 1024px long edge
MODEL_CANVAS_W = 1024
MODEL_CANVAS_H = int(1024 * SCREEN_H / SCREEN_W)

# Scale factors: real_px = canvas_px * SCALE
SCALE_X = SCREEN_W / MODEL_CANVAS_W
SCALE_Y = SCREEN_H / MODEL_CANVAS_H

# Grid drawn on screenshots at canvas resolution; every GRID_STEP canvas-px
GRID_STEP = 100

# =============================================================================
# ABORT FLAG — Ctrl+Q sets this to stop the current response
# =============================================================================
# A threading.Event that process_chat_turn checks at every loop iteration.
# When set, the turn is abandoned and control returns to the input prompt.
_abort_event = threading.Event()

# =============================================================================
# TESSERACT SETUP
# =============================================================================
# Windows: download from https://github.com/UB-Mannheim/tesseract/wiki
# Linux:   sudo apt install tesseract-ocr   OR   sudo dnf install tesseract
#          pip install pytesseract
#
_TESSERACT_AVAILABLE = False

try:
    import pytesseract
    if _IS_WINDOWS:
        TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
    # On Linux, tesseract is on PATH after apt/dnf install — no path needed
    pytesseract.get_tesseract_version()
    _TESSERACT_AVAILABLE = True
except Exception:
    pass


# ── UI Automation (UIA) SETUP — Windows only ─────────────────────────────────
# pip install uiautomation pywin32
#
# NOTE: The `uiautomation` package sometimes prints a generic
# "update your Windows" message when its internal COM initialization
# (via comtypes) fails. That message is almost always misleading — the real
# cause is usually one of:
#   1. pywin32 post-install script never ran (fixes most import-time COM issues)
#      Fix: run `python Scripts/pywin32_postinstall.py -install` as admin
#   2. COM apartment threading conflict — a prior library already called
#      CoInitialize with a different threading model in this process
#   3. Running inside a sandboxed/restricted session without COM access
#   4. UIA COM components genuinely missing (rare on any Windows 8+ system)
# We capture the REAL underlying exception below so the startup print shows
# the actual cause instead of the library's generic advice.
_UIA_AVAILABLE   = False
_UIA_INIT_ERROR  = None
win32gui = None
auto     = None

if _IS_WINDOWS:
    try:
        import win32gui
    except ImportError as e:
        _UIA_INIT_ERROR = f"pywin32 not installed: {e}"

    if win32gui is not None:
        try:
            import uiautomation as auto
            # Force a real COM call now (not just import) so init failures
            # surface here instead of silently later during the first real
            # UIA operation deep inside a tool call.
            _ = auto.GetRootControl()
            _UIA_AVAILABLE = True
        except Exception as e:
            _UIA_INIT_ERROR = f"{type(e).__name__}: {e}"
            auto = None

            # ── Retry once with explicit COM apartment initialization ────────
            # Fixes the most common cause: another library (or Ollama's own
            # threading) already initialized COM with an incompatible model
            # before uiautomation got a chance to.
            try:
                import comtypes
                comtypes.CoInitialize()
                import importlib
                if "uiautomation" in globals():
                    importlib.reload(auto)
                else:
                    import uiautomation as auto
                _ = auto.GetRootControl()
                _UIA_AVAILABLE  = True
                _UIA_INIT_ERROR = None
            except Exception as e2:
                _UIA_INIT_ERROR = (
                    f"{_UIA_INIT_ERROR} | Retry with CoInitialize also failed: "
                    f"{type(e2).__name__}: {e2}"
                )
                auto = None

# Control types that can host meaningful child content — used as scan roots
_CONTAINER_TYPES = {
    "PaneControl", "GroupControl", "CustomControl", "ToolbarControl",
    "DocumentControl", "WindowControl", "TabControl", "TabItemControl",
    "TreeControl", "ScrollBarControl", "MenuBarControl", "MenuControl",
}
# Control types that represent something a user can interact with
_ACTIONABLE_TYPES = {
    "ButtonControl", "EditControl", "CheckBoxControl", "HyperlinkControl",
    "ListItemControl", "MenuItemControl", "TabItemControl", "RadioButtonControl",
    "ComboBoxControl", "SliderControl", "ImageControl", "SplitButtonControl",
    "TextControl",
}

if _UIA_AVAILABLE:
    class AppMapNavigator:
        """
        UIA navigator tuned for deep Electron/Chromium UI trees (VS Code, Modrinth,
        Discord, etc). These apps wrap real content in many layers of generic
        Pane/Group/Custom containers, often 8-15 levels deep, so shallow depth
        cutoffs (3-4) never reach anything actionable.

        Strategy:
          - discover_ui_subtrees: walk up to MAX_DISCOVER_DEPTH levels, collect
            ANY named/automation-id'd container as a candidate subtree, not just
            top-level ones. Electron apps nest meaningfully-named panes deep.
          - inspect_subtree_controls: walk up to MAX_INSPECT_DEPTH levels below
            a chosen subtree root, collecting all actionable controls regardless
            of how deep they sit.
          - If a subtree yields nothing, automatically retry one level shallower
            in the tree (parent) so Midum doesn't have to manually backtrack.
        """

        MAX_DISCOVER_DEPTH = 12
        MAX_INSPECT_DEPTH  = 12
        MAX_RESULTS        = 60   # cap result lists so they don't blow the context

        def __init__(self, maps_dir="storage/app_maps"):
            self.maps_dir = os.path.abspath(maps_dir)
            os.makedirs(self.maps_dir, exist_ok=True)
            self._live_cache     = {}
            self._snapshot_cache = {}   # window_title → last snapshot element list

        def _get_map_path(self, window_title: str) -> str:
            safe_name = "".join([c if c.isalnum() else "_" for c in window_title.lower()]).strip("_")
            return os.path.join(self.maps_dir, f"{safe_name}.json")

        def load_app_blueprint(self, window_title: str) -> dict:
            path = self._get_map_path(window_title)
            if os.path.exists(path):
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        return json.load(f)
                except Exception:
                    pass
            return {"window_title": window_title, "subtrees": {}, "known_controls": {}}

        def save_app_blueprint(self, window_title: str, blueprint: dict):
            path = self._get_map_path(window_title)
            try:
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(blueprint, f, indent=2)
            except Exception:
                pass

        # ── Known Windows shell surface class names ───────────────────────────
        # These windows exist permanently but have no useful title and are never
        # the foreground window. They must be found by class name, not title.
        # Keys are the canonical aliases the model can pass as window_title.
        _SHELL_ALIASES: dict[str, list[str]] = {
            # Primary taskbar (Win10 + Win11)
            "taskbar":          ["Shell_TrayWnd"],
            "start":            ["Shell_TrayWnd"],   # Start button lives inside
            "start button":     ["Shell_TrayWnd"],
            "start menu":       ["Windows.UI.Core.CoreWindow",   # Win11 Start
                                 "Shell_TrayWnd"],
            # Secondary taskbars on multi-monitor setups
            "taskbar2":         ["Shell_SecondaryTrayWnd"],
            "secondary taskbar":["Shell_SecondaryTrayWnd"],
            # System tray / notification area
            "tray":             ["Shell_TrayWnd"],
            "system tray":      ["Shell_TrayWnd"],
            "notification area":["Shell_TrayWnd"],
            "clock":            ["Shell_TrayWnd"],
            # Overflow popup for hidden tray icons
            "tray overflow":    ["NotifyIconOverflowWindow"],
            "overflow":         ["NotifyIconOverflowWindow"],
            # Desktop itself
            "desktop":          ["Progman", "WorkerW"],
            # Action Center / Quick Settings (Win11)
            "action center":    ["Windows.UI.Core.CoreWindow"],
            "quick settings":   ["Windows.UI.Core.CoreWindow"],
            "notifications":    ["Windows.UI.Core.CoreWindow"],
            # Search bar / Search panel
            "search":           ["SearchHost",
                                 "Windows.UI.Core.CoreWindow"],
            "search bar":       ["Shell_TrayWnd"],
            # Task View / Virtual Desktops button
            "task view":        ["Shell_TrayWnd"],
            # Volume / Wi-Fi / Battery flyout
            "volume":           ["Windows.UI.Core.CoreWindow",
                                 "Shell_TrayWnd"],
        }

        def _find_shell_hwnd(self, alias: str) -> int | None:
            """
            Find a Windows shell surface by its canonical alias.
            Tries each class name in order and returns the first valid hwnd.
            """
            for cls in self._SHELL_ALIASES.get(alias.lower(), []):
                h = win32gui.FindWindow(cls, None)
                if h:
                    return h
            return None

        # ── App name → window class / title fragment mapping ──────────────────
        # Maps natural app names to (win32_class, title_fragment) pairs.
        # win32_class is used with FindWindow for a fast exact match.
        # title_fragment is the stable part of the title that never changes
        # regardless of what document/page/tab is currently open.
        _APP_ALIASES: dict[str, tuple[str, str]] = {
            "chrome":         ("Chrome_WidgetWin_1",  "Google Chrome"),
            "google chrome":  ("Chrome_WidgetWin_1",  "Google Chrome"),
            "brave":          ("Chrome_WidgetWin_1",  "Brave"),
            "brave browser":  ("Chrome_WidgetWin_1",  "Brave"),
            "firefox":        ("MozillaWindowClass",  "Firefox"),
            "edge":           ("Chrome_WidgetWin_1",  "Microsoft Edge"),
            "microsoft edge": ("Chrome_WidgetWin_1",  "Microsoft Edge"),
            "vscode":         ("Chrome_WidgetWin_1",  "Visual Studio Code"),
            "vs code":        ("Chrome_WidgetWin_1",  "Visual Studio Code"),
            "visual studio code": ("Chrome_WidgetWin_1", "Visual Studio Code"),
            "notepad":        ("Notepad",             "Notepad"),
            "notepad++":      ("Notepad++",           "Notepad++"),
            "explorer":       ("CabinetWClass",       "File Explorer"),
            "file explorer":  ("CabinetWClass",       "File Explorer"),
            "discord":        ("Chrome_WidgetWin_1",  "Discord"),
            "slack":          ("Chrome_WidgetWin_1",  "Slack"),
            "whatsapp":       ("Chrome_WidgetWin_1",  "WhatsApp"),
            "spotify":        ("Chrome_WidgetWin_1",  "Spotify"),
            "teams":          ("Chrome_WidgetWin_1",  "Microsoft Teams"),
            "obs":            ("Qt5152QWindowIcon",   "OBS"),
            "terminal":       ("CASCADIA_HOSTING_WINDOW_CLASS", "Windows Terminal"),
            "windows terminal": ("CASCADIA_HOSTING_WINDOW_CLASS", "Windows Terminal"),
            "powershell":     ("ConsoleWindowClass",  "Windows PowerShell"),
            "cmd":            ("ConsoleWindowClass",  "Command Prompt"),
            "paint":          ("MSPaintApp",          "Paint"),
            "word":           ("OpusApp",             "Word"),
            "excel":          ("XLMAIN",              "Excel"),
            "powerpoint":     ("PPTFrameClass",       "PowerPoint"),
            "outlook":        ("rctrl_renwnd32",      "Outlook"),
            "zoom":           ("zoom",                "Zoom"),
            "vlc":            ("Qt5QWindowIcon",      "VLC"),
            "calculator":     ("ApplicationFrameWindow", "Calculator"),
            "settings":       ("ApplicationFrameWindow", "Settings"),
            "task manager":   ("TaskManagerWindow",   "Task Manager"),
        }

        def _resolve_app_window(self, name: str) -> int | None:
            """
            Find the best window for a given app name, regardless of its
            current title (which changes as pages/documents change).

            Strategy:
            1. Check _APP_ALIASES for a known win32 class → FindWindow by class
               (fast, exact, title-independent).
            2. Enumerate all windows and find those whose title contains the
               stable title_fragment, preferring the foreground window.
            3. If multiple instances exist, prefer the one most recently
               brought to the foreground.
            """
            low = name.strip().lower()
            entry = self._APP_ALIASES.get(low)

            if entry:
                win32_class, title_frag = entry

                # Fast path: FindWindow by class (finds ANY window of this app)
                # This returns the topmost window of that class.
                hwnd = win32gui.FindWindow(win32_class, None)
                if hwnd and win32gui.IsWindowVisible(hwnd):
                    return hwnd

                # Enumerate all windows of this class (multi-window apps)
                matches: list[tuple[int, str]] = []
                def cb_class(h, _):
                    try:
                        cls = win32gui.GetClassName(h)
                        if cls == win32_class and win32gui.IsWindowVisible(h):
                            title = win32gui.GetWindowText(h).strip()
                            if title:
                                matches.append((h, title))
                    except Exception:
                        pass
                win32gui.EnumWindows(cb_class, None)

                if matches:
                    # Prefer foreground
                    try:
                        fg = win32gui.GetForegroundWindow()
                        for h, _ in matches:
                            if h == fg:
                                return h
                    except Exception:
                        pass
                    # Prefer the one whose title contains the stable fragment
                    frag_matches = [
                        (h, t) for h, t in matches
                        if title_frag.lower() in t.lower()
                    ]
                    if frag_matches:
                        return frag_matches[0][0]
                    return matches[0][0]

            # No alias entry — fall through to normal title search
            return None

        def _canonical_app_name(self, window_title: str) -> str:
            """
            Given any window title (possibly with dynamic content like page
            names), return the canonical app name for blueprint keying.

            "Python vs C++ - Google Search - Google Chrome" → "Google Chrome"
            "main.py - Midum - Visual Studio Code"         → "Visual Studio Code"
            "New tab - Google Chrome"                       → "Google Chrome"
            "Discord"                                       → "Discord"
            """
            # Check if any app alias's stable fragment appears in the title
            for alias, (_, frag) in self._APP_ALIASES.items():
                if frag.lower() in window_title.lower():
                    return frag   # return the stable display name
            # Fall back to stripping dynamic prefix: take the last " - " segment
            parts = [p.strip() for p in re.split(r"\s+[-–|]\s+", window_title)]
            return parts[-1] if parts else window_title

        def _find_window(self, window_title: str) -> int | None:
            """
            Find a window by title, with layered resolution:

            1. Shell alias (taskbar, start, tray, desktop…)
            2. App name alias (_APP_ALIASES) — class-based, title-independent.
               This is the key fix for dynamic titles: "New tab - Google Chrome",
               "Python vs C++ - Google Chrome" etc. all resolve to Chrome's hwnd
               by class name, not by matching the title string.
            3. Exact title match
            4. Raw Win32 class name
            5. Substring match across all windows
            """
            # ── 1. Shell alias ────────────────────────────────────────────────
            alias = window_title.strip().lower()
            shell_hwnd = self._find_shell_hwnd(alias)
            if shell_hwnd:
                return shell_hwnd

            # ── 2. App name alias (title-independent) ─────────────────────────
            app_hwnd = self._resolve_app_window(window_title)
            if app_hwnd:
                return app_hwnd

            # ── 3. Exact title match ──────────────────────────────────────────
            hwnd = win32gui.FindWindow(None, window_title)
            if hwnd:
                return hwnd

            # ── 4. Raw class name ─────────────────────────────────────────────
            hwnd = win32gui.FindWindow(window_title, None)
            if hwnd:
                return hwnd

            # ── 5. Substring match across ALL top-level windows ───────────────
            matches: list[tuple[int, str, bool]] = []

            def cb(h, _):
                title = win32gui.GetWindowText(h).strip()
                if not title:
                    return
                visible = bool(win32gui.IsWindowVisible(h))
                if window_title.lower() in title.lower():
                    try:
                        rect = win32gui.GetWindowRect(h)
                        if (rect[2] - rect[0]) <= 0 or (rect[3] - rect[1]) <= 0:
                            if visible:
                                return
                    except Exception:
                        pass
                    matches.append((h, title, visible))

            win32gui.EnumWindows(cb, None)

            if not matches:
                return None
            if len(matches) == 1:
                return matches[0][0]

            visible_matches = [m for m in matches if m[2]]
            pool = visible_matches if visible_matches else matches

            try:
                fg = win32gui.GetForegroundWindow()
                for h, _t, _v in pool:
                    if h == fg:
                        return h
            except Exception:
                pass

            pool.sort(key=lambda m: len(m[1]), reverse=True)
            return pool[0][0]

        def _resolve_shell_element(self, description: str) -> tuple[int | None, str | None]:
            """
            For descriptions that directly name a shell control ("Start button",
            "Search", "Show hidden icons", "Clock", etc.), return the hwnd of the
            shell surface that contains it, and the canonical alias to use.
            Returns (None, None) if description doesn't match any shell surface.
            """
            desc_l = description.strip().lower()
            # Map description keywords to shell surface aliases
            _DESC_TO_ALIAS = {
                "start":        "start",
                "search bar":   "taskbar",
                "search":       "taskbar",
                "task view":    "taskbar",
                "show desktop": "taskbar",
                "clock":        "taskbar",
                "date":         "taskbar",
                "time":         "taskbar",
                "tray":         "tray",
                "notification": "tray",
                "system tray":  "tray",
                "wifi":         "tray",
                "volume":       "tray",
                "battery":      "tray",
                "speaker":      "tray",
                "network":      "tray",
                "hidden icons": "tray overflow",
                "overflow":     "tray overflow",
                "desktop":      "desktop",
            }
            for keyword, alias in _DESC_TO_ALIAS.items():
                if keyword in desc_l:
                    hwnd = self._find_shell_hwnd(alias)
                    if hwnd:
                        return hwnd, alias
            return None, None

        def discover_ui_subtrees(self, window_title: str) -> list[dict]:
            hwnd = self._find_shell_hwnd(window_title.strip().lower()) or self._find_window(window_title)
            if not hwnd:
                return []

            root_element = auto.ControlFromHandle(hwnd)
            if window_title not in self._live_cache:
                self._live_cache[window_title] = {}

            containers = []
            seen_keys  = set()

            for element, depth in auto.WalkControl(root_element, maxDepth=self.MAX_DISCOVER_DEPTH):
                try:
                    type_name = element.ControlTypeName
                    name      = (element.Name or "").strip()
                    auto_id   = (element.AutomationId or "").strip()
                except Exception:
                    continue

                if type_name not in _CONTAINER_TYPES:
                    continue
                if not (name or auto_id):
                    continue

                key = f"{name or auto_id}::{type_name}::{depth}"
                if key in seen_keys:
                    continue
                seen_keys.add(key)

                self._live_cache[window_title][key] = element
                containers.append({
                    "subtree_key":  key,
                    "type":         type_name,
                    "name":         name,
                    "automation_id": auto_id,
                    "depth":        depth,
                })

                if len(containers) >= self.MAX_RESULTS:
                    break

            return containers

        def inspect_subtree_controls(self, window_title: str, subtree_key: str) -> list[dict]:
            app_cache = self._live_cache.get(window_title, {})
            subtree_root = app_cache.get(subtree_key)

            if not subtree_root:
                self.discover_ui_subtrees(window_title)
                subtree_root = self._live_cache.get(window_title, {}).get(subtree_key)
                if not subtree_root:
                    return []

            controls   = []
            seen_names = set()

            try:
                walker = auto.WalkControl(subtree_root, maxDepth=self.MAX_INSPECT_DEPTH)
            except Exception:
                return []

            for element, depth in walker:
                try:
                    type_name = element.ControlTypeName
                    name      = (element.Name or "").strip()
                    auto_id   = (element.AutomationId or "").strip()
                except Exception:
                    continue

                if type_name not in _ACTIONABLE_TYPES:
                    continue
                # TextControl is extremely common as static labels — only keep
                # it if it has an AutomationId (suggests it's a real interactive
                # element wrapper, common in Electron) or looks like a button label
                if type_name == "TextControl" and not auto_id:
                    continue
                if not (name or auto_id):
                    continue

                dedup_key = f"{name}::{auto_id}::{type_name}"
                if dedup_key in seen_names:
                    continue
                seen_names.add(dedup_key)

                controls.append({
                    "type": type_name, "name": name, "automation_id": auto_id, "depth": depth
                })

                if len(controls) >= self.MAX_RESULTS:
                    break

            return controls

        def safely_trigger_ui_element(self, window_title: str, control_type: str,
                                       search_property: str, property_value: str,
                                       action: str, text_to_type: str = "") -> str:
            hwnd = self._find_shell_hwnd(window_title.strip().lower()) or self._find_window(window_title)
            if not hwnd:
                return f"Error: Window '{window_title}' not found."

            root = auto.ControlFromHandle(hwnd)
            search_args = {}
            if search_property == "automation_id":
                search_args["AutomationId"] = property_value
            elif search_property == "name":
                search_args["Name"] = property_value
            elif search_property == "class_name":
                search_args["ClassName"] = property_value

            c_type = control_type.lower()
            # searchDepth=0 means unlimited in uiautomation — search the whole subtree
            search_args["searchDepth"] = 0

            if c_type == "button":
                element = root.ButtonControl(**search_args)
            elif c_type in ("edit", "input"):
                element = root.EditControl(**search_args)
            elif c_type == "text":
                element = root.TextControl(**search_args)
            elif c_type == "checkbox":
                element = root.CheckBoxControl(**search_args)
            elif c_type == "menuitem":
                element = root.MenuItemControl(**search_args)
            elif c_type == "listitem":
                element = root.ListItemControl(**search_args)
            elif c_type == "image":
                element = root.ImageControl(**search_args)
            else:
                element = root.Control(**search_args)

            if not element.Exists(1, 0.5):
                return (
                    f"Error: Could not locate {control_type} matching "
                    f"{search_property}='{property_value}' anywhere in '{window_title}'. "
                    f"Try manual_inspect_app_subtree on a different container, or fall back to "
                    f"fallback_click_text if this element has visible text."
                )

            try:
                if action == "click":
                    try:
                        element.GetInvokePattern().Invoke()
                        return f"Programmatically invoked '{property_value}'."
                    except Exception:
                        try:
                            element.SetFocus()
                        except Exception:
                            pass
                        element.Click(simulateMove=False)
                        return f"Background clicked '{property_value}'."
                elif action == "set_text":
                    try:
                        element.GetValuePattern().SetValue(text_to_type)
                    except Exception:
                        element.SetValue(text_to_type)
                    return f"Populated text field '{property_value}'."
                elif action == "get_text":
                    try:
                        return element.GetValuePattern().Value
                    except Exception:
                        return element.Name
            except Exception as e:
                return f"Action failed: {str(e)}"
            return "Unknown action."

        def _blueprint_key(self, window_title: str) -> str:
            """
            Normalise any window title to a stable blueprint key using
            _canonical_app_name, so dynamic titles like
            'Python vs C++ - Google Chrome' and 'New tab - Google Chrome'
            both map to 'Google Chrome' and share the same blueprint.
            """
            return self._canonical_app_name(window_title)

        def _lookup_blueprint(self, window_title: str, description: str) -> dict | None:
            """
            Return a saved control record for (app, description) if one exists,
            else None. Record shape:
              {"automation_id": str, "type": str, "name": str, "depth": int}
            """
            key  = self._blueprint_key(window_title)
            bp   = self.load_app_blueprint(key)
            desc = description.strip().lower()
            return bp.get("known_controls", {}).get(desc)

        def _save_to_blueprint(self, window_title: str, description: str,
                               element, type_name: str, depth: int):
            """
            Persist a successful element lookup so future calls skip the tree walk.
            Only saves elements that have an AutomationId — name-only lookups are
            less stable across sessions.
            """
            try:
                auto_id = (element.AutomationId or "").strip()
                name    = (element.Name or "").strip()
                if not auto_id:
                    return   # not worth caching without a stable ID
                key  = self._blueprint_key(window_title)
                bp   = self.load_app_blueprint(key)
                bp.setdefault("known_controls", {})[description.strip().lower()] = {
                    "automation_id": auto_id,
                    "type":          type_name,
                    "name":          name,
                    "depth":         depth,
                }
                self.save_app_blueprint(key, bp)
                print(f"   [Blueprint] Saved '{description}' → AutomationId='{auto_id}' for '{key}'")
            except Exception:
                pass

        def _find_by_automation_id(self, root, auto_id: str, type_name: str):
            """
            Fast direct lookup by AutomationId. Returns the element or None.
            Uses uiautomation's built-in search which is much faster than a
            manual tree walk.
            """
            try:
                search_args = {"AutomationId": auto_id, "searchDepth": 0}
                t = type_name.lower().replace("control", "").strip()
                ctrl_map = {
                    "button":      lambda: root.ButtonControl(**search_args),
                    "edit":        lambda: root.EditControl(**search_args),
                    "text":        lambda: root.TextControl(**search_args),
                    "checkbox":    lambda: root.CheckBoxControl(**search_args),
                    "menuitem":    lambda: root.MenuItemControl(**search_args),
                    "listitem":    lambda: root.ListItemControl(**search_args),
                    "image":       lambda: root.ImageControl(**search_args),
                    "hyperlink":   lambda: root.HyperlinkControl(**search_args),
                    "combobox":    lambda: root.ComboBoxControl(**search_args),
                    "radiobutton": lambda: root.RadioButtonControl(**search_args),
                    "tabitem":     lambda: root.TabItemControl(**search_args),
                    "splitbutton": lambda: root.SplitButtonControl(**search_args),
                }
                factory = ctrl_map.get(t, lambda: root.Control(**search_args))
                el = factory()
                if el.Exists(0.5, 0.1):
                    return el
            except Exception:
                pass
            return None

        def snapshot_ui(self, window_title: str, filter_type: str = "") -> str:
            """
            Walk the visible, actionable portion of a window's UIA tree and
            return a compact numbered table the model can use to pick elements
            by index via act_on_element_by_index.

            Visibility rules (all must pass):
              - BoundingRectangle width > 0 AND height > 0
              - Not fully off-screen
              - IsEnabled == True  (grayed-out controls excluded)
              - ControlType in _ACTIONABLE_TYPES (containers excluded)
              - Name or AutomationId not empty

            The snapshot is stored in _snapshot_cache[window_title] so
            act_on_element_by_index can find the element object later.
            """
            hwnd = (
                self._find_shell_hwnd(window_title.strip().lower())
                or self._resolve_app_window(window_title)
                or self._find_window(window_title)
            )
            if not hwnd:
                return (
                    f"Error: Window '{window_title}' not found. "
                    f"Call list_active_windows to see open windows."
                )

            root = auto.ControlFromHandle(hwnd)

            # Short type aliases for the table
            _TYPE_SHORT = {
                "ButtonControl":      "btn",
                "EditControl":        "edit",
                "CheckBoxControl":    "chk",
                "RadioButtonControl": "radio",
                "ComboBoxControl":    "combo",
                "ListItemControl":    "item",
                "MenuItemControl":    "menu",
                "TabItemControl":     "tab",
                "HyperlinkControl":   "link",
                "SliderControl":      "slider",
                "SplitButtonControl": "split",
                "ImageControl":       "img",
                "TextControl":        "text",
            }

            filter_low = filter_type.strip().lower()

            elements = []
            seen_ids  = set()

            try:
                for el, depth in auto.WalkControl(root, maxDepth=self.MAX_INSPECT_DEPTH):
                    try:
                        type_name = el.ControlTypeName
                        name      = (el.Name or "").strip()
                        auto_id   = (el.AutomationId or "").strip()
                    except Exception:
                        continue

                    if type_name not in _ACTIONABLE_TYPES:
                        continue
                    if not (name or auto_id):
                        continue

                    # ── Visibility + enabled check ────────────────────────────
                    try:
                        rect = el.BoundingRectangle
                        if rect.width() <= 0 or rect.height() <= 0:
                            continue
                        if rect.right < 0 or rect.bottom < 0:
                            continue
                        if rect.left > SCREEN_W * 2 or rect.top > SCREEN_H * 2:
                            continue
                    except Exception:
                        continue   # no rect = definitely not visible

                    try:
                        if not el.IsEnabled:
                            continue
                    except Exception:
                        pass   # can't check enabled state — include anyway

                    # ── Optional type filter ──────────────────────────────────
                    short = _TYPE_SHORT.get(type_name, type_name.replace("Control", "").lower())
                    if filter_low and filter_low not in short and filter_low not in type_name.lower():
                        continue

                    # ── Dedup: same name+type at the SAME depth — likely the same
                    # element appearing in multiple subtree walks.
                    # Different depths = different elements (e.g. two "Close" buttons
                    # in VS Code at depth 4 and depth 9 are different controls).
                    dedup_key = f"{name or auto_id}::{type_name}::{depth // 3}"
                    if dedup_key in seen_ids:
                        continue
                    seen_ids.add(dedup_key)

                    # ── State flags ───────────────────────────────────────────
                    flags = []
                    try:
                        if el.HasKeyboardFocus:
                            flags.append("focused")
                    except Exception:
                        pass
                    try:
                        state = el.CurrentToggleState   # 0=off 1=on 2=indeterminate
                        if state == 1:
                            flags.append("checked")
                        elif state == 0:
                            flags.append("unchecked")
                    except Exception:
                        pass
                    try:
                        if el.GetSelectionItemPattern().IsSelected:
                            flags.append("selected")
                    except Exception:
                        pass

                    label = name or auto_id
                    elements.append({
                        "element":  el,
                        "type":     type_name,
                        "short":    short,
                        "name":     name,
                        "auto_id":  auto_id,
                        "label":    label,
                        "flags":    flags,
                        "depth":    depth,
                    })

                    if len(elements) >= 120:   # hard cap — keeps output readable
                        break
            except Exception as e:
                return f"Error walking UI tree: {e}"

            if not elements:
                return (
                    f"No visible, enabled, actionable elements found in '{window_title}'. "
                    f"The window may render via canvas/WebGL — try fallback_click_text."
                )

            # Store for act_on_element_by_index
            self._snapshot_cache[window_title] = elements

            # ── Build compact table ───────────────────────────────────────────
            lines = [
                f"UI snapshot of '{window_title}' — {len(elements)} visible elements",
                f"Use act(target='{window_title}', index=N) to interact.",
                "",
                f"{'IDX':>4}  {'TYPE':<7}  {'STATUS':<12}  NAME",
                "─" * 72,
            ]
            for i, el in enumerate(elements):
                status = ", ".join(el["flags"]) if el["flags"] else "enabled"
                label  = el["label"][:48]   # truncate long names
                lines.append(f"{i:>4}  {el['short']:<7}  {status:<12}  {label}")

            return "\n".join(lines)

        def act_on_element_by_index(self, window_title: str, index: int,
                                     action: str = "click",
                                     text_to_type: str = "") -> str:
            """
            Act on an element from the last snapshot() call by its index.
            Exact — no scoring, no tree walk, no ambiguity.
            """
            cache = self._snapshot_cache.get(window_title)
            if not cache:
                return (
                    f"No snapshot found for '{window_title}'. "
                    f"Call snapshot_ui('{window_title}') first."
                )
            if index < 0 or index >= len(cache):
                return (
                    f"Index {index} out of range (snapshot has {len(cache)} elements, "
                    f"indices 0–{len(cache)-1})."
                )

            entry   = cache[index]
            element = entry["element"]
            label   = entry["label"]
            t_name  = entry["type"]

            # Verify the element is still alive (window may have changed)
            try:
                rect = element.BoundingRectangle
                if rect.width() <= 0 or rect.height() <= 0:
                    return (
                        f"Element #{index} ('{label}') is no longer visible. "
                        f"Call snapshot again to refresh."
                    )
            except Exception:
                return (
                    f"Element #{index} ('{label}') is stale — window may have changed. "
                    f"Call snapshot again to refresh."
                )

            hwnd = (
                self._find_shell_hwnd(window_title.strip().lower())
                or self._resolve_app_window(window_title)
                or self._find_window(window_title)
            )

            print(f"   [Snapshot act] #{index} '{label}' ({t_name}) → {action}")
            return self._act_on_element(element, label, t_name, action, text_to_type, hwnd or 0)

        def find_and_act(self, window_title: str, description: str, action: str = "click",
                          text_to_type: str = "") -> str:
            """
            ONE-CALL UI interaction — finds the best-matching element and acts on it.

            Lookup order:
              1. Blueprint cache  — direct AutomationId lookup, O(1), exact.
              2. Full tree search — scored walk, saves result to blueprint on success.

            Supports shell surfaces: pass window_title as "taskbar", "start",
            "tray", "desktop", "action center", etc.
            """
            # ── Window resolution ─────────────────────────────────────────────────
            hwnd = self._find_shell_hwnd(window_title.strip().lower())
            if not hwnd:
                hwnd = self._find_window(window_title)
            if not hwnd:
                shell_hwnd, _alias = self._resolve_shell_element(description)
                if shell_hwnd:
                    hwnd = shell_hwnd
            if not hwnd:
                return (
                    f"Error: Window '{window_title}' not found. "
                    f"Call list_active_windows to see exact titles. "
                    f"For shell controls use: 'taskbar', 'start', 'tray', "
                    f"'desktop', 'action center', 'search', 'tray overflow'."
                )

            root = auto.ControlFromHandle(hwnd)

            # ── 1. Blueprint fast path ────────────────────────────────────────────
            cached = self._lookup_blueprint(window_title, description)
            if cached:
                el = self._find_by_automation_id(
                    root, cached["automation_id"], cached["type"]
                )
                if el:
                    print(f"   [Blueprint] ✓ Hit for '{description}' "
                          f"(AutomationId='{cached['automation_id']}')")
                    result = self._act_on_element(
                        el, cached["name"], cached["type"],
                        action, text_to_type, hwnd
                    )
                    if "Error" not in result and "failed" not in result.lower():
                        return result
                    # Cache hit but action failed — element may have moved.
                    # Fall through to full search and update the blueprint.
                    print(f"   [Blueprint] Cache hit but action failed — refreshing.")
                else:
                    print(f"   [Blueprint] Cache miss (element gone) — doing full search.")

            # ── 2. Full tree search ───────────────────────────────────────────────
            # Keyword expansion
            _EXPANSIONS = {
                "close":        ["close", "close button", "x"],
                "minimize":     ["minimize", "minimise", "minimize button"],
                "maximize":     ["maximize", "maximise", "restore", "maximize button"],
                "send":         ["send", "send message", "send button", "submit"],
                "search":       ["search", "search box", "search bar", "find", "search field"],
                "address bar":  ["address bar", "address and search bar", "url", "location", "omnibox"],
                "new tab":      ["new tab", "open new tab", "add tab"],
                "back":         ["back", "back button", "navigate back", "go back"],
                "forward":      ["forward", "forward button", "navigate forward"],
                "settings":     ["settings", "preferences", "options", "gear"],
                "menu":         ["menu", "hamburger", "app menu", "main menu"],
                "ok":           ["ok", "okay", "confirm", "yes", "accept"],
                "cancel":       ["cancel", "no", "dismiss", "close"],
                "save":         ["save", "save file", "save document"],
                "open":         ["open", "open file", "browse"],
                "refresh":      ["refresh", "reload", "f5"],
                "new":          ["new", "new file", "new document", "create"],
                "delete":       ["delete", "remove", "trash"],
                "copy":         ["copy", "copy text"],
                "paste":        ["paste", "paste text"],
                "input":        ["input", "text field", "text box", "entry", "edit"],
                "chat input":   ["chat input", "message input", "type a message", "message box",
                                 "enter a prompt", "prompt", "compose"],
            }
            desc_lower    = description.strip().lower()
            desc_variants = {desc_lower}
            for key, variants in _EXPANSIONS.items():
                if key in desc_lower or desc_lower in key:
                    desc_variants.update(variants)
            desc_tokens = set(re.findall(r"[a-z0-9]+", desc_lower))

            candidates = []
            try:
                for element, depth in auto.WalkControl(root, maxDepth=self.MAX_INSPECT_DEPTH):
                    try:
                        type_name = element.ControlTypeName
                        name      = (element.Name or "").strip()
                        auto_id   = (element.AutomationId or "").strip()
                    except Exception:
                        continue
                    if type_name not in _ACTIONABLE_TYPES and type_name not in _CONTAINER_TYPES:
                        continue
                    if not (name or auto_id):
                        continue
                    try:
                        rect = element.BoundingRectangle
                        if rect.width() <= 0 or rect.height() <= 0:
                            continue
                        if rect.right < 0 or rect.bottom < 0:
                            continue
                        if rect.left > SCREEN_W * 2 or rect.top > SCREEN_H * 2:
                            continue
                        visible_area = rect.width() * rect.height()
                    except Exception:
                        visible_area = 0
                    candidates.append({
                        "element": element, "type": type_name,
                        "name": name, "automation_id": auto_id,
                        "depth": depth, "visible_area": visible_area,
                    })
                    if len(candidates) >= 600:
                        break
            except Exception as e:
                return f"Error walking UI tree for '{window_title}': {str(e)}"

            if not candidates:
                return (
                    f"No interactive elements found in '{window_title}'. "
                    f"This window likely renders via canvas/WebGL — UIA cannot see inside it. "
                    f"Try fallback_click_text or fallback_view_screen instead."
                )

            _STRIP_SUFFIXES = [" button", " tab", " field", " box", " bar",
                               " control", " panel", " window", " icon", " link"]

            def _strip_suffix(s: str) -> str:
                for sfx in _STRIP_SUFFIXES:
                    if s.endswith(sfx):
                        return s[:-len(sfx)].strip()
                return s

            _TYPE_WEIGHT = {
                "ButtonControl": 1.5, "EditControl": 1.4, "HyperlinkControl": 1.3,
                "CheckBoxControl": 1.3, "RadioButtonControl": 1.3, "ComboBoxControl": 1.2,
                "MenuItemControl": 1.2, "ListItemControl": 1.1, "TabItemControl": 1.1,
                "SplitButtonControl": 1.1, "SliderControl": 1.0,
                "ImageControl": 0.8, "TextControl": 0.6,
            }

            def score(c) -> float:
                name_l    = c["name"].lower()
                id_l      = c["automation_id"].lower()
                name_core = _strip_suffix(name_l)
                id_core   = _strip_suffix(id_l)
                s = 0.0
                if desc_lower == name_l or desc_lower == name_core:
                    s += 300
                elif desc_lower == id_l or desc_lower == id_core:
                    s += 220
                for variant in desc_variants:
                    if variant == name_l or variant == name_core:
                        s += 180; break
                    if variant == id_l or variant == id_core:
                        s += 140; break
                if len(desc_lower) >= 3:
                    if desc_lower in name_l:   s += 60
                    elif desc_lower in name_core: s += 55
                    if desc_lower in id_l:     s += 30
                    for variant in desc_variants:
                        if len(variant) >= 3 and variant in name_l:
                            s += 40; break
                name_tokens = set(re.findall(r"[a-z0-9]+", name_l))
                id_tokens   = set(re.findall(r"[a-z0-9]+", id_l))
                s += sum(min(len(t) * 3, 15) for t in desc_tokens & name_tokens)
                s += sum(min(len(t) * 1, 5)  for t in desc_tokens & id_tokens)
                type_w = _TYPE_WEIGHT.get(c["type"], 0.5)
                if s > 0:   s *= type_w
                elif c["type"] in _ACTIONABLE_TYPES: s += 5 * type_w
                else:        s -= 15
                if c["visible_area"] > 0:
                    s += min(c["visible_area"] / 5000, 8)
                if c["depth"] <= 2: s -= 10
                else:               s -= c["depth"] * 0.1
                return s

            scored     = sorted(candidates, key=score, reverse=True)
            best       = scored[0]
            best_score = score(best)

            if best_score <= 0:
                by_type: dict[str, list[str]] = {}
                for c in candidates:
                    label = c["name"] or c["automation_id"]
                    if label:
                        by_type.setdefault(c["type"], []).append(f"'{label}'")
                lines = []
                for t in sorted(by_type):
                    items = by_type[t][:8]
                    lines.append(f"  {t}: {', '.join(items)}"
                                 + (f" (+{len(by_type[t])-8} more)" if len(by_type[t]) > 8 else ""))
                return (
                    f"No element matched '{description}' in '{window_title}'.\n"
                    f"All visible elements by type:\n" + "\n".join(lines) + "\n"
                    f"Retry with an exact name from the list above, or use fallback_click_text."
                )

            close_matches = [c for c in scored[1:5] if best_score - score(c) <= 10]
            elem_desc = best["name"] or best["automation_id"] or best["type"]
            if close_matches and best_score < 150:
                alt_names = ", ".join(
                    f"'{c['name'] or c['automation_id']}' ({c['type']})"
                    for c in close_matches[:3]
                )
                print(f"   [UIA] ⚠ Ambiguous: chose '{elem_desc}' ({best['type']}), "
                      f"alternatives: {alt_names}")

            print(f"   [UIA] Best match: '{elem_desc}' ({best['type']}) "
                  f"score={best_score:.1f} depth={best['depth']}")

            # ── Act on the best match ─────────────────────────────────────────────
            result = self._act_on_element(
                best["element"], elem_desc, best["type"],
                action, text_to_type, hwnd
            )

            # ── Save to blueprint on success ──────────────────────────────────────
            if "Success" in result or "invoked" in result or "clicked" in result:
                self._save_to_blueprint(
                    window_title, description,
                    best["element"], best["type"], best["depth"]
                )

            return result

        def _act_on_element(self, element, elem_desc: str, type_name: str,
                             action: str, text_to_type: str, hwnd: int) -> str:
            """
            Perform the requested action on a resolved UIA element.
            Extracted from find_and_act so both the blueprint fast-path and
            the full search path share the same action logic.
            """
            # ── get_text ──────────────────────────────────────────────────────────
            if action == "get_text":
                try:
                    val = element.GetValuePattern().Value
                    if val:
                        return f"Text of '{elem_desc}': {val}"
                except Exception:
                    pass
                try:
                    txt = element.GetTextPattern().DocumentRange.GetText(-1)
                    if txt:
                        return f"Text of '{elem_desc}': {txt}"
                except Exception:
                    pass
                return f"Text of '{elem_desc}': {element.Name}"

            # ── set_text ──────────────────────────────────────────────────────────
            if action == "set_text":
                try:
                    element.GetValuePattern().SetValue(text_to_type)
                    return f"Success: set text of '{elem_desc}' to '{text_to_type[:40]}'"
                except Exception:
                    pass
                try:
                    element.SetValue(text_to_type)
                    return f"Success: set text of '{elem_desc}' (SetValue)."
                except Exception:
                    pass
                try:
                    rect = element.BoundingRectangle
                    cx = rect.left + rect.width() // 2
                    cy = rect.top  + rect.height() // 2
                    _do_click(cx, cy, "left_click", label=f"focus '{elem_desc}'")
                    time.sleep(0.15)
                    import pyperclip
                    pyperclip.copy(text_to_type)
                    ps = (
                        "Add-Type -AssemblyName System.Windows.Forms\n"
                        "[System.Windows.Forms.SendKeys]::SendWait('^a')\n"
                        "Start-Sleep -Milliseconds 50\n"
                        "[System.Windows.Forms.SendKeys]::SendWait('^v')"
                    )
                    execute_terminal_command(ps)
                    return f"Success: set text of '{elem_desc}' via clipboard paste."
                except Exception:
                    pass
                try:
                    element.Click(simulateMove=False)
                    time.sleep(0.1)
                    type_text(text_to_type)
                    return f"Set text via click+type on '{elem_desc}'."
                except Exception as e:
                    return f"Error: all set_text methods failed on '{elem_desc}': {e}"

            # ── click ─────────────────────────────────────────────────────────────
            # (a) InvokePattern
            try:
                element.GetInvokePattern().Invoke()
                return f"Success: invoked '{elem_desc}' ({type_name}) via UIA InvokePattern."
            except Exception:
                pass
            # (b) TogglePattern
            try:
                element.GetTogglePattern().Toggle()
                return f"Success: toggled '{elem_desc}' ({type_name}) via UIA TogglePattern."
            except Exception:
                pass
            # (c) SelectionItemPattern
            try:
                element.GetSelectionItemPattern().Select()
                return f"Success: selected '{elem_desc}' ({type_name}) via UIA SelectionItemPattern."
            except Exception:
                pass
            # (d) Coordinate click — universal fallback
            try:
                rect = element.BoundingRectangle
                if rect.width() > 0 and rect.height() > 0:
                    cx = rect.left + rect.width()  // 2
                    cy = rect.top  + rect.height() // 2
                    try:
                        import win32con
                        win32gui.ShowWindow(hwnd, win32con.SW_RESTORE)
                        win32gui.SetForegroundWindow(hwnd)
                        time.sleep(0.08)
                    except Exception:
                        pass
                    return _do_click(cx, cy, "left_click",
                                     label=f"UIA-coord '{elem_desc}'")
            except Exception:
                pass
            # (e) UIA Click — last resort
            try:
                try:
                    element.SetFocus()
                except Exception:
                    pass
                element.Click(simulateMove=False)
                return f"Success: clicked '{elem_desc}' ({type_name}) via UIA Click."
            except Exception as e:
                return (
                    f"Error: all click methods failed for '{elem_desc}' ({type_name}). "
                    f"Last error: {e}. Try fallback_click_text if it has visible text."
                )
        def read_aggregated_text(self, window_title: str, container_key: str = None) -> str:
            """
            Aggregates hundreds of small TextControl siblings into readable paragraphs.
            Resolves window by app name so dynamic titles don't break it.
            """
            hwnd = (
                self._find_shell_hwnd(window_title.strip().lower())
                or self._resolve_app_window(window_title)
                or self._find_window(window_title)
            )
            if not hwnd:
                # Tell the model what the window is actually called now
                current_titles = []
                def cb(h, _):
                    if win32gui.IsWindowVisible(h):
                        t = win32gui.GetWindowText(h).strip()
                        if t:
                            current_titles.append(t)
                win32gui.EnumWindows(cb, None)
                suggestions = [t for t in current_titles if window_title.split("-")[-1].strip().lower() in t.lower()]
                hint = f" Did you mean one of: {suggestions[:3]}?" if suggestions else ""
                return (
                    f"Window '{window_title}' not found (title may have changed).{hint} "
                    f"Call list_active_windows to see current titles, then retry."
                )
            
            root = auto.ControlFromHandle(hwnd)
            # If a specific container was requested, use that as the root
            search_root = self._live_cache.get(window_title, {}).get(container_key, root)
            
            aggregated_lines = []
            last_y = -1
            current_line = []
            
            # Walk the subtree, collecting TextControls
            for element, depth in auto.WalkControl(search_root, maxDepth=self.MAX_INSPECT_DEPTH):
                if element.ControlTypeName == "TextControl":
                    text = (element.Name or "").strip()
                    if not text: continue
                    
                    try:
                        rect = element.BoundingRectangle
                        # Logic: If elements are on the same Y-level (roughly), 
                        # treat them as part of the same line/paragraph
                        if abs(rect.top - last_y) < 15:
                            current_line.append(text)
                        else:
                            if current_line:
                                aggregated_lines.append(" ".join(current_line))
                            current_line = [text]
                            last_y = rect.top
                    except:
                        current_line.append(text)
            
            if current_line:
                aggregated_lines.append(" ".join(current_line))
                
            return "\n".join(aggregated_lines)

        def query_gemini_app(self, prompt: str, wait_for_response: int = 90) -> str:
            """
            Thin delegator kept ONLY so external dispatchers that call tools
            as getattr(ui_navigator, tool_name)(**args) (e.g. gui.pyw) keep
            working. Does NOT use UIA — routes straight to the module-level,
            CDP/browser-based query_gemini_app(), same as the CLI dispatcher
            in this file uses directly.
            """
            return query_gemini_app(prompt, wait_for_response=wait_for_response)

        def manage_gemini_chat(self, action: str, chat_name: str = None) -> str:
            window_title = "Gemini"
            hwnd = win32gui.FindWindow(None, window_title)
            if not hwnd: return "Error: Gemini window not found."
            
            root = auto.ControlFromHandle(hwnd)
            
            if action == "new_chat":
                # Search for the "New chat" button/icon
                for element, _ in auto.WalkControl(root):
                    if element.Name and "New chat" in element.Name:
                        element.Click()
                        return "Started a new Gemini chat."
                return "Error: Could not find 'New chat' button."
                
            elif action == "open_recent" and chat_name:
                # Search for the chat in the sidebar list
                for element, _ in auto.WalkControl(root):
                    if element.Name and chat_name.lower() in element.Name.lower():
                        element.Click()
                        return f"Opened recent chat: {chat_name}."
                return f"Error: Could not find recent chat named '{chat_name}'."
                
            return "Invalid action."
    ui_navigator = AppMapNavigator()
else:
    ui_navigator = None

# =============================================================================
# LINUX UI AUTOMATION — AppMapNavigatorLinux
# =============================================================================
# Mirrors every public method of AppMapNavigator but uses:
#   - AT-SPI2 via pyatspi for accessibility tree walking (replaces uiautomation)
#   - python-xlib / subprocess xdotool for mouse/keyboard (replaces win32gui + SendKeys)
#   - xdotool for window finding and focusing (replaces win32gui.FindWindow)
#   - xclip/xsel via subprocess for clipboard (replaces pyperclip on Wayland/X11)
#
# INSTALLATION (Debian/Ubuntu/Fedora):
#   sudo apt install python3-pyatspi xdotool xclip   # Debian/Ubuntu
#   sudo dnf install at-spi2-core xdotool xclip      # Fedora
#   pip install pyatspi                               # Python binding
#
# AT-SPI2 must be enabled in your desktop session. It is on by default in
# GNOME. For KDE/XFCE run: gsettings set org.gnome.desktop.interface
# toolkit-accessibility true  (or enable in Accessibility settings).
#
# Wayland note: xdotool works under XWayland for most apps. For native
# Wayland windows you may need ydotool (requires root or uinput group).

_PYATSPI_AVAILABLE = False
if _IS_LINUX:
    try:
        import pyatspi  # type: ignore[import-untyped]  # Linux-only package
        _PYATSPI_AVAILABLE = True
    except ImportError:
        pass

# AT-SPI role names that correspond to containers (mirrors _CONTAINER_TYPES)
_ATSPI_CONTAINER_ROLES = {
    "panel", "filler", "scroll pane", "split pane", "layered pane",
    "frame", "window", "dialog", "tool bar", "menu bar", "page tab list",
    "page tab", "tree", "tree table", "table",
}
# AT-SPI role names that are interactive (mirrors _ACTIONABLE_TYPES)
_ATSPI_ACTIONABLE_ROLES = {
    "push button", "toggle button", "check box", "radio button",
    "text", "entry", "password text", "combo box", "list item",
    "menu item", "check menu item", "radio menu item", "link",
    "slider", "spin button", "image", "label",
}


def _run(cmd: list, timeout: int = 10) -> tuple[str, str]:
    """Run a subprocess command, return (stdout, stderr)."""
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
        return r.stdout.strip(), r.stderr.strip()
    except Exception as e:
        return "", str(e)


class AppMapNavigatorLinux:
    """
    Linux equivalent of AppMapNavigator.

    Uses AT-SPI2 for accessibility tree inspection and xdotool for all
    mouse/keyboard simulation. The public API is identical to AppMapNavigator
    so the rest of Midum (tool dispatch, tool schemas) requires no changes —
    just swap which navigator is assigned to ui_navigator at startup.

    Window identification uses xdotool search --name (substring, case-insensitive)
    mirroring the Windows fuzzy title match.

    Coordinate system: all click coordinates are real screen pixels. There is no
    canvas scaling layer on Linux — xdotool takes absolute screen coordinates.
    """

    MAX_DISCOVER_DEPTH = 12
    MAX_INSPECT_DEPTH  = 12
    MAX_RESULTS        = 60

    def __init__(self, maps_dir="storage/app_maps_linux"):
        self.maps_dir = os.path.abspath(maps_dir)
        os.makedirs(self.maps_dir, exist_ok=True)
        self._live_cache: dict = {}   # window_title -> {key -> atspi node}

    # ── Window finding ────────────────────────────────────────────────────────

    def _find_window_id(self, window_title: str) -> str | None:
        """
        Return the first xdotool window ID (as string) whose name contains
        window_title (case-insensitive). Returns None if not found.
        """
        stdout, _ = _run(["xdotool", "search", "--name", window_title])
        ids = [line.strip() for line in stdout.splitlines() if line.strip()]
        if not ids:
            return None
        # Prefer the focused window among matches
        focused_stdout, _ = _run(["xdotool", "getactivewindow"])
        focused = focused_stdout.strip()
        if focused in ids:
            return focused
        return ids[0]

    def _find_atspi_window(self, window_title: str):
        """
        Walk the AT-SPI desktop tree looking for a top-level window whose
        name contains window_title (case-insensitive). Returns the AT-SPI
        Accessible or None.
        """
        if not _PYATSPI_AVAILABLE:
            return None
        try:
            desktop = pyatspi.Registry.getDesktop(0)
            low = window_title.lower()
            # Collect all matches
            matches = []
            for app in desktop:
                if app is None:
                    continue
                try:
                    for win in app:
                        if win is None:
                            continue
                        name = (win.name or "").lower()
                        if low in name:
                            matches.append(win)
                except Exception:
                    continue
            if not matches:
                return None
            if len(matches) == 1:
                return matches[0]
            # Prefer focused window
            try:
                focused = pyatspi.Registry.getDesktop(0)   # re-fetch for state
                for w in matches:
                    try:
                        state = w.getState()
                        if state.contains(pyatspi.STATE_ACTIVE):
                            return w
                    except Exception:
                        pass
            except Exception:
                pass
            return matches[0]
        except Exception:
            return None

    # ── AT-SPI tree walking ───────────────────────────────────────────────────

    def _walk_atspi(self, node, max_depth: int, _depth: int = 0):
        """Generator: yield (node, depth) for the entire subtree."""
        if node is None or _depth > max_depth:
            return
        yield node, _depth
        try:
            for i in range(node.childCount):
                child = node.getChildAtIndex(i)
                yield from self._walk_atspi(child, max_depth, _depth + 1)
        except Exception:
            pass

    def _role_name(self, node) -> str:
        try:
            return node.getRole().name.lower().replace("_", " ")
        except Exception:
            return ""

    # ── Public API ────────────────────────────────────────────────────────────

    def discover_ui_subtrees(self, window_title: str) -> list[dict]:
        """
        Scan the AT-SPI tree of window_title and return a list of named
        container nodes — equivalent to AppMapNavigator.discover_ui_subtrees.
        """
        win = self._find_atspi_window(window_title)
        if not win:
            return []

        if window_title not in self._live_cache:
            self._live_cache[window_title] = {}

        containers = []
        seen_keys  = set()

        for node, depth in self._walk_atspi(win, self.MAX_DISCOVER_DEPTH):
            role = self._role_name(node)
            if role not in _ATSPI_CONTAINER_ROLES:
                continue
            name    = (node.name or "").strip()
            desc    = ""
            try:
                desc = (node.description or "").strip()
            except Exception:
                pass
            if not (name or desc):
                continue

            key = f"{name or desc}::{role}::{depth}"
            if key in seen_keys:
                continue
            seen_keys.add(key)
            self._live_cache[window_title][key] = node
            containers.append({
                "subtree_key":   key,
                "type":          role,
                "name":          name,
                "description":   desc,
                "depth":         depth,
            })
            if len(containers) >= self.MAX_RESULTS:
                break

        return containers

    def inspect_subtree_controls(self, window_title: str, subtree_key: str) -> list[dict]:
        """
        Return all actionable controls under subtree_key — equivalent to
        AppMapNavigator.inspect_subtree_controls.
        """
        app_cache    = self._live_cache.get(window_title, {})
        subtree_root = app_cache.get(subtree_key)
        if not subtree_root:
            self.discover_ui_subtrees(window_title)
            subtree_root = self._live_cache.get(window_title, {}).get(subtree_key)
            if not subtree_root:
                return []

        controls   = []
        seen_names = set()

        for node, depth in self._walk_atspi(subtree_root, self.MAX_INSPECT_DEPTH):
            role = self._role_name(node)
            if role not in _ATSPI_ACTIONABLE_ROLES:
                continue
            name = (node.name or "").strip()
            desc = ""
            try:
                desc = (node.description or "").strip()
            except Exception:
                pass
            if not (name or desc):
                continue
            dedup = f"{name}::{desc}::{role}"
            if dedup in seen_names:
                continue
            seen_names.add(dedup)
            controls.append({"type": role, "name": name, "description": desc, "depth": depth})
            if len(controls) >= self.MAX_RESULTS:
                break

        return controls

    def _score_node(self, node, desc_lower: str, desc_tokens: set, depth: int) -> float:
        """Score an AT-SPI node against a plain-English description."""
        name_l = (node.name or "").lower()
        try:
            node_desc_l = (node.description or "").lower()
        except Exception:
            node_desc_l = ""
        role = self._role_name(node)
        s = 0.0

        if desc_lower == name_l:
            s += 200
        elif desc_lower == node_desc_l:
            s += 150

        if len(desc_lower) >= 3:
            if desc_lower in name_l:
                s += 40
            if desc_lower in node_desc_l:
                s += 20

        name_tokens = set(re.findall(r"[a-z0-9]+", name_l))
        desc_tok    = set(re.findall(r"[a-z0-9]+", node_desc_l))
        s += 8 * len(desc_tokens & name_tokens)
        s += 3 * len(desc_tokens & desc_tok)

        if role in _ATSPI_ACTIONABLE_ROLES:
            s += 30
        else:
            s -= 20

        s -= depth * 0.2
        return s

    def _get_node_center(self, node) -> tuple[int, int] | None:
        """Return screen (x, y) center of node's bounding box, or None."""
        try:
            comp = node.queryComponent()
            box  = comp.getExtents(pyatspi.DESKTOP_COORDS)
            if box.width <= 0 or box.height <= 0:
                return None
            return box.x + box.width // 2, box.y + box.height // 2
        except Exception:
            return None

    def _do_click_linux(self, x: int, y: int, click_type: str = "left_click") -> str:
        """Click at absolute screen coordinates using xdotool."""
        button = {"left_click": "1", "right_click": "3", "double_click": "1"}.get(click_type, "1")
        cmds = (
            ["xdotool", "mousemove", "--sync", str(x), str(y)],
            ["xdotool", "click", "--clearmodifiers", button],
        )
        if click_type == "double_click":
            cmds = (
                ["xdotool", "mousemove", "--sync", str(x), str(y)],
                ["xdotool", "click", "--clearmodifiers", "--repeat", "2", button],
            )
        for cmd in cmds:
            _, err = _run(cmd)
            if err:
                return f"{click_type} at ({x},{y}) — warning: {err[:100]}"
        return f"Success: {click_type} at screen ({x},{y})"

    def find_and_act(self, window_title: str, description: str,
                     action: str = "click", text_to_type: str = "") -> str:
        """
        One-call UI interaction — mirrors AppMapNavigator.find_and_act.

        1. Find the AT-SPI window.
        2. Walk the entire tree, score every node against description.
        3. Act on the best match:
           click     → try AT-SPI DoAction("click"), fall back to xdotool coordinate click
           set_text  → try AT-SPI SetValue, fall back to focus + xdotool type
           get_text  → AT-SPI GetText or Name
        """
        if not _PYATSPI_AVAILABLE:
            return "Error: pyatspi not installed. Run: pip install pyatspi"

        win = self._find_atspi_window(window_title)
        if not win:
            return (
                f"Error: Window '{window_title}' not found. "
                f"Run list_active_windows to see open window titles."
            )

        desc_lower  = description.strip().lower()
        desc_tokens = set(re.findall(r"[a-z0-9]+", desc_lower))

        candidates = []
        try:
            for node, depth in self._walk_atspi(win, self.MAX_INSPECT_DEPTH):
                role = self._role_name(node)
                if role not in _ATSPI_ACTIONABLE_ROLES and role not in _ATSPI_CONTAINER_ROLES:
                    continue
                name = (node.name or "").strip()
                try:
                    ndesc = (node.description or "").strip()
                except Exception:
                    ndesc = ""
                if not (name or ndesc):
                    continue
                candidates.append((node, depth))
                if len(candidates) >= 400:
                    break
        except Exception as e:
            return f"Error walking AT-SPI tree for '{window_title}': {e}"

        if not candidates:
            return (
                f"No interactive elements found in '{window_title}'. "
                f"The app may not expose an AT-SPI tree. "
                f"Try fallback_click_text instead."
            )

        scored = sorted(candidates, key=lambda nd: self._score_node(nd[0], desc_lower, desc_tokens, nd[1]), reverse=True)
        best_node, best_depth = scored[0]
        best_score = self._score_node(best_node, desc_lower, desc_tokens, best_depth)

        if best_score <= 0:
            sample = ", ".join(
                f"'{n.name}' ({self._role_name(n)})"
                for n, _ in scored[:25]
                if (n.name or "").strip()
            )
            return (
                f"No element matched '{description}' in '{window_title}'. "
                f"Available elements include: {sample}. "
                f"Retry with one of these exact names."
            )

        elem_desc = (best_node.name or "").strip() or self._role_name(best_node)

        # ── get_text ──────────────────────────────────────────────────────────
        if action == "get_text":
            try:
                text_iface = best_node.queryText()
                return f"Text of '{elem_desc}': {text_iface.getText(0, -1)}"
            except Exception:
                return f"Text of '{elem_desc}': {best_node.name}"

        # ── set_text ──────────────────────────────────────────────────────────
        if action == "set_text":
            # Try AT-SPI EditableText interface first
            try:
                edit = best_node.queryEditableText()
                edit.setTextContents(text_to_type)
                return f"Success: set text of '{elem_desc}' to '{text_to_type[:40]}'"
            except Exception:
                pass
            # Fall back: focus element, clear, type via xdotool
            center = self._get_node_center(best_node)
            if center:
                self._do_click_linux(*center)
                time.sleep(0.1)
                # Select all + delete existing text
                _run(["xdotool", "key", "--clearmodifiers", "ctrl+a"])
                time.sleep(0.05)
                _run(["xdotool", "key", "--clearmodifiers", "Delete"])
                time.sleep(0.05)
                _, err = _run(["xdotool", "type", "--clearmodifiers", "--delay", "20", text_to_type])
                if err:
                    return f"set_text via xdotool on '{elem_desc}' — warning: {err[:100]}"
                return f"Success: typed into '{elem_desc}' via xdotool"
            return f"Error: Could not set text on '{elem_desc}' — no bounding box."

        # ── click ─────────────────────────────────────────────────────────────
        # (a) AT-SPI DoAction "click"
        try:
            actions = best_node.queryAction()
            for i in range(actions.nActions):
                if actions.getName(i).lower() in ("click", "activate", "press"):
                    actions.doAction(i)
                    return f"Success: invoked '{elem_desc}' (AT-SPI DoAction)."
        except Exception:
            pass

        # (b) Coordinate click via xdotool
        center = self._get_node_center(best_node)
        if center:
            # Focus the window first
            win_id = self._find_window_id(window_title)
            if win_id:
                _run(["xdotool", "windowfocus", "--sync", win_id])
                time.sleep(0.1)
            return self._do_click_linux(*center, click_type=action if action in ("left_click", "right_click", "double_click") else "left_click")

        return (
            f"Found '{elem_desc}' but could not determine its screen position. "
            f"Try fallback_click_text instead."
        )

    def safely_trigger_ui_element(self, window_title: str, control_type: str,
                                   search_property: str, property_value: str,
                                   action: str, text_to_type: str = "") -> str:
        """
        Manual precise interaction — mirrors AppMapNavigator.safely_trigger_ui_element.
        Searches by name or description match, then acts.
        """
        win = self._find_atspi_window(window_title)
        if not win:
            return f"Error: Window '{window_title}' not found."

        target = property_value.lower()
        best   = None

        for node, _ in self._walk_atspi(win, self.MAX_INSPECT_DEPTH):
            name = (node.name or "").lower()
            try:
                ndesc = (node.description or "").lower()
            except Exception:
                ndesc = ""
            if search_property == "name" and target in name:
                best = node
                break
            if search_property in ("automation_id", "class_name") and target in ndesc:
                best = node
                break

        if not best:
            return (
                f"Error: Could not locate {control_type} matching "
                f"{search_property}='{property_value}' in '{window_title}'."
            )

        return self.find_and_act(window_title, best.name or property_value, action, text_to_type)

    def read_aggregated_text(self, window_title: str, container_key: str = None) -> str:
        """
        Aggregate all text from a window's AT-SPI tree into readable lines —
        mirrors AppMapNavigator.read_aggregated_text.
        """
        win = self._find_atspi_window(window_title)
        if not win:
            return "Window not found."

        root = win
        if container_key:
            root = self._live_cache.get(window_title, {}).get(container_key, win)

        lines_by_y: dict[int, list[str]] = {}

        for node, _ in self._walk_atspi(root, self.MAX_INSPECT_DEPTH):
            role = self._role_name(node)
            if role not in ("label", "text", "static text"):
                continue
            text = (node.name or "").strip()
            if not text:
                continue
            y = 0
            try:
                comp = node.queryComponent()
                box  = comp.getExtents(pyatspi.DESKTOP_COORDS)
                y    = box.y
            except Exception:
                pass
            # Group text on roughly the same Y line (within 12px)
            bucket = next((k for k in lines_by_y if abs(k - y) < 12), y)
            lines_by_y.setdefault(bucket, []).append(text)

        if not lines_by_y:
            return ""
        return "\n".join(" ".join(words) for _, words in sorted(lines_by_y.items()))

    def query_gemini_app(self, prompt: str, wait_for_response: int = 90) -> str:
        """
        Thin delegator kept ONLY so external dispatchers that call tools as
        getattr(ui_navigator, tool_name)(**args) (e.g. gui.pyw) keep working.
        Does NOT use AT-SPI/xdotool — routes straight to the module-level,
        CDP/browser-based query_gemini_app(), same as the CLI dispatcher in
        this file uses directly.
        """
        return query_gemini_app(prompt, wait_for_response=wait_for_response)

    def manage_gemini_chat(self, action: str, chat_name: str = None) -> str:
        """Mirrors AppMapNavigator.manage_gemini_chat."""
        window_title = "Gemini"
        win = self._find_atspi_window(window_title)
        if not win:
            return "Error: Gemini window not found."

        if action == "new_chat":
            result = self.find_and_act(window_title, "New chat", action="click")
            return result if "Success" in result else "Error: Could not find 'New chat' button."

        if action == "open_recent" and chat_name:
            result = self.find_and_act(window_title, chat_name, action="click")
            return result if "Success" in result else f"Error: Could not find recent chat '{chat_name}'."

        return "Invalid action."

    def _get_map_path(self, window_title: str) -> str:
        safe = "".join(c if c.isalnum() else "_" for c in window_title.lower()).strip("_")
        return os.path.join(self.maps_dir, f"{safe}.json")

    def load_app_blueprint(self, window_title: str) -> dict:
        path = self._get_map_path(window_title)
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return {"window_title": window_title, "subtrees": {}, "known_controls": {}}

    def save_app_blueprint(self, window_title: str, blueprint: dict):
        path = self._get_map_path(window_title)
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(blueprint, f, indent=2)
        except Exception:
            pass


# ── Pick the right navigator based on OS ─────────────────────────────────────
if _IS_LINUX and _PYATSPI_AVAILABLE:
    ui_navigator = AppMapNavigatorLinux()
    print("🐧 [Linux UI navigator: AppMapNavigatorLinux (AT-SPI2 + xdotool)]")
elif _IS_LINUX and not _PYATSPI_AVAILABLE:
    ui_navigator = AppMapNavigatorLinux()   # still usable for xdotool fallbacks
    print("🐧 [Linux UI navigator: AppMapNavigatorLinux (xdotool only — install pyatspi for full tree access)]")
elif _UIA_AVAILABLE:
    pass   # already set above: ui_navigator = AppMapNavigator()
else:
    ui_navigator = None

# =============================================================================
# GEMINI SETUP (UPDATED FOR GOOGLE-GENAI & FREE TIER PROTECTION)
# =============================================================================
#
# INSTALLATION:
#   pip install google-genai
#
# API KEY SETUP (one-time):
#   1. Get a free key at https://aistudio.google.com/app/apikey
#   2. Create the secrets file at the path shown in SECRETS_FILE above.
#      Example:  { "GEMINI_API_KEY": "AIza..." }
#
# CREDIT DEFENSE MECHANISM:
#   To prevent running out of requests or hitting tight Token Per Minute (TPM) 
#   limits immediately on the free tier, we default routing to gemini-2.0-flash 
#   and window large text contexts safely.
#
_GEMINI_AVAILABLE = False
_gemini_client    = None

def _load_gemini():
    """
    Load API key from secrets file and initialise the unified GenAI client.

    NOTE: as of this rewrite, the Gemini API is NOT used by consult_gemini
    or query_gemini_app anymore — both route exclusively through the free
    web chat interface (gemini.google.com/app) via browser/CDP automation,
    since the API's free-tier request/token-per-minute limits are far
    tighter than the web chat UI's own limits. This loader is kept only
    so _GEMINI_AVAILABLE can still be reported in the startup banner if a
    key happens to be configured; nothing in the active code path calls
    into _gemini_client.
    """
    global _GEMINI_AVAILABLE, _gemini_client
    try:
        from google import genai
        secrets_path = os.path.abspath(SECRETS_FILE)
        if not os.path.exists(secrets_path):
            return False, f"Secrets file not found: {secrets_path}"
        with open(secrets_path, "r", encoding="utf-8") as f:
            secrets = json.load(f)
        key = secrets.get("GEMINI_API_KEY", "").strip()
        if not key:
            return False, "GEMINI_API_KEY is empty in secrets file."
        
        # Initialize modern unified client
        _gemini_client = genai.Client(api_key=key)
        _GEMINI_AVAILABLE = True
        return True, "OK"
    except ImportError:
        return False, "google-genai not installed. Run: pip install google-genai"
    except Exception as e:
        return False, str(e)


# Attempt load at module import time; failure is non-fatal
_gemini_load_ok, _gemini_load_msg = _load_gemini()


def consult_gemini(prompt, task_type="auto", context=""):
    """
    Send a prompt to Gemini via the actual WEB CHAT INTERFACE
    (https://gemini.google.com/app), using the community `gemini_webapi`
    library in query_gemini_app() — cookie-based, no browser automation
    or UIA involved. This intentionally does NOT fall back to the metered
    Gemini API — the API's free-tier request/token-per-minute limits are
    far tighter than what the web chat UI itself allows, so a silent
    fallback would burn through API credits without the caller ever
    knowing. If the web route fails, this returns a clear error instead
    so you can retry rather than unknowingly hitting the API.

    task_type is accepted for backward compatibility but has no effect
    now that there's no model to route between — it's always whatever
    model gemini.google.com/app is currently serving.
    """
    full_prompt = prompt.strip()
    if context.strip():
        if len(context) > 60000:
            context = context[:30000] + "\n... [TRUNCATED] ...\n" + context[-30000:]
        full_prompt = "[CONTEXT]\n" + context.strip() + "\n\n[TASK]\n" + full_prompt

    if not _GEMINI_WEBAPI_AVAILABLE:
        return f"Error: Gemini web chat is unavailable — {_gemini_webapi_load_msg}"

    print("   [Gemini web] Sending consult request...")
    try:
        result = query_gemini_app(full_prompt)
    except Exception as e:
        return f"Error: Gemini web chat request failed: {e}"

    if result and not result.startswith("Error"):
        print(f"   [Gemini web] Response received ({len(result)} chars)")
        return "[Gemini/web]\n" + result

    return f"Error: Gemini web chat failed — {result}"


# =============================================================================
# OPENROUTER SETUP
# =============================================================================
#
# INSTALLATION:
#   pip install requests   (already required elsewhere in this file)
#
# API KEY SETUP (one-time):
#   1. Get a key at https://openrouter.ai/keys
#   2. Add it to the SAME secrets file used for Gemini:
#      { "GEMINI_API_KEY": "...", "OPENROUTER_API_KEY": "sk-or-v1-..." }
#
# OpenRouter exposes an OpenAI-compatible /chat/completions endpoint, so this
# uses plain `requests` rather than a dedicated SDK — one less dependency.
#
_OPENROUTER_AVAILABLE = False
_OPENROUTER_API_KEY   = None

def _load_openrouter():
    """Load OpenRouter API key from the shared secrets file."""
    global _OPENROUTER_AVAILABLE, _OPENROUTER_API_KEY
    try:
        secrets_path = os.path.abspath(SECRETS_FILE)
        if not os.path.exists(secrets_path):
            return False, f"Secrets file not found: {secrets_path}"
        with open(secrets_path, "r", encoding="utf-8") as f:
            secrets = json.load(f)
        key = secrets.get("OPENROUTER_API_KEY", "").strip()
        if not key:
            return False, "OPENROUTER_API_KEY is empty in secrets file."
        _OPENROUTER_API_KEY   = key
        _OPENROUTER_AVAILABLE = True
        return True, "OK"
    except Exception as e:
        return False, str(e)


_openrouter_load_ok, _openrouter_load_msg = _load_openrouter()


def _openrouter_chat(messages: list, model: str = None, tools_schema: list = None,
                      timeout: int = 60, _retries: int = 4) -> dict:
    """
    Call OpenRouter's OpenAI-compatible /chat/completions endpoint.
    Returns a dict normalised to the SAME shape _call_ollama produces:
        {"message": {"role": "assistant", "content": str, "tool_calls": [...] }}
    so process_chat_turn can treat Ollama and OpenRouter identically.

    Retries on 429 (rate limit) and 502/503 (upstream overloaded) — both are
    common on OpenRouter's free-tier models, which is the primary use case
    for OPENROUTER_CONSULT_MODE="always". Raises RuntimeError with the
    ACTUAL error message from OpenRouter's response body (not just a generic
    HTTP status), since that body usually says exactly what went wrong
    (rate-limited, model overloaded, invalid key, etc).
    """
    if not _OPENROUTER_AVAILABLE or not _OPENROUTER_API_KEY:
        raise RuntimeError(
            f"OpenRouter not available: {_openrouter_load_msg}. "
            f"Add OPENROUTER_API_KEY to {SECRETS_FILE}"
        )
    if requests is None:
        raise RuntimeError("requests library not installed: pip install requests")

    model = model or OPENROUTER_MODEL

    payload = {
        "model": model,
        "messages": messages,
    }
    if tools_schema:
        payload["tools"] = tools_schema

    headers = {
        "Authorization": f"Bearer {_OPENROUTER_API_KEY}",
        "Content-Type":  "application/json",
        "HTTP-Referer":  "https://github.com/jarvis-local-agent",
        "X-Title":       "Midum Desktop Agent",
    }

    last_err = None
    for attempt in range(_retries + 1):
        try:
            resp = requests.post(
                f"{OPENROUTER_API_BASE}/chat/completions",
                headers=headers, json=payload, timeout=timeout
            )

            # Pull the response body regardless of status — OpenRouter puts
            # the real error reason in JSON even on 4xx/5xx responses.
            try:
                data = resp.json()
            except Exception:
                data = {}

            if resp.status_code != 200:
                err_obj = data.get("error", {}) if isinstance(data, dict) else {}
                err_msg = err_obj.get("message") or resp.text[:200] or f"HTTP {resp.status_code}"

                retryable = resp.status_code in (429, 502, 503, 504)
                if retryable and attempt < _retries:
                    # Exponential backoff with a cap, plus small jitter, since
                    # free-tier 429s are a SHARED rate-limit pool — a fixed
                    # short wait often isn't enough for it to clear.
                    import random as _random
                    wait_s = min(2.0 * (2 ** attempt), 20.0) + _random.uniform(0, 0.75)
                    print(f"   [OpenRouter] {resp.status_code} ({err_msg[:80]}) — "
                          f"retrying in {wait_s:.1f}s... (attempt {attempt + 1}/{_retries + 1})")
                    time.sleep(wait_s)
                    last_err = err_msg
                    continue

                raise RuntimeError(f"OpenRouter API error ({resp.status_code}): {err_msg}")

            # 200 OK but the API-level payload can still carry an error object
            if isinstance(data, dict) and data.get("error"):
                raise RuntimeError(f"OpenRouter error: {data['error'].get('message', data['error'])}")

            break   # success

        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            last_err = str(e)
            if attempt < _retries:
                wait_s = 1.5 * (attempt + 1)
                print(f"   [OpenRouter] Connection issue — retrying in {wait_s:.1f}s...")
                time.sleep(wait_s)
                continue
            raise RuntimeError(f"OpenRouter connection failed after {_retries + 1} attempts: {last_err}")
    else:
        raise RuntimeError(f"OpenRouter failed after {_retries + 1} attempts: {last_err}")

    choice  = (data.get("choices") or [{}])[0]
    msg     = choice.get("message", {}) or {}
    content = msg.get("content") or ""
    raw_tool_calls = msg.get("tool_calls") or []

    # Normalise tool_calls to the same shape Ollama's client produces:
    # [{"function": {"name": str, "arguments": dict}}]
    normalised_calls = []
    for tc in raw_tool_calls:
        fn = tc.get("function", {})
        args = fn.get("arguments", {})
        if isinstance(args, str):
            try:
                args = json.loads(args)
            except Exception:
                pass
        normalised_calls.append({"function": {"name": fn.get("name", ""), "arguments": args}})

    return {
        "message": {
            "role": "assistant",
            "content": content,
            "tool_calls": normalised_calls,
        }
    }


def _openrouter_chat_with_fallback(messages: list, model: str = None, tools_schema: list = None,
                                    timeout: int = 60) -> dict:
    """
    Wraps _openrouter_chat() with automatic fallback across
    OPENROUTER_FALLBACK_MODELS. A single free-tier model can 429 for
    minutes at a time because the rate limit is shared across ALL
    OpenRouter users of that model, not just you — retrying the SAME
    model harder doesn't help. This tries the requested model first
    (with its own internal retries), and if it still fails with a
    retryable error, moves on to the next model in the fallback list
    before giving up.
    """
    requested = model or OPENROUTER_MODEL
    chain = [requested] + [m for m in OPENROUTER_FALLBACK_MODELS if m != requested]

    last_err = None
    for i, m in enumerate(chain):
        try:
            resp = _openrouter_chat(messages, model=m, tools_schema=tools_schema, timeout=timeout)
            if i > 0:
                print(f"   [OpenRouter] Recovered using fallback model: {m}")
            return resp
        except RuntimeError as e:
            last_err = e
            msg = str(e)
            retryable = any(code in msg for code in ("429", "502", "503", "504")) or "connection failed" in msg.lower()
            if retryable and i < len(chain) - 1:
                print(f"   [OpenRouter] {m} unavailable ({msg[:100]}) — trying next fallback model...")
                continue
            raise
    raise last_err or RuntimeError("OpenRouter: all fallback models failed.")


def consult_openrouter(prompt: str, context: str = "", model: str = None) -> str:
    """
    Send a plain reasoning/planning prompt to OpenRouter (no tool schema —
    this is for text generation, same role as consult_gemini). Used as the
    secondary planning brain, consulted far more often than Gemini since
    it's a direct cheap/free API call with no desktop-app UI cost.
    """
    if not _OPENROUTER_AVAILABLE:
        return f"OpenRouter is not available: {_openrouter_load_msg}"

    full_prompt = prompt.strip()
    if context.strip():
        if len(context) > 60000:
            context = context[:30000] + "\n... [TRUNCATED] ...\n" + context[-30000:]
        full_prompt = "[CONTEXT]\n" + context.strip() + "\n\n[TASK]\n" + full_prompt

    try:
        use_model = model or OPENROUTER_MODEL
        print(f"   [OpenRouter] Model: {use_model}")
        resp   = _openrouter_chat_with_fallback(
            [{"role": "user", "content": full_prompt}],
            model=use_model
        )
        result = (resp["message"]["content"] or "").strip()
        print(f"   [OpenRouter] Response received ({len(result)} chars)")
        return f"[OpenRouter/{use_model}]\n" + result
    except Exception as e:
        return f"OpenRouter error: {str(e)}"


def delegate_to_openrouter(task: str, context: str = "", model: str = None,
                            max_steps: int = 10) -> str:
    """
    Turn OpenRouter into an actual coworker instead of just a text consultant.
    Hands `task` off to a fresh, FULLY TOOL-CAPABLE agent loop running on
    OpenRouter — it can call UIA, CDP, filesystem, terminal, or any other
    tool Midum has, exactly like the primary loop, then reports back a
    final summary that gets relayed to the user.

    Architecturally this spins up process_chat_turn on a brand-new isolated
    conversation seeded with the task, with force_provider="openrouter" so
    it runs on OpenRouter regardless of the global MODEL_PROVIDER. The outer
    Midum loop and this delegated sub-loop do NOT share conversation
    history — only the final summary comes back to the caller.

    The shared response-memory scratchpad file is saved and restored around
    the delegated call, since process_chat_turn wipes it at the start of
    every invocation (including this nested one) and we don't want a
    delegated sub-task to clobber the outer task's in-progress notes.
    """
    if not _OPENROUTER_AVAILABLE:
        return f"OpenRouter is not available: {_openrouter_load_msg}"

    use_model = model or OPENROUTER_MODEL
    print(f"   [Delegate → OpenRouter/{use_model}] Task: {task[:80]}")

    _saved_scratchpad = None
    try:
        if os.path.exists(RESPONSE_MEMORY):
            with open(RESPONSE_MEMORY, "r", encoding="utf-8") as f:
                _saved_scratchpad = f.read()
    except Exception:
        pass

    try:
        sub_system_prompt = get_system_prompt(
            effective_provider="openrouter", effective_model=use_model
        )
        sub_system_prompt += (
            "\n\n━━━ DELEGATED TASK MODE ━━━\n"
            "You have been handed a specific task by Midum (the primary agent) to "
            "complete independently. You have FULL access to every tool listed above — "
            "act autonomously to complete it, calling tools directly rather than asking "
            "anyone for permission. When finished, reply with a clear plain-text summary "
            "of what you did and the result — this summary is relayed directly to the "
            "user, so make it complete and readable on its own."
        )

        task_message = task.strip()
        if context.strip():
            task_message = f"[CONTEXT FROM MIDUM]\n{context.strip()}\n\n[TASK]\n{task_message}"

        sub_history = [
            {"role": "system", "content": sub_system_prompt},
            {
                "role": "user",
                "content": (
                    f"{task_message}\n\n"
                    "[SYSTEM]: Act immediately. Execute the first tool call now. "
                    "Do not explain — just act. Give a final plain-text summary when done."
                ),
            },
        ]

        summary, sub_tool_outputs = process_chat_turn(
            sub_history,
            user_request=task,
            force_provider="openrouter",
            force_model=use_model,
            max_steps=max_steps,
        )

        step_note = f" ({len(sub_tool_outputs)} tool call(s) executed)" if sub_tool_outputs else ""
        return f"[OpenRouter coworker/{use_model} — task complete{step_note}]\n{summary}"

    except Exception as e:
        return f"Delegation to OpenRouter failed: {e}"

    finally:
        # Restore the outer task's scratchpad exactly as it was before we
        # ran a nested process_chat_turn call that would have wiped it.
        try:
            if _saved_scratchpad is not None:
                with open(RESPONSE_MEMORY, "w", encoding="utf-8") as f:
                    f.write(_saved_scratchpad)
        except Exception:
            pass


def set_openrouter_model(model_id: str) -> str:
    """
    Switch OPENROUTER_MODEL at runtime — no restart needed. Applies
    immediately to consult_openrouter, delegate_to_openrouter, and (if
    MODEL_PROVIDER == "openrouter") the primary execution loop, since all
    of those read the global fresh on every call rather than caching it.
    """
    global OPENROUTER_MODEL
    old = OPENROUTER_MODEL
    new = model_id.strip()
    if not new:
        return "Error: empty model ID."
    OPENROUTER_MODEL = new
    print(f"🔀 [OpenRouter model switched: '{old}' → '{new}']")
    return (
        f"OpenRouter model switched: '{old}' → '{new}'. "
        f"This is in-memory only for the current session — edit OPENROUTER_MODEL "
        f"in main.py to make it the default on next launch."
    )


def list_openrouter_models(free_only: bool = True) -> str:
    """
    Fetch the live list of models available on OpenRouter and return a
    numbered indexed table. Follow up with set_openrouter_model_by_index(N)
    to switch — consistent with the choose-over-search pattern used
    elsewhere (list_paths_indexed, list_skills_indexed, etc).
    """
    if requests is None:
        return "requests library not installed: pip install requests"
    try:
        headers = {}
        if _OPENROUTER_API_KEY:
            headers["Authorization"] = f"Bearer {_OPENROUTER_API_KEY}"
        resp = requests.get(f"{OPENROUTER_API_BASE}/models", headers=headers, timeout=15)
        resp.raise_for_status()
        data = resp.json().get("data", [])
    except Exception as e:
        return f"Error fetching OpenRouter model list: {e}"

    items = []
    for m in data:
        model_id = m.get("id", "")
        pricing  = m.get("pricing", {}) or {}
        is_free  = model_id.endswith(":free") or (
            str(pricing.get("prompt", "1")) == "0" and str(pricing.get("completion", "1")) == "0"
        )
        if free_only and not is_free:
            continue
        items.append({
            "id":      model_id,
            "name":    m.get("name", model_id),
            "context": m.get("context_length", "?"),
            "free":    is_free,
        })

    if not items:
        return "No models found matching the filter."

    _store_index("openrouter_models", items)

    lines = [
        f"OpenRouter models ({len(items)}{' free' if free_only else ''}) — "
        f"currently active: {OPENROUTER_MODEL}",
        "Use set_openrouter_model_by_index(index) to switch.",
        "",
        f"{'IDX':>4}  {'CTX':>8}  {'FREE':<5}  MODEL ID",
        "─" * 72,
    ]
    for i, m in enumerate(items):
        lines.append(f"{i:>4}  {str(m['context']):>8}  {'yes' if m['free'] else 'no':<5}  {m['id']}")
    return "\n".join(lines)


def set_openrouter_model_by_index(index: int) -> str:
    """Switch OPENROUTER_MODEL to the entry at `index` from list_openrouter_models."""
    entry = _get_indexed("openrouter_models", index)
    if entry is None:
        return f"Index {index} not found. Call list_openrouter_models() first."
    return set_openrouter_model(entry["id"])


# =============================================================================
# GEMINI API (OFFICIAL) SETUP — real API key, no browser/cookie hacks
# =============================================================================
#
# This is a SEPARATE code path from both consult_gemini (web-chat scraping)
# and the unused google-genai SDK client (_gemini_client) loaded above. It
# talks to Google's OFFICIAL OpenAI-compatible endpoint for the Gemini API:
#     https://ai.google.dev/gemini-api/docs/openai
# which means we can reuse the exact same plain-`requests` pipeline as
# OpenRouter (_openrouter_chat) — same request shape, same normalised
# response shape, same retry logic — instead of hand-rolling anything
# genai-SDK-specific. Structured tool calling (the `tools` schema already
# used everywhere else in this file) works natively here.
#
# API KEY SETUP (one-time):
#   1. Get a free key at https://aistudio.google.com/app/apikey
#   2. Add it to the SAME secrets file used everywhere else:
#      { "GEMINI_API_KEY": "AIza...", "OPENROUTER_API_KEY": "sk-or-v1-..." }
#
_GEMINI_API_AVAILABLE = False
_GEMINI_API_KEY        = None

def _load_gemini_api():
    """Load GEMINI_API_KEY from the shared secrets file (official API path)."""
    global _GEMINI_API_AVAILABLE, _GEMINI_API_KEY
    try:
        secrets_path = os.path.abspath(SECRETS_FILE)
        if not os.path.exists(secrets_path):
            return False, f"Secrets file not found: {secrets_path}"
        with open(secrets_path, "r", encoding="utf-8") as f:
            secrets = json.load(f)
        key = secrets.get("GEMINI_API_KEY", "").strip()
        if not key:
            return False, "GEMINI_API_KEY is empty in secrets file."
        _GEMINI_API_KEY       = key
        _GEMINI_API_AVAILABLE = True
        return True, "OK"
    except Exception as e:
        return False, str(e)


_gemini_api_load_ok, _gemini_api_load_msg = _load_gemini_api()


def _sanitize_messages_for_gemini_api(messages: list) -> list:
    """
    Midum's internal conversation_history uses a loose, Ollama-shaped
    message format that OpenRouter's backend tolerates/auto-repairs, but
    Google's OFFICIAL OpenAI-compatible endpoint validates strictly and will
    400 INVALID_ARGUMENT on it. Two specific things need fixing on the way
    out, without touching the shared internal format used by every other
    provider:

    1. Assistant `tool_calls` entries here have no "id" / "type": "function"
       — added by _openrouter_chat/_call_ollama's normaliser, but OpenAI's
       schema requires both, and the FOLLOWING "tool" role message must
       carry a matching "tool_call_id".
    2. `function.arguments` is stored as a Python dict internally, but the
       OpenAI/Gemini schema requires it to be a JSON-encoded STRING.

    This walks the message list once, assigns synthetic ids to any
    assistant tool_calls that are missing them, re-serialises arguments to
    strings, and threads matching tool_call_id values onto the immediately
    following "tool" messages (FIFO, since Midum only ever emits one tool
    call per step — see "tool_calls[:1]" in process_chat_turn).
    """
    sanitized  = []
    pending_ids = []
    counter    = 0
    for m in messages:
        m = dict(m)
        role = m.get("role")

        if role == "assistant" and m.get("tool_calls"):
            new_calls = []
            for tc in m["tool_calls"]:
                fn   = (tc or {}).get("function", {}) or {}
                name = fn.get("name", "")
                args = fn.get("arguments", {})
                if not isinstance(args, str):
                    try:
                        args = json.dumps(args)
                    except Exception:
                        args = "{}"
                call_id = tc.get("id") or f"call_{counter}"
                counter += 1
                pending_ids.append(call_id)
                new_calls.append({
                    "id": call_id,
                    "type": "function",
                    "function": {"name": name, "arguments": args},
                })
            m["tool_calls"] = new_calls
            if not m.get("content"):
                m["content"] = None   # OpenAI schema allows null content alongside tool_calls
            sanitized.append(m)

        elif role == "tool":
            if not m.get("tool_call_id"):
                m["tool_call_id"] = pending_ids.pop(0) if pending_ids else f"call_{counter}"
            sanitized.append(m)

        else:
            sanitized.append(m)

    return sanitized


def _gemini_api_chat(messages: list, model: str = None, tools_schema: list = None,
                      timeout: int = 60, _retries: int = 2) -> dict:
    """
    Call Google's OFFICIAL Gemini API through its OpenAI-compatible
    /chat/completions endpoint. Returns a dict normalised to the SAME shape
    _call_ollama / _openrouter_chat produce:
        {"message": {"role": "assistant", "content": str, "tool_calls": [...] }}
    so process_chat_turn can treat every provider identically.

    Retries on 429 (rate limit) and 502/503 (upstream overloaded), mirroring
    _openrouter_chat exactly. Raises RuntimeError with the actual error
    message from Gemini's response body on failure.
    """
    if not _GEMINI_API_AVAILABLE or not _GEMINI_API_KEY:
        raise RuntimeError(
            f"Gemini API not available: {_gemini_api_load_msg}. "
            f"Add GEMINI_API_KEY to {SECRETS_FILE}"
        )
    if requests is None:
        raise RuntimeError("requests library not installed: pip install requests")

    model = model or GEMINI_API_MODEL

    payload = {
        "model": model,
        "messages": _sanitize_messages_for_gemini_api(messages),
    }
    if tools_schema:
        payload["tools"] = tools_schema

    headers = {
        "Authorization": f"Bearer {_GEMINI_API_KEY}",
        "Content-Type":  "application/json",
    }

    last_err = None
    for attempt in range(_retries + 1):
        try:
            resp = requests.post(
                f"{GEMINI_API_BASE}/chat/completions",
                headers=headers, json=payload, timeout=timeout
            )

            try:
                data = resp.json()
            except Exception:
                data = {}

            if resp.status_code != 200:
                # Google's error body is sometimes a dict ({"error": {...}}) and
                # sometimes a single-element list ([{"error": {...}}]) — handle both.
                err_container = data
                if isinstance(err_container, list) and err_container:
                    err_container = err_container[0]
                err_obj = err_container.get("error", {}) if isinstance(err_container, dict) else {}
                err_msg = err_obj.get("message") or resp.text[:300] or f"HTTP {resp.status_code}"

                retryable = resp.status_code in (429, 502, 503, 504)
                if retryable and attempt < _retries:
                    wait_s = 1.5 * (attempt + 1)
                    print(f"   [Gemini API] {resp.status_code} ({err_msg[:80]}) — "
                          f"retrying in {wait_s:.1f}s...")
                    time.sleep(wait_s)
                    last_err = err_msg
                    continue

                raise RuntimeError(f"Gemini API error ({resp.status_code}): {err_msg}")

            if isinstance(data, dict) and data.get("error"):
                raise RuntimeError(f"Gemini API error: {data['error'].get('message', data['error'])}")

            break   # success

        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            last_err = str(e)
            if attempt < _retries:
                wait_s = 1.5 * (attempt + 1)
                print(f"   [Gemini API] Connection issue — retrying in {wait_s:.1f}s...")
                time.sleep(wait_s)
                continue
            raise RuntimeError(f"Gemini API connection failed after {_retries + 1} attempts: {last_err}")
    else:
        raise RuntimeError(f"Gemini API failed after {_retries + 1} attempts: {last_err}")

    choice  = (data.get("choices") or [{}])[0]
    msg     = choice.get("message", {}) or {}
    content = msg.get("content") or ""
    raw_tool_calls = msg.get("tool_calls") or []

    normalised_calls = []
    for tc in raw_tool_calls:
        fn = tc.get("function", {})
        args = fn.get("arguments", {})
        if isinstance(args, str):
            try:
                args = json.loads(args)
            except Exception:
                pass
        normalised_calls.append({"function": {"name": fn.get("name", ""), "arguments": args}})

    return {
        "message": {
            "role": "assistant",
            "content": content,
            "tool_calls": normalised_calls,
        }
    }


def consult_gemini_api(prompt: str, context: str = "", model: str = None) -> str:
    """
    Send a plain reasoning/planning prompt to the OFFICIAL Gemini API (API
    key, not web chat) — no tool schema, text generation only. Same role as
    consult_openrouter/consult_gemini, but using the metered official API.
    """
    if not _GEMINI_API_AVAILABLE:
        return f"Gemini API is not available: {_gemini_api_load_msg}"

    full_prompt = prompt.strip()
    if context.strip():
        if len(context) > 60000:
            context = context[:30000] + "\n... [TRUNCATED] ...\n" + context[-30000:]
        full_prompt = "[CONTEXT]\n" + context.strip() + "\n\n[TASK]\n" + full_prompt

    try:
        use_model = model or GEMINI_API_MODEL
        print(f"   [Gemini API] Model: {use_model}")
        resp   = _gemini_api_chat(
            [{"role": "user", "content": full_prompt}],
            model=use_model
        )
        result = (resp["message"]["content"] or "").strip()
        print(f"   [Gemini API] Response received ({len(result)} chars)")
        return f"[Gemini-API/{use_model}]\n" + result
    except Exception as e:
        return f"Gemini API error: {str(e)}"


def delegate_to_gemini_api(task: str, context: str = "", model: str = None,
                            max_steps: int = 10) -> str:
    """
    Mirror of delegate_to_openrouter(): hand `task` off to a fresh, FULLY
    TOOL-CAPABLE agent loop running on the official Gemini API — it can call
    UIA, CDP, filesystem, terminal, MCP tools, or any other tool Midum has,
    exactly like the primary loop, then reports back a final summary.

    Runs process_chat_turn on a brand-new isolated conversation seeded with
    the task, with force_provider="gemini_api" so it runs on the Gemini API
    regardless of the global MODEL_PROVIDER. Does NOT share conversation
    history with the outer loop — only the final summary comes back.
    """
    if not _GEMINI_API_AVAILABLE:
        return f"Gemini API is not available: {_gemini_api_load_msg}"

    use_model = model or GEMINI_API_MODEL
    print(f"   [Delegate → Gemini-API/{use_model}] Task: {task[:80]}")

    _saved_scratchpad = None
    try:
        if os.path.exists(RESPONSE_MEMORY):
            with open(RESPONSE_MEMORY, "r", encoding="utf-8") as f:
                _saved_scratchpad = f.read()
    except Exception:
        pass

    try:
        sub_system_prompt = get_system_prompt(
            effective_provider="gemini_api", effective_model=use_model
        )
        sub_system_prompt += (
            "\n\n━━━ DELEGATED TASK MODE ━━━\n"
            "You have been handed a specific task by Midum (the primary agent) to "
            "complete independently. You have FULL access to every tool listed above — "
            "act autonomously to complete it, calling tools directly rather than asking "
            "anyone for permission. When finished, reply with a clear plain-text summary "
            "of what you did and the result — this summary is relayed directly to the "
            "user, so make it complete and readable on its own."
        )

        task_message = task.strip()
        if context.strip():
            task_message = f"[CONTEXT FROM MIDUM]\n{context.strip()}\n\n[TASK]\n{task_message}"

        sub_history = [
            {"role": "system", "content": sub_system_prompt},
            {
                "role": "user",
                "content": (
                    f"{task_message}\n\n"
                    "[SYSTEM]: Act immediately. Execute the first tool call now. "
                    "Do not explain — just act. Give a final plain-text summary when done."
                ),
            },
        ]

        summary, sub_tool_outputs = process_chat_turn(
            sub_history,
            user_request=task,
            force_provider="gemini_api",
            force_model=use_model,
            max_steps=max_steps,
        )

        step_note = f" ({len(sub_tool_outputs)} tool call(s) executed)" if sub_tool_outputs else ""
        return f"[Gemini-API coworker/{use_model} — task complete{step_note}]\n{summary}"

    except Exception as e:
        return f"Delegation to Gemini API failed: {e}"

    finally:
        try:
            if _saved_scratchpad is not None:
                with open(RESPONSE_MEMORY, "w", encoding="utf-8") as f:
                    f.write(_saved_scratchpad)
        except Exception:
            pass


def set_gemini_api_model(model_id: str) -> str:
    """
    Switch GEMINI_API_MODEL at runtime — no restart needed. Applies
    immediately to consult_gemini_api, delegate_to_gemini_api, and (if
    MODEL_PROVIDER == "gemini_api") the primary execution loop.
    """
    global GEMINI_API_MODEL
    old = GEMINI_API_MODEL
    new = model_id.strip()
    if not new:
        return "Error: empty model ID."
    GEMINI_API_MODEL = new
    print(f"🔀 [Gemini API model switched: '{old}' → '{new}']")
    return (
        f"Gemini API model switched: '{old}' → '{new}'. "
        f"This is in-memory only for the current session — edit GEMINI_API_MODEL "
        f"in main.py to make it the default on next launch."
    )


# =============================================================================
# GROQCLOUD SETUP
# =============================================================================
# GroqCloud (https://console.groq.com) runs open models (Llama, Qwen,
# DeepSeek, Kimi, ...) on custom LPU hardware, and exposes a genuinely free
# tier — no credit card required — through the SAME OpenAI-compatible
# /chat/completions endpoint shape as OpenRouter and the official Gemini
# API. That means we can reuse the exact same plain-`requests` pipeline
# (same request shape, same normalised response shape, same retry/fallback
# logic) instead of hand-rolling anything Groq-specific. Structured tool
# calling works natively and is reliable on Groq's supported models.
#
# API KEY SETUP (one-time):
#   1. Get a free key at https://console.groq.com/keys
#   2. Add it to the SAME secrets file used everywhere else:
#      { "GEMINI_API_KEY": "...", "OPENROUTER_API_KEY": "...", "GROQ_API_KEY": "gsk_..." }
#
_GROQ_AVAILABLE = False
_GROQ_API_KEY    = None

def _load_groq():
    """Load GROQ_API_KEY from the shared secrets file."""
    global _GROQ_AVAILABLE, _GROQ_API_KEY
    try:
        secrets_path = os.path.abspath(SECRETS_FILE)
        if not os.path.exists(secrets_path):
            return False, f"Secrets file not found: {secrets_path}"
        with open(secrets_path, "r", encoding="utf-8") as f:
            secrets = json.load(f)
        key = secrets.get("GROQ_API_KEY", "").strip()
        if not key:
            return False, "GROQ_API_KEY is empty in secrets file."
        _GROQ_API_KEY   = key
        _GROQ_AVAILABLE = True
        return True, "OK"
    except Exception as e:
        return False, str(e)


_groq_load_ok, _groq_load_msg = _load_groq()


def _sanitize_messages_for_groq(messages: list) -> list:
    """
    Same fixup as _sanitize_messages_for_gemini_api(): Midum's internal
    conversation_history uses a loose, Ollama-shaped message format.
    GroqCloud's OpenAI-compatible endpoint validates strictly and will 400
    on missing tool_call ids / non-string function.arguments, so we repair
    those on the way out without touching the shared internal format used
    by every other provider.
    """
    sanitized   = []
    pending_ids = []
    counter     = 0
    for m in messages:
        m = dict(m)
        role = m.get("role")

        if role == "assistant" and m.get("tool_calls"):
            new_calls = []
            for tc in m["tool_calls"]:
                fn   = (tc or {}).get("function", {}) or {}
                name = fn.get("name", "")
                args = fn.get("arguments", {})
                if not isinstance(args, str):
                    try:
                        args = json.dumps(args)
                    except Exception:
                        args = "{}"
                call_id = tc.get("id") or f"call_{counter}"
                counter += 1
                pending_ids.append(call_id)
                new_calls.append({
                    "id": call_id,
                    "type": "function",
                    "function": {"name": name, "arguments": args},
                })
            m["tool_calls"] = new_calls
            if not m.get("content"):
                m["content"] = None
            sanitized.append(m)

        elif role == "tool":
            if not m.get("tool_call_id"):
                m["tool_call_id"] = pending_ids.pop(0) if pending_ids else f"call_{counter}"
            sanitized.append(m)

        else:
            sanitized.append(m)

    return sanitized


def _groq_chat(messages: list, model: str = None, tools_schema: list = None,
               timeout: int = 60, _retries: int = 2) -> dict:
    """
    Call GroqCloud's OpenAI-compatible /chat/completions endpoint. Returns a
    dict normalised to the SAME shape _call_ollama / _openrouter_chat /
    _gemini_api_chat produce:
        {"message": {"role": "assistant", "content": str, "tool_calls": [...] }}
    so process_chat_turn can treat every provider identically.

    Retries on 429 (rate limit) and 502/503 (upstream overloaded), mirroring
    _openrouter_chat / _gemini_api_chat exactly. Raises RuntimeError with the
    actual error message from Groq's response body on failure.
    """
    if not _GROQ_AVAILABLE or not _GROQ_API_KEY:
        raise RuntimeError(
            f"GroqCloud not available: {_groq_load_msg}. "
            f"Add GROQ_API_KEY to {SECRETS_FILE}"
        )
    if requests is None:
        raise RuntimeError("requests library not installed: pip install requests")

    model = model or GROQ_MODEL

    payload = {
        "model": model,
        "messages": _sanitize_messages_for_groq(messages),
    }
    if tools_schema:
        payload["tools"] = tools_schema

    headers = {
        "Authorization": f"Bearer {_GROQ_API_KEY}",
        "Content-Type":  "application/json",
    }

    last_err = None
    for attempt in range(_retries + 1):
        try:
            resp = requests.post(
                f"{GROQ_API_BASE}/chat/completions",
                headers=headers, json=payload, timeout=timeout
            )

            try:
                data = resp.json()
            except Exception:
                data = {}

            if resp.status_code != 200:
                err_obj = data.get("error", {}) if isinstance(data, dict) else {}
                err_msg = err_obj.get("message") or resp.text[:300] or f"HTTP {resp.status_code}"

                retryable = resp.status_code in (429, 502, 503, 504)
                if retryable and attempt < _retries:
                    wait_s = 1.5 * (attempt + 1)
                    print(f"   [Groq] {resp.status_code} ({err_msg[:80]}) — "
                          f"retrying in {wait_s:.1f}s...")
                    time.sleep(wait_s)
                    last_err = err_msg
                    continue

                raise RuntimeError(f"Groq error ({resp.status_code}): {err_msg}")

            if isinstance(data, dict) and data.get("error"):
                raise RuntimeError(f"Groq error: {data['error'].get('message', data['error'])}")

            break   # success

        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            last_err = str(e)
            if attempt < _retries:
                wait_s = 1.5 * (attempt + 1)
                print(f"   [Groq] Connection issue — retrying in {wait_s:.1f}s...")
                time.sleep(wait_s)
                continue
            raise RuntimeError(f"Groq connection failed after {_retries + 1} attempts: {last_err}")
    else:
        raise RuntimeError(f"Groq failed after {_retries + 1} attempts: {last_err}")

    choice  = (data.get("choices") or [{}])[0]
    msg     = choice.get("message", {}) or {}
    content = msg.get("content") or ""
    raw_tool_calls = msg.get("tool_calls") or []

    normalised_calls = []
    for tc in raw_tool_calls:
        fn = tc.get("function", {})
        args = fn.get("arguments", {})
        if isinstance(args, str):
            try:
                args = json.loads(args)
            except Exception:
                pass
        normalised_calls.append({"function": {"name": fn.get("name", ""), "arguments": args}})

    return {
        "message": {
            "role": "assistant",
            "content": content,
            "tool_calls": normalised_calls,
        }
    }


def _groq_chat_with_fallback(messages: list, model: str = None, tools_schema: list = None,
                              timeout: int = 60) -> dict:
    """
    Wraps _groq_chat() with automatic fallback across GROQ_FALLBACK_MODELS,
    mirroring _openrouter_chat_with_fallback(). If the requested (or
    default) model repeatedly fails with a retryable error, try the next
    model in the fallback chain instead of failing the whole turn.
    """
    requested = model or GROQ_MODEL
    chain = [requested] + [m for m in GROQ_FALLBACK_MODELS if m != requested]

    last_exc = None
    for m in chain:
        try:
            resp = _groq_chat(messages, model=m, tools_schema=tools_schema, timeout=timeout)
            if m != requested:
                print(f"   [Groq] Fell back to model: {m}")
            return resp
        except Exception as e:
            last_exc = e
            print(f"   [Groq] Model '{m}' failed ({e}) — trying next fallback...")
            continue
    raise RuntimeError(f"All Groq models failed. Last error: {last_exc}")


def consult_groq(prompt: str, context: str = "", model: str = None) -> str:
    """
    Send a plain reasoning/planning prompt to GroqCloud — no tool schema,
    text generation only. Same role as consult_openrouter/consult_gemini_api,
    but using GroqCloud's free-tier fast inference.
    """
    if not _GROQ_AVAILABLE:
        return f"GroqCloud is not available: {_groq_load_msg}"

    full_prompt = prompt.strip()
    if context.strip():
        if len(context) > 60000:
            context = context[:30000] + "\n... [TRUNCATED] ...\n" + context[-30000:]
        full_prompt = "[CONTEXT]\n" + context.strip() + "\n\n[TASK]\n" + full_prompt

    try:
        use_model = model or GROQ_MODEL
        print(f"   [Groq] Model: {use_model}")
        resp   = _groq_chat_with_fallback(
            [{"role": "user", "content": full_prompt}],
            model=use_model
        )
        result = (resp["message"]["content"] or "").strip()
        print(f"   [Groq] Response received ({len(result)} chars)")
        return f"[Groq/{use_model}]\n" + result
    except Exception as e:
        return f"GroqCloud error: {str(e)}"


def delegate_to_groq(task: str, context: str = "", model: str = None,
                      max_steps: int = 10) -> str:
    """
    Mirror of delegate_to_openrouter()/delegate_to_gemini_api(): hand `task`
    off to a fresh, FULLY TOOL-CAPABLE agent loop running on GroqCloud — it
    can call UIA, CDP, filesystem, terminal, MCP tools, or any other tool
    Midum has, exactly like the primary loop, then reports back a final
    summary.

    Runs process_chat_turn on a brand-new isolated conversation seeded with
    the task, with force_provider="groq" so it runs on GroqCloud regardless
    of the global MODEL_PROVIDER. Does NOT share conversation history with
    the outer loop — only the final summary comes back.
    """
    if not _GROQ_AVAILABLE:
        return f"GroqCloud is not available: {_groq_load_msg}"

    use_model = model or GROQ_MODEL
    print(f"   [Delegate → Groq/{use_model}] Task: {task[:80]}")

    _saved_scratchpad = None
    try:
        if os.path.exists(RESPONSE_MEMORY):
            with open(RESPONSE_MEMORY, "r", encoding="utf-8") as f:
                _saved_scratchpad = f.read()
    except Exception:
        pass

    try:
        sub_system_prompt = get_system_prompt(
            effective_provider="groq", effective_model=use_model
        )
        sub_system_prompt += (
            "\n\n━━━ DELEGATED TASK MODE ━━━\n"
            "You have been handed a specific task by Midum (the primary agent) to "
            "complete independently. You have FULL access to every tool listed above — "
            "act autonomously to complete it, calling tools directly rather than asking "
            "anyone for permission. When finished, reply with a clear plain-text summary "
            "of what you did and the result — this summary is relayed directly to the "
            "user, so make it complete and readable on its own."
        )

        task_message = task.strip()
        if context.strip():
            task_message = f"[CONTEXT FROM MIDUM]\n{context.strip()}\n\n[TASK]\n{task_message}"

        sub_history = [
            {"role": "system", "content": sub_system_prompt},
            {
                "role": "user",
                "content": (
                    f"{task_message}\n\n"
                    "[SYSTEM]: Act immediately. Execute the first tool call now. "
                    "Do not explain — just act. Give a final plain-text summary when done."
                ),
            },
        ]

        summary, sub_tool_outputs = process_chat_turn(
            sub_history,
            user_request=task,
            force_provider="groq",
            force_model=use_model,
            max_steps=max_steps,
        )

        step_note = f" ({len(sub_tool_outputs)} tool call(s) executed)" if sub_tool_outputs else ""
        return f"[Groq coworker/{use_model} — task complete{step_note}]\n{summary}"

    except Exception as e:
        return f"Delegation to GroqCloud failed: {e}"

    finally:
        try:
            if _saved_scratchpad is not None:
                with open(RESPONSE_MEMORY, "w", encoding="utf-8") as f:
                    f.write(_saved_scratchpad)
        except Exception:
            pass


def set_groq_model(model_id: str) -> str:
    """
    Switch GROQ_MODEL at runtime — no restart needed. Applies immediately to
    consult_groq, delegate_to_groq, and (if MODEL_PROVIDER == "groq") the
    primary execution loop, since all three read the module-level global.
    """
    global GROQ_MODEL
    old = GROQ_MODEL
    new = model_id.strip()
    if not new:
        return "Error: empty model ID."
    GROQ_MODEL = new
    print(f"🔀 [Groq model switched: '{old}' → '{new}']")
    return (
        f"Groq model switched: '{old}' → '{new}'. "
        f"This is in-memory only for the current session — edit GROQ_MODEL "
        f"in main.py to make it the default on next launch."
    )


def list_groq_models() -> str:
    """
    Fetch the live list of models currently available on GroqCloud from
    /models, format it as a numbered indexed table. Follow up with
    set_groq_model_by_index(N) or set_groq_model(model_id) to switch.
    """
    if not _GROQ_AVAILABLE:
        return f"GroqCloud is not available: {_groq_load_msg}"
    if requests is None:
        return "requests library not installed: pip install requests"

    try:
        headers = {"Authorization": f"Bearer {_GROQ_API_KEY}"}
        resp = requests.get(f"{GROQ_API_BASE}/models", headers=headers, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        items = data.get("data", [])
    except Exception as e:
        return f"Failed to fetch Groq model list: {e}"

    if not items:
        return "No models returned by GroqCloud."

    items.sort(key=lambda m: m.get("id", ""))
    _store_index("groq_models", items)

    lines = [
        f"currently active: {GROQ_MODEL}",
        "Use set_groq_model_by_index(index) to switch.",
        ""
    ]
    for i, m in enumerate(items):
        lines.append(f"[{i}] {m.get('id', '?')} (owner: {m.get('owned_by', '?')})")
    return "\n".join(lines)


def set_groq_model_by_index(index: int) -> str:
    """Switch GROQ_MODEL to the entry at `index` from list_groq_models."""
    entry = _get_indexed("groq_models", index)
    if not entry:
        return f"Index {index} not found. Call list_groq_models() first."
    return set_groq_model(entry["id"])


# =============================================================================
# MCP (MODEL CONTEXT PROTOCOL) SERVER SUPPORT
# =============================================================================
#
# Lets Midum connect to external MCP servers and use their tools WITHOUT
# dumping every registered tool's JSON schema into Midum's context window
# (which is exactly what would happen if each MCP tool were added directly
# to the `tools` list below — with several servers connected that can blow
# past the context budget fast). Instead, Midum gets three uniform tools:
#
#   list_mcp_servers()                       — which servers are connected (no tool detail)
#   show_server_tools(server)                — on demand: tools + schemas for ONE server
#   call_mcp_tool(server, tool_name, args)    — uniform invocation for ANY tool on ANY server
#
# plus two management tools to make connecting a server easy:
#
#   connect_mcp_server(name, transport, ...) — connect now, and remember it for next time
#   disconnect_mcp_server(server, forget)    — disconnect, optionally forgetting it too
#
# CONFIG FILE (auto-loaded and auto-connected at startup):
#   storage/mcp_servers.json — a JSON list of server configs, e.g.:
#   [
#     {"name": "filesystem", "transport": "stdio", "command": "npx",
#      "args": ["-y", "@modelcontextprotocol/server-filesystem", "/home/user"]},
#     {"name": "weather", "transport": "http", "url": "https://example.com/mcp"}
#   ]
#   Editing this file by hand works too — it's just what connect_mcp_server writes to.
#
# INSTALL:
#   pip install mcp
#
_MCP_SERVER_ORDER: list = []     # server names, in the order they were connected (gives indices)
_MCP_SERVERS: dict      = {}     # name -> _MCPServerHandle


class _MCPServerHandle:
    """Everything Midum knows about one connected (or attempted) MCP server."""
    def __init__(self, name: str, config: dict):
        self.name        = name
        self.config      = config
        self.session     = None     # mcp.ClientSession, once connected
        self.tools       = []       # [{"name","description","input_schema"}, ...]
        self.exit_stack  = None     # contextlib.AsyncExitStack keeping the transport open
        self.connected   = False
        self.error       = None


def _load_mcp_config() -> list:
    """Read storage/mcp_servers.json. Returns [] if missing or invalid."""
    try:
        if not os.path.exists(MCP_SERVERS_FILE):
            return []
        with open(MCP_SERVERS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception as e:
        print(f"⚠️ [MCP] Could not read {MCP_SERVERS_FILE}: {e}")
        return []


def _save_mcp_config(configs: list):
    """Write the given list of server configs to storage/mcp_servers.json."""
    try:
        os.makedirs(os.path.dirname(MCP_SERVERS_FILE), exist_ok=True)
        with open(MCP_SERVERS_FILE, "w", encoding="utf-8") as f:
            json.dump(configs, f, indent=2)
    except Exception as e:
        print(f"⚠️ [MCP] Could not save {MCP_SERVERS_FILE}: {e}")


def _mcp_upsert_config(new_entry: dict):
    """Add or replace (by name) one server entry in the persisted config file."""
    configs = _load_mcp_config()
    configs = [c for c in configs if c.get("name") != new_entry.get("name")]
    configs.append(new_entry)
    _save_mcp_config(configs)


def _mcp_remove_config(name: str):
    configs = [c for c in _load_mcp_config() if c.get("name") != name]
    _save_mcp_config(configs)


class _MCPManager:
    """
    Owns a single background asyncio event loop that MCP client sessions run
    on for their whole lifetime (they're async context managers that need to
    stay entered between calls), and exposes plain synchronous methods so
    the rest of Midum — which is entirely sync/threaded — never has to
    touch asyncio directly.
    """
    def __init__(self):
        self._loop   = None
        self._thread = None
        self._ready  = threading.Event()

    def _ensure_loop(self):
        if self._loop is not None:
            return
        def _runner():
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            self._loop = loop
            self._ready.set()
            loop.run_forever()
        self._thread = threading.Thread(target=_runner, daemon=True, name="mcp-loop")
        self._thread.start()
        self._ready.wait(timeout=5)

    def _run(self, coro, timeout: float = 30):
        self._ensure_loop()
        fut = asyncio.run_coroutine_threadsafe(coro, self._loop)
        return fut.result(timeout=timeout)

    async def _connect_async(self, config: dict):
        from contextlib import AsyncExitStack
        from mcp import ClientSession

        transport = (config.get("transport") or "stdio").lower()
        stack = AsyncExitStack()
        try:
            if transport == "stdio":
                from mcp import StdioServerParameters
                from mcp.client.stdio import stdio_client
                env = None
                if config.get("env"):
                    env = {**os.environ, **config["env"]}
                params = StdioServerParameters(
                    command=config["command"],
                    args=config.get("args", []),
                    env=env,
                )
                read, write = await stack.enter_async_context(stdio_client(params))

            elif transport in ("http", "streamable_http", "streamable-http"):
                from mcp.client.streamable_http import streamablehttp_client
                read, write, _ = await stack.enter_async_context(
                    streamablehttp_client(config["url"], headers=config.get("headers"))
                )

            elif transport == "sse":
                from mcp.client.sse import sse_client
                read, write = await stack.enter_async_context(
                    sse_client(config["url"], headers=config.get("headers"))
                )

            else:
                raise ValueError(
                    f"Unknown transport '{transport}'. Use 'stdio', 'http', or 'sse'."
                )

            session = await stack.enter_async_context(ClientSession(read, write))
            await session.initialize()
            listed = await session.list_tools()
            tool_list = [
                {
                    "name": t.name,
                    "description": t.description or "",
                    "input_schema": t.inputSchema or {"type": "object", "properties": {}},
                }
                for t in listed.tools
            ]
            return stack, session, tool_list
        except Exception:
            await stack.aclose()
            raise

    def connect(self, name: str, config: dict) -> tuple:
        """Connect (or reconnect) to a server. Returns (ok: bool, message: str)."""
        if not _MCP_SDK_AVAILABLE:
            return False, "The 'mcp' package is not installed. Run: pip install mcp"

        # Reconnecting: tear down any existing connection under this name first.
        existing = _MCP_SERVERS.get(name)
        if existing and existing.exit_stack:
            try:
                self._run(existing.exit_stack.aclose(), timeout=10)
            except Exception:
                pass

        handle = _MCPServerHandle(name, config)
        _MCP_SERVERS[name] = handle
        if name not in _MCP_SERVER_ORDER:
            _MCP_SERVER_ORDER.append(name)

        try:
            stack, session, tool_list = self._run(
                self._connect_async(config), timeout=config.get("connect_timeout", 30)
            )
            handle.exit_stack = stack
            handle.session    = session
            handle.tools      = tool_list
            handle.connected  = True
            handle.error      = None
            return True, f"Connected to '{name}' — {len(tool_list)} tool(s) available."
        except Exception as e:
            handle.connected = False
            handle.error     = str(e)
            return False, f"Failed to connect to '{name}': {e}"

    def disconnect(self, name: str) -> str:
        handle = _MCP_SERVERS.get(name)
        if not handle:
            return f"No connected server named '{name}'."
        if handle.exit_stack:
            try:
                self._run(handle.exit_stack.aclose(), timeout=10)
            except Exception:
                pass
        del _MCP_SERVERS[name]
        if name in _MCP_SERVER_ORDER:
            _MCP_SERVER_ORDER.remove(name)
        return f"Disconnected '{name}'."

    def call_tool(self, name: str, tool_name: str, arguments: dict) -> str:
        handle = _MCP_SERVERS.get(name)
        if not handle:
            return f"Error: no connected server named '{name}'. Call list_mcp_servers() first."
        if not handle.connected:
            return f"Error: server '{name}' is not connected ({handle.error})."
        if not any(t["name"] == tool_name for t in handle.tools):
            known = ", ".join(t["name"] for t in handle.tools) or "(none)"
            return (f"Error: '{tool_name}' is not a tool on server '{name}'. "
                    f"Known tools: {known}. Call show_server_tools() to double-check.")

        async def _call():
            return await handle.session.call_tool(tool_name, arguments or {})

        try:
            result = self._run(_call(), timeout=handle.config.get("call_timeout", 60))
        except Exception as e:
            return f"Error calling '{tool_name}' on '{name}': {e}"

        parts = []
        for block in getattr(result, "content", []) or []:
            if getattr(block, "type", None) == "text":
                parts.append(block.text)
            else:
                parts.append(str(block))
        text = "\n".join(parts) if parts else "(no output)"
        if getattr(result, "isError", False):
            return f"[Tool error from '{tool_name}' on '{name}'] {text}"
        return text


_mcp_manager = _MCPManager()


def _mcp_resolve_name(server) -> str:
    """
    Accepts either an index (int/numeric string, matching list_mcp_servers
    order) or a server name directly, and returns the resolved name — or
    None if it doesn't match anything currently connected.
    """
    if server is None:
        return None
    server_str = str(server).strip()
    if server_str.isdigit():
        idx = int(server_str)
        if 0 <= idx < len(_MCP_SERVER_ORDER):
            return _MCP_SERVER_ORDER[idx]
        return None
    return server_str if server_str in _MCP_SERVERS else None


def _mcp_normalize_name(s: str) -> str:
    """Lowercase and strip everything but letters/digits — so 'get_weather',
    'getWeather', 'get-weather', and 'getweather' all compare equal."""
    return re.sub(r'[^a-z0-9]', '', (s or '').lower())


def _mcp_find_tool_matches(name: str) -> list:
    """
    Search every connected MCP server's tool list for a tool matching `name`
    once separators/case are normalized away. Weak local models frequently
    call an MCP tool directly by its own name (skipping call_mcp_tool) and
    mangle underscores/hyphens/case while doing it — this is what lets that
    still resolve. Returns a list of (server_name, real_tool_name) — usually
    0 or 1 matches, but can be 2+ if the same tool name exists on multiple
    connected servers (caller should treat that as ambiguous, not pick one).
    """
    target = _mcp_normalize_name(name)
    if not target:
        return []
    matches = []
    for server_name in _MCP_SERVER_ORDER:
        handle = _MCP_SERVERS.get(server_name)
        if not handle or not handle.connected:
            continue
        for t in handle.tools:
            if _mcp_normalize_name(t["name"]) == target:
                matches.append((server_name, t["name"]))
    return matches


def _mcp_autoroute_tool_call(name: str, args):
    """
    If `name` isn't a registered top-level tool but matches exactly one
    tool on exactly one connected MCP server, transparently rewrite the
    call into the equivalent call_mcp_tool(...) invocation. This is what
    keeps tool-calling uniform even when a weak model shortcuts straight
    to calling the MCP tool by its own (possibly underscore-mangled) name
    instead of going through call_mcp_tool as instructed.

    Returns (name, args) — unchanged if `name` is already a known
    top-level tool, or if there's no unambiguous MCP match.
    """
    if name in _known_tool_names():
        return name, args
    matches = _mcp_find_tool_matches(name)
    if len(matches) == 1:
        server_name, real_tool_name = matches[0]
        print(f"   [MCP autoroute] '{name}' → call_mcp_tool("
              f"server='{server_name}', tool_name='{real_tool_name}')")
        return "call_mcp_tool", {
            "server": server_name,
            "tool_name": real_tool_name,
            "arguments": args or {}
        }
    return name, args


def init_mcp_servers_from_config():
    """
    Called once at startup: auto-connects every server saved in
    storage/mcp_servers.json. Failures are non-fatal and left visible in
    list_mcp_servers() so Midum (or the user) can see what went wrong.
    """
    configs = _load_mcp_config()
    if not configs:
        return
    if not _MCP_SDK_AVAILABLE:
        print("⚠️ [MCP] mcp_servers.json has entries but the 'mcp' package "
              "isn't installed — run: pip install mcp")
        return
    for cfg in configs:
        name = cfg.get("name")
        if not name:
            continue
        ok, msg = _mcp_manager.connect(name, cfg)
        icon = "🔌" if ok else "⚠️"
        print(f"{icon} [MCP] {msg}")


def list_mcp_servers() -> str:
    """
    Lists connected MCP servers ONLY — not their tools — to keep this cheap
    on context. Use show_server_tools(server) to see what a given server offers.
    """
    if not _MCP_SERVER_ORDER:
        hint = "" if _MCP_SDK_AVAILABLE else " (install the 'mcp' package first: pip install mcp)"
        return f"No MCP servers connected.{hint} Use connect_mcp_server(...) to add one."

    lines = []
    for i, name in enumerate(_MCP_SERVER_ORDER):
        handle = _MCP_SERVERS[name]
        if handle.connected:
            transport = handle.config.get("transport", "stdio")
            lines.append(f"[{i}] {name} — connected ({transport}) — {len(handle.tools)} tool(s)")
        else:
            lines.append(f"[{i}] {name} — connection failed: {handle.error}")
    return "\n".join(lines)


def show_server_tools(server) -> str:
    """
    Returns the tool names, descriptions, and JSON input schemas for ONE
    server, identified by index (from list_mcp_servers) or name. This is
    the only place full tool schemas are shown to Midum — deliberately
    on-demand, per server, instead of always-loaded.
    """
    name = _mcp_resolve_name(server)
    if name is None:
        return (f"Unknown server '{server}'. Call list_mcp_servers() first "
                f"to see valid indices/names.")
    handle = _MCP_SERVERS[name]
    if not handle.connected:
        return f"Server '{name}' is not connected ({handle.error})."
    if not handle.tools:
        return f"Server '{name}' is connected but exposes no tools."
    return json.dumps(
        {"server": name, "tools": handle.tools},
        indent=2
    )


def call_mcp_tool(server, tool_name: str, arguments) -> str:
    """
    Uniform invocation for ANY tool on ANY connected MCP server — this is
    the ONLY tool actually needed to use MCP servers, once you've checked
    show_server_tools() for the right tool_name/arguments shape.
    """
    name = _mcp_resolve_name(server)
    if name is None:
        return (f"Unknown server '{server}'. Call list_mcp_servers() first "
                f"to see valid indices/names.")
    if isinstance(arguments, str):
        try:
            arguments = json.loads(arguments) if arguments.strip() else {}
        except Exception:
            return f"Error: 'arguments' was not valid JSON: {arguments!r}"
    return _mcp_manager.call_tool(name, tool_name, arguments or {})


def list_native_tools() -> str:
    """
    Lists every built-in Midum tool by name + one-line description ONLY —
    no parameter schemas — to keep this cheap on context. Use
    show_native_tool_schema(tool_name) to get the full JSON parameter
    schema for a specific tool before calling it.

    This exists so the full native tool catalogue never has to be inlined
    into a persistent system prompt/Gem: providers that support real native
    function-calling (Ollama/OpenRouter/Gemini-API/Groq) already get the
    full `tools` schema for free via the API's tools= parameter, so this
    on-demand pair (list_native_tools / show_native_tool_schema) is mainly
    for Gemini-web, which has no native tool-calling protocol and would
    otherwise need everything inlined into one big Gem prompt.
    """
    lines = []
    for i, t in enumerate(tools):
        fn = t.get("function", {})
        name = fn.get("name", "?")
        desc = (fn.get("description") or "").strip().splitlines()[0] if fn.get("description") else ""
        lines.append(f"[{i}] {name} — {desc}")
    return "\n".join(lines)


def show_native_tool_schema(tool_name: str) -> str:
    """
    Returns the full {name, description, parameters} JSON schema for ONE
    native Midum tool, identified by index (from list_native_tools) or
    exact name. Call this right before using a tool you haven't already
    seen the schema for this session.
    """
    match = None
    if isinstance(tool_name, str) and tool_name.strip().isdigit():
        idx = int(tool_name.strip())
        if 0 <= idx < len(tools):
            match = tools[idx]
    if match is None:
        for t in tools:
            if t.get("function", {}).get("name") == tool_name:
                match = t
                break
    if match is None:
        return (f"Unknown native tool '{tool_name}'. Call list_native_tools() first "
                f"to see valid indices/names.")
    fn = match.get("function", {})
    return json.dumps(
        {
            "name": fn.get("name"),
            "description": fn.get("description", ""),
            "parameters": fn.get("parameters", {}),
        },
        indent=2,
    )


def connect_mcp_server(name: str, transport: str = "stdio", command: str = None,
                        args: list = None, url: str = None, env: dict = None,
                        headers: dict = None, persist: bool = True) -> str:
    """
    Connect to a new MCP server right now. If persist=True (default), it's
    also saved to storage/mcp_servers.json so it auto-connects on every
    future startup — this is the "easy to connect" path: one tool call and
    it's remembered.

    transport='stdio' needs command (+ optional args, env) — for local
    servers launched as a subprocess, e.g. command='npx',
    args=['-y', '@modelcontextprotocol/server-filesystem', '/home/user'].

    transport='http' or 'sse' needs url (+ optional headers) — for remote
    MCP servers reachable over the network.
    """
    if not name or not name.strip():
        return "Error: a server 'name' is required."
    name = name.strip()

    config = {"name": name, "transport": transport}
    if transport == "stdio":
        if not command:
            return "Error: transport='stdio' requires 'command'."
        config["command"] = command
        if args:
            config["args"] = args
        if env:
            config["env"] = env
    elif transport in ("http", "streamable_http", "sse"):
        if not url:
            return f"Error: transport='{transport}' requires 'url'."
        config["url"] = url
        if headers:
            config["headers"] = headers
    else:
        return "Error: transport must be 'stdio', 'http', or 'sse'."

    ok, msg = _mcp_manager.connect(name, config)
    if ok and persist:
        _mcp_upsert_config(config)
        msg += " Saved — will auto-connect on future startups."
    return msg


def disconnect_mcp_server(server, forget: bool = False) -> str:
    """
    Disconnect a currently-connected MCP server. If forget=True, also
    remove it from storage/mcp_servers.json so it stops auto-connecting.
    """
    name = _mcp_resolve_name(server)
    if name is None:
        return f"Unknown server '{server}'. Call list_mcp_servers() first."
    msg = _mcp_manager.disconnect(name)
    if forget:
        _mcp_remove_config(name)
        msg += " Removed from saved config."
    return msg


# =============================================================================
# 0b. GUI USER-PROMPT TOOLS
# =============================================================================
# Lets Midum pop up a small native GUI dialog to get something from the user
# without guessing — a missing file path, an approval, a disambiguating
# choice, or arbitrary free text. Every call BLOCKS until the user responds
# (or dismisses the window), then the answer is fed back as a normal tool
# result so Midum can continue the same turn. The model can request several
# of these in a single turn (e.g. ask a question AND request a file path);
# by default it requests none — these are opt-in, situational tools, not a
# forced step at the end of every turn.
try:
    import tkinter as _tk
    from tkinter import filedialog as _tk_filedialog
    _TKINTER_AVAILABLE = True
except ImportError:
    _tk = None
    _tk_filedialog = None
    _TKINTER_AVAILABLE = False

# When Midum is running inside the CustomTkinter desktop GUI (gui.pyw), the
# GUI installs a callable here at startup. If present, every ask_user_*
# tool below routes through it instead of popping up a separate native
# tkinter window — the request/response instead renders as an inline card
# in the main chat, styled to match the rest of the app. The hook's
# signature is: hook(kind: str, payload: dict) -> str, and it BLOCKS the
# calling thread (the engine worker thread) until the user responds, so
# every ask_user_* call below still behaves exactly as documented from the
# model's point of view. When running main.py standalone (no GUI attached),
# this stays None and the original native tkinter dialogs are used as-is.
_gui_ask_hook = None


def _gui_root():
    """Create a hidden, always-on-top root window to anchor a dialog to."""
    root = _tk.Tk()
    root.withdraw()
    try:
        root.attributes("-topmost", True)
    except Exception:
        pass
    return root


def ask_user_text(prompt: str, title: str = "Midum needs input") -> str:
    """
    Pop up a GUI textbox asking the user to type a free-form answer — e.g. a
    missing file path, a name, a value, or anything else with no fixed set
    of options. Returns the typed text, '[USER SUBMITTED EMPTY TEXT]' if
    they submitted blank, or '[USER CANCELLED]' if they closed the dialog.
    """
    if _gui_ask_hook is not None:
        return _gui_ask_hook("text", {"prompt": prompt, "title": title})
    if not _TKINTER_AVAILABLE:
        return "[GUI ERROR] tkinter not available on this system — ask the user in plain text instead."
    result = {"value": None}
    root = _gui_root()
    win = _tk.Toplevel(root)
    win.title(title)
    try: win.attributes("-topmost", True)
    except Exception: pass

    _tk.Label(win, text=prompt, wraplength=420, justify="left", padx=16, pady=12).pack()
    entry = _tk.Entry(win, width=52)
    entry.pack(padx=16, pady=(0, 8))
    entry.focus_set()

    def submit(event=None):
        result["value"] = entry.get()
        win.destroy()

    def cancel():
        result["value"] = None
        win.destroy()

    entry.bind("<Return>", submit)
    btn_frame = _tk.Frame(win)
    btn_frame.pack(pady=(0, 12))
    _tk.Button(btn_frame, text="Submit", command=submit, width=10).pack(side="left", padx=6)
    _tk.Button(btn_frame, text="Cancel", command=cancel, width=10).pack(side="left", padx=6)

    win.protocol("WM_DELETE_WINDOW", cancel)
    win.grab_set()
    root.wait_window(win)
    root.destroy()

    if result["value"] is None:
        return "[USER CANCELLED]"
    return result["value"] if result["value"] else "[USER SUBMITTED EMPTY TEXT]"


def ask_user_file_path(prompt: str = "Select a file", must_exist: bool = True) -> str:
    """
    Pop up a native file-picker dialog. Use this whenever the user tells you
    to open/read/edit/save something involving a file but does not give you
    a path — instead of guessing, ask. Returns the absolute path chosen, or
    '[USER CANCELLED]' if they closed the dialog without picking anything.
    """
    if _gui_ask_hook is not None:
        return _gui_ask_hook("file", {"prompt": prompt, "must_exist": must_exist})
    if not _TKINTER_AVAILABLE:
        return "[GUI ERROR] tkinter not available on this system — ask the user for the path in plain text instead."
    root = _gui_root()
    try:
        if must_exist:
            path = _tk_filedialog.askopenfilename(title=prompt, parent=root)
        else:
            path = _tk_filedialog.asksaveasfilename(title=prompt, parent=root)
    finally:
        root.destroy()
    return path if path else "[USER CANCELLED]"


def ask_user_approval(message: str, details: str = "") -> str:
    """
    Pop up an Approve / Decline GUI dialog with two buttons. Use this before
    doing something the user should explicitly sign off on — deleting or
    overwriting files, sending a message, spending money, running a risky
    or irreversible command, etc. Returns 'APPROVED' or 'DECLINED'.
    """
    if _gui_ask_hook is not None:
        return _gui_ask_hook("approval", {"message": message, "details": details})
    if not _TKINTER_AVAILABLE:
        return "[GUI ERROR] tkinter not available on this system — ask the user to approve in plain text instead."
    result = {"value": "DECLINED"}
    root = _gui_root()
    win = _tk.Toplevel(root)
    win.title("Midum requests approval")
    try: win.attributes("-topmost", True)
    except Exception: pass

    _tk.Label(win, text=message, wraplength=420, justify="left",
              padx=16, pady=(16, 4), font=("Segoe UI", 10, "bold")).pack()
    if details:
        _tk.Label(win, text=details, wraplength=420, justify="left", padx=16, pady=(0, 8)).pack()

    def approve():
        result["value"] = "APPROVED"
        win.destroy()

    def decline():
        result["value"] = "DECLINED"
        win.destroy()

    btn_frame = _tk.Frame(win)
    btn_frame.pack(pady=12)
    _tk.Button(btn_frame, text="\u2705 Approve", command=approve, width=12, bg="#d7f5d7").pack(side="left", padx=8)
    _tk.Button(btn_frame, text="\u274c Decline", command=decline, width=12, bg="#f5d7d7").pack(side="left", padx=8)

    win.protocol("WM_DELETE_WINDOW", decline)
    win.grab_set()
    root.wait_window(win)
    root.destroy()
    return result["value"]


def ask_user_choice(question: str, choice_1: str = "", choice_2: str = "",
                     choice_3: str = "", choice_4: str = "", allow_custom: bool = True) -> str:
    """
    Pop up a multiple-choice GUI dialog: your question plus up to 4 options
    you define, and (unless allow_custom=False) a 5th free-text box so the
    user can type something else entirely. Use this to disambiguate what the
    user wants with a couple of taps instead of a back-and-forth in text.
    Returns the exact text of the option the user picked, their custom text,
    or '[USER CANCELLED]' if they closed the dialog.
    """
    if _gui_ask_hook is not None:
        return _gui_ask_hook("choice", {
            "question": question,
            "options": [c for c in (choice_1, choice_2, choice_3, choice_4) if c],
            "allow_custom": allow_custom,
        })
    if not _TKINTER_AVAILABLE:
        return "[GUI ERROR] tkinter not available on this system — ask the user in plain text instead."
    options = [c for c in (choice_1, choice_2, choice_3, choice_4) if c]
    result = {"value": None}
    root = _gui_root()
    win = _tk.Toplevel(root)
    win.title("Midum has a question")
    try: win.attributes("-topmost", True)
    except Exception: pass

    _tk.Label(win, text=question, wraplength=420, justify="left",
              padx=16, pady=(16, 8), font=("Segoe UI", 10, "bold")).pack()

    def choose(opt):
        result["value"] = opt
        win.destroy()

    for opt in options:
        _tk.Button(win, text=opt, width=48, anchor="w",
                   command=lambda o=opt: choose(o)).pack(padx=16, pady=3, fill="x")

    if allow_custom:
        row = _tk.Frame(win)
        row.pack(padx=16, pady=(10, 4), fill="x")
        custom_entry = _tk.Entry(row, width=36)
        custom_entry.pack(side="left", fill="x", expand=True)

        def submit_custom(event=None):
            txt = custom_entry.get().strip()
            if txt:
                choose(txt)

        custom_entry.bind("<Return>", submit_custom)
        _tk.Button(row, text="Other...", command=submit_custom).pack(side="left", padx=(6, 0))

    win.protocol("WM_DELETE_WINDOW", lambda: choose("[USER CANCELLED]"))
    win.grab_set()
    root.wait_window(win)
    root.destroy()
    return result["value"] if result["value"] is not None else "[USER CANCELLED]"


# =============================================================================
# 1. TOOL SCHEMAS
# =============================================================================

tools = [
    {
        "type": "function",
        "function": {
            "name": "manual_scan_app_layouts",
            "description": "Scan an active window to find its major layout containers (subtrees). Use this first when exploring a new app.",
            "parameters": {
                "type": "object",
                "properties": {
                    "window_title": {"type": "string", "description": "Exact title of the window (e.g., 'Gemini')."}
                },
                "required": ["window_title"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "manual_inspect_app_subtree",
            "description": "Scan a specific layout container (found via manual_scan_app_layouts) to reveal the interactive buttons and text fields inside it.",
            "parameters": {
                "type": "object",
                "properties": {
                    "window_title": {"type": "string"},
                    "subtree_key": {"type": "string", "description": "The name or automation_id of the container."}
                },
                "required": ["window_title", "subtree_key"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "click_ui_element",
            "description": (
                "PREFERRED tool for interacting with desktop app UI elements (buttons, "
                "fields, menu items, close/minimize buttons, etc). ONE call does "
                "everything: finds the window, searches its ENTIRE UI tree for an "
                "element matching your plain-English description, and acts on it. "
                "Supports normal app windows AND Windows shell surfaces — pass "
                "window_title as 'taskbar', 'start', 'tray', 'desktop', "
                "'action center', 'search', or 'tray overflow' to interact with "
                "the Start button, system tray, notification area, clock, Wi-Fi, "
                "volume, battery, and other shell controls that have no normal title. "
                "Examples: click_ui_element('taskbar', 'Start') clicks the Start button. "
                "click_ui_element('tray', 'Wi-Fi') clicks the Wi-Fi tray icon. "
                "click_ui_element('tray overflow', 'Show hidden icons') opens the overflow. "
                "If this returns 'No element matched', it shows all available element "
                "names grouped by type — retry with one of those exact names. "
                "If it says canvas/WebGL, use fallback_click_text instead."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "window_title": {
                        "type": "string",
                        "description": (
                            "Window title or substring, OR a shell surface alias: "
                            "'taskbar', 'start', 'tray', 'desktop', 'action center', "
                            "'search', 'tray overflow', 'secondary taskbar'."
                        )
                    },
                    "description": {
                        "type": "string",
                        "description": "Plain-English description of the element, e.g. 'Close button', 'Send message', 'Start button', 'Wi-Fi'."
                    },
                    "action": {
                        "type": "string",
                        "description": "'click' (default), 'set_text', or 'get_text'.",
                        "enum": ["click", "set_text", "get_text"]
                    },
                    "text_to_type": {
                        "type": "string",
                        "description": "Required if action is 'set_text' — the text to enter."
                    }
                },
                "required": ["window_title", "description"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "snapshot",
            "description": (
                "ONE tool that snapshots ALL visible, interactive elements into a numbered "
                "indexed table — works IDENTICALLY on a desktop app window (Discord, VS Code, "
                "Settings, any app) AND a browser tab. It is NOT browser-only. "
                "Returns: IDX | TYPE | NAME/LABEL | STATUS. "
                "Follow up with act(index, action) to interact by index — exact, never wrong. "
                "\n\n"
                "For DESKTOP WINDOWS: pass target='<window title>' e.g. target='Google Chrome', "
                "target='Visual Studio Code', target='taskbar', target='tray'. "
                "Uses UIA (Windows UI Automation) to read the app's accessibility tree. "
                "\n\n"
                "For BROWSER PAGES: pass target='browser' (reads the active tab) or "
                "target='browser:N' (reads tab N). "
                "Uses CDP (Chrome DevTools Protocol) to read the live DOM. "
                "Requires Chrome with --remote-debugging-port=9222. "
                "\n\n"
                "filter_type: optional, e.g. 'button', 'edit', 'link', 'input', 'tab', 'menuitem'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "target": {
                        "type": "string",
                        "description": (
                            "What to snapshot. "
                            "Desktop: canonical window title or shell alias ('Google Chrome', 'taskbar', 'tray'). "
                            "Browser page: 'browser' for active tab, 'browser:1' for tab 1."
                        )
                    },
                    "filter_type": {
                        "type": "string",
                        "description": "Optional type filter: 'button', 'edit', 'link', 'input', 'tab', 'menuitem', 'checkbox'."
                    }
                },
                "required": ["target"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "act",
            "description": (
                "Interact with an element from the last snapshot() call by its index. "
                "Works for both desktop windows and browser pages — same tool, same syntax. "
                "The index is exact: no scoring, no fuzzy matching, no wrong element. "
                "Always call snapshot(target) first, then act(index). "
                "\n\n"
                "action: 'click' (default), 'set_text' (type into a field), 'get_text'. "
                "target: same value you passed to snapshot() — needed to look up the right cache."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "target": {
                        "type": "string",
                        "description": "Same target you passed to snapshot(). E.g. 'Google Chrome' or 'browser'."
                    },
                    "index": {
                        "type": "integer",
                        "description": "Element index from the snapshot() output."
                    },
                    "action": {
                        "type": "string",
                        "enum": ["click", "set_text", "get_text"],
                        "description": "'click' (default), 'set_text', or 'get_text'."
                    },
                    "text_to_type": {
                        "type": "string",
                        "description": "Required if action is 'set_text'."
                    }
                },
                "required": ["target", "index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "manual_interact_with_ui",
            "description": (
                "Most tasks should use click_ui_element instead — it's a single call and "
                "has automatic coordinate fallback. Only use this if click_ui_element "
                "failed and you have an EXACT automation_id from manual_inspect_app_subtree."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "window_title": {"type": "string"},
                    "control_type": {"type": "string", "description": "'button', 'edit', 'text', etc."},
                    "search_property": {"type": "string", "description": "'automation_id', 'name', or 'class_name'"},
                    "property_value": {"type": "string", "description": "The target identifier value."},
                    "action": {"type": "string", "description": "'click', 'set_text', or 'get_text'"},
                    "text_to_type": {"type": "string", "description": "Only required if action is 'set_text'."}
                },
                "required": ["window_title", "control_type", "search_property", "property_value", "action"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_gemini_app",
            "description": (
                "Sends a prompt to Gemini via the actual WEB CHAT INTERFACE "
                "(https://gemini.google.com/app) using the community `gemini_webapi` "
                "library — NOT a browser automation, NOT a desktop application, and NOT "
                "the metered developer API. It talks directly to Gemini's internal web "
                "endpoints using session cookies, so there's no page to load or button "
                "to click — just a plain request. Uses the free web app's own usage "
                "limits. Returns ONLY the single newest reply, never the full "
                "conversation history. Prefer consult_gemini for most cases — it already "
                "routes through this same mechanism, with no fallback to the metered API."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "The complex query, data processing prompt, or reasoning task to send to Gemini."
                    }
                },
                "required": ["prompt"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "manage_gemini_chat",
            "description": "Manage the Gemini application state by performing actions like starting a new chat or selecting a recent chat.",
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["new_chat", "open_recent"],
                        "description": "The action to perform."
                    },
                    "chat_name": {
                        "type": "string",
                        "description": "Required if action is 'open_recent'. The specific name of the chat to open."
                    }
                },
                "required": ["action"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "list_active_windows",
            "description": "List all currently open and visible window titles on the desktop. Use this if you are unsure of the exact window_title to pass to the UI interaction tools.",
            "parameters": {
                "type": "object",
                "properties": {}
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_aggregated_text",
            "description": "Read text from a window or specific container by merging sibling TextControl elements into readable paragraphs.",
            "parameters": {
                "type": "object",
                "properties": {
                    "window_title": {"type": "string"},
                    "container_key": {"type": "string", "description": "Optional: Specific subtree key to read from."}
                },
                "required": ["window_title"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_local_file",
            "description": "Read the contents of a local file.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Absolute or relative path to the file."}
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_local_file",
            "description": "Completely overwrite or clear a file. Pass '' to wipe it clean.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path":    {"type": "string"},
                    "content": {"type": "string"}
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "append_local_file",
            "description": "Add content to the end of a file without modifying existing content.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path":    {"type": "string"},
                    "content": {"type": "string"}
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "execute_terminal_command",
            "description": (
                "Run a PowerShell command on the local Windows machine. "
                "To open an app: ALWAYS use Start-Process with the full path from paths.md. "
                "Example: Start-Process 'C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe'. "
                "NEVER use 'start chrome', 'start firefox', or any shorthand — always the full path. "
                "If you do not know the path, call read_paths first."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "command":           {"type": "string", "description": "The exact PowerShell command to execute."},
                    "working_directory": {"type": "string", "description": "Optional: absolute folder path to run from."}
                },
                "required": ["command"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_internet",
            "description": "Search the internet for real-time information or documentation.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string"}
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fallback_view_screen",
            "description": (
                "Capture a live screenshot of the desktop, downscaled to canvas size with a "
                "coordinate grid burned in. Use this when you need to visually inspect the screen "
                "or when you need to identify coordinates for fallback_click_grid. "
                "For text-heavy GUIs, prefer fallback_find_text over reading the grid manually."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fallback_find_text",
            "description": (
                "Use OCR (Tesseract) to locate a text string on the current screen. "
                "Returns the canvas coordinates of the best match and a full list of all detected "
                "text with their positions. Use this instead of reading the grid image when you "
                "want to click a button, label, or menu item that has visible text — it is faster "
                "and more accurate than visual grid estimation. "
                "Pass the exact text or a substring of it. Case-insensitive."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {
                        "type": "string",
                        "description": "The text string to search for on screen. Case-insensitive substring match."
                    }
                },
                "required": ["text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fallback_click_grid",
            "description": (
                "Simulate a mouse click at canvas coordinates from the grid screenshot. "
                "Python scales these to real screen pixels automatically. "
                "Use this when you have read coordinates from the fallback_view_screen grid. "
                "For clicking text elements, prefer fallback_find_text which gives you "
                "precise coordinates without needing to read the grid."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "x": {"type": "integer", "description": "Canvas x-coordinate from the grid screenshot."},
                    "y": {"type": "integer", "description": "Canvas y-coordinate from the grid screenshot."},
                    "click_type": {
                        "type": "string",
                        "description": "'left_click' (default), 'right_click', or 'double_click'.",
                        "enum": ["left_click", "right_click", "double_click"]
                    }
                },
                "required": ["x", "y"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fallback_click_text",
            "description": (
                "Find a text element on screen using OCR and click it in one step. "
                "This is the most accurate way to click buttons, menu items, and labels. "
                "Use this whenever the element you want to click has readable text. "
                "If multiple matches exist, clicks the one with the highest OCR confidence."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {
                        "type": "string",
                        "description": "The visible text of the element to click. Case-insensitive substring match."
                    },
                    "click_type": {
                        "type": "string",
                        "description": "'left_click' (default), 'right_click', or 'double_click'.",
                        "enum": ["left_click", "right_click", "double_click"]
                    }
                },
                "required": ["text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "type_text",
            "description": (
                "Type text at the current cursor position using keyboard simulation. "
                "Always set expected_window to the title of the window you just clicked — "
                "this prevents accidentally typing into the wrong app (e.g. Midum's own terminal). "
                "Use special_key for Enter, Tab, Escape, F-keys etc."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "text":            {"type": "string", "description": "The text to type."},
                    "special_key":     {"type": "string", "description": "Optional key to press after typing: 'Enter', 'Tab', 'Escape', 'F5', etc."},
                    "expected_window": {"type": "string", "description": "Title substring of the window that should be in the foreground. Typing is aborted if a different window is active."}
                },
                "required": ["text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "update_memory",
            "description": "Persist important information. 'target' must be 'master', 'project', or 'session'.",
            "parameters": {
                "type": "object",
                "properties": {
                    "target":  {"type": "string"},
                    "content": {"type": "string"}
                },
                "required": ["target", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "set_current_goal",
            "description": "Update the current goal. Use goal='none' to clear when a task is done.",
            "parameters": {
                "type": "object",
                "properties": {
                    "goal":   {"type": "string"},
                    "reason": {"type": "string"}
                },
                "required": ["goal"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "load_skill",
            "description": "Load a skill file by name (without .md). Call list_skills first if unsure.",
            "parameters": {
                "type": "object",
                "properties": {
                    "skill_name": {"type": "string"}
                },
                "required": ["skill_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_skills",
            "description": "List all available skills with descriptions.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_instructions",
            "description": (
                "Read instructions.md — user preferences and behavioural rules "
                "(e.g. preferred command style, formatting, workflow habits). "
                "Consult before any task where HOW matters, not just what to do."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "add_instruction",
            "description": (
                "Add a preference or behavioural rule to instructions.md. "
                "Call when the user states a preference or corrects your behaviour."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "instruction": {"type": "string", "description": "The rule to record."}
                },
                "required": ["instruction"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_paths",
            "description": (
                "Read paths.md — absolute paths to apps, folders, and files. "
                "Consult when you need a path you are not certain of."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "add_path",
            "description": "Add a labelled path entry to paths.md.",
            "parameters": {
                "type": "object",
                "properties": {
                    "label": {"type": "string", "description": "Short name, e.g. 'Blender'."},
                    "path":  {"type": "string", "description": "Absolute path on disk."},
                    "note":  {"type": "string", "description": "Optional extra context."}
                },
                "required": ["label", "path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_domain_knowledge",
            "description": (
                "Create a domain-specific knowledge file (like commands.md but for a specific "
                "tool, e.g. blender_commands.md). Registered in domain_index.md."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "name":            {"type": "string", "description": "snake_case name, no extension."},
                    "description":     {"type": "string", "description": "One-line description."},
                    "initial_content": {"type": "string", "description": "Optional seed content."}
                },
                "required": ["name", "description"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_domain_knowledge",
            "description": "List all registered domain knowledge files with descriptions.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_domain_knowledge",
            "description": "Read a domain knowledge file by name (without .md extension).",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Filename without extension."}
                },
                "required": ["name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_domain_skill",
            "description": (
                "Create a domain-specific skill file for a tool or workflow. "
                "Stored in skills dir, registered in domain_skills_index.md and skills.md."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "name":        {"type": "string", "description": "snake_case filename, no extension."},
                    "domain":      {"type": "string", "description": "Tool this belongs to, e.g. 'blender'."},
                    "description": {"type": "string", "description": "One-line description."},
                    "content":     {"type": "string", "description": "Full Markdown skill instructions."}
                },
                "required": ["name", "domain", "description", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_domain_skills",
            "description": "List all domain-specific skills grouped by domain.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "consult_gemini",
            "description": (
                "DEFAULT consult tool — try this FIRST for reasoning/analysis/research tasks. "
                "Send a reasoning, analysis, or research task to a Gemini model via Google AI Studio. "
                "Use this when: (1) the user explicitly says 'ask Gemini' or 'consult Gemini', "
                "(2) the task requires deep reasoning, complex analysis, code review, architectural "
                "decisions, or multi-step planning that exceeds your own confident ability, "
                "(3) you need a second opinion or want to cross-check your own reasoning. "
                "Midum selects the most appropriate model automatically based on task complexity "
                "unless you specify task_type. "
                "Models available (free tier): "
                "quick=gemini-2.0-flash-lite (fast, simple tasks), "
                "balanced=gemini-2.0-flash (default, multi-step reasoning), "
                "hard=gemini-2.5-flash (complex analysis, long context), "
                "expert=gemini-2.5-pro (hardest problems, slowest)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "The full question or task for Gemini. Be specific and complete."
                    },
                    "task_type": {
                        "type": "string",
                        "description": "Model tier: 'auto' (default), 'quick', 'balanced', 'hard', or 'expert'."
                    },
                    "context": {
                        "type": "string",
                        "description": "Optional: relevant context to include (file contents, memory, prior results)."
                    }
                },
                "required": ["prompt"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "consult_openrouter",
            "description": (
                "Send a reasoning, analysis, or research task to a model on OpenRouter "
                "(currently configured: see OPENROUTER_MODEL). "
                "This is a FALLBACK/explicit-request tool, not a default choice — prefer "
                "consult_gemini first. Only use consult_openrouter when: "
                "(1) the user explicitly says 'ask OpenRouter' or names OpenRouter, or "
                "(2) consult_gemini just failed/errored and you still need a second-brain "
                "answer. Do not reach for this casually as an alternative to Gemini."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "The full question or task for the model. Be specific and complete."
                    },
                    "context": {
                        "type": "string",
                        "description": "Optional: relevant context to include (file contents, memory, prior results)."
                    },
                    "model": {
                        "type": "string",
                        "description": "Optional: override the default OpenRouter model ID for this call."
                    }
                },
                "required": ["prompt"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "delegate_to_openrouter",
            "description": (
                "Hand off an entire task to OpenRouter as a real COWORKER, not just a "
                "text consultant. Unlike consult_openrouter (which only returns text), "
                "the OpenRouter model spun up here gets FULL access to every tool Midum "
                "has — it can click UI elements, run terminal commands, read/write files, "
                "browse the web, everything — and works through the task independently, "
                "then reports back a final summary that you relay to the user. "
                "\n\n"
                "Use this when a sub-task is complex enough to benefit from a stronger "
                "model actually DOING the work end-to-end (not just planning it), or "
                "when the user explicitly asks to delegate/offload something. "
                "\n\n"
                "This runs in an ISOLATED conversation — the delegate does not see your "
                "conversation history, only what you put in `task` and `context`. Be "
                "complete and specific in both."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "task": {
                        "type": "string",
                        "description": (
                            "The complete task to hand off, written as a self-contained "
                            "instruction — the delegate has no other context except this "
                            "and whatever you put in `context`."
                        )
                    },
                    "context": {
                        "type": "string",
                        "description": "Optional: relevant background the delegate needs (file contents, prior findings, constraints)."
                    },
                    "model": {
                        "type": "string",
                        "description": "Optional: override the default OpenRouter model ID for this delegated task."
                    },
                    "max_steps": {
                        "type": "integer",
                        "description": "Optional: cap on tool-call steps the delegate can take before it must report back. Default 10."
                    }
                },
                "required": ["task"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_openrouter_models",
            "description": (
                "List models available on OpenRouter as a numbered indexed table "
                "(IDX | context length | free? | model ID). Defaults to free-tier models only. "
                "Follow up with set_openrouter_model_by_index(index) to switch the active model — "
                "this is the CHOOSE-pattern way to pick a model instead of typing an exact ID."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "free_only": {
                        "type": "boolean",
                        "description": "If true (default), only show free-tier models. Set false to see all models including paid."
                    }
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "set_openrouter_model_by_index",
            "description": "Switch the active OpenRouter model to the entry at `index` from list_openrouter_models.",
            "parameters": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer", "description": "Index from list_openrouter_models output."}
                },
                "required": ["index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "set_openrouter_model",
            "description": (
                "Switch the active OpenRouter model directly by its exact model ID "
                "(e.g. 'anthropic/claude-3.7-sonnet', 'meta-llama/llama-3.3-70b-instruct:free'). "
                "Use list_openrouter_models + set_openrouter_model_by_index instead if you don't "
                "know the exact ID. Takes effect immediately, no restart needed."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "model_id": {"type": "string", "description": "Exact OpenRouter model ID."}
                },
                "required": ["model_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "consult_gemini_api",
            "description": (
                "Send a plain reasoning/planning prompt to Gemini via the OFFICIAL "
                "Google Gemini API (real API key, structured request — not the web-chat "
                "scraping used by consult_gemini). Returns text only, no tool access — "
                "use delegate_to_gemini_api instead if the task needs real actions taken."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "The question or planning prompt."},
                    "context": {"type": "string", "description": "Optional background context."},
                    "model": {
                        "type": "string",
                        "description": "Optional: override the default Gemini API model ID for this call."
                    }
                },
                "required": ["prompt"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "delegate_to_gemini_api",
            "description": (
                "Hand off an entire task to the OFFICIAL Gemini API as a real COWORKER, "
                "not just a text consultant — same idea as delegate_to_openrouter but "
                "running on Gemini through a real API key with native structured tool "
                "calling (not the web-chat session used by delegate_to_gemini_web). The "
                "Gemini sub-agent gets FULL access to every tool Midum has — UI "
                "automation, terminal, files, browser, MCP servers, everything — and "
                "works through the task independently in an ISOLATED conversation, then "
                "reports back a final summary."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "task": {
                        "type": "string",
                        "description": (
                            "The complete task to hand off, written as a self-contained "
                            "instruction — the delegate has no other context except this "
                            "and whatever you put in `context`."
                        )
                    },
                    "context": {
                        "type": "string",
                        "description": "Optional: relevant background the delegate needs (file contents, prior findings, constraints)."
                    },
                    "model": {
                        "type": "string",
                        "description": "Optional: override the default Gemini API model ID for this delegated task."
                    },
                    "max_steps": {
                        "type": "integer",
                        "description": "Optional: cap on tool-call steps the delegate can take before it must report back. Default 10."
                    }
                },
                "required": ["task"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "set_gemini_api_model",
            "description": (
                "Switch the active Gemini API model directly by its exact model ID "
                "(e.g. 'gemini-3.1-flash-lite'). Takes effect immediately, no restart needed."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "model_id": {"type": "string", "description": "Exact Gemini API model ID."}
                },
                "required": ["model_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "consult_groq",
            "description": (
                "Send a reasoning, analysis, or research task to a model on GroqCloud "
                "(currently configured: see GROQ_MODEL) — GroqCloud's free tier gives "
                "fast inference with no credit card required. "
                "This is a FALLBACK/explicit-request tool, not a default choice — prefer "
                "consult_gemini first. Only use consult_groq when: "
                "(1) the user explicitly says 'ask Groq' or names GroqCloud, or "
                "(2) consult_gemini just failed/errored and you still need a second-brain "
                "answer. Do not reach for this casually as an alternative to Gemini."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "The full question or task for the model. Be specific and complete."
                    },
                    "context": {
                        "type": "string",
                        "description": "Optional: relevant context to include (file contents, memory, prior results)."
                    },
                    "model": {
                        "type": "string",
                        "description": "Optional: override the default Groq model ID for this call."
                    }
                },
                "required": ["prompt"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "delegate_to_groq",
            "description": (
                "Hand off an entire task to GroqCloud as a real COWORKER, not just a "
                "text consultant. Unlike consult_groq (which only returns text), "
                "the Groq model spun up here gets FULL access to every tool Midum "
                "has — it can click UI elements, run terminal commands, read/write files, "
                "browse the web, everything — and works through the task independently, "
                "then reports back a final summary that you relay to the user. GroqCloud "
                "runs on fast LPU hardware, so this is a good pick when the user wants "
                "quick free-tier turnaround on a delegated sub-task. "
                "\n\n"
                "This runs in an ISOLATED conversation — the delegate does not see your "
                "conversation history, only what you put in `task` and `context`. Be "
                "complete and specific in both."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "task": {
                        "type": "string",
                        "description": (
                            "The complete task to hand off, written as a self-contained "
                            "instruction — the delegate has no other context except this "
                            "and whatever you put in `context`."
                        )
                    },
                    "context": {
                        "type": "string",
                        "description": "Optional: relevant background the delegate needs (file contents, prior findings, constraints)."
                    },
                    "model": {
                        "type": "string",
                        "description": "Optional: override the default Groq model ID for this delegated task."
                    },
                    "max_steps": {
                        "type": "integer",
                        "description": "Optional: cap on tool-call steps the delegate can take before it must report back. Default 10."
                    }
                },
                "required": ["task"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_groq_models",
            "description": (
                "List models currently available on GroqCloud as a numbered indexed "
                "table (IDX | model ID | owner). Follow up with "
                "set_groq_model_by_index(index) to switch the active model — this is "
                "the CHOOSE-pattern way to pick a model instead of typing an exact ID."
            ),
            "parameters": {
                "type": "object",
                "properties": {},
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "set_groq_model_by_index",
            "description": "Switch the active Groq model to the entry at `index` from list_groq_models.",
            "parameters": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer", "description": "Index from list_groq_models output."}
                },
                "required": ["index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "set_groq_model",
            "description": (
                "Switch the active Groq model directly by its exact model ID "
                "(e.g. 'llama-3.3-70b-versatile', 'qwen/qwen3-32b'). "
                "Use list_groq_models + set_groq_model_by_index instead if you don't "
                "know the exact ID. Takes effect immediately, no restart needed."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "model_id": {"type": "string", "description": "Exact Groq model ID."}
                },
                "required": ["model_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "delegate_to_gemini_web",
            "description": (
                "Hand off an entire task to Gemini (via the web-app account session, "
                "gemini_webapi) as a real COWORKER, not just a text consultant — "
                "same idea as delegate_to_openrouter but running on Gemini's own "
                "account session instead. The Gemini sub-agent gets FULL access to "
                "every tool Midum has and works through the task independently in "
                "an ISOLATED conversation (its own ChatSession), then reports back a "
                "final summary. Slower per-step than delegate_to_openrouter (each "
                "step is a real gemini.google.com round trip) — prefer "
                "delegate_to_openrouter unless the user specifically asks for Gemini."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "task": {
                        "type": "string",
                        "description": (
                            "The complete task to hand off, written as a self-contained "
                            "instruction — the delegate has no other context except this "
                            "and whatever you put in `context`."
                        )
                    },
                    "context": {
                        "type": "string",
                        "description": "Optional: relevant background the delegate needs (file contents, prior findings, constraints)."
                    },
                    "max_steps": {
                        "type": "integer",
                        "description": "Optional: cap on tool-call steps the delegate can take before it must report back. Default 10."
                    }
                },
                "required": ["task"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "set_gemini_web_model",
            "description": (
                "Pin the model Gemini-web uses (when MODEL_PROVIDER=='gemini_web', or "
                "for delegate_to_gemini_web) to an exact model_name/display_name string "
                "(e.g. 'gemini-3-flash'). Pass an empty string to go back to "
                "auto-selecting the fastest available model at runtime."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "model_name": {"type": "string", "description": "Exact Gemini model name/display name, or '' for auto."}
                },
                "required": ["model_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "wait",
            "description": "Pause execution for a specific number of seconds. Use this when waiting for an application to launch, a web page to load, or a background process to complete.",
            "parameters": {
                "type": "object",
                "properties": {
                    "seconds": {
                        "type": "number",
                        "description": "The number of seconds to pause (can be a decimal, e.g., 1.5)."
                    }
                },
                "required": ["seconds"]
            }
        }
    },
    {"type":"function","function":{"name":"read_file_smart","description":"Read any file: txt/md/py/json/csv/html/.pdf(requires pymupdf)/.docx(requires mammoth). Returns chunk 1 for large files with chunk count — call read_file_chunk for rest.","parameters":{"type":"object","properties":{"path":{"type":"string","description":"Absolute path to the file."}},"required":["path"]}}},
    {"type":"function","function":{"name":"read_file_chunk","description":"Read chunk N (1-based) of a large file after read_file_smart reports multiple chunks.","parameters":{"type":"object","properties":{"path":{"type":"string"},"chunk_index":{"type":"integer","description":"1-based chunk number."}},"required":["path","chunk_index"]}}},
    {"type":"function","function":{"name":"write_docx_file","description":"Write a .docx Word document from Markdown-style text (# headings, **bold**). Requires python-docx: pip install python-docx.","parameters":{"type":"object","properties":{"path":{"type":"string","description":"Absolute path ending in .docx."},"content":{"type":"string","description":"Markdown-style text content."}},"required":["path","content"]}}},
    {
        "type": "function",
        "function": {
            "name": "generate_image",
            "description": (
                "Generate one or more images from a text prompt using Gemini's web app "
                "(gemini.google.com) via the free gemini_webapi session — the same session "
                "used by query_gemini_app/delegate_to_gemini_web. No image API key, no "
                "per-image metering, and the image is NOT saved to disk automatically — it's "
                "kept in memory and shown inline in the GUI chat with Download/Copy buttons "
                "the user can click if they want to keep it. Use this whenever the user asks "
                "you to create/generate/draw/make a picture or image."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "Detailed description of the image to generate, e.g. 'A watercolor painting of a lighthouse at sunset'."
                    },
                    "count": {
                        "type": "integer",
                        "description": "How many image variations to request (default 1)."
                    }
                },
                "required": ["prompt"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_flowchart",
            "description": (
                "Build and render a full explanatory flowchart in the chat (a real box-and-arrow "
                "diagram in the GUI, plus ASCII and Mermaid fallbacks in plain terminal mode). "
                "Use this any time you need to explain a process, decision tree, algorithm, or "
                "step-by-step system visually instead of (or in addition to) prose. "
                "Give every node a short unique 'id', a human-readable 'label', a 'type' "
                "(start | process | decision | io | end), and a 'next' list of the ids it flows "
                "into — optionally with an edge 'label' (e.g. 'yes'/'no' out of a decision node)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {
                        "type": "string",
                        "description": "Short title for the flowchart."
                    },
                    "steps": {
                        "type": "array",
                        "description": "The nodes of the flowchart, in any order (at least one should be type='start').",
                        "items": {
                            "type": "object",
                            "properties": {
                                "id": {
                                    "type": "string",
                                    "description": "Short unique identifier for this node, e.g. 'check_login'."
                                },
                                "label": {
                                    "type": "string",
                                    "description": "The text shown inside the node's box."
                                },
                                "type": {
                                    "type": "string",
                                    "enum": ["start", "process", "decision", "io", "end"],
                                    "description": "Node shape/role. 'start' = entry point, 'decision' = branching question, 'io' = input/output, 'end' = terminal node."
                                },
                                "next": {
                                    "type": "array",
                                    "description": "IDs (or {to, label} objects) this node flows into. Use edge labels like 'yes'/'no' on decision branches.",
                                    "items": {
                                        "type": "object",
                                        "properties": {
                                            "to": {"type": "string", "description": "id of the next node."},
                                            "label": {"type": "string", "description": "Optional label for this edge, e.g. 'yes', 'no', 'on error'."}
                                        },
                                        "required": ["to"]
                                    }
                                }
                            },
                            "required": ["id", "label", "type"]
                        }
                    }
                },
                "required": ["title", "steps"]
            }
        }
    },
    {"type":"function","function":{"name":"write_response_memory","description":"Overwrite the response scratchpad (response_memory.md). Call FIRST for any multi-step task with a numbered plan. Wiped automatically when set_current_goal(none) fires.","parameters":{"type":"object","properties":{"content":{"type":"string","description":"Plan, checklist, or notes."}},"required":["content"]}}},
    {"type":"function","function":{"name":"append_response_memory","description":"Append a note or partial result to the response scratchpad. Use to log progress and accumulate partial outputs during a task.","parameters":{"type":"object","properties":{"content":{"type":"string","description":"Note or partial result."}},"required":["content"]}}},
    {"type":"function","function":{"name":"read_response_memory","description":"Read the current response scratchpad to check your plan or assemble a final answer from accumulated notes.","parameters":{"type":"object","properties":{}}}},
    {
        "type": "function",
        "function": {
            "name": "say",
            "description": (
                "Print a message to the user mid-turn, then continue acting. "
                "Use this to narrate what you are doing WHILE doing it — e.g. say('Opening Chrome...') "
                "then immediately call execute_terminal_command. "
                "Do NOT use this as a substitute for acting. Always follow a say() with a real tool call."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "message": {"type": "string", "description": "The text to show the user right now."}
                },
                "required": ["message"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_directory",
            "description": (
                "List a directory as a numbered indexed table (IDX | TYPE | SIZE | NAME). "
                "PREFERRED over explore_path — returns an index you can act on. "
                "Follow up with open_path(path, index) to open a file or enter a subfolder."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Absolute path to the directory."}
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "open_path",
            "description": (
                "Act on a directory entry from list_directory by index. "
                "Directories: drills in and returns another indexed listing. "
                "Files: reads and returns the file content."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path":  {"type": "string",  "description": "The same path you passed to list_directory."},
                    "index": {"type": "integer", "description": "Entry index from list_directory output."}
                },
                "required": ["path", "index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "find_file",
            "description": (
                "Search for a file by name under a root directory. "
                "Returns a numbered indexed list of all matches. "
                "Follow up with open_path_by_index(index) to open the chosen file."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "filename":    {"type": "string", "description": "Filename or substring to search for."},
                    "search_root": {"type": "string", "description": "Optional: directory to search under. Defaults to home directory."}
                },
                "required": ["filename"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "open_path_by_index",
            "description": "Open/read a file from the last find_file() result by index.",
            "parameters": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer", "description": "Index from find_file output."}
                },
                "required": ["index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_skills_indexed",
            "description": (
                "List all available skills as a numbered indexed table. "
                "PREFERRED over list_skills — returns an index. "
                "Follow up with load_skill_by_index(index)."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "load_skill_by_index",
            "description": "Load a skill from the list_skills_indexed snapshot by index.",
            "parameters": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer", "description": "Index from list_skills_indexed output."}
                },
                "required": ["index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_paths_indexed",
            "description": (
                "Parse paths.md into a numbered indexed table (IDX | LABEL | PATH). "
                "PREFERRED over read_paths — returns an index. "
                "Follow up with get_path(index) to retrieve the exact path string."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_path",
            "description": "Return the full path string for an entry from list_paths_indexed by index.",
            "parameters": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer", "description": "Index from list_paths_indexed output."}
                },
                "required": ["index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_domain_knowledge_indexed",
            "description": (
                "List domain knowledge files as a numbered indexed table. "
                "PREFERRED over list_domain_knowledge. "
                "Follow up with read_domain_by_index(index)."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_domain_by_index",
            "description": "Read a domain knowledge file by its index from list_domain_knowledge_indexed.",
            "parameters": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer", "description": "Index from list_domain_knowledge_indexed output."}
                },
                "required": ["index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_domain_skills_indexed",
            "description": (
                "List all domain skills as a numbered indexed table grouped by domain. "
                "PREFERRED over list_domain_skills. "
                "Follow up with load_skill(name) using the name shown."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "open_search_result",
            "description": (
                "Open a web search result from the last search_internet call by index. "
                "Much faster than typing the URL manually. "
                "Call search_internet first, then open_search_result(index) to open the chosen result."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "index":   {"type": "integer", "description": "Result index from search_internet output (0-based)."},
                    "browser": {"type": "string",  "description": "Browser to open in: 'chrome' (default), 'brave', 'firefox', 'edge'."}
                },
                "required": ["index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ocr_snapshot",
            "description": (
                "Take a screenshot and OCR the entire screen into a numbered indexed table "
                "(IDX | CONF | CX | CY | TEXT). "
                "PREFERRED over fallback_find_text when you want to see all text at once. "
                "Follow up with click_ocr_index(index) to click any element by index. "
                "Use when UIA returns no elements (canvas/WebGL apps, games, etc)."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "click_ocr_index",
            "description": "Click a text element from the last ocr_snapshot() by index. Exact — no text matching needed.",
            "parameters": {
                "type": "object",
                "properties": {
                    "index":      {"type": "integer", "description": "Index from ocr_snapshot output."},
                    "click_type": {"type": "string",  "description": "'left_click' (default), 'right_click', or 'double_click'.", "enum": ["left_click", "right_click", "double_click"]}
                },
                "required": ["index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_browser_page",
            "description": (
                "Read the full text content of the currently active browser tab using Chrome DevTools Protocol (CDP). "
                "Returns the page title, URL, and clean readable text extracted from the DOM — "
                "NO HTML tags, just the actual text content. "
                "Works on ANY page including Google results, articles, YouTube descriptions, etc. "
                "Use this instead of read_aggregated_text for browser content — UIA cannot read web page text. "
                "Requires Chrome to be running with remote debugging enabled (see setup note)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "tab_index": {
                        "type": "integer",
                        "description": "Which tab to read (0 = first/active tab, 1 = second tab, etc). Default 0."
                    }
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_browser_tabs",
            "description": (
                "List all open tabs in Chrome as a numbered indexed table (IDX | TITLE | URL). "
                "Follow up with read_browser_page(tab_index=N) or open_url to navigate."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_js_in_browser",
            "description": (
                "Execute arbitrary JavaScript in the current browser tab via CDP and return the result. "
                "Use for reading page data that other tools can't reach: "
                "document.title, element values, computed content, etc. "
                "Example: run_js_in_browser(\"document.title\") returns the page title. "
                "Example: run_js_in_browser(\"document.querySelector('h1').textContent\") reads the first heading."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "script":    {"type": "string",  "description": "JavaScript expression or statement to execute."},
                    "tab_index": {"type": "integer", "description": "Tab index (0 = active). Default 0."}
                },
                "required": ["script"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "open_url",
            "description": (
                "Open a URL directly in Chrome (or the default browser). "
                "This is the fastest way to navigate — ONE call replaces: "
                "click address bar → type URL → press Enter. "
                "If Chrome is already open, the URL opens in a new tab. "
                "If Chrome is not open, it launches Chrome and opens the URL. "
                "Prefer this over manually clicking the address bar and typing."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": (
                            "The full URL to open, e.g. 'https://youtube.com', "
                            "'https://google.com/search?q=python'. "
                            "http:// or https:// prefix is added automatically if missing."
                        )
                    },
                    "browser": {
                        "type": "string",
                        "description": (
                            "Which browser to use: 'chrome' (default), 'brave', 'firefox', 'edge', "
                            "or 'default' to use the system default browser."
                        ),
                        "enum": ["chrome", "brave", "firefox", "edge", "default"]
                    }
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_mcp_servers",
            "description": (
                "List connected MCP (Model Context Protocol) servers ONLY — names, "
                "connection status, and tool COUNT, not the tools themselves. "
                "Always call this first before using any MCP server. Follow up with "
                "show_server_tools(server) to see what a specific server offers."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "show_server_tools",
            "description": (
                "Show the tool names, descriptions, and JSON input schemas for ONE "
                "connected MCP server. Call list_mcp_servers() first to get the "
                "server's index or name. Use this right before call_mcp_tool so you "
                "know the exact tool_name and arguments shape to send."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "server": {
                        "type": "string",
                        "description": "Server index (e.g. '0') or name, from list_mcp_servers()."
                    }
                },
                "required": ["server"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "call_mcp_tool",
            "description": (
                "Call ANY tool on ANY connected MCP server. This is the single, "
                "uniform way to invoke MCP tools — check show_server_tools(server) "
                "first for the exact tool_name and the arguments it expects."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "server": {
                        "type": "string",
                        "description": "Server index (e.g. '0') or name, from list_mcp_servers()."
                    },
                    "tool_name": {
                        "type": "string",
                        "description": "Exact tool name, from show_server_tools()."
                    },
                    "arguments": {
                        "type": "object",
                        "description": (
                            "Arguments object matching that tool's input schema "
                            "(from show_server_tools()). Use {} if it takes none."
                        )
                    }
                },
                "required": ["server", "tool_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "connect_mcp_server",
            "description": (
                "Connect a new MCP server and (by default) remember it for future "
                "startups — the easy way to add a server. Use transport='stdio' for "
                "a local server launched as a subprocess (command + args), or "
                "transport='http'/'sse' for a remote server reachable by URL."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "A short unique name for this server, e.g. 'filesystem'."},
                    "transport": {
                        "type": "string",
                        "enum": ["stdio", "http", "sse"],
                        "description": "How to connect. 'stdio' for a local subprocess, 'http' or 'sse' for a remote URL."
                    },
                    "command": {"type": "string", "description": "Executable to launch (stdio only), e.g. 'npx' or 'python'."},
                    "args": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Command-line arguments (stdio only), e.g. ['-y', '@modelcontextprotocol/server-filesystem', '/home/user']."
                    },
                    "url": {"type": "string", "description": "Server URL (http/sse only)."},
                    "env": {"type": "object", "description": "Extra environment variables for the subprocess (stdio only)."},
                    "headers": {"type": "object", "description": "Extra HTTP headers, e.g. an Authorization token (http/sse only)."},
                    "persist": {
                        "type": "boolean",
                        "description": "Save to storage/mcp_servers.json so it auto-connects next startup. Default true."
                    }
                },
                "required": ["name", "transport"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "disconnect_mcp_server",
            "description": "Disconnect a connected MCP server. Optionally forget it so it stops auto-connecting on future startups.",
            "parameters": {
                "type": "object",
                "properties": {
                    "server": {"type": "string", "description": "Server index (e.g. '0') or name, from list_mcp_servers()."},
                    "forget": {"type": "boolean", "description": "Also remove it from storage/mcp_servers.json. Default false."}
                },
                "required": ["server"],
            },
        },
    },
    # ── GUI user-interaction tools ─────────────────────────────────────────────
    # Optional, situational. Default is to call NONE of these. Use one only
    # when you actually need something specific from the user that you can't
    # reasonably infer or find yourself. Several can be called in the same turn.
    {
        "type": "function",
        "function": {
            "name": "ask_user_text",
            "description": (
                "Pop up a GUI textbox and ask the user to type a free-form answer. "
                "Use for anything with no fixed set of options — a name, a value, "
                "clarifying detail, etc. Blocks until the user submits or cancels."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "The question/instruction shown above the textbox."},
                    "title": {"type": "string", "description": "Dialog window title. Optional."}
                },
                "required": ["prompt"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ask_user_file_path",
            "description": (
                "Pop up a native file-picker dialog. Use this whenever the user tells "
                "you to open/read/edit/save a file but does not give you a path — ask "
                "instead of guessing. Blocks until the user picks a file or cancels."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "Dialog title, e.g. 'Select the file to open'. Optional."},
                    "must_exist": {
                        "type": "boolean",
                        "description": "True (default) for picking an existing file to open. False for a save/new-file dialog."
                    }
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ask_user_approval",
            "description": (
                "Pop up an Approve / Decline GUI dialog with two buttons. Use before "
                "anything the user should explicitly sign off on — deleting or "
                "overwriting files, sending a message on their behalf, spending money, "
                "running a risky or irreversible command. Blocks until the user clicks "
                "a button. Returns 'APPROVED' or 'DECLINED'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "message": {"type": "string", "description": "Short, bold headline of what's being approved, e.g. 'Delete 12 files in Downloads?'"},
                    "details": {"type": "string", "description": "Optional extra context/detail shown below the headline."}
                },
                "required": ["message"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ask_user_choice",
            "description": (
                "Pop up a multiple-choice GUI dialog: your question plus up to 4 "
                "options you define as buttons, and (by default) a 5th free-text box "
                "so the user can type something else. Use this to disambiguate what "
                "the user wants with a tap instead of back-and-forth text. Blocks "
                "until the user picks a button or submits custom text."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "question": {"type": "string", "description": "The question shown at the top of the dialog."},
                    "choice_1": {"type": "string", "description": "First option button label."},
                    "choice_2": {"type": "string", "description": "Second option button label. Optional."},
                    "choice_3": {"type": "string", "description": "Third option button label. Optional."},
                    "choice_4": {"type": "string", "description": "Fourth option button label. Optional."},
                    "allow_custom": {
                        "type": "boolean",
                        "description": "Whether to show a free-text 'Other...' box in addition to the buttons. Default true."
                    }
                },
                "required": ["question", "choice_1"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_native_tools",
            "description": (
                "List every built-in Midum tool by name + one-line description "
                "ONLY — not full parameter schemas. Follow up with "
                "show_native_tool_schema(tool_name) to get the exact arguments a "
                "tool expects before calling it. Mainly useful when you were not "
                "given the full native tool catalogue up front."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "show_native_tool_schema",
            "description": (
                "Show the full JSON parameter schema for ONE native Midum tool. "
                "Call list_native_tools() first to get the tool's index or exact "
                "name, then call this right before using it so you know the exact "
                "arguments shape to send."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "tool_name": {
                        "type": "string",
                        "description": "Tool index (e.g. '0') or exact name, from list_native_tools()."
                    }
                },
                "required": ["tool_name"],
            },
        },
    },
]

# =============================================================================
# GROQ LAZY TOOL LOADING (Option A)
# =============================================================================
# Groq's free tier is TPM-capped (as low as 6000 tokens/minute on some
# models). The full `tools` schema above is ~73KB / ~18k tokens on its own —
# more than the entire budget before a single message is sent. Instead of
# sending all ~88 tool schemas on every Groq call, we send a small CORE set
# covering the actions Midum needs most often, plus two meta-tools
# (list_more_tools / load_tool_by_index) that let Midum pull in any other
# tool's full schema on demand, mirroring the existing
# "list indexed, then load by index" pattern used for skills/paths/MCP.
#
# Loaded extra tools persist for the rest of the session (reset on
# "new session") so Midum doesn't have to reload the same tool repeatedly.

_TOOLS_BY_NAME = {t["function"]["name"]: t for t in tools}

GROQ_CORE_TOOL_NAMES = [
    "execute_terminal_command",
    "list_paths_indexed", "get_path",
    "list_directory", "open_path", "find_file", "open_path_by_index",
    "read_local_file", "write_local_file", "append_local_file",
    "search_internet", "open_url",
    "type_text", "wait", "say",
    "ask_user_text", "ask_user_approval", "ask_user_choice", "ask_user_file_path",
]
GROQ_CORE_TOOL_NAMES = [n for n in GROQ_CORE_TOOL_NAMES if n in _TOOLS_BY_NAME]

_GROQ_META_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "list_more_tools",
            "description": (
                "List ALL tools NOT currently loaded, as a numbered index with a "
                "one-line description each (e.g. GUI automation, browser DOM/CDP, "
                "MCP servers, domain knowledge, model delegation/consulting, memory "
                "and goal tools). The core tools already cover terminal, files, "
                "paths, search, URLs, and user prompts — only call this when none of "
                "those fit. Follow up with load_tool_by_index to actually load one."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "load_tool_by_index",
            "description": (
                "Load a tool from the most recent list_more_tools() index so it "
                "becomes callable on your NEXT turn. Loaded tools persist for the "
                "rest of the session and stack — previously loaded tools stay "
                "available."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer", "description": "Index number from list_more_tools()."}
                },
                "required": ["index"],
            },
        },
    },
]

# Tools loaded on-demand for the Groq provider this session.
_groq_extra_tools: list = []
# Numbered index cache backing the last list_more_tools() call, so
# load_tool_by_index(N) knows what N refers to.
_groq_more_tools_index: list = []


def list_more_tools() -> str:
    """Return a numbered index of every tool not already loaded for Groq."""
    global _groq_more_tools_index
    loaded_names = set(GROQ_CORE_TOOL_NAMES) | {t["function"]["name"] for t in _groq_extra_tools}
    remaining = [n for n in _TOOLS_BY_NAME if n not in loaded_names]
    _groq_more_tools_index = remaining
    if not remaining:
        return "All tools are already loaded."
    lines = []
    for i, name in enumerate(remaining):
        desc = _TOOLS_BY_NAME[name]["function"].get("description", "") or ""
        short = desc.split(". ")[0].split("\n")[0][:110]
        lines.append(f"{i}  {name}  — {short}")
    return "\n".join(lines)


def load_tool_by_index(index: int) -> str:
    """Load a tool by index from the last list_more_tools() call."""
    if not _groq_more_tools_index:
        return "Call list_more_tools() first to see what's available."
    if not (0 <= index < len(_groq_more_tools_index)):
        return f"Invalid index. Valid range: 0-{len(_groq_more_tools_index) - 1}."
    name = _groq_more_tools_index[index]
    if any(t["function"]["name"] == name for t in _groq_extra_tools):
        return f"'{name}' is already loaded."
    _groq_extra_tools.append(_TOOLS_BY_NAME[name])
    return f"Loaded '{name}'. It's now available for your next tool call."


def reset_groq_loaded_tools():
    """Clear session-loaded extra tools — called on 'new session'."""
    global _groq_extra_tools, _groq_more_tools_index
    _groq_extra_tools = []
    _groq_more_tools_index = []


def _get_groq_tools_schema() -> list:
    """Core tool subset + anything loaded on-demand + the meta-tools, kept
    small deliberately to fit Groq's tight TPM budget."""
    core = [_TOOLS_BY_NAME[n] for n in GROQ_CORE_TOOL_NAMES]
    return core + _groq_extra_tools + _GROQ_META_TOOLS


def open_url(url: str, browser: str = "chrome") -> str:
    """
    Open a URL in the specified browser.

    Strategy:
      Windows — uses Start-Process with the browser executable + URL argument,
                which opens a new tab if the browser is already running.
      Linux   — uses xdg-open for 'default', otherwise nohup + browser binary.

    The URL is normalised (http:// prefix added if missing).
    """
    # Normalise URL
    url = url.strip()
    if url and not re.match(r"^https?://", url, re.IGNORECASE):
        url = "https://" + url

    browser = (browser or "chrome").strip().lower()
    print(f"   [open_url] {browser} → {url}")

    if _IS_LINUX:
        if browser == "default":
            cmd = f"xdg-open '{url}'"
        else:
            _LINUX_BROWSERS = {
                "chrome":   ["/usr/bin/google-chrome",   "/usr/bin/google-chrome-stable",
                             "/usr/bin/chromium",         "/usr/bin/chromium-browser"],
                "brave":    ["/usr/bin/brave-browser",   "/usr/bin/brave"],
                "firefox":  ["/usr/bin/firefox"],
                "edge":     ["/usr/bin/microsoft-edge",  "/usr/bin/microsoft-edge-stable"],
            }
            candidates = _LINUX_BROWSERS.get(browser, _LINUX_BROWSERS["chrome"])
            exe = next((p for p in candidates if os.path.exists(p)), None)
            if not exe:
                # Fall back to xdg-open
                cmd = f"xdg-open '{url}'"
            else:
                cmd = f"nohup '{exe}' '{url}' &>/dev/null &"
        result = execute_terminal_command(cmd)
        if "error" in result.lower() and "nohup" not in result.lower():
            return f"Error opening URL: {result}"
        return f"Success: opened '{url}' in {browser}."

    else:
        # Windows — use Start-Process with the browser name; Chrome/Brave accept
        # a URL as the first positional argument and open it in a new tab.
        _WIN_BROWSERS = {
            "chrome":  "chrome.exe",
            "brave":   "brave.exe",
            "firefox": "firefox.exe",
            "edge":    "msedge.exe",
            "default": None,
        }
        exe = _WIN_BROWSERS.get(browser)

        if browser == "default" or exe is None:
            cmd = f"Start-Process '{url}'"
        else:
            cmd = f"Start-Process '{exe}' -ArgumentList '{url}'"

        result = execute_terminal_command(cmd)
        stderr = result.split("STDERR:")[-1].strip() if "STDERR:" in result else ""

        # If exe not found on PATH, try to find it via paths.md
        if stderr and ("cannot find" in stderr.lower() or "not recognized" in stderr.lower()):
            # Try common hardcoded paths as a last resort
            _FALLBACK_PATHS = {
                "chrome.exe":  r"C:\Program Files\Google\Chrome\Application\chrome.exe",
                "brave.exe":   r"C:\Program Files\BraveSoftware\Brave-Browser\Application\brave.exe",
                "firefox.exe": r"C:\Program Files\Mozilla Firefox\firefox.exe",
                "msedge.exe":  r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
            }
            fallback = _FALLBACK_PATHS.get(exe)
            if fallback and os.path.exists(fallback):
                cmd2    = f"Start-Process '{fallback}' -ArgumentList '{url}'"
                result2 = execute_terminal_command(cmd2)
                stderr2 = result2.split("STDERR:")[-1].strip() if "STDERR:" in result2 else ""
                if not stderr2 or len(stderr2) < 5:
                    return f"Success: opened '{url}' in {browser} (via fallback path)."
            return (
                f"Could not find '{exe}'. "
                f"Call list_paths_indexed() to find the browser path, "
                f"then use execute_terminal_command(\"Start-Process 'C:\\\\...\\\\{exe}' "
                f"-ArgumentList '{url}'\")."
            )

        if stderr and len(stderr) > 5:
            return f"open_url warning: {stderr[:150]}"

        return f"Success: opened '{url}' in {browser}."


# =============================================================================
# BROWSER DOM — Chrome DevTools Protocol (CDP)
# =============================================================================
#
# HOW TO ENABLE (one-time setup):
#
#   Windows: Create a shortcut to chrome.exe with the extra flag:
#     --remote-debugging-port=9222
#   Or launch from terminal:
#     "C:\...\chrome.exe" --remote-debugging-port=9222
#
#   Linux:
#     google-chrome --remote-debugging-port=9222 &
#
#   The port only needs to be open once per Chrome session.
#   After enabling, all CDP tools work instantly with no extra setup.
#
# INSTALLATION:
#   pip install websocket-client requests
#
# =============================================================================

CDP_PORT    = 9222
CDP_HOST    = "localhost"
CDP_BASE    = f"http://{CDP_HOST}:{CDP_PORT}"

_cdp_ws_cache: dict[int, str] = {}   # tab_index → websocket URL

try:
    import websocket as _websocket
    import json as _json_mod
    _requests = requests   # reuse the module-level import (already loaded above)
    _CDP_AVAILABLE = requests is not None
except ImportError:
    _CDP_AVAILABLE = False

def _cdp_get_tabs() -> list[dict]:
    """Return list of open tab descriptors from Chrome's /json endpoint."""
    if not _CDP_AVAILABLE:
        return []
    try:
        resp = _requests.get(f"{CDP_BASE}/json", timeout=3)
        tabs = [t for t in resp.json() if t.get("type") == "page"]
        return tabs
    except Exception:
        return []

def _cdp_call(ws_url: str, method: str, params: dict = None, _retries: int = 2) -> dict:
    """
    Send a single CDP command over a fresh WebSocket connection and return result.
    Retries on connection errors. Raises on JS exceptions so callers can handle them.
    """
    if not _CDP_AVAILABLE:
        raise RuntimeError("websocket-client not installed: pip install websocket-client")
    last_err = None
    for attempt in range(_retries + 1):
        ws = None
        try:
            # Chrome 111+ enforces an Origin allowlist on CDP WebSocket
            # connections and rejects handshakes with no matching Origin
            # header (this is the "origin policy" rejection). Sending an
            # explicit localhost Origin header satisfies the default
            # allowlist so --remote-allow-origins=* is not required.
            ws = _websocket.create_connection(
                ws_url, timeout=12,
                origin=f"http://{CDP_HOST}:{CDP_PORT}"
            )
            msg = _json_mod.dumps({"id": 1, "method": method, "params": params or {}})
            ws.send(msg)
            # Drain until we get our response (id==1); skip CDP events
            deadline = time.time() + 12
            while time.time() < deadline:
                raw  = ws.recv()
                data = _json_mod.loads(raw)
                if data.get("id") == 1:
                    # Propagate runtime exceptions as Python errors
                    exc = (data.get("result") or {}).get("exceptionDetails")
                    if exc:
                        msg_text = exc.get("text") or exc.get("exception", {}).get("description", "JS error")
                        raise RuntimeError(f"CDP JS exception: {msg_text}")
                    return data
            raise TimeoutError("CDP response timed out")
        except (RuntimeError, TimeoutError):
            raise   # don't retry logic errors
        except Exception as e:
            last_err = e
            if attempt < _retries:
                time.sleep(0.3 * (attempt + 1))
        finally:
            if ws:
                try:
                    ws.close()
                except Exception:
                    pass
    final_err = str(last_err) if last_err else "unknown"
    if "403" in final_err or "handshake" in final_err.lower() or "forbidden" in final_err.lower():
        raise ConnectionError(
            f"CDP WebSocket rejected (likely an origin policy issue): {final_err}\n"
            f"Fix: relaunch Chrome/Brave with BOTH flags: "
            f"--remote-debugging-port={CDP_PORT} --remote-allow-origins=*"
        )
    raise ConnectionError(f"CDP call failed after {_retries + 1} attempts: {last_err}")

def _cdp_ws_for_tab(tab_index: int = 0) -> str | None:
    """Return the WebSocket debugger URL for a tab."""
    tabs = _cdp_get_tabs()
    if not tabs or tab_index >= len(tabs):
        return None
    return tabs[tab_index].get("webSocketDebuggerUrl")

def _cdp_require() -> str | None:
    """Return error string if CDP is unavailable, else None."""
    if not _CDP_AVAILABLE:
        return (
            "CDP not available. Install: pip install websocket-client requests\n"
            "Then launch Chrome with: --remote-debugging-port=9222"
        )
    tabs = _cdp_get_tabs()
    if not tabs:
        return (
            "Chrome is not running with remote debugging enabled.\n"
            "Launch Chrome with: --remote-debugging-port=9222\n"
            "  Windows: Create a shortcut with that flag appended to the Target.\n"
            "  Linux:   google-chrome --remote-debugging-port=9222 &\n"
            "If tabs ARE listed but WebSocket calls still fail with an origin/policy "
            "error, add --remote-allow-origins=* as well (Midum already sends a "
            "matching Origin header, so this should rarely be needed)."
        )
    return None


# =============================================================================
# GEMINI WEB APP CLIENT — via community `gemini_webapi` library
# =============================================================================
#
# query_gemini_app() no longer drives a real Chrome tab over CDP. Instead it
# uses the community-maintained `gemini_webapi` library
# (https://github.com/HanaokaYuzu/Gemini-API), which talks directly to
# gemini.google.com's internal endpoints using your browser's session
# cookies. No page load, no clicking into a prompt box, no clipboard —
# just a plain async HTTP call. It still uses the FREE web app's own
# usage limits, not the metered developer API.
#
# INSTALLATION:
#   pip install -U gemini_webapi
#   pip install -U browser-cookie3      (optional, see below)
#
# AUTHENTICATION (one-time), whichever is easier:
#   A) Automatic — install browser-cookie3 and be logged into
#      https://gemini.google.com in a supported browser (Chrome, Firefox,
#      Edge, Brave, etc — see the project README for the full list).
#      gemini_webapi will pull the session cookies straight from your
#      browser's cookie store; no further config needed.
#   B) Manual — log into https://gemini.google.com, open DevTools (F12)
#      -> Network tab -> refresh -> find the __Secure-1PSID and
#      __Secure-1PSIDTS cookies (Application/Storage tab, or a request's
#      Cookie header), then add them to the secrets file:
#        { "GEMINI_SECURE_1PSID": "...", "GEMINI_SECURE_1PSIDTS": "..." }
# =============================================================================

_GEMINI_WEBAPI_AVAILABLE = False
_gemini_webapi_load_msg  = ""
try:
    from gemini_webapi import GeminiClient as _GeminiWebClient
    from gemini_webapi.constants import Model as _GeminiModelEnum
    _GEMINI_WEBAPI_AVAILABLE = True

    # gemini_webapi logs internally via loguru at DEBUG/WARNING level,
    # which by default dumps raw HTTP response bodies (the ")]}'" /
    # array-of-arrays payloads you saw printed) straight to stderr on
    # every request — none of that is an actual Midum error, it's just
    # unfiltered library-internal logging. Cap it to ERROR so only real
    # failures surface; auth/session recovery is handled by our own
    # retry logic below, not by reading these log lines.
    try:
        from loguru import logger as _gemini_loguru_logger
        _gemini_loguru_logger.remove()
        _gemini_loguru_logger.add(lambda msg: None, level="ERROR")
    except Exception:
        pass
except ImportError as _e:
    _GeminiWebClient = None
    _GeminiModelEnum = None
    _gemini_webapi_load_msg = f"gemini_webapi not installed: {_e}. Run: pip install -U gemini_webapi"

_gemini_web_client      = None   # lazily-initialised GeminiClient (singleton)
_gemini_web_client_lock = threading.Lock()

_gemini_async_loop   = None       # persistent background asyncio loop
_gemini_async_thread = None


def _get_gemini_async_loop():
    """
    Start (once) a background thread running a persistent asyncio event
    loop. gemini_webapi is async; this lets the rest of this otherwise-sync
    script call into it and block for a result, while still letting the
    library's background cookie-refresh task keep running between calls.
    """
    global _gemini_async_loop, _gemini_async_thread
    if _gemini_async_loop is not None:
        return _gemini_async_loop

    _gemini_async_loop = asyncio.new_event_loop()

    def _runner():
        asyncio.set_event_loop(_gemini_async_loop)
        _gemini_async_loop.run_forever()

    _gemini_async_thread = threading.Thread(target=_runner, daemon=True, name="gemini-webapi-loop")
    _gemini_async_thread.start()
    return _gemini_async_loop


def _run_gemini_coro(coro, timeout: float = 90.0):
    """Run a gemini_webapi coroutine on the background loop and block for the result."""
    loop = _get_gemini_async_loop()
    fut  = asyncio.run_coroutine_threadsafe(coro, loop)
    return fut.result(timeout=timeout)


def _get_gemini_web_client():
    """
    Lazily create (and cache) the GeminiClient singleton.

    Cookie resolution order:
      1. GEMINI_SECURE_1PSID / GEMINI_SECURE_1PSIDTS from the secrets file,
         if present.
      2. Otherwise, pass no cookies and let gemini_webapi try to pull them
         automatically via browser-cookie3 (if installed) from whatever
         browser you're logged into gemini.google.com in.

    Returns (client, None) on success, or (None, error_message) on failure.
    """
    global _gemini_web_client
    if _gemini_web_client is not None:
        return _gemini_web_client, None
    if not _GEMINI_WEBAPI_AVAILABLE:
        return None, _gemini_webapi_load_msg

    with _gemini_web_client_lock:
        if _gemini_web_client is not None:
            return _gemini_web_client, None

        secure_1psid, secure_1psidts = "", ""
        secrets_path = os.path.abspath(SECRETS_FILE)
        try:
            if os.path.exists(secrets_path):
                with open(secrets_path, "r", encoding="utf-8") as f:
                    secrets = json.load(f)
                secure_1psid   = secrets.get("GEMINI_SECURE_1PSID", "").strip()
                secure_1psidts = secrets.get("GEMINI_SECURE_1PSIDTS", "").strip()
        except Exception:
            pass   # fall through to browser-cookie3 auto-detection

        try:
            client = _GeminiWebClient(secure_1psid or None, secure_1psidts or None)
            _run_gemini_coro(
                client.init(timeout=30, auto_close=False, auto_refresh=True),
                timeout=35,
            )
            _gemini_web_client = client
            return _gemini_web_client, None
        except Exception as e:
            return None, (
                f"Failed to initialise Gemini web client: {e}. Either add "
                f"GEMINI_SECURE_1PSID/GEMINI_SECURE_1PSIDTS to the secrets file "
                f"({secrets_path}), or install browser-cookie3 (pip install -U "
                f"browser-cookie3) and make sure you're logged into "
                f"https://gemini.google.com in a supported browser."
            )




# =============================================================================
# GEMINI-WEB PRIMARY EXECUTION BACKEND
# =============================================================================
# Implements MODEL_PROVIDER == "gemini_web": Gemini, driven through the same
# gemini_webapi client/session plumbing set up above, as a full tool-calling
# execution brain sitting alongside Ollama and OpenRouter. See the module
# docstring / implementation brief for the full architecture. Summary:
#
#   - NO GEM. Gems proved problematic (flaky create/update/fetch round trips,
#     an extra persistent server-side object to get out of sync, silent
#     "no-gem mode" degradation) for what they bought us. Instead the same
#     persona + JSON tool-call-format instructions are sent as a single
#     plain-text priming message — the FIRST message of a fresh ChatSession
#     — exactly once per conversation, no create_gem/update_gem/fetch_gems
#     calls anywhere in this backend anymore.
#   - The native tool schema is NOT inlined into that priming message either.
#     Instead Gemini gets two native tools of its own — list_native_tools()
#     and show_native_tool_schema(tool_name) — and discovers native tools on
#     demand in-turn, the exact same pattern already used for MCP tools via
#     list_mcp_servers()/show_server_tools(). This keeps every hop's prompt
#     small and means the tool catalogue can never drift out of sync with a
#     stale cached Gem prompt, since there's nothing cached server-side to
#     go stale.
#   - One persistent gemini_webapi ChatSession per conversation_history
#     object carries multi-hop tool-loop context across round trips; new
#     tool results are injected into that SAME session as new messages
#     (prefixed with GEMINI_WEB_TOOL_RESULT_MARKER), never as independent
#     generate_content() calls (which would lose the session).
#   - Output is parsed with the exact same legacy JSON tool-call parser used
#     for Ollama/OpenRouter free-text tool calls (_extract_legacy_tool_calls),
#     scanning ONLY output.text — never output.thoughts.
#   - A `source`/`server` field on the emitted JSON (see _try_parse_tool_json)
#     disambiguates native vs. MCP dispatch; this is threaded through the
#     shared parser so it benefits every provider, not just Gemini.
# =============================================================================

# Concurrency guard: the scraped-cookie session backing _gemini_web_client is
# a SINGLE account session. Multiple simultaneous tool-loop tasks hammering
# it concurrently would interleave ChatSession state unpredictably. Rather
# than silently allowing that, every Gemini-web primary call is serialized
# through this lock — a second concurrent task simply waits its turn instead
# of corrupting the first task's in-progress tool loop.
_gemini_primary_call_lock = threading.Lock()

# One gemini_webapi ChatSession per conversation_history object, keyed by
# id() (conversation_history is a plain list owned/mutated by the caller,
# never replaced mid-conversation in this codebase, so id() is stable for
# the conversation's lifetime). Also tracks how much of that list has
# already been delivered to Gemini, so each hop only sends the NEW
# messages (the session already has everything sent previously) rather
# than re-sending the whole transcript as text every step.
#   id(conversation_history) -> {
#       "chat": ChatSession,
#       "sent": list[dict]   # exact message objects already delivered
#   }
_gemini_web_sessions: dict = {}

# Cached "fastest available model" resolution, refreshed once per process
# (or whenever GEMINI_WEB_MODEL changes) rather than on every hop.
_gemini_web_model_cache = None


def _gemini_web_persona_prompt() -> str:
    """
    Builds the plain-text persona + JSON tool-call output convention that
    used to live in a Gem. Sent as the FIRST message of every fresh
    ChatSession (see _gemini_web_get_session/_call_gemini_web_primary) —
    no create_gem/update_gem/fetch_gems round trip, nothing pinned
    server-side, nothing that can go stale or silently fall back to
    "no-gem mode".

    Deliberately does NOT inline the native tool schema. Instead Gemini
    gets list_native_tools()/show_native_tool_schema(tool_name) — the same
    on-demand discovery pattern already used for MCP tools via
    list_mcp_servers()/show_server_tools() — so this priming message stays
    small and the tool catalogue can never drift out of sync.
    """
    return (
        "You are Midum, a desktop AI agent. You act by emitting tool calls "
        "as JSON — you have no native function-calling protocol here, so "
        "this JSON convention IS your only way to act.\n"
        "\n"
        "\u2501\u2501\u2501 TOOL-CALL OUTPUT FORMAT \u2501\u2501\u2501\n"
        "When you want to call a tool, your ENTIRE reply must be ONE JSON "
        "object and nothing else \u2014 no commentary before or after, no markdown "
        "fences required (but tolerated if you include them):\n"
        '  {"name": "<tool_name>", "arguments": {<args>}, "source": "<source>"}\n'
        "\n"
        "`source` tells the harness which catalogue the tool comes from:\n"
        '  - "native"   \u2014 a built-in Midum tool. This is almost every call.\n'
        '  - "mcp:<server>" \u2014 the tool came from show_server_tools(<server>) for\n'
        "                 a connected MCP server (e.g. \"mcp:jira\"). Only use this\n"
        "                 after you've actually called show_server_tools for that\n"
        "                 server this session \u2014 don't guess an MCP tool exists.\n"
        "If you omit `source`, the harness assumes \"native\".\n"
        "\n"
        "Call exactly ONE tool per reply. Never batch multiple tool calls into one "
        "JSON object or array \u2014 one at a time, wait for the [TOOL_RESULT], then "
        "decide the next step.\n"
        "\n"
        "When you are done and have a final answer for the user (no more tools "
        "needed), reply with PLAIN TEXT \u2014 no JSON, no `name`/`arguments` keys.\n"
        "\n"
        "\u2501\u2501\u2501 WHERE TOOLS COME FROM (nothing is pre-loaded \u2014 discover, don't guess) \u2501\u2501\u2501\n"
        "- NATIVE tools: every tool Midum has built in. You are NOT given their "
        "full schemas up front. Call list_native_tools() (a native tool) to see "
        "names + one-line descriptions, then show_native_tool_schema(tool_name) "
        "for the exact arguments a specific tool expects, before calling it for "
        "the first time this session.\n"
        "- MCP tools: external servers connected to Midum, which can change "
        "independently of this prompt. Call list_mcp_servers() to see what's "
        "connected, then show_server_tools(server) to get that ONE server's tool "
        "schemas before calling any of its tools. Never assume an MCP tool's "
        "name or arguments without having called show_server_tools for it first "
        "this session.\n"
        "- call_mcp_tool(server, tool_name, arguments) is the uniform way to "
        "invoke any MCP tool once you know its schema, if you'd rather route "
        "through it explicitly than emit source=\"mcp:<server>\" directly \u2014 both "
        "work identically, the harness normalizes to the same dispatch.\n"
        "- Once you've discovered a tool's schema this session, you don't need "
        "to look it up again \u2014 just call it.\n"
        "\n"
        "\u2501\u2501\u2501 ACT, DON'T NARRATE \u2501\u2501\u2501\n"
        "If you have a next step to take, emit the tool-call JSON for it. Never "
        "write out what you're 'about to do' as plain text instead of doing it. "
        "Plain-text replies with no tool call are ONLY for a genuine final "
        "answer, or a question you need the user to answer before continuing.\n"
    )


def _gemini_web_pick_model(client):
    """
    Resolve the model to drive the primary execution loop with.

      - If GEMINI_WEB_MODEL is set, use it verbatim (matched against
        client.list_models() by model_name/display_name).
      - Otherwise auto-pick the fastest suitable model from whatever the
        account currently has available (never hardcode a specific fast/
        lite model name — the web app's lineup isn't stable over time).
        Heuristic: prefer a "flash"-tier model over "pro"/"advanced"/
        "thinking" variants, since latency is a real cost here (§2.4/§3.5).

    Returns a Model/AvailableModel/str suitable for the `model=` kwarg on
    ChatSession/generate_content, or Model.UNSPECIFIED (let Gemini pick)
    if nothing usable was found.
    """
    global _gemini_web_model_cache
    if _gemini_web_model_cache is not None:
        return _gemini_web_model_cache

    if GEMINI_WEB_MODEL:
        _gemini_web_model_cache = GEMINI_WEB_MODEL
        return _gemini_web_model_cache

    try:
        available = client.list_models() or []
    except Exception:
        available = []

    if not available:
        _gemini_web_model_cache = _GeminiModelEnum.UNSPECIFIED if _GeminiModelEnum else None
        return _gemini_web_model_cache

    def _rank(m):
        name = (m.model_name or "").lower() + " " + (m.display_name or "").lower()
        if "flash" in name and "thinking" not in name and "advanced" not in name and "plus" not in name:
            return 0
        if "flash" in name:
            return 1
        return 2

    available.sort(key=_rank)
    chosen = available[0]
    print(f"🧠 [Gemini-web] Auto-selected model: {chosen.model_name} ({chosen.display_name})")
    _gemini_web_model_cache = chosen
    return _gemini_web_model_cache


def _gemini_web_render_message(msg: dict) -> str | None:
    """
    Render one OpenAI-style conversation_history message into Gemini-web
    prompt text. Returns None for messages that shouldn't be sent at all
    (Gemini's own prior assistant turns are already in the ChatSession
    server-side — re-sending them would duplicate context).
    """
    role    = msg.get("role", "")
    content = msg.get("content", "")
    if isinstance(content, list):
        # Some call sites build OpenAI-style multi-part content; flatten to text.
        content = "\n".join(
            part.get("text", "") for part in content
            if isinstance(part, dict) and part.get("type") == "text"
        )
    content = (content or "").strip()

    if role == "system":
        # The persona/system prompt is primed once as the first turn of a
        # fresh ChatSession (see _gemini_web_persona_prompt /
        # _call_gemini_web_primary) — resending conversation_history's own
        # system entries on every hop would just bloat each round trip for
        # no benefit. Skip entirely.
        return None
    if role == "assistant":
        # Gemini's own prior turns already live server-side in the
        # ChatSession; nothing to inject back.
        return None
    if role == "tool":
        if not content:
            return None
        return f"{GEMINI_WEB_TOOL_RESULT_MARKER} {content}"
    # role == "user" (or unknown, treated as a plain user turn)
    return content if content else None


def _gemini_web_get_session(conversation_history: list, client, model):
    """
    Get-or-create the persistent ChatSession for this conversation, and
    return (chat_session, delta_messages, is_new) where delta_messages is
    the list of conversation_history entries not yet delivered to Gemini,
    and is_new is True the very first time this session is created (so the
    caller knows to prime it with the persona prompt as the first turn —
    no Gem, so nothing is pinned server-side automatically).

    If conversation_history has been rewritten out from under us in a way
    that isn't a clean append (e.g. the sliding HISTORY_WINDOW in
    process_chat_turn dropped older entries, or this is a genuinely new
    task reusing the same list object), the safe move is to start a FRESH
    ChatSession and resend the full current window as one prompt — silent
    corruption of an old session's dangling state is worse than one extra
    full-context turn.
    """
    key   = id(conversation_history)
    state = _gemini_web_sessions.get(key)

    if state is not None:
        sent = state["sent"]
        if len(conversation_history) >= len(sent) and conversation_history[:len(sent)] == sent:
            delta = conversation_history[len(sent):]
            state["sent"] = list(conversation_history)
            return state["chat"], delta, False
        # History diverged from what we last sent (window slid, or history
        # was mutated) — treat as a fresh task on a fresh session.
        print("   [Gemini-web] Conversation history diverged from session state — "
              "starting a fresh ChatSession for this task.")

    chat = client.start_chat(model=model)
    _gemini_web_sessions[key] = {"chat": chat, "sent": list(conversation_history)}
    return chat, list(conversation_history), True


def _gemini_web_reset_session(conversation_history: list):
    """Drop the cached session for this conversation (used on unrecoverable
    session/auth failures, so the next call starts clean instead of
    repeatedly hammering a dead session)."""
    _gemini_web_sessions.pop(id(conversation_history), None)


def _call_gemini_web_primary(messages, result_q, model_override: str = None):
    """
    Drives ONE step of the primary tool-calling loop via Gemini through
    gemini_webapi, on a persistent ChatSession. Populates result_q with the
    exact same ("ok", resp) / ("err", exc) contract as
    _call_ollama/_call_openrouter_primary, so process_chat_turn is
    completely unaware Gemini is driving instead of a local/OpenRouter
    model — this is the entire point of matching the existing dispatch
    pattern instead of a parallel one.

    `messages` here is the SAME trimmed conversation_history list
    process_chat_turn already threads through the other two backends
    (system messages + a sliding window of the rest) — used both to derive
    the ChatSession delta (§2.4) and, on a fresh/diverged session, as the
    full resend content.

    Session-recovery (§3.5): on an auth/session failure mid-loop, one
    reinit + one retry is attempted before surfacing a clear recoverable
    error — never silently corrupting the loop by pretending the call
    succeeded.
    """
    try:
        from gemini_webapi.exceptions import AuthError, GeminiError, TimeoutError as _GeminiTimeoutError
    except ImportError:
        AuthError = GeminiError = _GeminiTimeoutError = Exception  # pragma: no cover

    def _attempt():
        client, err = _get_gemini_web_client()
        if err:
            raise RuntimeError(f"Gemini web client unavailable: {err}")

        model = model_override or _gemini_web_pick_model(client)

        chat, delta, is_new = _gemini_web_get_session(messages, client, model)

        rendered = [r for r in (_gemini_web_render_message(m) for m in delta) if r]
        if is_new:
            # No Gem here — prime the fresh ChatSession with the persona +
            # tool-call-format instructions as the very first turn, exactly
            # once per conversation, instead of a pinned server-side object.
            rendered = [_gemini_web_persona_prompt()] + rendered
        if not rendered:
            # Nothing new to say (can happen right after a fresh-session
            # full resend where everything was system/assistant messages).
            # Fall back to a neutral nudge so we never send an empty prompt.
            prompt = "[SYSTEM]: Continue with the task."
        else:
            prompt = "\n\n".join(rendered)

        with _gemini_primary_call_lock:
            output = _run_gemini_coro(
                chat.send_message(prompt),
                timeout=GEMINI_WEB_HOP_TIMEOUT,
            )
        return output

    try:
        try:
            output = _attempt()
        except (AuthError, _GeminiTimeoutError, GeminiError) as e:
            # Session fragility is a first-class risk (§ context) — one
            # recovery attempt: drop the dead client/session and retry
            # exactly once before giving up.
            print(f"⚠️  [Gemini-web] Session/auth error mid-loop ({e}) — "
                  f"reinitializing and retrying once...")
            global _gemini_web_client
            _gemini_web_client = None
            _gemini_web_reset_session(messages)
            output = _attempt()

        raw_text = output.text or ""   # NEVER read output.thoughts here (§3.3)

        legacy_calls, cleaned_content = _extract_legacy_tool_calls(raw_text)
        if legacy_calls:
            legacy_calls = legacy_calls[:1]   # one tool at a time (§ non-goals)

        resp = {
            "message": {
                "role": "assistant",
                "content": cleaned_content,
                "tool_calls": legacy_calls,
            }
        }
        result_q.put(("ok", resp))

    except Exception as e:
        _gemini_web_reset_session(messages)
        result_q.put(("err", (
            f"Gemini-web session/request failed: {e}. If this persists, the "
            f"scraped cookie session likely expired — refresh "
            f"GEMINI_SECURE_1PSID/GEMINI_SECURE_1PSIDTS in the secrets file, "
            f"or re-login in the browser gemini_webapi pulls cookies from."
        )))


def set_gemini_web_model(model_name: str) -> str:
    """
    Pin GEMINI_WEB_MODEL to an exact model name/display name, or pass ""
    to go back to auto-selecting the fastest available model. Takes effect
    on the next primary-loop call — no restart needed.
    """
    global GEMINI_WEB_MODEL, _gemini_web_model_cache
    old = GEMINI_WEB_MODEL
    GEMINI_WEB_MODEL = (model_name or "").strip()
    _gemini_web_model_cache = None
    return (f"Gemini-web model changed: '{old or '(auto)'}' → "
            f"'{GEMINI_WEB_MODEL or '(auto)'}'. Takes effect on the next call.")


def delegate_to_gemini_web(task: str, context: str = "", max_steps: int = 10) -> str:
    """
    Mirror of delegate_to_openrouter(): hand `task` off to a fresh,
    FULLY TOOL-CAPABLE agent loop running on Gemini-web — full access to
    every tool Midum has, in an ISOLATED conversation (its own
    ChatSession, not shared with the caller's), reporting back a final
    summary.
    """
    if not _GEMINI_WEBAPI_AVAILABLE:
        return f"Gemini web client is not available: {_gemini_webapi_load_msg}"

    print(f"   [Delegate → Gemini-web] Task: {task[:80]}")

    _saved_scratchpad = None
    try:
        if os.path.exists(RESPONSE_MEMORY):
            with open(RESPONSE_MEMORY, "r", encoding="utf-8") as f:
                _saved_scratchpad = f.read()
    except Exception:
        pass

    try:
        sub_system_prompt = get_system_prompt(effective_provider="gemini_web")
        sub_system_prompt += (
            "\n\n━━━ DELEGATED TASK MODE ━━━\n"
            "You have been handed a specific task by Midum (the primary agent) to "
            "complete independently. Act autonomously to complete it. When finished, "
            "reply with a clear plain-text summary of what you did and the result — "
            "this summary is relayed directly to the user."
        )
        task_message = task.strip()
        if context.strip():
            task_message = f"[CONTEXT FROM MIDUM]\n{context.strip()}\n\n[TASK]\n{task_message}"

        sub_history = [
            {"role": "system", "content": sub_system_prompt},
            {
                "role": "user",
                "content": (
                    f"{task_message}\n\n"
                    "[SYSTEM]: Act immediately. Emit the first tool-call JSON now. "
                    "Do not explain — just act. Give a final plain-text summary when done."
                ),
            },
        ]

        summary, sub_tool_outputs = process_chat_turn(
            sub_history,
            user_request=task,
            force_provider="gemini_web",
            max_steps=max_steps,
        )
        step_note = f" ({len(sub_tool_outputs)} tool call(s) executed)" if sub_tool_outputs else ""
        return f"[Gemini-web coworker — task complete{step_note}]\n{summary}"

    except Exception as e:
        return f"Delegation to Gemini-web failed: {e}"

    finally:
        try:
            if _saved_scratchpad is not None:
                with open(RESPONSE_MEMORY, "w", encoding="utf-8") as f:
                    f.write(_saved_scratchpad)
        except Exception:
            pass




# JS helpers injected into pages
_JS_GET_TEXT = """
(function() {
    // Remove script/style/noscript nodes then return innerText
    var clone = document.body.cloneNode(true);
    ['script','style','noscript','nav','footer','header'].forEach(function(tag) {
        Array.from(clone.querySelectorAll(tag)).forEach(function(el) { el.remove(); });
    });
    var text = clone.innerText || clone.textContent || '';
    // Collapse excess whitespace
    return text.replace(/[ \\t]+/g, ' ').replace(/\\n{3,}/g, '\\n\\n').trim().slice(0, 15000);
})()
"""

_JS_GET_ELEMENTS = """
(function(filterType) {
    var selectors = {
        'button':   'button, [role="button"], input[type="button"], input[type="submit"]',
        'link':     'a[href]',
        'input':    'input:not([type="hidden"]), textarea',
        'select':   'select',
        'textarea': 'textarea',
        '':         'a[href], button, input:not([type="hidden"]), select, textarea, [role="button"], [role="link"], [role="tab"], [role="menuitem"], [role="option"], [tabindex]:not([tabindex="-1"])'
    };
    var sel = selectors[filterType] || selectors[''];
    var els = Array.from(document.querySelectorAll(sel));
    var results = [];
    els.forEach(function(el, i) {
        var rect = el.getBoundingClientRect();
        if (rect.width <= 0 || rect.height <= 0) return;   // hidden
        if (rect.top < -100 || rect.bottom > window.innerHeight + 100) return;  // off-screen
        var type = el.tagName.toLowerCase();
        if (el.getAttribute('role')) type = el.getAttribute('role');
        var label = el.innerText || el.value || el.placeholder ||
                    el.getAttribute('aria-label') || el.getAttribute('title') ||
                    el.getAttribute('name') || el.getAttribute('id') || '';
        label = label.trim().replace(/\\s+/g, ' ').slice(0, 80);
        results.push({
            idx: results.length,
            type: type,
            label: label,
            tag: el.tagName.toLowerCase(),
            href: el.href || '',
            x: Math.round(rect.left + rect.width/2),
            y: Math.round(rect.top  + rect.height/2),
        });
    });
    return results;
})(FILTER_TYPE_PLACEHOLDER)
"""

# Per-call DOM element snapshot cache: tab_index → list of element dicts
_browser_element_cache: dict[int, list] = {}


def list_browser_tabs() -> str:
    err = _cdp_require()
    if err:
        return err
    tabs = _cdp_get_tabs()
    if not tabs:
        return "No open tabs found."
    _store_index("browser_tabs", [{"title": t.get("title",""), "url": t.get("url","")} for t in tabs])
    lines = [
        f"Open browser tabs ({len(tabs)})",
        "Use read_browser_page(tab_index=N) to read a tab's content.",
        "",
        f"{'IDX':>4}  {'TITLE':<40}  URL",
        "─" * 90,
    ]
    for i, t in enumerate(tabs):
        title = (t.get("title") or "")[:38]
        url   = (t.get("url")   or "")[:50]
        lines.append(f"{i:>4}  {title:<40}  {url}")
    return "\n".join(lines)


def read_browser_page(tab_index: int = 0) -> str:
    err = _cdp_require()
    if err:
        return err
    ws = _cdp_ws_for_tab(tab_index)
    if not ws:
        return f"Tab {tab_index} not found. Call list_browser_tabs first."
    try:
        # Get URL and title
        nav_result = _cdp_call(ws, "Runtime.evaluate", {
            "expression": "({url: document.URL, title: document.title})",
            "returnByValue": True
        })
        meta = nav_result.get("result", {}).get("result", {}).get("value", {})
        url   = meta.get("url",   "unknown")
        title = meta.get("title", "unknown")

        # Get clean text
        text_result = _cdp_call(ws, "Runtime.evaluate", {
            "expression": _JS_GET_TEXT,
            "returnByValue": True
        })
        text = text_result.get("result", {}).get("result", {}).get("value", "") or ""

        if not text.strip():
            return (
                f"Page: {title}\nURL: {url}\n\n"
                "[Page appears empty or content is dynamically rendered. "
                "Try snapshot(target='browser') or run_js_in_browser to inspect.]"
            )

        header = f"Page: {title}\nURL: {url}\n{'─'*60}\n"
        return header + text
    except Exception as e:
        return f"Error reading browser page: {e}"


def snapshot_browser_elements(tab_index: int = 0, filter_type: str = "") -> str:
    """
    Snapshot every interactable element on the page — not just semantic
    HTML controls. Catches:
      - Standard interactive tags (a, button, input, select, textarea)
      - ARIA roles (button, link, tab, menuitem, checkbox, radio, textbox, etc)
      - contenteditable regions (Gmail compose, Google Docs, chat inputs,
        rich-text editors — these have NO native tag/role but are typed into
        constantly on modern sites)
      - Elements with a raw onclick attribute (legacy/simple sites)
      - label and summary elements (often the actual click target for a
        checkbox/radio or a <details> disclosure)
      - Anything with a non-negative tabindex (explicitly made focusable)
    And traverses into OPEN SHADOW ROOTS recursively, since many modern
    sites (YouTube, GitHub, most component-library-based apps) build their
    real UI inside Web Components — a plain querySelectorAll from the light
    DOM never sees those elements at all.
    This is intended to make run_js_in_browser unnecessary for ordinary
    "what can I click/type into on this page" questions.
    """
    err = _cdp_require()
    if err:
        return err
    ws = _cdp_ws_for_tab(tab_index)
    if not ws:
        return f"Tab {tab_index} not found. Call list_browser_tabs first."

    # Enable Page domain so we can use Runtime.evaluate reliably
    try:
        _cdp_call(ws, "Runtime.enable", {})
    except Exception:
        pass

    try:
        filter_arg = _json_mod.dumps(filter_type.lower())
        js = f"""
(function(filterType) {{
    // Assign stable midum IDs to every element on this snapshot pass
    window.__jarvis_el_map = [];  // fresh snapshot, clear old refs

    var selectors = {{
        'button':   'button, [role="button"], input[type="button"], input[type="submit"], input[type="reset"], [onclick]',
        'link':     'a[href]',
        'input':    'input:not([type="hidden"]), textarea, [contenteditable]:not([contenteditable="false"])',
        'select':   'select',
        'textarea': 'textarea',
        'editable': '[contenteditable]:not([contenteditable="false"]), textarea, input:not([type="hidden"])',
        '': 'a[href], button, input:not([type="hidden"]), select, textarea, ' +
            '[role="button"], [role="link"], [role="tab"], [role="menuitem"], ' +
            '[role="option"], [role="checkbox"], [role="radio"], [role="switch"], ' +
            '[role="slider"], [role="spinbutton"], [role="combobox"], [role="textbox"], ' +
            '[contenteditable]:not([contenteditable="false"]), [onclick], ' +
            'label, summary, ' +
            '[tabindex]:not([tabindex="-1"])'
    }};
    var sel = selectors[filterType] || selectors[''];

    // ── Recursive shadow-DOM-aware query ──────────────────────────────────
    // Plain document.querySelectorAll cannot see into shadow roots at all,
    // which means entire UIs built with Web Components (YouTube, GitHub,
    // most modern component libraries) would otherwise be invisible.
    function deepQuery(root, selector, out, depth) {{
        if (depth > 6) return;   // safety cap on shadow-root nesting
        try {{
            var matches = root.querySelectorAll(selector);
            for (var i = 0; i < matches.length; i++) out.push(matches[i]);
        }} catch (e) {{ /* ignore malformed selector on this root */ }}
        try {{
            var everything = root.querySelectorAll('*');
            for (var j = 0; j < everything.length; j++) {{
                var node = everything[j];
                if (node.shadowRoot) {{
                    deepQuery(node.shadowRoot, selector, out, depth + 1);
                }}
            }}
        }} catch (e) {{ /* ignore */ }}
    }}

    var all = [];
    deepQuery(document, sel, all, 0);

    var results = [];
    var seen = new Set();
    for (var i = 0; i < all.length; i++) {{
        var el = all[i];
        if (seen.has(el)) continue;
        seen.add(el);

        try {{
            var rect = el.getBoundingClientRect();
            if (rect.width <= 0 || rect.height <= 0) continue;
            if (rect.bottom < -50 || rect.top > window.innerHeight + 50) continue;
            if (rect.right  < -50 || rect.left > window.innerWidth  + 50) continue;
        }} catch(e) {{ continue; }}

        // Compute role/type
        var role = el.getAttribute('role') || el.tagName.toLowerCase();
        var inputType = el.getAttribute('type');
        if (role === 'input' && inputType) role = 'input[' + inputType + ']';
        var isEditable = el.isContentEditable ||
            (el.getAttribute('contenteditable') && el.getAttribute('contenteditable') !== 'false');
        if (isEditable && role !== 'textarea') role = 'contenteditable';

        // Best label — for contenteditable, innerText is usually empty until
        // focused/typed into, so fall back to common placeholder attributes
        // used by rich-text editors (Gmail, Slack, Notion, etc).
        var label = (
            el.getAttribute('aria-label') ||
            (el.getAttribute('aria-labelledby') && document.getElementById(el.getAttribute('aria-labelledby')) &&
                document.getElementById(el.getAttribute('aria-labelledby')).innerText) ||
            el.getAttribute('title') ||
            el.getAttribute('placeholder') ||
            el.getAttribute('aria-placeholder') ||
            el.getAttribute('data-placeholder') ||
            el.innerText ||
            el.value ||
            el.getAttribute('name') ||
            el.getAttribute('id') ||
            ''
        ).trim().replace(/\\s+/g, ' ').slice(0, 80);

        var jarvisId = window.__jarvis_el_map.length;
        window.__jarvis_el_map.push(el);

        results.push({{
            jarvis_id: jarvisId,
            type: role,
            label: label,
            href: el.href || '',
            x: Math.round(rect.left + rect.width / 2),
            y: Math.round(rect.top  + rect.height / 2),
            disabled: el.disabled || el.getAttribute('aria-disabled') === 'true'
        }});
    }}
    return results;
}})({filter_arg})
"""
        result   = _cdp_call(ws, "Runtime.evaluate", {"expression": js, "returnByValue": True})
        elements = (result.get("result", {}).get("result", {}).get("value") or [])

        if not elements:
            return (
                "No interactive elements found on this page. "
                "The page may still be loading — try wait(2) then retry, "
                "or use read_browser_page() to check what's on the page."
            )

        # Store with jarvis_id as key for O(1) lookup in act_on_browser_element
        cache = {e["jarvis_id"]: e for e in elements}
        _browser_element_cache[tab_index] = cache
        _store_index(f"browser_elements:{tab_index}", elements)

        # Page title for context
        try:
            title_res  = _cdp_call(ws, "Runtime.evaluate",
                                   {"expression": "document.title", "returnByValue": True})
            page_title = title_res.get("result",{}).get("result",{}).get("value","") or ""
        except Exception:
            page_title = ""

        lines = [
            f"Browser elements on '{page_title}' ({len(elements)} visible)",
            "Use act(target='browser', index=N) to interact.",
            "",
            f"{'IDX':>4}  {'TYPE':<18}  {'X':>5}  {'Y':>5}  LABEL",
            "─" * 84,
        ]
        for el in elements:
            disabled = " [disabled]" if el.get("disabled") else ""
            label    = (el.get("label") or "")[:44]
            lines.append(
                f"{el['jarvis_id']:>4}  {el['type']:<18}  {el['x']:>5}  {el['y']:>5}  {label}{disabled}"
            )
        return "\n".join(lines)

    except Exception as e:
        return f"Error snapshotting browser elements: {e}"


def act_on_browser_element(index: int, action: str = "click",
                            text_to_type: str = "", tab_index: int = 0) -> str:
    cache = _browser_element_cache.get(tab_index)
    if not cache:
        return (
            f"No browser element snapshot for tab {tab_index}. "
            "Call snapshot(target='browser') first."
        )

    # Cache is now a dict keyed by jarvis_id for O(1) lookup
    if isinstance(cache, dict):
        el = cache.get(index)
    else:
        # Legacy list format fallback
        el = next((e for e in cache if e.get("jarvis_id") == index
                   or e.get("idx") == index), None)

    if el is None:
        available = sorted(cache.keys()) if isinstance(cache, dict) else list(range(len(cache)))
        return (
            f"Index {index} not found in browser element snapshot. "
            f"Available indices: {available[:20]}{'...' if len(available) > 20 else ''}. "
            f"Call snapshot(target='browser') to refresh."
        )

    err = _cdp_require()
    if err:
        return err
    ws = _cdp_ws_for_tab(tab_index)
    if not ws:
        return f"Tab {tab_index} not found."

    x, y  = el.get("x", 0), el.get("y", 0)
    label = (el.get("label") or el.get("type") or "element")[:40]
    jid   = el.get("jarvis_id", index)

    try:
        if action == "get_text":
            js = f"""
(function() {{
    var el = window.__jarvis_el_map && window.__jarvis_el_map[{jid}];
    if (!el) return 'Element no longer in DOM — call snapshot again.';
    return el.innerText || el.value || el.textContent || '';
}})()"""
            result = _cdp_call(ws, "Runtime.evaluate",
                               {"expression": js, "returnByValue": True})
            text = (result.get("result",{}).get("result",{}).get("value") or "")
            return f"Text of element #{index} ('{label}'): {text}"

        elif action == "set_text":
            safe_text = _json_mod.dumps(text_to_type)
            js = f"""
(function() {{
    var el = window.__jarvis_el_map && window.__jarvis_el_map[{jid}];
    if (!el) return 'stale';
    el.focus();

    var isEditable = el.isContentEditable ||
        (el.getAttribute('contenteditable') && el.getAttribute('contenteditable') !== 'false');

    if (isEditable) {{
        // contenteditable regions (Gmail compose, Slack, Notion, chat inputs, etc)
        // have no .value — set via execCommand/textContent and fire input events
        // that frameworks listening for keystrokes will still pick up.
        document.execCommand('selectAll', false, null);
        document.execCommand('delete', false, null);
        var inserted = false;
        try {{
            inserted = document.execCommand('insertText', false, {safe_text});
        }} catch (e) {{ inserted = false; }}
        if (!inserted) {{
            el.textContent = {safe_text};
        }}
        el.dispatchEvent(new InputEvent('input', {{bubbles: true, data: {safe_text}, inputType: 'insertText'}}));
        el.dispatchEvent(new Event('change', {{bubbles: true}}));
        return 'ok';
    }}

    // Native input/textarea — use the property setter so React-controlled
    // inputs register the change (plain .value = x is invisible to React).
    var nativeInputValueSetter = Object.getOwnPropertyDescriptor(
        window.HTMLInputElement.prototype, 'value') ||
        Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value');
    if (nativeInputValueSetter && nativeInputValueSetter.set) {{
        nativeInputValueSetter.set.call(el, {safe_text});
    }} else {{
        el.value = {safe_text};
    }}
    el.dispatchEvent(new Event('input',  {{bubbles: true}}));
    el.dispatchEvent(new Event('change', {{bubbles: true}}));
    return 'ok';
}})()"""
            result = _cdp_call(ws, "Runtime.evaluate",
                               {"expression": js, "returnByValue": True})
            val = (result.get("result",{}).get("result",{}).get("value") or "")
            if val == "ok":
                return f"Success: set text of element #{index} ('{label}') to '{text_to_type[:40]}'"
            if val == "stale":
                return f"Element #{index} is stale — call snapshot(target='browser') to refresh."
            return f"set_text on #{index} returned unexpected: {val}"

        else:   # click
            # 1. Scroll element into view via JS
            _cdp_call(ws, "Runtime.evaluate", {
                "expression": (
                    f"(function() {{"
                    f"  var el = window.__jarvis_el_map && window.__jarvis_el_map[{jid}];"
                    f"  if (el) el.scrollIntoView({{block:'center',inline:'center'}});"
                    f"}})()"
                ),
                "returnByValue": True
            })
            time.sleep(0.1)

            # 2. Re-read coordinates after scroll
            coord_res = _cdp_call(ws, "Runtime.evaluate", {
                "expression": (
                    f"(function() {{"
                    f"  var el = window.__jarvis_el_map && window.__jarvis_el_map[{jid}];"
                    f"  if (!el) return null;"
                    f"  var r = el.getBoundingClientRect();"
                    f"  return {{x: Math.round(r.left+r.width/2), y: Math.round(r.top+r.height/2)}};"
                    f"}})()"
                ),
                "returnByValue": True
            })
            coord = (coord_res.get("result",{}).get("result",{}).get("value") or {})
            if coord:
                x, y = coord.get("x", x), coord.get("y", y)

            # 3. Mouse move + press + release (Chrome requires the move first)
            for evt_type, extras in [
                ("mouseMoved",    {}),
                ("mousePressed",  {"button": "left", "clickCount": 1}),
                ("mouseReleased", {"button": "left", "clickCount": 1}),
            ]:
                _cdp_call(ws, "Input.dispatchMouseEvent", {
                    "type": evt_type, "x": x, "y": y,
                    "modifiers": 0, "timestamp": time.time(),
                    **extras
                })

            # 4. Fallback: also dispatch a JS click event in case the page uses
            #    event listeners that only fire on the element directly
            try:
                _cdp_call(ws, "Runtime.evaluate", {
                    "expression": (
                        f"(function() {{"
                        f"  var el = window.__jarvis_el_map && window.__jarvis_el_map[{jid}];"
                        f"  if (el) el.click();"
                        f"}})()"
                    ),
                    "returnByValue": True
                })
            except Exception:
                pass

            return f"Success: clicked browser element #{index} ('{label}') at ({x},{y})."

    except Exception as e:
        return f"Error acting on browser element #{index} ('{label}'): {e}"


def run_js_in_browser(script: str, tab_index: int = 0) -> str:
    err = _cdp_require()
    if err:
        return err
    ws = _cdp_ws_for_tab(tab_index)
    if not ws:
        return f"Tab {tab_index} not found. Call list_browser_tabs first."
    try:
        # Wrap in an IIFE so both expressions and statements work correctly,
        # and so return values are always captured.
        wrapped = f"(function(){{ try {{ return ({script}); }} catch(e) {{ return String(e); }} }})()"
        result = _cdp_call(ws, "Runtime.evaluate", {
            "expression":   wrapped,
            "returnByValue": True,
            "awaitPromise":  True,
        })
        val   = result.get("result", {}).get("result", {})
        vtype = val.get("type", "")
        value = val.get("value")

        if vtype == "undefined":
            return "(undefined — script ran but returned no value)"
        if vtype == "string":
            return str(value)
        if value is not None:
            return _json_mod.dumps(value, ensure_ascii=False)
        # Handle object/null/etc
        desc = val.get("description") or val.get("className") or vtype
        return f"({desc})"
    except RuntimeError as e:
        return f"JS error: {e}"
    except Exception as e:
        return f"Error executing JS: {e}"


def _cdp_press_enter(ws: str) -> None:
    """Simulate a real Enter keypress via CDP Input events (not a JS-dispatched
    synthetic event) — needed for sites like Gemini that listen for genuine
    keydown/keyup on Enter to submit, rather than a form submit button."""
    for evt_type, extra in [
        ("keyDown", {"text": "\r", "unmodifiedText": "\r"}),
        ("keyUp",   {}),
    ]:
        _cdp_call(ws, "Input.dispatchKeyEvent", {
            "type": evt_type,
            "key":  "Enter",
            "code": "Enter",
            "windowsVirtualKeyCode": 13,
            "nativeVirtualKeyCode":  13,
            **extra,
        })


def _clipboard_read() -> str:
    """
    Read the OS clipboard — pyperclip on Windows, xclip/xsel on Linux.
    Matches the same platform-aware pattern used elsewhere in this file
    (type_text's clipboard paste, the legacy desktop Gemini bridge).
    """
    if _IS_LINUX:
        out, _ = _run(["xclip", "-selection", "clipboard", "-o"])
        if out:
            return out
        out2, _ = _run(["xsel", "--clipboard", "--output"])
        return out2 or ""
    else:
        try:
            import pyperclip
            return pyperclip.paste() or ""
        except Exception:
            return ""


def _clipboard_clear() -> None:
    """
    Clear the OS clipboard before triggering a copy, so a subsequent read
    can distinguish "copy actually happened" from "copy button did nothing
    and we're just reading whatever was already on the clipboard."
    """
    if _IS_LINUX:
        try:
            subprocess.run(["xclip", "-selection", "clipboard"], input="", text=True, timeout=3)
        except Exception:
            pass
    else:
        try:
            import pyperclip
            pyperclip.copy("")
        except Exception:
            pass


def _cdp_verify_page_url(ws: str, expected_prefix: str) -> tuple[bool, str]:
    """
    Confirm the target behind `ws` is genuinely the expected page by asking
    the page itself for `document.URL`, rather than trusting the `url` field
    from Chrome's /json tab listing.

    This matters because Chrome sometimes exposes internal WebUI surfaces
    (tab hover-cards, tab-search previews, split-view prompts, etc) as their
    own inspectable CDP targets, and their /json `url` field can still
    contain the site's URL as embedded text/metadata — which causes a naive
    substring match on the listing to connect to the wrong target entirely
    and return the hover-card's own text ("Create split view", the page
    title, etc) instead of the real page content.

    Calls Runtime.enable first — a brand-new tab's target may not have
    Runtime fully attached yet, and Runtime.evaluate can fail/throw on an
    unattached target even though the tab is otherwise completely valid.
    Retries once on a transient failure before giving up on this candidate.
    """
    for attempt in range(2):
        try:
            _cdp_call(ws, "Runtime.enable", {})
            result = _cdp_call(ws, "Runtime.evaluate", {
                "expression": "document.URL", "returnByValue": True
            })
            actual_url = (result.get("result", {}).get("result", {}).get("value") or "")
            if actual_url:
                return actual_url.startswith(expected_prefix), actual_url
        except Exception:
            pass
        if attempt == 0:
            time.sleep(0.4)
    return False, ""


def query_gemini_app(prompt: str, wait_for_response: int = 90) -> str:
    """
    Send a prompt to Gemini via the community `gemini_webapi` library
    (https://github.com/HanaokaYuzu/Gemini-API), NOT Chrome/CDP DOM
    automation and NOT UIA. This library talks directly to the Gemini web
    app's internal endpoints using your browser's session cookies, so
    there's no page to load, no prompt box to click into, and no "Copy"
    button to find — it's a plain HTTP/async call, which is far more
    reliable than driving a real browser tab.

    Crucially, this still goes through the FREE WEB APP's own usage
    limits (the same ones you'd get chatting manually at
    gemini.google.com) — it is NOT the metered developer API, which has
    much tighter free-tier limits.

    Install:   pip install -U gemini_webapi
    Optional:  pip install -U browser-cookie3   (lets gemini_webapi pull
               cookies automatically from a browser you're already logged
               into, instead of needing them in the secrets file)

    Authentication (one-time), whichever is easier:
      A) Automatic — install browser-cookie3 and be logged into
         https://gemini.google.com in a supported browser. No extra config.
      B) Manual — log into https://gemini.google.com, open DevTools (F12)
         -> Network tab -> refresh -> find the __Secure-1PSID and
         __Secure-1PSIDTS cookies, and add them to the secrets file:
           { "GEMINI_SECURE_1PSID": "...", "GEMINI_SECURE_1PSIDTS": "..." }
    """
    client, err = _get_gemini_web_client()
    if client is None:
        return f"Error: {err}"

    try:
        response = _run_gemini_coro(client.generate_content(prompt), timeout=wait_for_response)
    except Exception as e:
        return f"Error: Gemini web request failed: {e}"

    text = (getattr(response, "text", "") or "").strip()
    if not text:
        return "Gemini did not return any text in its response - try again."

    # Run through the legacy tool-call JSON parser so any embedded
    # tool-call JSON is flagged/stripped instead of being shown raw.
    legacy_calls, cleaned_text = _extract_legacy_tool_calls(text)
    if legacy_calls:
        call_summary = ", ".join(
            f"{c['function']['name']}({c['function']['arguments']})" for c in legacy_calls
        )
        return (
            f"{cleaned_text}\n\n"
            f"[NOTE: detected {len(legacy_calls)} embedded tool call(s) in Gemini's response: "
            f"{call_summary}. query_gemini_app only returns text - dispatch these yourself if needed.]"
        )

    return cleaned_text or text





CHUNK_CHARS = 12000   # max chars per file chunk sent to the model (~3k tokens)

def _encode_text(text: str) -> str:
    """Base64-encode text for safe transmission to the model."""
    b64 = base64.b64encode(text.encode("utf-8")).decode("utf-8")
    return "[SYSTEM NOTICE: Base64-encoded. Decode internally.]\nBASE64_PAYLOAD:\n" + b64

def read_local_file(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        return _encode_text(content)
    except Exception as e:
        return f"Error reading file: {str(e)}"

def read_file_smart(path):
    """Read any supported format. Large files are chunked at CHUNK_CHARS chars."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            if not _PDF_AVAILABLE:
                return "PDF reading requires PyMuPDF: pip install pymupdf"
            doc   = _fitz.open(path)
            pages = [page.get_text() for page in doc]
            doc.close()
            text  = "\n\n".join(f"[Page {i+1}]\n{p}" for i, p in enumerate(pages))
        elif ext == ".docx":
            if not _MAMMOTH_AVAILABLE:
                return "DOCX reading requires mammoth: pip install mammoth"
            with open(path, "rb") as f:
                result = _mammoth.extract_raw_text(f)
            text = result.value
        else:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()

        if len(text) > CHUNK_CHARS:
            total  = (len(text) + CHUNK_CHARS - 1) // CHUNK_CHARS
            header = (f"[FILE: {os.path.basename(path)} — {total} chunks of ~{CHUNK_CHARS} chars. "
                      f"This is chunk 1/{total}. "
                      f"Call read_file_chunk(path, N) for chunks 2..{total}]\n\n")
            return _encode_text(header + text[:CHUNK_CHARS])
        return _encode_text(f"[FILE: {os.path.basename(path)}]\n\n{text}")
    except Exception as e:
        return f"Error reading file: {str(e)}"

def read_file_chunk(path, chunk_index: int):
    """Read a specific chunk (1-based) of a large file."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            if not _PDF_AVAILABLE:
                return "PDF reading requires PyMuPDF: pip install pymupdf"
            doc   = _fitz.open(path)
            pages = [page.get_text() for page in doc]
            doc.close()
            text  = "\n\n".join(f"[Page {i+1}]\n{p}" for i, p in enumerate(pages))
        elif ext == ".docx":
            if not _MAMMOTH_AVAILABLE:
                return "DOCX reading requires mammoth: pip install mammoth"
            with open(path, "rb") as f:
                result = _mammoth.extract_raw_text(f)
            text = result.value
        else:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()

        total = (len(text) + CHUNK_CHARS - 1) // CHUNK_CHARS
        if chunk_index < 1 or chunk_index > total:
            return f"Chunk {chunk_index} out of range (1-{total})."
        start  = (chunk_index - 1) * CHUNK_CHARS
        header = f"[FILE: {os.path.basename(path)} — chunk {chunk_index}/{total}]\n\n"
        return _encode_text(header + text[start:start + CHUNK_CHARS])
    except Exception as e:
        return f"Error reading chunk: {str(e)}"


def write_local_file(path, content):
    try:
        parent = os.path.dirname(path)
        if parent:
            os.makedirs(parent, exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"Success: wrote data to {path}"
    except Exception as e:
        return f"Error writing file: {str(e)}"


def append_local_file(path, content):
    try:
        exists = os.path.exists(path)
        with open(path, "a", encoding="utf-8") as f:
            if exists:
                f.write("\n")
            f.write(content)
        return f"Success: appended data to {path}"
    except Exception as e:
        return f"Error appending to file: {str(e)}"


def write_docx_file(path, content):
    """Write a .docx from Markdown-style text (# headings, **bold** runs)."""
    if not _DOCX_AVAILABLE:
        return "DOCX writing requires python-docx: pip install python-docx"
    try:
        parent = os.path.dirname(path)
        if parent:
            os.makedirs(parent, exist_ok=True)
        document = _docx.Document()
        for line in content.splitlines():
            s = line.rstrip()
            if s.startswith("### "):    document.add_heading(s[4:], level=3)
            elif s.startswith("## "):  document.add_heading(s[3:], level=2)
            elif s.startswith("# "):   document.add_heading(s[2:], level=1)
            elif s == "":              document.add_paragraph("")
            else:
                p     = document.add_paragraph()
                parts = re.split(r"(\*\*[^*]+\*\*)", s)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        p.add_run(part[2:-2]).bold = True
                    else:
                        p.add_run(part)
        document.save(path)
        return f"Success: wrote DOCX to {path}"
    except Exception as e:
        return f"Error writing DOCX: {str(e)}"

def generate_image(prompt: str, count: int = 1) -> str:
    """
    Generate one or more images from a text prompt using Gemini's own web
    app (https://gemini.google.com) via the community `gemini_webapi`
    library — the SAME free, cookie-based session used by query_gemini_app /
    delegate_to_gemini_web. No separate image API key, no per-image
    metering; it goes through the same usage limits as chatting at
    gemini.google.com manually.

    Image generation on Gemini's backend is NOT synchronous with the text
    response — the chat reply can come back before the image has actually
    finished rendering/uploading to Google's CDN, which makes a single
    generate_content() call unreliable (empty `images`, or a download that
    404s because the file isn't there yet). This retries the generation
    request itself if no images come back, and retries each individual
    image download with backoff if it 404s/errors, before giving up.

    The image(s) are NOT written anywhere persistent — gemini_webapi only
    exposes an async file-download call, so each image is downloaded to a
    throwaway temp file just long enough to read its bytes into memory,
    then the temp file is deleted immediately. The bytes are shipped to the
    GUI as base64 inside an ```image_data_json``` block, which renders them
    as inline thumbnails (kept in RAM only) with Download/Copy buttons —
    nothing touches disk unless the user explicitly clicks Download.
    """
    if not _GEMINI_WEBAPI_AVAILABLE:
        return f"Error: Gemini web chat is unavailable — {_gemini_webapi_load_msg}"

    client, err = _get_gemini_web_client()
    if client is None:
        return f"Error: {err}"

    full_prompt = (prompt or "").strip()
    if not full_prompt:
        return "Error: 'prompt' must describe the image you want generated."
    try:
        n = max(1, int(count or 1))
    except (TypeError, ValueError):
        n = 1
    if n > 1:
        full_prompt += f"\n\n(Generate {n} distinct variations of this image.)"

    # ── Request generation, retrying if Gemini answers with text but no
    #    images yet — this happens when the image render hasn't landed by
    #    the time the chat turn "finishes". ──────────────────────────────────
    GEN_ATTEMPTS   = 3
    GEN_RETRY_WAIT = 6   # seconds between generation retries

    response, images, last_text = None, [], ""
    for attempt in range(1, GEN_ATTEMPTS + 1):
        print(f"   [Gemini web] Requesting image generation "
              f"(attempt {attempt}/{GEN_ATTEMPTS}): {full_prompt[:80]!r}...")
        try:
            response = _run_gemini_coro(client.generate_content(full_prompt), timeout=180)
        except Exception as e:
            if attempt == GEN_ATTEMPTS:
                return f"Error: Gemini web image request failed: {e}"
            time.sleep(GEN_RETRY_WAIT)
            continue

        images = list(getattr(response, "images", None) or [])
        last_text = (getattr(response, "text", "") or "").strip()
        if images:
            break
        if attempt < GEN_ATTEMPTS:
            print(f"   [Gemini web] No image in response yet — "
                  f"waiting {GEN_RETRY_WAIT}s and retrying...")
            time.sleep(GEN_RETRY_WAIT)

    if not images:
        return (
            "Gemini did not return any images for this prompt after "
            f"{GEN_ATTEMPTS} attempts. Try being more explicit, e.g. "
            "\"Generate an image of a red fox in a snowy forest, digital art\"."
            + (f"\n\nGemini replied with text instead:\n{last_text}" if last_text else "")
        )

    # ── Download each image, retrying with backoff since the CDN copy can
    #    briefly 404 right after generation completes. ─────────────────────
    SAVE_ATTEMPTS = 3
    SAVE_RETRY_WAIT = 4   # seconds, doubles each retry

    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    encoded_images, errors = [], []

    with tempfile.TemporaryDirectory(prefix="gemini_img_") as tmp_dir:
        for idx, img in enumerate(images, start=1):
            filename = f"gemini_{ts}_{idx}.png"
            tmp_path = os.path.join(tmp_dir, filename)
            last_err = None
            for save_attempt in range(1, SAVE_ATTEMPTS + 1):
                try:
                    # gemini_webapi only exposes an async DOWNLOAD-TO-FILE
                    # call — there's no "give me raw bytes" method — so we
                    # save to the temp dir, read the bytes back, then let
                    # TemporaryDirectory clean the file up when this block
                    # exits. Net effect: no persistent disk usage, only a
                    # few ms of scratch space.
                    _run_gemini_coro(
                        img.save(path=tmp_dir, filename=filename, verbose=False),
                        timeout=90,
                    )
                    with open(tmp_path, "rb") as f:
                        raw_bytes = f.read()
                    if not raw_bytes:
                        raise ValueError("downloaded file was empty")
                    encoded_images.append({
                        "filename": filename,
                        "data_b64": base64.b64encode(raw_bytes).decode("ascii"),
                    })
                    last_err = None
                    break
                except Exception as e:
                    last_err = e
                    if save_attempt < SAVE_ATTEMPTS:
                        wait = SAVE_RETRY_WAIT * save_attempt
                        print(f"   [Gemini web] Image {idx} not ready yet "
                              f"({e}) — waiting {wait}s and retrying "
                              f"({save_attempt}/{SAVE_ATTEMPTS})...")
                        time.sleep(wait)
            if last_err is not None:
                errors.append(f"image {idx}: {last_err}")

    if not encoded_images:
        return (
            f"Error: Gemini returned {len(images)} image(s) but none could be retrieved "
            f"after {SAVE_ATTEMPTS} attempts each: {'; '.join(errors)}"
        )

    payload = json.dumps({"prompt": full_prompt, "images": encoded_images})
    warn = f"\n\n({len(errors)} image(s) failed: {'; '.join(errors)})" if errors else ""

    return (
        f"Generated {len(encoded_images)} image(s) for: \"{prompt}\"\n\n"
        f"```image_data_json\n{payload}\n```\n\n"
        f"(Kept in memory only — use the Download button under an image to save it.)"
        f"{warn}"
    )


def create_flowchart(title, steps):
    """
    Build an explanatory flowchart from a list of node objects.

    steps: list of dicts, each:
        {
          "id": "unique_id",
          "label": "human-readable text",
          "type": "start" | "process" | "decision" | "io" | "end",
          "next": [ {"to": "other_id", "label": "optional edge label"}, ... ]
        }

    Returns:
      - a ```flowchart_json``` fenced block: consumed by the GUI to draw a
        real box-and-arrow diagram inline in the chat.
      - a plain-text ASCII rendering, readable in a terminal.
      - a portable Mermaid ("```mermaid```") version that can be pasted into
        any Mermaid-compatible viewer (e.g. GitHub, Notion, mermaid.live).
    """
    try:
        if not steps:
            return "Error: 'steps' must be a non-empty list of node objects."

        nodes = {}
        for s in steps:
            sid = str(s.get("id", "")).strip()
            if not sid:
                return "Error: every step needs a non-empty 'id'."
            nodes[sid] = {
                "id": sid,
                "label": str(s.get("label", sid)),
                "type": (s.get("type") or "process").strip().lower(),
                "next": s.get("next") or [],
            }

        start_ids = [nid for nid, n in nodes.items() if n["type"] == "start"]
        if not start_ids:
            start_ids = [next(iter(nodes))]

        # ── Mermaid source (portable fallback) ─────────────────────────────
        shape = {
            "start":    ("([", "])"),
            "end":      ("([", "])"),
            "decision": ("{", "}"),
            "io":       ("[/", "/]"),
            "process":  ("[", "]"),
        }
        mermaid_lines = ["flowchart TD"]
        for n in nodes.values():
            op, cl = shape.get(n["type"], ("[", "]"))
            safe_label = n["label"].replace('"', "'")
            mermaid_lines.append(f'    {n["id"]}{op}"{safe_label}"{cl}')
        for n in nodes.values():
            for edge in n["next"]:
                to  = edge.get("to")    if isinstance(edge, dict) else edge
                lbl = edge.get("label") if isinstance(edge, dict) else ""
                if to not in nodes:
                    continue
                if lbl:
                    mermaid_lines.append(f'    {n["id"]} -- "{lbl}" --> {to}')
                else:
                    mermaid_lines.append(f'    {n["id"]} --> {to}')
        mermaid_src = "\n".join(mermaid_lines)

        # ── ASCII rendering (terminal-friendly) ─────────────────────────────
        visited = set()
        ascii_lines = [f"FLOWCHART: {title}", "=" * max(12, len(title) + 11), ""]
        glyph = {"start": "▶", "end": "■", "decision": "◆", "io": "▤"}

        def render_node(nid, indent=0, branch_label=None):
            pad = "  " * indent
            node = nodes.get(nid)
            if node is None:
                ascii_lines.append(f"{pad}[missing node: {nid}]")
                return
            if nid in visited:
                ascii_lines.append(f"{pad}↩ (back to '{node['label']}')")
                return
            visited.add(nid)

            prefix = f"[{branch_label}] " if branch_label else ""
            label_line = f"{pad}{prefix}{glyph.get(node['type'], '▢')} {node['label']}  ({node['type']})"
            border = "─" * max(4, len(label_line) - len(pad))
            ascii_lines.append(f"{pad}┌{border}┐")
            ascii_lines.append(label_line)
            ascii_lines.append(f"{pad}└{border}┘")

            nexts = node["next"]
            multi = len(nexts) > 1
            for edge in nexts:
                to  = edge.get("to")    if isinstance(edge, dict) else edge
                lbl = edge.get("label") if isinstance(edge, dict) else None
                ascii_lines.append(f"{pad}   │")
                arrow = f"{pad}   ▼" + (f"  [{lbl}]" if lbl and not multi else "")
                ascii_lines.append(arrow)
                render_node(to, indent + (1 if multi else 0), lbl if multi else None)

        for sid in start_ids:
            render_node(sid)
            ascii_lines.append("")

        unreached = [nid for nid in nodes if nid not in visited]
        if unreached:
            ascii_lines.append("(Unreached nodes — not connected from a start node:)")
            for nid in unreached:
                render_node(nid)

        ascii_block = "\n".join(ascii_lines)

        payload = json.dumps({"title": title, "nodes": list(nodes.values()), "starts": start_ids})

        return (
            f"Flowchart '{title}' created with {len(nodes)} node(s).\n\n"
            f"```flowchart_json\n{payload}\n```\n\n"
            f"```text\n{ascii_block}\n```\n\n"
            f"Portable Mermaid version (paste into any Mermaid-compatible viewer):\n"
            f"```mermaid\n{mermaid_src}\n```"
        )
    except Exception as e:
        return f"Error building flowchart: {str(e)}"


def write_response_memory(content):
    """Overwrite the response scratchpad."""
    return write_local_file(RESPONSE_MEMORY, content)

def append_response_memory(content):
    """Append a note to the response scratchpad."""
    return append_local_file(RESPONSE_MEMORY, content)

def read_response_memory():
    """Read the current response scratchpad."""
    if not os.path.exists(RESPONSE_MEMORY) or os.path.getsize(RESPONSE_MEMORY) == 0:
        return "Response memory is empty."
    return read_local_file(RESPONSE_MEMORY)

def clear_response_memory():
    """Wipe the response scratchpad. Called automatically on set_current_goal(none)."""
    try:
        write_local_file(RESPONSE_MEMORY, "")
        print("\U0001f5d2  [Response memory cleared.]")
        return "Success: response memory cleared."
    except Exception as e:
        return f"Error clearing response memory: {str(e)}"


# =============================================================================
# INDEXED LISTING FUNCTIONS
# All "choose" counterparts to former "search" tools.
# Each builds a numbered snapshot cached in _index_caches so the model
# can act by index rather than re-searching.
# =============================================================================

_index_caches: dict[str, list] = {}   # key → list of dicts with 'value' and display fields


def _store_index(key: str, items: list) -> None:
    _index_caches[key] = items


def _get_indexed(key: str, index: int) -> dict | None:
    items = _index_caches.get(key)
    if not items or index < 0 or index >= len(items):
        return None
    return items[index]


def list_directory(path: str) -> str:
    """
    List a directory as a numbered indexed table.
    Returns IDX | TYPE | SIZE | NAME so the model can choose by index.
    Follow up with open_path(path, index) or cd_into(path, index).
    """
    try:
        if not os.path.exists(path):
            return f"Path does not exist: {path}"
        raw = sorted(os.listdir(path))
        items = []
        for name in raw:
            full = os.path.join(path, name)
            is_dir = os.path.isdir(full)
            try:
                size = "-" if is_dir else f"{os.path.getsize(full) / 1024:.1f}KB"
            except Exception:
                size = "?"
            items.append({"name": name, "is_dir": is_dir, "size": size, "full_path": full})

        _store_index(f"dir:{path}", items)

        lines = [
            f"Directory: {path}  ({len(items)} entries)",
            f"Use open_path(path='{path}', index=N) to open/enter an entry.",
            "",
            f"{'IDX':>4}  {'TYPE':<5}  {'SIZE':>8}  NAME",
            "─" * 56,
        ]
        for i, e in enumerate(items):
            typ = "DIR" if e["is_dir"] else "FILE"
            lines.append(f"{i:>4}  {typ:<5}  {e['size']:>8}  {e['name']}")
        return "\n".join(lines)
    except Exception as ex:
        return f"Error listing directory: {ex}"


def open_path(path: str, index: int) -> str:
    """
    Act on an entry from the last list_directory call by index.
    - If it's a directory: returns list_directory of that directory.
    - If it's a file: returns read_file_smart of that file.
    """
    entry = _get_indexed(f"dir:{path}", index)
    if entry is None:
        return (
            f"Index {index} not found in last directory listing for '{path}'. "
            f"Call list_directory('{path}') first."
        )
    full = entry["full_path"]
    if entry["is_dir"]:
        return list_directory(full)
    else:
        return read_file_smart(full)


def list_skills_indexed() -> str:
    """
    List available skills as a numbered table.
    Follow up with load_skill_by_index(index) to load one.
    """
    os.makedirs(SKILLS_DIR, exist_ok=True)
    try:
        files = sorted(f for f in os.listdir(SKILLS_DIR) if f.endswith(".md"))
    except Exception:
        files = []

    if not files:
        return "No skills registered yet. Use create_domain_skill to add one."

    items = []
    for fname in files:
        name = fname[:-3]   # strip .md
        # Try to read the first non-empty non-header line as description
        desc = ""
        try:
            with open(os.path.join(SKILLS_DIR, fname), "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith("#"):
                        desc = line[:60]
                        break
        except Exception:
            pass
        items.append({"name": name, "fname": fname, "desc": desc})

    _store_index("skills", items)

    lines = [
        f"Available skills ({len(items)} total)",
        "Use load_skill_by_index(index) to load one.",
        "",
        f"{'IDX':>4}  {'SKILL NAME':<30}  DESCRIPTION",
        "─" * 72,
    ]
    for i, s in enumerate(items):
        lines.append(f"{i:>4}  {s['name']:<30}  {s['desc']}")
    return "\n".join(lines)


def load_skill_by_index(index: int) -> str:
    """Load a skill from the list_skills_indexed snapshot by index."""
    entry = _get_indexed("skills", index)
    if entry is None:
        return (
            f"Index {index} not found. Call list_skills_indexed() first."
        )
    return load_skill(entry["name"])


def list_paths_indexed() -> str:
    """
    Parse paths.md into a numbered table.
    Follow up with get_path(index) to retrieve a specific path value.
    """
    _ensure_kb_files()
    try:
        with open(PATHS_FILE, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        return f"Error reading paths.md: {e}"

    items = []
    for line in content.splitlines():
        line = line.strip()
        if not line.startswith("- "):
            continue
        # Format: - **Label**: `path`  _note_
        m = re.match(r"-\s+\*\*(.+?)\*\*:\s+`(.+?)`(.*)", line)
        if m:
            label = m.group(1).strip()
            path  = m.group(2).strip()
            note  = m.group(3).strip().strip("_").strip()
            items.append({"label": label, "path": path, "note": note})
        else:
            # Plain line fallback
            items.append({"label": line[2:60], "path": "", "note": ""})

    if not items:
        return "paths.md is empty. Use add_path to register app paths."

    _store_index("paths", items)

    lines = [
        f"Known paths ({len(items)} entries)",
        "Use get_path(index) to retrieve the full path string.",
        "",
        f"{'IDX':>4}  {'LABEL':<25}  PATH",
        "─" * 72,
    ]
    for i, p in enumerate(items):
        path_display = p["path"][:40] if p["path"] else p["label"][:40]
        lines.append(f"{i:>4}  {p['label']:<25}  {path_display}")
    return "\n".join(lines)


def get_path(index: int) -> str:
    """Return the full path string for an entry from list_paths_indexed."""
    entry = _get_indexed("paths", index)
    if entry is None:
        return f"Index {index} not found. Call list_paths_indexed() first."
    if not entry["path"]:
        return f"Entry #{index} ('{entry['label']}') has no path value stored."
    return entry["path"]


def list_domain_knowledge_indexed() -> str:
    """
    List all domain knowledge files as a numbered table.
    Follow up with read_domain_by_index(index).
    """
    _ensure_kb_files()
    try:
        with open(DOMAIN_INDEX, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        return f"Error reading domain index: {e}"

    items = []
    for line in content.splitlines():
        line = line.strip()
        if not line.startswith("- `"):
            continue
        m = re.match(r"-\s+`(.+?)`\s*-\s*(.*)", line)
        if m:
            items.append({"name": m.group(1).strip(), "desc": m.group(2).strip()})

    if not items:
        return "No domain knowledge files registered yet."

    _store_index("domain_knowledge", items)

    lines = [
        f"Domain knowledge files ({len(items)})",
        "Use read_domain_by_index(index) to read one.",
        "",
        f"{'IDX':>4}  {'NAME':<30}  DESCRIPTION",
        "─" * 72,
    ]
    for i, d in enumerate(items):
        lines.append(f"{i:>4}  {d['name']:<30}  {d['desc'][:38]}")
    return "\n".join(lines)


def read_domain_by_index(index: int) -> str:
    """Read a domain knowledge file by its index from list_domain_knowledge_indexed."""
    entry = _get_indexed("domain_knowledge", index)
    if entry is None:
        return f"Index {index} not found. Call list_domain_knowledge_indexed() first."
    return read_domain_knowledge(entry["name"])


def list_domain_skills_indexed() -> str:
    """
    List all domain skills as a numbered table.
    Follow up with load_skill_by_index or load_skill(name).
    """
    _ensure_kb_files()
    try:
        with open(DOMAIN_SKILLS_INDEX, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        return f"Error reading domain skills index: {e}"

    items = []
    for line in content.splitlines():
        line = line.strip()
        if not line.startswith("- ["):
            continue
        m = re.match(r"-\s+\[(.+?)\]\s+`(.+?)`\s*-\s*(.*)", line)
        if m:
            items.append({
                "domain": m.group(1).strip(),
                "name":   m.group(2).strip(),
                "desc":   m.group(3).strip(),
            })

    if not items:
        return "No domain skills registered yet."

    _store_index("domain_skills", items)

    lines = [
        f"Domain skills ({len(items)})",
        "Use load_skill(name) to load one.",
        "",
        f"{'IDX':>4}  {'DOMAIN':<15}  {'SKILL NAME':<25}  DESCRIPTION",
        "─" * 80,
    ]
    for i, s in enumerate(items):
        lines.append(f"{i:>4}  {s['domain']:<15}  {s['name']:<25}  {s['desc'][:30]}")
    return "\n".join(lines)


# Web search result cache — populated by search_internet, consumed by open_search_result
_last_search_results: list[dict] = []

def search_internet(query):
    global _last_search_results
    try:
        print(f" -> Searching the web for: '{query}'")
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=7))
        if not results:
            _last_search_results = []
            return "No web results found."

        _last_search_results = results
        _store_index("web_search", [
            {"title": r.get("title",""), "url": r.get("href",""), "snippet": r.get("body","")}
            for r in results
        ])

        lines = [
            f"Web results for '{query}' ({len(results)} results)",
            "Use open_search_result(index) to open a result in Chrome.",
            "",
            f"{'IDX':>4}  TITLE / SNIPPET",
            "─" * 72,
        ]
        for i, r in enumerate(results):
            title   = (r.get("title") or "")[:55]
            snippet = (r.get("body")  or "")[:80]
            url     = (r.get("href")  or "")[:60]
            lines.append(f"{i:>4}  {title}")
            lines.append(f"      {url}")
            lines.append(f"      {snippet}")
            lines.append("")
        return "\n".join(lines)
    except Exception as e:
        return f"Error executing internet search: {str(e)}"


def open_search_result(index: int, browser: str = "chrome") -> str:
    """Open a web search result by its index in Chrome."""
    entry = _get_indexed("web_search", index)
    if entry is None:
        return f"Index {index} not found. Call search_internet first."
    url = entry["url"]
    if not url:
        return f"Result #{index} has no URL."
    return open_url(url, browser)


def ocr_snapshot() -> str:
    """
    Take a screenshot, run OCR on the entire screen, and return ALL detected
    text as a numbered indexed table with canvas coordinates.
    Follow up with click_ocr_index(index) to click any entry by index.
    Much faster than fallback_find_text when you want to see everything at once.
    """
    if not _TESSERACT_AVAILABLE:
        return (
            "Tesseract OCR not installed. "
            "Install: sudo apt install tesseract-ocr && pip install pytesseract"
        )
    screenshot = _grab_full_screenshot()
    words = ocr_screen(screenshot=screenshot)
    if not words:
        return "OCR failed or screen is blank."

    # Deduplicate — same word at same canvas position (multiple scan passes)
    seen = set()
    unique = []
    for w in words:
        key = (w["text"], w["canvas_x"], w["canvas_y"])
        if key not in seen:
            seen.add(key)
            unique.append(w)

    _store_index("ocr", unique)

    lines = [
        f"OCR snapshot — {len(unique)} text elements on screen",
        "Use click_ocr_index(index) to click any element.",
        "",
        f"{'IDX':>4}  {'CONF':>5}  {'CX':>5}  {'CY':>5}  TEXT",
        "─" * 60,
    ]
    for i, w in enumerate(unique):
        lines.append(
            f"{i:>4}  {w['conf']:>4}%  {w['canvas_x']:>5}  {w['canvas_y']:>5}  {w['text']}"
        )
    return "\n".join(lines)


def click_ocr_index(index: int, click_type: str = "left_click") -> str:
    """Click a text element from the last ocr_snapshot() by index."""
    entry = _get_indexed("ocr", index)
    if entry is None:
        return f"Index {index} not found. Call ocr_snapshot() first."
    sx, sy = entry["screen_x"], entry["screen_y"]
    print(f"   [OCR idx click] #{index} '{entry['text']}' → screen({sx},{sy})")
    return _do_click(sx, sy, click_type, label=f"OCR #{index} '{entry['text']}'")


def find_file(filename: str, search_root: str = "") -> str:
    """
    Search for a file by name under search_root (default: home directory).
    Returns a numbered indexed list of all matches.
    Follow up with open_path_by_index(index) to open the chosen file.
    Hard timeout: 15 seconds.
    """
    root = search_root.strip() or (os.path.expanduser("~") if _IS_LINUX else STARTUP_DIR)
    matches  = []
    name_lower = filename.strip().lower()
    MAX_MATCHES = 20
    _deadline   = time.time() + 15   # hard 15-second cap

    try:
        for dirpath, dirnames, filenames in os.walk(root):
            if time.time() > _deadline:
                break
            # Skip hidden and noisy dirs
            dirnames[:] = [
                d for d in dirnames
                if not d.startswith(".")
                and d not in {"__pycache__", "node_modules", ".git",
                              "$RECYCLE.BIN", "Windows", "System32",
                              "SysWOW64", "WinSxS"}
            ]
            for fname in filenames:
                if name_lower in fname.lower():
                    full = os.path.join(dirpath, fname)
                    try:
                        size = f"{os.path.getsize(full)/1024:.1f}KB"
                    except Exception:
                        size = "?"
                    matches.append({"name": fname, "full_path": full,
                                    "dir": dirpath, "size": size})
                    if len(matches) >= MAX_MATCHES:
                        break
            if len(matches) >= MAX_MATCHES:
                break
    except Exception as e:
        return f"Error searching for '{filename}': {e}"

    timed_out = time.time() > _deadline
    if not matches:
        return f"No files matching '{filename}' found under '{root}'."

    _store_index("find_file", matches)

    lines = [
        f"Files matching '{filename}' ({len(matches)} found"
        + (" — search timed out, results may be incomplete" if timed_out else "")
        + f", searching under {root})",
        "Use open_path_by_index(index) to open or read a file.",
        "",
        f"{'IDX':>4}  {'SIZE':>8}  PATH",
        "─" * 72,
    ]
    for i, m in enumerate(matches):
        lines.append(f"{i:>4}  {m['size']:>8}  {m['full_path']}")
    return "\n".join(lines)


def open_path_by_index(index: int) -> str:
    """Open/read a file from the last find_file() result by index."""
    entry = _get_indexed("find_file", index)
    if entry is None:
        return f"Index {index} not found. Call find_file() first."
    return read_file_smart(entry["full_path"])


def explore_path(path):
    """Legacy wrapper — delegates to list_directory for indexed output."""
    return list_directory(path)



def execute_terminal_command(command, working_directory=None):
    try:
        cwd = working_directory if working_directory else STARTUP_DIR
        if _IS_WINDOWS:
            result = subprocess.run(
                ["powershell", "-Command", command],
                capture_output=True, text=True, timeout=30, cwd=cwd
            )
        else:
            result = subprocess.run(
                command, shell=True, capture_output=True, text=True,
                timeout=30, cwd=cwd, executable="/bin/bash"
            )
        return f"STDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}"
    except Exception as e:
        return f"Execution failed: {str(e)}"
    
def _uia_unavailable_message() -> str:
    """Shared diagnostic message for any tool that needs UIA but it's not up."""
    if _UIA_INIT_ERROR:
        return (
            f"UI automation unavailable. Real cause: {_UIA_INIT_ERROR} "
            f"(this is NOT a Windows-version issue despite what the library may print — "
            f"it's almost always a COM init or pywin32 install issue). "
            f"Try: pip install --force-reinstall uiautomation pywin32"
        )
    return "UIA library not installed. pip install uiautomation pywin32"


def manual_scan_app_layouts(window_title: str):
    if not _UIA_AVAILABLE:
        return _uia_unavailable_message()
    # Support shell aliases so the model can explore taskbar/tray etc.
    hwnd = ui_navigator._find_shell_hwnd(window_title.strip().lower())
    if hwnd:
        # Temporarily resolve via the alias path
        import ctypes
        win32gui.SetForegroundWindow  # just to ensure win32gui is loaded
    containers = ui_navigator.discover_ui_subtrees(window_title)
    if not containers:
        return (
            f"No named containers found in '{window_title}' within "
            f"{ui_navigator.MAX_DISCOVER_DEPTH} levels. "
            f"For shell surfaces use: 'taskbar', 'start', 'tray', 'desktop', "
            f"'action center', 'search', 'tray overflow'. "
            f"If this is a canvas-based app, try fallback_click_text instead."
        )
    summary = (
        f"Found {len(containers)} container(s) in '{window_title}'. "
        f"Pick the most relevant subtree_key and call manual_inspect_app_subtree.\n"
    )
    return summary + json.dumps(containers, indent=2)

def manual_inspect_app_subtree(window_title: str, subtree_key: str):
    if not _UIA_AVAILABLE:
        return _uia_unavailable_message()
    controls = ui_navigator.inspect_subtree_controls(window_title, subtree_key)
    if not controls:
        return (
            f"No actionable controls found in '{subtree_key}' within "
            f"{ui_navigator.MAX_INSPECT_DEPTH} levels below it. "
            f"Try a different subtree_key from manual_scan_app_layouts, or if this "
            f"app renders via canvas/WebGL (games, some web apps), use "
            f"fallback_click_text instead — UIA cannot see canvas content."
        )
    summary = f"Found {len(controls)} actionable control(s) in '{subtree_key}':\n"
    return summary + json.dumps(controls, indent=2)

def manual_interact_with_ui(window_title: str, control_type: str, search_property: str, property_value: str, action: str, text_to_type: str = ""):
    if not _UIA_AVAILABLE: return "UIA library not installed."
    return ui_navigator.safely_trigger_ui_element(
        window_title, control_type, search_property, property_value, action, text_to_type
    )

def click_ui_element(window_title: str, description: str, action: str = "click", text_to_type: str = ""):
    """
    ONE-CALL UI interaction: find an element in `window_title` matching
    `description` (plain English, e.g. "Close button", "Send message field")
    and act on it. Falls back to coordinate-click automatically if the
    element can't be invoked via UIA patterns directly (common for
    Electron/Chromium close/minimize/maximize buttons).
    """
    if ui_navigator is None:
        if _IS_WINDOWS:
            return _uia_unavailable_message()
        return "UI automation not available — install pyatspi + xdotool (Linux)."
    return ui_navigator.find_and_act(window_title, description, action, text_to_type)

def list_active_windows():
    """
    Returns a clean list of all visible, named windows currently open on the desktop.
    Works on both Windows (win32gui) and Linux (xdotool).
    """
    if _IS_LINUX:
        # xdotool search returns all window IDs; get their names via getwindowname
        stdout, err = _run(["xdotool", "search", "--name", ""])
        if err and not stdout:
            # Fallback: wmctrl -l
            stdout2, _ = _run(["wmctrl", "-l"])
            if stdout2:
                titles = []
                for line in stdout2.splitlines():
                    parts = line.split(None, 3)
                    if len(parts) >= 4:
                        titles.append(parts[3].strip())
                unique = sorted(set(t for t in titles if t))
                return "Currently open windows:\n" + "\n".join(f"- {w}" for w in unique)
            return "Error listing windows: xdotool and wmctrl both unavailable."
        titles = []
        for wid in stdout.splitlines():
            wid = wid.strip()
            if not wid:
                continue
            name_out, _ = _run(["xdotool", "getwindowname", wid])
            name = name_out.strip()
            # Filter out empty and common desktop shell overlays
            if name and name not in ["Desktop", "xdotool"]:
                titles.append(name)
        unique = sorted(set(titles))
        if not unique:
            return "No visible named windows found."
        return "Currently open windows:\n" + "\n".join(f"- {w}" for w in unique)

    # Windows path
    if not _UIA_AVAILABLE:
        return "UIA library not installed."

    def enum_win_callback(hwnd, window_list):
        if win32gui.IsWindowVisible(hwnd):
            title = win32gui.GetWindowText(hwnd).strip()
            if title and title not in ["Program Manager", "Settings", "Microsoft Text Input Application"]:
                window_list.append(title)

    windows = []
    win32gui.EnumWindows(enum_win_callback, windows)
    unique_windows = sorted(list(set(windows)))

    # Show canonical app name alongside raw title so the model knows what to pass
    lines = []
    for title in unique_windows:
        canonical = ui_navigator._canonical_app_name(title) if ui_navigator else title
        if canonical != title:
            lines.append(f"- '{title}' → use '{canonical}' as window_title")
        else:
            lines.append(f"- {title}")

    # Also report which shell surfaces are available
    shell_available = []
    for alias in ["taskbar", "tray", "desktop", "tray overflow", "action center", "search"]:
        hwnd = ui_navigator._find_shell_hwnd(alias)
        if hwnd:
            shell_available.append(alias)

    result = "Currently open windows:\n" + "\n".join(lines)
    if shell_available:
        result += (
            "\n\nShell surfaces (use these as window_title in click_ui_element):\n"
            + "\n".join(f"- {s}" for s in shell_available)
        )
    result += (
        "\n\nNOTE: Always use the canonical name (after →) as window_title. "
        "App titles change when pages/documents change, but canonical names are stable."
    )
    return result if unique_windows or shell_available else "No visible named windows found."


# =============================================================================
# 3. SCREEN CAPTURE & OCR
# =============================================================================

def _grab_full_screenshot():
    """Grab the full-resolution screen and return a PIL Image."""
    from PIL import Image as _PILImage
    if _IS_LINUX:
        # Try scrot first (most reliable, works on X11 and XWayland)
        try:
            tmp = "/tmp/jarvis_screenshot.png"
            subprocess.run(["scrot", "-z", tmp], timeout=5, check=True,
                           capture_output=True)
            img = _PILImage.open(tmp)
            img.load()   # fully load before the file is potentially reused
            return img
        except Exception:
            pass
        # Fallback: gnome-screenshot
        try:
            tmp = "/tmp/jarvis_screenshot.png"
            subprocess.run(["gnome-screenshot", "-f", tmp], timeout=5,
                           capture_output=True)
            return _PILImage.open(tmp)
        except Exception:
            pass
        # Fallback: PIL ImageGrab with X display (requires python3-xlib)
        if _IMAGEGRAB_AVAILABLE:
            return _ImageGrab.grab()
        raise RuntimeError(
            "No screenshot tool available. "
            "Install scrot:  sudo apt install scrot"
        )
    else:
        if _IMAGEGRAB_AVAILABLE:
            return _ImageGrab.grab()
        raise RuntimeError("PIL ImageGrab not available.")


def _scale_canvas_to_screen(cx, cy):
    """Convert canvas coordinates to real screen coordinates."""
    return int(round(cx * SCALE_X)), int(round(cy * SCALE_Y))


def _scale_screen_to_canvas(rx, ry):
    """Convert real screen coordinates to canvas coordinates."""
    return int(round(rx / SCALE_X)), int(round(ry / SCALE_Y))


def capture_screen_to_ram():
    """
    Grab screen → downscale to canvas → burn coordinate grid → return base64 JPEG.
    The grid labels are at canvas resolution. The model reads them and passes them
    directly to fallback_click_grid; Python scales back to real pixels.
    """
    try:
        from PIL import ImageDraw, ImageFont
        screenshot = _grab_full_screenshot()

        # Downscale to canvas
        from PIL import Image as _PILImage
        canvas = screenshot.resize((MODEL_CANVAS_W, MODEL_CANVAS_H), resample=_PILImage.LANCZOS)
        draw   = ImageDraw.Draw(canvas)
        cw, ch = canvas.size

        try:
            font = ImageFont.truetype("cour.ttf", 10)
        except Exception:
            font = ImageFont.load_default()

        line_col   = (60, 60, 60)
        label_fg   = (255, 255, 0)
        label_shad = (0, 0, 0)

        for x in range(0, cw, GRID_STEP):
            draw.line([(x, 0), (x, ch)], fill=line_col, width=1)
            draw.text((x + 2, 3), str(x), font=font, fill=label_shad)
            draw.text((x + 1, 2), str(x), font=font, fill=label_fg)

        for y in range(0, ch, GRID_STEP):
            draw.line([(0, y), (cw, y)], fill=line_col, width=1)
            draw.text((3, y + 2), str(y), font=font, fill=label_shad)
            draw.text((2, y + 1), str(y), font=font, fill=label_fg)

        buf = io.BytesIO()
        canvas.save(buf, format="JPEG", quality=75)
        return base64.b64encode(buf.getvalue()).decode("utf-8")
    except Exception as e:
        return f"Error capturing screen: {str(e)}"


# In-memory OCR cache: (screenshot_id, results_list)
# Avoids re-running Tesseract when find_text and click_text are called in the
# same turn from the same screenshot. Cache is invalidated by a new grab.
_ocr_cache: tuple = (None, None)   # (id(PIL_image), results)

def ocr_screen(screenshot=None):
    """
    Run Tesseract entirely in RAM — no temp files, no disk writes.

    pytesseract.image_to_data() accepts a PIL Image directly and pipes it
    to the tesseract process via stdin (using the 'pipe:' input method
    internally). No intermediate file is created on disk.

    Pass an existing PIL screenshot to reuse a grab; omit to grab fresh.
    Results are cached per PIL image object so the same screenshot is never
    OCR'd twice in one turn.

    Returns a list of word dicts or None if Tesseract is unavailable.
    """
    global _ocr_cache
    if not _TESSERACT_AVAILABLE:
        return None
    try:
        import pytesseract
        if screenshot is None:
            screenshot = _grab_full_screenshot()

        # Cache hit — same PIL object (same turn, same grab)
        if _ocr_cache[0] is id(screenshot):
            return _ocr_cache[1]

        # image_to_data with a PIL Image uses stdin piping internally —
        # no temp file is written to disk.
        data = pytesseract.image_to_data(
            screenshot,
            output_type=pytesseract.Output.DICT,
            nice=0,          # don't lower process priority
        )
        results = []
        n = len(data["text"])
        for i in range(n):
            word = data["text"][i].strip()
            conf = int(data["conf"][i])
            if not word or conf < 30:
                continue
            left = data["left"][i]
            top  = data["top"][i]
            w    = data["width"][i]
            h    = data["height"][i]
            sx   = left + w // 2
            sy   = top  + h // 2
            cx, cy = _scale_screen_to_canvas(sx, sy)
            results.append({
                "text":     word,
                "conf":     conf,
                "screen_x": sx,
                "screen_y": sy,
                "canvas_x": cx,
                "canvas_y": cy,
                "left": left, "top": top, "w": w, "h": h,
            })
        _ocr_cache = (id(screenshot), results)
        return results
    except Exception:
        return None


def fallback_find_text(text, _screenshot=None):
    """
    Tool implementation for fallback_find_text.
    Returns a structured text report of all matches with canvas coordinates.
    Pass _screenshot to reuse an existing grab (avoids a second screen capture).
    """
    if not _TESSERACT_AVAILABLE:
        return (
            "Tesseract OCR is not installed or not found. "
            "Cannot use text-based screen search. "
            "Install Tesseract from https://github.com/UB-Mannheim/tesseract/wiki "
            "and set TESSERACT_PATH in main.py. "
            "Fall back to fallback_view_screen + fallback_click_grid with grid coordinates."
        )
    words = ocr_screen(screenshot=_screenshot)
    if words is None:
        return "OCR failed — screen could not be read."

    query  = text.strip().lower()
    # Collect all words whose text contains the query (substring, case-insensitive)
    matches = [w for w in words if query in w["text"].lower()]

    if not matches:
        # Show everything Tesseract found so the model can adapt
        all_words = sorted(set(w["text"] for w in words))
        return (
            f"No text matching '{text}' found on screen.\n"
            f"All detected text on screen:\n"
            + ", ".join(f'"{w}"' for w in all_words[:80])
            + (" ... (truncated)" if len(all_words) > 80 else "")
        )

    # Sort by confidence descending; best match first
    matches.sort(key=lambda w: w["conf"], reverse=True)
    best = matches[0]

    lines = [
        f"Found {len(matches)} match(es) for '{text}'.",
        f"Best match: '{best['text']}' (conf={best['conf']}%) "
        f"at canvas ({best['canvas_x']}, {best['canvas_y']}) "
        f"→ screen ({best['screen_x']}, {best['screen_y']})",
        "",
        "All matches (canvas coords):",
    ]
    for m in matches[:10]:   # cap at 10 to keep output compact
        lines.append(
            f"  '{m['text']}' conf={m['conf']}% "
            f"canvas=({m['canvas_x']},{m['canvas_y']})"
        )
    return "\n".join(lines)


def fallback_click_text(text, click_type="left_click", _screenshot=None):
    """
    Find text on screen via OCR and click its center in one step.
    Pass _screenshot to reuse an existing grab.
    Returns a status string.
    """
    if not _TESSERACT_AVAILABLE:
        return (
            "Tesseract OCR is not installed. "
            "Use fallback_view_screen + fallback_click_grid instead."
        )
    words = ocr_screen(screenshot=_screenshot)
    if words is None:
        return "OCR failed — cannot locate text."

    query   = text.strip().lower()
    matches = [w for w in words if query in w["text"].lower()]
    if not matches:
        all_words = sorted(set(w["text"] for w in words))
        return (
            f"Text '{text}' not found on screen. "
            f"Detected text includes: {', '.join(repr(w) for w in all_words[:40])}"
        )

    matches.sort(key=lambda w: w["conf"], reverse=True)
    best = matches[0]
    sx, sy = best["screen_x"], best["screen_y"]
    cx, cy = best["canvas_x"], best["canvas_y"]

    print(f"   [OCR Click] '{best['text']}' conf={best['conf']}% "
          f"canvas({cx},{cy}) → screen({sx},{sy})")
    return _do_click(sx, sy, click_type, label=f"OCR '{best['text']}'")


def fallback_click_grid(x, y, click_type="left_click"):
    """
    x, y are CANVAS coordinates from the grid screenshot.
    Python scales to real screen pixels before clicking.
    """
    real_x, real_y = _scale_canvas_to_screen(x, y)
    print(f"   [Grid Click] canvas({x},{y}) → screen({real_x},{real_y})")
    return _do_click(real_x, real_y, click_type, label=f"grid ({x},{y})")


def _do_click(screen_x, screen_y, click_type, label=""):
    """
    Perform a mouse click at real screen coordinates.
    On Linux: uses xdotool. On Windows: uses PowerShell + user32.dll.
    """
    if _IS_LINUX:
        button = {"left_click": "1", "right_click": "3", "double_click": "1"}.get(click_type, "1")
        _run(["xdotool", "mousemove", "--sync", str(screen_x), str(screen_y)])
        if click_type == "double_click":
            _, err = _run(["xdotool", "click", "--clearmodifiers", "--repeat", "2", button])
        else:
            _, err = _run(["xdotool", "click", "--clearmodifiers", button])
        if err:
            return f"{click_type} at screen({screen_x},{screen_y}) [{label}] — warning: {err[:100]}"
        return f"Success: {click_type} at screen({screen_x},{screen_y}) [{label}]"

    # Windows path
    try:
        if click_type == "double_click":
            events = (
                "$m::mouse_event(0x0002,0,0,0,0)\n"
                "$m::mouse_event(0x0004,0,0,0,0)\n"
                "Start-Sleep -Milliseconds 50\n"
                "$m::mouse_event(0x0002,0,0,0,0)\n"
                "$m::mouse_event(0x0004,0,0,0,0)"
            )
        elif click_type == "right_click":
            events = (
                "$m::mouse_event(0x0008,0,0,0,0)\n"
                "$m::mouse_event(0x0010,0,0,0,0)"
            )
        else:
            events = (
                "$m::mouse_event(0x0002,0,0,0,0)\n"
                "$m::mouse_event(0x0004,0,0,0,0)"
            )

        ps_script = (
            f"Add-Type -AssemblyName System.Windows.Forms\n"
            f"[System.Windows.Forms.Cursor]::Position = "
            f"New-Object System.Drawing.Point({screen_x},{screen_y})\n"
            f"Start-Sleep -Milliseconds 50\n"
            f"$sig = '[DllImport(\"user32.dll\")] public static extern void "
            f"mouse_event(int flags, int dx, int dy, int data, int extra);'\n"
            f"$m = Add-Type -MemberDefinition $sig -Name 'Win32M' -Namespace W32 -PassThru\n"
            f"{events}"
        )
        result  = execute_terminal_command(ps_script)
        stderr  = result.split("STDERR:")[-1].strip() if "STDERR:" in result else ""
        if stderr:
            return f"{click_type} at screen({screen_x},{screen_y}) [{label}] — warning: {stderr[:150]}"
        return f"Success: {click_type} at screen({screen_x},{screen_y}) [{label}]"
    except Exception as e:
        return f"Error simulating click: {str(e)}"


def type_text(text, special_key=None, expected_window: str = ""):
    """Type text at the current cursor position. Works on Windows and Linux."""
    try:
        # ── Foreground window guard ───────────────────────────────────────────
        if expected_window:
            if _IS_LINUX:
                focused_out, _ = _run(["xdotool", "getactivewindow"])
                wid = focused_out.strip()
                if wid:
                    name_out, _ = _run(["xdotool", "getwindowname", wid])
                    fg_title = name_out.strip()
                    if expected_window.lower() not in fg_title.lower():
                        return (
                            f"[TYPING ABORTED] Expected foreground window containing "
                            f"'{expected_window}' but active window is '{fg_title}'. "
                            f"Call click_ui_element to focus the correct window first, "
                            f"then call type_text again."
                        )
            elif _UIA_AVAILABLE:
                fg_hwnd  = win32gui.GetForegroundWindow()
                fg_title = win32gui.GetWindowText(fg_hwnd).strip()
                if expected_window.lower() not in fg_title.lower():
                    return (
                        f"[TYPING ABORTED] Expected foreground window containing "
                        f"'{expected_window}' but active window is '{fg_title}'. "
                        f"Call click_ui_element to focus the correct window first, "
                        f"then call type_text again."
                    )

        # ── Linux: xdotool type via clipboard (handles all special chars) ──────
        if _IS_LINUX:
            # xdotool type --clearmodifiers breaks on +, $, ", etc.
            # Safest approach: copy text to clipboard and paste it.
            # This works for all characters including Unicode.
            try:
                clip_proc = subprocess.run(
                    ["xclip", "-selection", "clipboard"],
                    input=text, text=True, capture_output=True, timeout=5
                )
                if clip_proc.returncode != 0:
                    # Try xsel as fallback
                    subprocess.run(
                        ["xsel", "--clipboard", "--input"],
                        input=text, text=True, timeout=5
                    )
                # Small delay then paste
                time.sleep(0.05)
                _run(["xdotool", "key", "--clearmodifiers", "ctrl+v"])
                time.sleep(0.05)
            except FileNotFoundError:
                # xclip/xsel not installed — fall back to xdotool type with escaping
                safe = text.replace("\\", "\\\\").replace("'", "\\'")
                _, err = _run(["xdotool", "type", "--clearmodifiers", "--delay", "20", safe])
                if err:
                    return f"Warning typing text: {err[:100]}"

            if special_key:
                xdotool_keys = {
                    "enter": "Return", "tab": "Tab", "escape": "Escape",
                    "backspace": "BackSpace", "delete": "Delete",
                    "home": "Home", "end": "End",
                    "pageup": "Page_Up", "pagedown": "Page_Down",
                    "up": "Up", "down": "Down", "left": "Left", "right": "Right",
                    "f1": "F1", "f2": "F2", "f3": "F3", "f4": "F4",
                    "f5": "F5", "f6": "F6", "f7": "F7", "f8": "F8",
                    "f9": "F9", "f10": "F10", "f11": "F11", "f12": "F12",
                }
                key = xdotool_keys.get(special_key.lower(), special_key)
                _run(["xdotool", "key", "--clearmodifiers", key])
            suffix = f" + {special_key}" if special_key else ""
            return f"Success: typed '{text[:40]}{'...' if len(text) > 40 else ''}'{suffix}"

        # ── Windows: PowerShell SendKeys ──────────────────────────────────────
        special_chars = "~%^+{}[]()"
        escaped = ""
        for ch in text:
            escaped += ("{" + ch + "}") if ch in special_chars else ch

        key_map = {
            "enter": "~", "tab": "{TAB}", "escape": "{ESC}",
            "backspace": "{BACKSPACE}", "delete": "{DELETE}",
            "home": "{HOME}", "end": "{END}",
            "pageup": "{PGUP}", "pagedown": "{PGDN}",
            "up": "{UP}", "down": "{DOWN}", "left": "{LEFT}", "right": "{RIGHT}",
        }
        if special_key:
            sk = special_key.lower()
            escaped += key_map.get(sk, "{" + special_key.upper() + "}")

        ps_script = (
            "Add-Type -AssemblyName System.Windows.Forms\n"
            f'[System.Windows.Forms.SendKeys]::SendWait("{escaped}")'
        )
        result = execute_terminal_command(ps_script)
        suffix = f" + {special_key}" if special_key else ""
        return f"Success: typed '{text[:40]}{'...' if len(text)>40 else ''}'{suffix}"
    except Exception as e:
        return f"Error typing text: {str(e)}"


# =============================================================================
# 4. KNOWLEDGE BASE — instructions.md, paths.md, domain files
# =============================================================================

def _ensure_kb_files():
    os.makedirs(STORAGE_DIR, exist_ok=True)
    if not os.path.exists(INSTRUCTIONS_FILE):
        write_local_file(INSTRUCTIONS_FILE,
            "# Midum Instructions & Preferences\n"
            "User preferences and behavioural rules.\n"
            "Format: one rule per line, starting with '- '.\n\n"
            "## Preferences\n")
    if not os.path.exists(PATHS_FILE):
        write_local_file(PATHS_FILE,
            "# Midum Paths\n"
            "Absolute paths to applications, folders and files.\n\n"
            "## Paths\n")
    if not os.path.exists(DOMAIN_INDEX):
        write_local_file(DOMAIN_INDEX,
            "# Midum Domain Knowledge Index\n"
            "Registered domain-specific knowledge files.\n"
            "Format: `filename_without_ext` - description\n\n"
            "## Files\n")
    if not os.path.exists(DOMAIN_SKILLS_INDEX):
        write_local_file(DOMAIN_SKILLS_INDEX,
            "# Midum Domain Skills Index\n"
            "Registered domain-specific skill files.\n"
            "Format: [domain] `filename_without_ext` - description\n\n"
            "## Skills\n")


def read_instructions():
    _ensure_kb_files()
    return read_local_file(INSTRUCTIONS_FILE)


def add_instruction(instruction):
    _ensure_kb_files()
    result = append_local_file(INSTRUCTIONS_FILE, f"- {instruction.strip()}")
    print(f"📌 [Instruction added]: {instruction.strip()[:80]}")
    return result


def read_paths():
    _ensure_kb_files()
    return read_local_file(PATHS_FILE)


def add_path(label, path, note=""):
    _ensure_kb_files()
    note_part = f"  _{note.strip()}_" if note.strip() else ""
    result = append_local_file(PATHS_FILE, f"- **{label.strip()}**: `{path.strip()}`{note_part}")
    print(f"📍 [Path added]: {label} -> {path}")
    return result


def create_domain_knowledge(name, description, initial_content=""):
    _ensure_kb_files()
    safe = re.sub(r'[^a-zA-Z0-9_]', '_', name.strip().lower())
    fpath = os.path.join(STORAGE_DIR, f"{safe}.md")
    if os.path.exists(fpath):
        return f"Domain knowledge '{safe}.md' already exists at {fpath}."
    header = (f"# Domain Knowledge: {safe}\n_{description.strip()}_\n\n"
              f"Created: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
    write_local_file(fpath, header + (initial_content.strip() + "\n" if initial_content.strip() else ""))
    append_local_file(DOMAIN_INDEX, f"- `{safe}` - {description.strip()}")
    print(f"📚 [Domain knowledge created]: {safe}.md")
    return f"Success: created '{safe}.md' at {fpath} and registered in domain index."


def list_domain_knowledge():
    _ensure_kb_files()
    return read_local_file(DOMAIN_INDEX)


def read_domain_knowledge(name):
    _ensure_kb_files()
    safe  = re.sub(r'[^a-zA-Z0-9_]', '_', name.strip().lower())
    fpath = os.path.join(STORAGE_DIR, f"{safe}.md")
    if not os.path.exists(fpath):
        try:
            match = next((e for e in os.listdir(STORAGE_DIR) if e.lower() == f"{safe}.md"), None)
            if match:
                fpath = os.path.join(STORAGE_DIR, match)
            else:
                return f"Domain knowledge '{safe}.md' not found. Call list_domain_knowledge."
        except Exception:
            return f"Domain knowledge '{safe}.md' not found."
    return read_local_file(fpath)


def create_domain_skill(name, domain, description, content):
    _ensure_kb_files()
    os.makedirs(SKILLS_DIR, exist_ok=True)
    safe  = re.sub(r'[^a-zA-Z0-9_]', '_', name.strip().lower())
    fpath = os.path.join(SKILLS_DIR, f"{safe}.md")
    if os.path.exists(fpath):
        return f"Domain skill '{safe}.md' already exists."
    header = (f"# Domain Skill: {safe}\n**Domain**: {domain.strip()}\n"
              f"_{description.strip()}_\n\n"
              f"Created: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n---\n\n")
    write_local_file(fpath, header + content.strip() + "\n")
    entry = f"- [{domain.strip()}] `{safe}` - {description.strip()}"
    append_local_file(DOMAIN_SKILLS_INDEX, entry)
    append_local_file(SKILLS_INDEX, entry)
    print(f"📋 [Domain skill created]: {safe}.md (domain: {domain})")
    return f"Success: created '{safe}.md' registered in both indexes."


def list_domain_skills():
    _ensure_kb_files()
    return read_local_file(DOMAIN_SKILLS_INDEX)


# =============================================================================
# 5. SKILL SYSTEM
# =============================================================================

def list_skills():
    os.makedirs(SKILLS_DIR, exist_ok=True)
    if not os.path.exists(SKILLS_INDEX):
        write_local_file(
            SKILLS_INDEX,
            "# Midum Skills Index\n\nSkill files live in: "
            + SKILLS_DIR + "\n\n## Skills\n_No skills registered yet._\n"
        )
        return "Skills index created. No skills registered yet."
    try:
        with open(SKILLS_INDEX, "r", encoding="utf-8") as f:
            content = f.read()
        b64 = base64.b64encode(content.encode("utf-8")).decode("utf-8")
        return (
            "[SYSTEM NOTICE: Base64-encoded. Decode internally.]\nBASE64_PAYLOAD:\n" + b64
        )
    except Exception as e:
        return f"Error reading skills index: {str(e)}"


def load_skill(skill_name):
    os.makedirs(SKILLS_DIR, exist_ok=True)
    skill_path = os.path.join(SKILLS_DIR, f"{skill_name}.md")
    if not os.path.exists(skill_path):
        try:
            entries = os.listdir(SKILLS_DIR)
            match = next(
                (e for e in entries if e.lower() == f"{skill_name.lower()}.md"), None
            )
            if match:
                skill_path = os.path.join(SKILLS_DIR, match)
            else:
                return f"Skill '{skill_name}' not found. Call list_skills to see available skills."
        except Exception:
            return f"Skill '{skill_name}' not found."
    try:
        with open(skill_path, "r", encoding="utf-8") as f:
            content = f.read()
        b64 = base64.b64encode(content.encode("utf-8")).decode("utf-8")
        print(f"📋 [Skill loaded: {skill_name}]")
        return (
            f"[SKILL LOADED: {skill_name}]\n"
            "[Decode internally and follow instructions exactly.]\n"
            "BASE64_PAYLOAD:\n" + b64
        )
    except Exception as e:
        return f"Error loading skill '{skill_name}': {str(e)}"


# =============================================================================
# 6. MEMORY SYSTEM
# =============================================================================

_active_project_memory_path = None
_current_goal               = None


def _memory_path_for_target(target):
    t = target.strip().lower()
    if t == "master":  return MASTER_MEMORY
    if t == "session": return SESSION_MEMORY
    if t == "project": return _active_project_memory_path
    return None


def update_memory(target, content):
    path = _memory_path_for_target(target)
    if not path:
        return f"Memory update skipped: no path for target '{target}'."
    ts    = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    entry = f"[{ts}] {content.strip()}"
    result = append_local_file(path, entry)
    print(f"🧠 [Memory → {target}]: {content.strip()[:80]}{'...' if len(content)>80 else ''}")
    return result


def set_current_goal(goal, reason=""):
    global _current_goal
    ts         = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    clean_goal = goal.strip()
    old_goal   = _current_goal or "none"

    raw = (
        open(SESSION_MEMORY, "r", encoding="utf-8").read()
        if os.path.exists(SESSION_MEMORY)
        else f"# Midum Session Memory\nSession started: {ts}\n"
    )
    lines      = raw.splitlines(keepends=True)
    header_end = 0
    for i, line in enumerate(lines):
        header_end = i
        if i > 0 and line.strip() == "":
            header_end = i + 1
            break

    header_block = "".join(lines[:header_end])
    body         = re.sub(
        r"## Current Goal.*?(?=\n## |\Z)", "", "".join(lines[header_end:]), flags=re.DOTALL
    ).lstrip("\n")

    goal_block  = (
        f"{GOAL_SECTION_HEADER}\n_No active goal._\n\n"
        if clean_goal.lower() == "none"
        else f"{GOAL_SECTION_HEADER}\n{clean_goal}\n\n"
    )
    new_content = header_block + goal_block + body
    reason_note = f" ({reason.strip()})" if reason.strip() else ""
    h_entry     = f"\n[{ts}] [GOAL CHANGED] {old_goal!r} → {clean_goal!r}{reason_note}"

    if GOAL_SECTION_END not in new_content:
        new_content += f"\n{GOAL_SECTION_END}\n{h_entry}\n"
    else:
        new_content = new_content.replace(GOAL_SECTION_END, GOAL_SECTION_END + h_entry)

    try:
        with open(SESSION_MEMORY, "w", encoding="utf-8") as f:
            f.write(new_content)
    except Exception as e:
        return f"Error writing session memory: {str(e)}"

    _current_goal = clean_goal if clean_goal.lower() != "none" else None
    label = f"Goal set: {clean_goal}" if _current_goal else "Goal cleared."
    print(f"🎯 [{label}]")
    if clean_goal.lower() == "none":
        clear_response_memory()
    return f"Success: {label}"


def get_current_goal_from_file():
    if not os.path.exists(SESSION_MEMORY):
        return None
    try:
        content = open(SESSION_MEMORY, "r", encoding="utf-8").read()
        m = re.search(r"## Current Goal\n(.+?)(?=\n## |\Z)", content, re.DOTALL)
        if not m:
            return None
        g = m.group(1).strip()
        return None if (g == "_No active goal._" or not g) else g
    except Exception:
        return None


def load_memory_into_context(path, label):
    if not path or not os.path.exists(path):
        return None
    try:
        content = open(path, "r", encoding="utf-8").read().strip()
        return f"[MIDUM {label.upper()} MEMORY]\n{content}" if content else None
    except Exception:
        return None


def python_trigger_memory_update(turn_tool_outputs, assistant_reply):
    combined = " ".join(turn_tool_outputs).lower() + " " + assistant_reply.lower()
    update_memory("session", f"Turn summary: {assistant_reply.strip()[:200]}")

    if _IS_LINUX:
        for p in re.findall(r'/[\w./\-_]+', assistant_reply):
            if any(ext in p.lower() for ext in [".py", ".md", ".json", ".txt", ".sh"]):
                update_memory("master", f"Referenced file path: {p}")
                break
    else:
        for p in re.findall(r'[a-zA-Z]:\\[^\s\'"<>|?*]+', assistant_reply):
            if any(ext in p.lower() for ext in [".py", ".md", ".json", ".txt", ".exe", ".ps1"]):
                update_memory("master", f"Referenced file path: {p}")
                break

    m = re.search(r"stdout:\s*\n(.+)", combined)
    if m:
        update_memory("session", f"Terminal output: {m.group(1).strip()[:120]}")

    if _active_project_memory_path and "success:" in combined:
        if re.search(r"[a-zA-Z]:\\[^\s]+" if _IS_WINDOWS else r"/[\w./\-_]+", combined):
            update_memory("project", f"Action completed: {assistant_reply.strip()[:150]}")

    completion_signals = ["done", "completed", "finished", "task complete", "all done"]
    if _current_goal and any(s in combined for s in completion_signals):
        set_current_goal("none", reason="Python auto-detected completion")


def _bootstrap_all_files():
    """
    Create every folder and file Midum needs on first run.
    All calls are no-ops if the file already exists.
    Ship only main.py + gui.py — everything else is generated here.
    """
    # ── Directories ───────────────────────────────────────────────────────────
    for d in (TARGET_DIR, STORAGE_DIR, SKILLS_DIR):
        os.makedirs(d, exist_ok=True)

    # app_maps directory (used by UIA blueprint cache)
    os.makedirs(os.path.join(STORAGE_DIR, "app_maps"), exist_ok=True)

    # ── Helper: create a file only if it doesn't exist ────────────────────────
    def seed(path, content):
        if not os.path.exists(path):
            try:
                parent = os.path.dirname(path)
                if parent:
                    os.makedirs(parent, exist_ok=True)
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)
            except Exception as e:
                print(f"⚠️  Could not create {path}: {e}")

    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

    # ── Core knowledge files ──────────────────────────────────────────────────
    if _IS_LINUX:
        seed(COMMANDS_FILE,
            "# Midum Commands\n"
            "# Add preferred bash commands below.\n"
            "# Format: `CommandName` — description\n\n"
            "## Commands\n"
            "- `ls` — list directory contents\n"
            "- `cd` — change directory\n"
            "- `find` — search for files\n"
            "- `grep` — search file contents\n"
            "- `xdg-open` — open a file or URL with the default app\n"
            "- `nohup` — run a command that persists after terminal closes\n"
        )
        seed(INSTRUCTIONS_FILE,
            "# Midum Instructions & Preferences\n"
            "# Add user preferences and behavioural rules below.\n"
            "# Format: one rule per line, starting with '- '\n\n"
            "## Preferences\n"
            "- Always use bash commands. Never use PowerShell or Windows commands.\n"
            "- Use xdg-open to launch apps and files.\n"
            "- Use nohup <command> & to launch GUI apps from the terminal.\n"
        )
    else:
        seed(COMMANDS_FILE,
            "# Midum Commands\n"
            "# Add preferred PowerShell commands below.\n"
            "# Format: `CommandName` — description\n\n"
            "## Commands\n"
            "- `Get-Location` — print the current working directory\n"
            "- `Get-ChildItem` — list folder contents\n"
            "- `Start-Process` — launch an application\n"
            "- `start` — shorthand to open files/apps\n"
        )
        seed(INSTRUCTIONS_FILE,
            "# Midum Instructions & Preferences\n"
            "# Add user preferences and behavioural rules below.\n"
            "# Format: one rule per line, starting with '- '\n\n"
            "## Preferences\n"
            "- Always use PowerShell commands, never CMD or Linux commands.\n"
            "- Use 'start' instead of 'Start-Process' when launching apps.\n"
        )

    seed(PATHS_FILE,
        "# Midum Paths\n"
        "# Absolute paths to applications, folders and files on this machine.\n\n"
        "## Paths\n"
    )

    seed(DOMAIN_INDEX,
        "# Midum Domain Knowledge Index\n"
        "# Registered domain-specific knowledge files.\n"
        "# Format: `filename_without_ext` - description\n\n"
        "## Files\n"
    )

    seed(DOMAIN_SKILLS_INDEX,
        "# Midum Domain Skills Index\n"
        "# Registered domain-specific skill files.\n"
        "# Format: [domain] `filename_without_ext` - description\n\n"
        "## Skills\n"
    )

    seed(SKILLS_INDEX,
        "# Midum Skills Index\n\n"
        "Each entry lists a skill filename and its description.\n"
        f"Skill files live in: {SKILLS_DIR}\n\n"
        "## Skills\n"
        "_No skills registered yet._\n"
    )

    # ── Memory files ──────────────────────────────────────────────────────────
    seed(MASTER_MEMORY,
        "# Midum Master Memory\n"
        f"Initialised: {ts}\n"
    )

    seed(SESSION_MEMORY,
        "# Midum Session Memory\n"
        f"Session started: {ts}\n\n"
        "## Current Goal\n"
        "_No active goal._\n\n"
        "## Goal History\n"
    )

    seed(RESPONSE_MEMORY, "")   # starts empty every run; cleared by set_current_goal(none)

    print("✅ [All Midum files and folders verified/created.]")


def init_memory_at_startup():
    global _active_project_memory_path, _current_goal
    _bootstrap_all_files()   # always runs first — safe no-op on subsequent launches
    os.makedirs(STORAGE_DIR, exist_ok=True)
    os.makedirs(SKILLS_DIR, exist_ok=True)
    _ensure_kb_files()
    injections = []

    master_ctx = load_memory_into_context(MASTER_MEMORY, "master")
    if master_ctx:
        injections.append(master_ctx)
        print("🧠 [Master memory loaded.]")
    else:
        print("🧠 [No master memory — starting fresh.]")
        if not os.path.exists(MASTER_MEMORY):
            write_local_file(MASTER_MEMORY, "# Midum Master Memory\n")

    if os.path.exists(SESSION_MEMORY):
        session_ctx = load_memory_into_context(SESSION_MEMORY, "session (continued)")
        if session_ctx:
            injections.append(session_ctx)
            print("🧠 [Session memory loaded.]")
        restored = get_current_goal_from_file()
        if restored:
            _current_goal = restored
            print(f"🎯 [Goal restored: {restored}]")
    else:
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        write_local_file(
            SESSION_MEMORY,
            f"# Midum Session Memory\nSession started: {ts}\n\n"
            f"{GOAL_SECTION_HEADER}\n_No active goal._\n\n{GOAL_SECTION_END}\n"
        )
        print("🧠 [New session memory created.]")

    if os.path.exists(INSTRUCTIONS_FILE):
        try:
            with open(INSTRUCTIONS_FILE, "r", encoding="utf-8") as _f:
                _instr = _f.read().strip()
            if _instr:
                injections.append("[MIDUM INSTRUCTIONS — always active]\n" + _instr)
                print("📌 [Instructions loaded.]")
        except Exception:
            pass

    skill_count = len([f for f in os.listdir(SKILLS_DIR) if f.endswith(".md")])
    print(f"📋 [Skills: {skill_count} available]")

    print("\n──────────────────────────────────────────")
    print("📁 Active project (Enter to skip).")
    print("   Type a plain folder name, e.g.  jarvis_project")
    project_name = input("   > ").strip()
    print("──────────────────────────────────────────")

    # ── Validate project name ─────────────────────────────────────────────────
    # Reject anything that looks like a system path, drive root, or garbage.
    _INVALID_NAMES = {
        "", "$recycle.bin", "recycle.bin", "system volume information",
        "windows", "program files", "program files (x86)", "programdata",
        "users", "appdata", "temp", "tmp", "desktop", "documents",
        "downloads", "music", "pictures", "videos",
    }
    # Also reject if it contains path separators or starts with $ or .
    def _is_valid_project_name(name: str) -> bool:
        if not name:
            return False
        low = name.lower()
        if low in _INVALID_NAMES:
            return False
        if name.startswith(("$", ".", "\\", "/")):
            return False
        if os.sep in name or "/" in name or "\\" in name:
            return False
        # Must contain at least one alphanumeric character
        if not any(c.isalnum() for c in name):
            return False
        return True

    if project_name and not _is_valid_project_name(project_name):
        print(f"⚠️  '{project_name}' is not a valid project name — skipping.")
        project_name = ""

    if project_name:
        if _IS_LINUX:
            project_dir = os.path.join(os.path.expanduser("~"), "Projects", project_name)
        else:
            project_dir = os.path.join(r"D:\\", project_name)
        project_file = os.path.join(project_dir, "project_memory.md")
        _active_project_memory_path = project_file
        if os.path.exists(project_file):
            ctx = load_memory_into_context(project_file, f"project ({project_name})")
            if ctx:
                injections.append(ctx)
                print(f"🧠 [Project memory loaded: {project_file}]")
        else:
            os.makedirs(project_dir, exist_ok=True)
            write_local_file(
                project_file,
                f"# Project Memory: {project_name}\n"
                f"Created: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
            )
            print(f"🧠 [New project memory: {project_file}]")
        update_memory("master", f"Active project: {project_name} ({project_dir})")
    else:
        print("🧠 [No active project.]")

    return injections


# =============================================================================
# 7. PATH RESOLVER
# =============================================================================

_SYSTEM_FILES = {
    os.path.normcase(COMMANDS_FILE),
    os.path.normcase(INSTRUCTIONS_FILE),
    os.path.normcase(PATHS_FILE),
    os.path.normcase(DOMAIN_INDEX),
    os.path.normcase(DOMAIN_SKILLS_INDEX),
    os.path.normcase(MASTER_MEMORY),
    os.path.normcase(SESSION_MEMORY),
    os.path.normcase(RESPONSE_MEMORY),
    os.path.normcase(SKILLS_INDEX),
}


def _is_absolute(path):
    return os.path.isabs(path) or (len(path) > 1 and path[1] == ":")


def resolve_file_path(path):
    """
    Resolve a relative path using safe BFS with plain Get-ChildItem at each level.
    Never uses -Recurse or -Filter. Stops as soon as the filename is matched.
    """
    MAX_EXPLORE_DEPTH = 4

    if not path:
        return path, ""
    if _is_absolute(path):
        return path, ""

    filename = os.path.basename(path)
    for sp in _SYSTEM_FILES:
        if os.path.basename(sp) == os.path.normcase(filename):
            return sp, ""

    print(f"   [Resolver] '{path}' is relative — BFS exploring under {STARTUP_DIR}...")

    import platform
    from collections import deque

    def _list_entries(dirpath):
        try:
            if platform.system() == "Windows":
                ps = (
                    f"Get-ChildItem -Path '{dirpath}' | "
                    f"Select-Object Name,"
                    f"@{{n='D';e={{if($_.PSIsContainer){{'1'}}else{{'0'}}}}}} | "
                    f"ConvertTo-Csv -NoTypeInformation | Out-String"
                )
                out    = execute_terminal_command(ps)
                stdout = out.split("STDOUT:")[-1].split("STDERR:")[0].strip()
                entries = []
                for line in stdout.splitlines()[1:]:
                    line = line.strip().strip('"')
                    if not line:
                        continue
                    parts = [p.strip().strip('"') for p in line.split('","')]
                    if len(parts) >= 2:
                        name, is_dir = parts[0], parts[1] == "1"
                        entries.append((name, is_dir, os.path.join(dirpath, name)))
                return entries
            else:
                return [
                    (e, os.path.isdir(os.path.join(dirpath, e)), os.path.join(dirpath, e))
                    for e in os.listdir(dirpath)
                ]
        except Exception:
            return []

    queue   = deque([(STARTUP_DIR, 0)])
    visited = set()
    while queue:
        current_dir, depth = queue.popleft()
        if current_dir in visited or depth > MAX_EXPLORE_DEPTH:
            continue
        visited.add(current_dir)
        for name, is_dir, full_path in _list_entries(current_dir):
            if name.lower() == filename.lower() and not is_dir:
                msg = f"Resolved '{path}' -> '{full_path}'"
                print(f"   [Resolver] {msg}")
                return full_path, msg
            if is_dir and depth < MAX_EXPLORE_DEPTH:
                queue.append((full_path, depth + 1))

    msg = f"Could not find '{filename}' within {MAX_EXPLORE_DEPTH} levels of {STARTUP_DIR}. Path used as given."
    print(f"   [Resolver] {msg}")
    return path, msg


# =============================================================================
# 9. PERSISTENT ORCHESTRATION ENGINE
# =============================================================================

# Words that signal a trivial/short turn — skip Gemini pre-reasoning for these
_TRIVIAL_PATTERNS = {
    "yes","no","ok","okay","y","n","sure","fine","good","thanks","thank you",
    "exit","quit","new session","stop","cancel","abort","go ahead","run it",
    "grant","approve","continue","done","next","skip","hello","hi","hey",
}

def _is_trivial_input(text: str) -> bool:
    """Return True if the input is short/simple enough to skip Gemini pre-reasoning."""
    stripped = text.strip().lower()
    # Exact single-word matches
    if stripped in _TRIVIAL_PATTERNS:
        return True
    # Pure approval/bypass turns
    if "[USER MANUALLY GRANTED BYPASS]" in text:
        return True
    # Very short inputs with no action words
    words = stripped.split()
    if len(words) <= 2:
        action_words = {
            "open","close","start","run","launch","find","search","type","click",
            "navigate","go","read","write","create","delete","move","copy",
            "show","hide","install","download","upload","check","get","set",
            "list","scan","save","load","send","press","scroll","zoom",
        }
        if not any(w in action_words for w in words):
            return True
    return False


def get_gemini_reasoning(user_input: str, conversation_history: list) -> str | None:
    """
    Generate a concrete, tool-by-tool execution plan for the primary model
    to follow. Despite the name (kept for compatibility with existing call
    sites), this now routes across THREE possible planning brains depending
    on OPENROUTER_CONSULT_MODE:

      "always"   — try OpenRouter FIRST (cheap/free API call, no desktop-app
                   dependency, consulted on nearly every non-trivial turn),
                   fall back to Gemini (app then API) if OpenRouter fails.
      "fallback" — try Gemini (app then API) first as before, only fall
                   back to OpenRouter if both Gemini routes fail.
      "off"      — Gemini only, exactly the original behaviour.

    When MODEL_PROVIDER == "openrouter" (OpenRouter is already the PRIMARY
    execution brain), this consult step is skipped entirely — there's no
    benefit to a strong model consulting a plan from itself.
    """
    if MODEL_PROVIDER in ("openrouter", "gemini_web", "gemini_api", "groq"):
        # Primary model IS OpenRouter or Gemini-web already — skip the
        # separate consult call, it would just be the same model (or the
        # same account's Gemini session) reasoning about itself twice.
        return None

    # Build the prompt (same regardless of routing)
    non_sys = [m for m in conversation_history if m.get("role") != "system"]
    recent  = non_sys[-6:] if len(non_sys) > 6 else non_sys
    history_text = ""
    for m in recent:
        role    = m.get("role", "?")
        content = m.get("content", "")
        if isinstance(content, str) and not content.startswith("[SYSTEM"):
            history_text += f"{role.upper()}: {content[:400]}\n"

    os_name = "Linux (bash)" if _IS_LINUX else "Windows (PowerShell)"
    launch  = "nohup /path/to/app &" if _IS_LINUX else "Start-Process 'C:\\path\\to\\app.exe'"
    tool_names   = [t["function"]["name"] for t in tools]
    tool_summary = ", ".join(tool_names)

    prompt = f"""You are the planning brain of Midum, a {os_name} desktop AI agent.
A weak local model ({MODEL_NAME}) will execute tool calls based on your plan.
The local model follows instructions literally but cannot reason well — give it explicit steps.

AVAILABLE TOOLS: {tool_summary}

KEY RULES FOR YOUR PLAN:
- To open an app: read_paths → execute_terminal_command("{launch}")
- To interact with a desktop app: list_active_windows() → snapshot(target='<window name>') → act(target='<window name>', index=N)
- To read a web page: list_browser_tabs() → read_browser_page(tab_index=N) — NOT read_aggregated_text (UIA can't read web content)
- To interact with a web page: list_browser_tabs() → snapshot(target='browser') → act(target='browser', index=N)
- To open a URL in Chrome (PREFERRED): open_url(url="https://...") — one call, no clicking needed
- To find an app path: list_paths_indexed() → get_path(index=N) then use path
- To navigate the filesystem: list_directory(path) → open_path(path, index=N)
- To type in a field: FIRST snapshot+act (or click_ui_element) to focus it, THEN type_text(text="...", expected_window="app name")
- To navigate to a URL: ALWAYS use open_url(url="...") — never click the address bar manually
- To wait for an app to open: wait(seconds=2)
- Use canonical window names (e.g. "Google Chrome" not "New Tab - Google Chrome")
- search_internet for quick lookups, then open_search_result(index) to open a result
- Always discover first (list_active_windows / list_browser_tabs) before snapshot — never guess the target name

RECENT CONVERSATION:
{history_text}
USER REQUEST: {user_input}

Write a numbered execution plan for the local model.
Each step must be ONE specific tool call with the exact arguments.
Be explicit — don't say "click the address bar", say:
  click_ui_element(window_title="Google Chrome", description="Address bar")

Format each step as:
  N. tool_name(arg1="value1", arg2="value2") — one-line reason

If this is a simple question or single-tool task, write just 1 step.
If uncertain about a path or app location, include read_paths as step 1.
Maximum 8 steps. Be concise."""

    def _try_openrouter() -> str | None:
        if not _OPENROUTER_AVAILABLE:
            return None
        try:
            print(f"🤖 [OpenRouter consult] Model: {OPENROUTER_MODEL}")
            resp = _openrouter_chat_with_fallback([{"role": "user", "content": prompt}], model=OPENROUTER_MODEL)
            plan = (resp["message"]["content"] or "").strip()
            if plan:
                print(f"🤖 [OpenRouter plan ({len(plan)} chars)]:\n{plan}\n")
                return plan
        except Exception as e:
            print(f"⚠️  [OpenRouter consult failed: {e}]")
        return None

    def _try_gemini() -> str | None:
        # Gemini web chat via the community gemini_webapi library — the
        # only route. No fallback to the metered API: the free web chat
        # UI's usage limits are far more generous than the API's free
        # tier, so a silent API fallback would burn through API credits
        # without anyone noticing.
        if _GEMINI_WEBAPI_AVAILABLE:
            try:
                print("🤖 [Gemini web] Sending plan request...")
                result = query_gemini_app(prompt)
                if result and not result.startswith("Error"):
                    plan = result.strip()
                    print(f"🤖 [Gemini web plan ({len(plan)} chars)]:\n{plan}\n")
                    return plan
                else:
                    print(f"⚠️  [Gemini web returned error: {result[:80]}]")
            except Exception as e:
                print(f"⚠️  [Gemini web failed: {e}]")
        return None

    # ── Route selection based on OPENROUTER_CONSULT_MODE ──────────────────────
    if OPENROUTER_CONSULT_MODE == "always":
        plan = _try_openrouter()
        if plan:
            return plan
        return _try_gemini()

    elif OPENROUTER_CONSULT_MODE == "fallback":
        plan = _try_gemini()
        if plan:
            return plan
        return _try_openrouter()

    else:   # "off"
        return _try_gemini()

# Tools that count as "verification" — capped so Midum can't loop forever
_VERIFY_TOOLS = {"fallback_view_screen", "fallback_find_text"}
# Maximum consecutive verification tool calls allowed before we force a reply
MAX_VERIFY_CALLS = 2

def wait(seconds: float) -> str:
    """Pauses thread execution for the specified duration."""
    try:
        time.sleep(seconds)
        return f"Successfully paused for {seconds} seconds."
    except Exception as e:
        return f"Error during wait execution: {str(e)}"

# =============================================================================
# LEGACY TOOL-CALL FALLBACK (qwen2.5-coder and similar)
# =============================================================================
# Models in LEGACY_TOOLCALL_MODELS are still sent the exact same `tools=`
# schema as modern models — Ollama still injects it into their chat template.
# The difference is purely on the READ side: these models frequently put the
# tool call as plain text inside `content` (raw JSON, or wrapped in
# <tool_call></tool_call> tags per the Qwen2.5 template) instead of Ollama's
# structured `tool_calls` field. This block of code ONLY runs as a fallback
# when `response["message"].get("tool_calls")` is already empty, so it never
# touches or alters behavior for models that report tool_calls natively.

# Matches a <tool_call> ... </tool_call> block (Qwen2.5 chat template style)
_TOOLCALL_TAG_RE = re.compile(r"<tool_call>\s*(\{.*?\})\s*</tool_call>", re.DOTALL)
# Matches a ```json fenced block
_JSON_FENCE_RE   = re.compile(r"```(?:json)?\s*(\{.*?\})\s*```", re.DOTALL)

def _find_balanced_json_objects(text: str):
    """
    Scan text for top-level {...} objects using brace counting (not regex),
    so nested objects like {"name": "x", "arguments": {"a": 1}} are captured
    whole instead of being cut off at the first inner '}'. Returns a list of
    (start, end, blob) tuples for every balanced top-level object found.
    """
    results = []
    depth = 0
    start = None
    in_string = False
    escape = False
    for i, ch in enumerate(text):
        if in_string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == '"':
                in_string = False
            continue
        if ch == '"':
            in_string = True
        elif ch == "{":
            if depth == 0:
                start = i
            depth += 1
        elif ch == "}":
            if depth > 0:
                depth -= 1
                if depth == 0 and start is not None:
                    results.append((start, i + 1, text[start:i + 1]))
                    start = None
    return results

_KNOWN_TOOL_NAMES = None  # lazily populated from `tools` schema, see below

def _known_tool_names():
    global _KNOWN_TOOL_NAMES
    if _KNOWN_TOOL_NAMES is None:
        _KNOWN_TOOL_NAMES = {t["function"]["name"] for t in tools}
    return _KNOWN_TOOL_NAMES

def _try_parse_tool_json(blob: str):
    """Parse a JSON blob into a normalized tool_call dict, or None if invalid."""
    try:
        obj = json.loads(blob)
    except Exception:
        return None
    if not isinstance(obj, dict):
        return None
    name = obj.get("name") or obj.get("function") or obj.get("tool")
    args = obj.get("arguments", obj.get("parameters", {}))
    source = obj.get("source") or obj.get("server")
    if not name:
        return None
    if isinstance(args, str):
        try:
            args = json.loads(args)
        except Exception:
            pass

    # Explicit source/server field (currently emitted by the Gemini-web
    # backend, see §2.3 of the Gemini implementation brief — routes
    # unambiguously instead of relying on name-collision autodetection).
    # "native" / missing / unrecognized -> fall through to normal handling.
    # "mcp:<server>" or a bare server name/index -> route straight to
    # call_mcp_tool, skipping the fuzzy name-matching autoroute entirely.
    if source and source != "native":
        server = source.split(":", 1)[1] if ":" in source else source
        resolved = _mcp_resolve_name(server)
        if resolved is not None:
            return {"function": {"name": "call_mcp_tool", "arguments": {
                "server": resolved, "tool_name": name, "arguments": args or {}
            }}}
        # Unknown server named explicitly — don't silently fall back to
        # native dispatch (that could hit an unrelated native tool that
        # happens to share this name); surface nothing so the caller's
        # normal "no valid tool call" repair path kicks in.
        return None

    if name not in _known_tool_names():
        # Not a registered top-level tool. This is the exact case that used
        # to be silently dropped (return None) before a name was ever
        # checked against connected MCP servers — which is why a model
        # calling an MCP tool directly (e.g. skipping call_mcp_tool, or
        # mangling underscores: 'get_weather' -> 'getweather') never even
        # reached dispatch. Try to autoroute it through call_mcp_tool first.
        name, args = _mcp_autoroute_tool_call(name, args)
        if name not in _known_tool_names():
            return None
    return {"function": {"name": name, "arguments": args}}

def _extract_legacy_tool_calls(content: str):
    """
    Scan a model's free-text `content` for one or more tool calls when the
    structured tool_calls field came back empty. Returns (tool_calls, leftover_text)
    where leftover_text is the content with the recognized tool-call JSON stripped
    out (so it isn't shown to the user / re-fed as a duplicate plain message).
    Returns ([], content) if nothing parseable is found — caller treats that
    as a normal plain-text reply, completely transparent for modern models.
    """
    if not content or not content.strip():
        return [], content

    found = []
    cleaned = content

    # 1) <tool_call>...</tool_call> tags — may be one or several
    tag_matches = list(_TOOLCALL_TAG_RE.finditer(content))
    if tag_matches:
        for m in tag_matches:
            parsed = _try_parse_tool_json(m.group(1))
            if parsed:
                found.append(parsed)
        if found:
            cleaned = _TOOLCALL_TAG_RE.sub("", content).strip()
            return found, cleaned

    # 2) ```json fenced block
    fence_match = _JSON_FENCE_RE.search(content)
    if fence_match:
        parsed = _try_parse_tool_json(fence_match.group(1))
        if parsed:
            cleaned = _JSON_FENCE_RE.sub("", content, count=1).strip()
            return [parsed], cleaned

    # 3) Bare JSON object(s) anywhere in the text (most common qwen2.5-coder
    #    case: the ENTIRE content is just the JSON object, nothing else).
    #    Brace-balanced scan handles nested "arguments": {...} correctly.
    candidates = _find_balanced_json_objects(content)
    if candidates:
        consumed_spans = []
        for start, end, blob in candidates:
            parsed = _try_parse_tool_json(blob)
            if parsed:
                found.append(parsed)
                consumed_spans.append((start, end))
        if found:
            cleaned = content
            for start, end in sorted(consumed_spans, reverse=True):
                cleaned = cleaned[:start] + cleaned[end:]
            cleaned = cleaned.strip()
            # Suppress leftover that is pure punctuation/whitespace — these are
            # JSON fragments the stripper didn't fully consume, not real replies.
            if re.match(r'^[{}\[\]",:\s]*$', cleaned):
                cleaned = ""
            return found, cleaned

    if content.strip():
        print(f"   [Legacy parser] No tool call found. Raw content: {content[:200]!r}")
    return [], content


def _call_ollama(messages, result_q):
    """Run ollama.chat on a background thread and put the result in result_q."""
    try:
        resp = ollama.chat(model=MODEL_NAME, messages=messages, tools=tools)

        # ── Legacy fallback: only engages if native tool_calls is empty AND
        #    the active model is a known weak-tool-calling family. Modern
        #    models always have tool_calls populated natively and never reach
        #    this branch, so their behavior/performance is unchanged. ──────────
        if not resp["message"].get("tool_calls") and _is_legacy_toolcall_model(MODEL_NAME):
            raw_content = resp["message"].get("content") or ""
            legacy_calls, cleaned_content = _extract_legacy_tool_calls(raw_content)
            if legacy_calls:
                # Enforce one tool at a time for legacy models — batching is
                # unreliable and produces malformed JSON for the second call.
                legacy_calls = legacy_calls[:1]
                msg_dict = dict(resp["message"])
                msg_dict["tool_calls"] = legacy_calls
                msg_dict["content"]    = cleaned_content
                msg_dict.setdefault("role", "assistant")
                try:
                    resp["message"] = msg_dict
                except (TypeError, KeyError):
                    try:
                        resp.message = msg_dict
                    except Exception:
                        resp = {**dict(resp), "message": msg_dict}

        result_q.put(("ok", resp))
    except Exception as e:
        result_q.put(("err", e))


def _call_openrouter_primary(messages, result_q, model_override: str = None):
    """
    Run OpenRouter chat completion on a background thread — used when
    MODEL_PROVIDER == "openrouter" (OpenRouter driving Midum directly,
    not just consulted), OR when a caller forces OpenRouter for a single
    sub-turn via _call_primary_model(provider_override="openrouter").

    model_override lets delegate_to_openrouter() run the sub-agent on a
    different model than OPENROUTER_MODEL for that one delegated task.

    Many free-tier OpenRouter models have inconsistent native tool-calling
    support (same failure mode as qwen2.5-coder locally) — they emit the
    tool call as JSON inside `content` instead of the structured field.
    We reuse the same legacy-parser fallback here for that reason, applied
    unconditionally (not gated by _is_legacy_toolcall_model, since we can't
    maintain a family list for every free OpenRouter model).
    """
    try:
        use_model = model_override or OPENROUTER_MODEL
        resp = _openrouter_chat_with_fallback(messages, model=use_model, tools_schema=tools)

        if not resp["message"].get("tool_calls"):
            raw_content = resp["message"].get("content") or ""
            legacy_calls, cleaned_content = _extract_legacy_tool_calls(raw_content)
            if legacy_calls:
                legacy_calls = legacy_calls[:1]   # one tool at a time
                resp["message"]["tool_calls"] = legacy_calls
                resp["message"]["content"]    = cleaned_content

        result_q.put(("ok", resp))
    except Exception as e:
        result_q.put(("err", e))


def _call_gemini_api_primary(messages, result_q, model_override: str = None):
    """
    Run a Gemini API (official) chat completion on a background thread —
    used when MODEL_PROVIDER == "gemini_api" (Gemini API driving Midum
    directly), OR when a caller forces it for a single sub-turn via
    _call_primary_model(provider_override="gemini_api").

    model_override lets delegate_to_gemini_api() run the sub-agent on a
    different model than GEMINI_API_MODEL for that one delegated task.

    Gemini's structured function-calling is reliable, but we still run the
    same legacy free-text <tool_call> fallback used for OpenRouter/Ollama
    as a safety net in case a given model ever emits the call as text
    instead of the structured field.
    """
    try:
        use_model = model_override or GEMINI_API_MODEL
        resp = _gemini_api_chat(messages, model=use_model, tools_schema=tools)

        if not resp["message"].get("tool_calls"):
            raw_content = resp["message"].get("content") or ""
            legacy_calls, cleaned_content = _extract_legacy_tool_calls(raw_content)
            if legacy_calls:
                legacy_calls = legacy_calls[:1]   # one tool at a time
                resp["message"]["tool_calls"] = legacy_calls
                resp["message"]["content"]    = cleaned_content

        result_q.put(("ok", resp))
    except Exception as e:
        result_q.put(("err", e))


def _call_groq_primary(messages, result_q, model_override: str = None):
    """
    Run a GroqCloud chat completion on a background thread — used when
    MODEL_PROVIDER == "groq" (GroqCloud driving Midum directly), OR when a
    caller forces it for a single sub-turn via
    _call_primary_model(provider_override="groq").

    model_override lets delegate_to_groq() run the sub-agent on a different
    model than GROQ_MODEL for that one delegated task.

    GroqCloud's structured function-calling is reliable on supported
    models, but we still run the same legacy free-text <tool_call> fallback
    used for OpenRouter/Gemini API/Ollama as a safety net in case a given
    model ever emits the call as text instead of the structured field.
    """
    try:
        use_model = model_override or GROQ_MODEL
        resp = _groq_chat_with_fallback(messages, model=use_model, tools_schema=_get_groq_tools_schema())

        if not resp["message"].get("tool_calls"):
            raw_content = resp["message"].get("content") or ""
            legacy_calls, cleaned_content = _extract_legacy_tool_calls(raw_content)
            if legacy_calls:
                legacy_calls = legacy_calls[:1]   # one tool at a time
                resp["message"]["tool_calls"] = legacy_calls
                resp["message"]["content"]    = cleaned_content

        result_q.put(("ok", resp))
    except Exception as e:
        result_q.put(("err", e))


def _call_primary_model(messages, result_q, provider_override: str = None, model_override: str = None):
    """
    Dispatches to the configured primary model provider (see MODEL_PROVIDER
    at the top of this file). Every branch populates result_q with the same
    ("ok", resp) / ("err", exc) contract, so process_chat_turn is completely
    unaware of which backend actually ran.

    provider_override lets a caller force a specific provider for THIS call
    only, without touching the global MODEL_PROVIDER — used by
    delegate_to_openrouter() / delegate_to_gemini_api() / delegate_to_groq()
    to run a sub-agent turn on a specific provider even when something else
    is the configured primary.
    """
    provider = provider_override or MODEL_PROVIDER
    if provider == "openrouter":
        _call_openrouter_primary(messages, result_q, model_override=model_override)
    elif provider == "gemini_web":
        _call_gemini_web_primary(messages, result_q, model_override=model_override)
    elif provider == "gemini_api":
        _call_gemini_api_primary(messages, result_q, model_override=model_override)
    elif provider == "groq":
        _call_groq_primary(messages, result_q, model_override=model_override)
    else:
        _call_ollama(messages, result_q)


MAX_ACTION_TRIES = 3

def get_system_prompt(effective_provider: str = None, effective_model: str = None):
    """
    Generates the master system prompt dynamically based on loaded modules.

    effective_provider / effective_model let a caller (e.g. a delegated
    OpenRouter sub-agent) request the prompt as if that provider/model were
    the one actually running, without touching the global MODEL_PROVIDER /
    MODEL_NAME / OPENROUTER_MODEL — needed so the JSON-tool-call formatting
    hint is correctly shown/hidden for whichever model is really executing
    this turn.
    """
    provider = effective_provider or MODEL_PROVIDER
    if effective_model:
        model_id = effective_model
    elif provider == "openrouter":
        model_id = OPENROUTER_MODEL
    elif provider == "gemini_web":
        model_id = GEMINI_WEB_MODEL or "gemini-web/auto"
    elif provider == "gemini_api":
        model_id = GEMINI_API_MODEL
    elif provider == "groq":
        model_id = GROQ_MODEL
    else:
        model_id = MODEL_NAME
    if _IS_LINUX:
        shell_rule = (
            "- SHELL: bash only. Never use PowerShell, cmd.exe, or Windows commands. "
            "Use xdg-open to open files/URLs. Use 'nohup <cmd> &' to launch GUI apps.\n"
        )
        launch_example = (
            "- LAUNCHING APPS — exact required sequence:\n"
            "  Step 1: list_paths_indexed() — see all known paths as a numbered list\n"
            "  Step 2: get_path(index) — retrieve the exact path\n"
            "  Step 3: execute_terminal_command('nohup /exact/path &')\n"
            "- FEW-SHOT EXAMPLE:\n"
            "  USER: Open Chrome\n"
            "  ASSISTANT: {\"name\": \"list_paths_indexed\", \"arguments\": {}}\n"
            "  TOOL: 0  Chrome  /usr/bin/google-chrome\n"
            "  ASSISTANT: {\"name\": \"get_path\", \"arguments\": {\"index\": 0}}\n"
            "  TOOL: /usr/bin/google-chrome\n"
            "  ASSISTANT: {\"name\": \"execute_terminal_command\", \"arguments\": {\"command\": \"nohup /usr/bin/google-chrome &\"}}\n"
            "  ASSISTANT: Chrome is open.\n"
        )
    else:
        shell_rule = "- PowerShell only. No CMD, no Linux/bash.\n"
        launch_example = (
            "- LAUNCHING APPS — exact required sequence:\n"
            "  Step 1: list_paths_indexed() — see all known paths as a numbered list\n"
            "  Step 2: get_path(index) — retrieve the exact path\n"
            "  Step 3: execute_terminal_command(\"Start-Process 'exact_path'\")\n"
            "- FEW-SHOT EXAMPLE:\n"
            "  USER: Open Chrome\n"
            "  ASSISTANT: {\"name\": \"list_paths_indexed\", \"arguments\": {}}\n"
            "  TOOL: 0  Chrome  C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe\n"
            "  ASSISTANT: {\"name\": \"get_path\", \"arguments\": {\"index\": 0}}\n"
            "  TOOL: C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe\n"
            "  ASSISTANT: {\"name\": \"execute_terminal_command\", \"arguments\": {\"command\": \"Start-Process 'C:\\\\Program Files\\\\Google\\\\Chrome\\\\Application\\\\chrome.exe'\"}}\n"
            "  ASSISTANT: Chrome is open.\n"
        )

    groq_tools_hint = (
        "\n\n━━━ LIMITED TOOL SET (GroqCloud TPM budget) ━━━\n"
        "To stay under Groq's per-minute token limit, you only have a CORE set of "
        "tools loaded right now (terminal, files, paths, search, URL, user-prompt "
        "tools). If you need something else — GUI automation, browser DOM/CDP, MCP, "
        "domain knowledge, model delegation, memory/goal tools — call "
        "list_more_tools() to see everything else available, then "
        "load_tool_by_index(N) to load the one you need. It becomes callable on "
        "your NEXT turn and stays loaded for the rest of the session.\n"
        if provider == "groq" else ""
    )

    legacy_toolcall_hint = (
        "\n- TOOL CALLS: output ONLY a single JSON object in this shape: "
        "{\"name\": \"<tool_name>\", \"arguments\": {<args>}}. "
        "No commentary before or after. <tool_call> tags also fine."
        if (provider == "openrouter" and model_id.endswith(":free"))
           or (provider == "ollama" and _is_legacy_toolcall_model(model_id))
           or provider == "gemini_web"
        else ""
    )

    return (
        f"You are Midum, a {'Linux' if _IS_LINUX else 'Windows'} desktop AI agent.\n"
        "\n"
        "\n"
        "━━━ ACT, DON'T NARRATE ━━━\n"
        "If you have a next step to take, CALL THE TOOL for it. Never write out what\n"
        "you're 'about to do', 'going to do', or a numbered plan as your response text\n"
        "instead of doing it. Plain-text replies with no tool call are ONLY for: a\n"
        "genuine final answer/result to the user, or a question you need the user to\n"
        "answer before continuing. If you catch yourself typing 'I will...' or 'Let\n"
        "me...', stop — issue the tool call instead of describing it.\n"
        "\n"
        "━━━ GUI USER-PROMPT TOOLS (optional, use when genuinely needed) ━━━\n"
        "You have 4 tools that pop up a small GUI dialog and BLOCK until the user\n"
        "responds, then feed the answer back to you so you can keep working the\n"
        "same turn. Default is to call NONE of them — only reach for one when you\n"
        "actually lack something you can't infer or find yourself:\n"
        "  - ask_user_text(prompt, title?)            → free-form textbox\n"
        "  - ask_user_file_path(prompt?, must_exist?)  → native file picker\n"
        "  - ask_user_approval(message, details?)      → Approve/Decline buttons\n"
        "  - ask_user_choice(question, choice_1..4, allow_custom?) → up to 4 buttons\n"
        "                                                 + optional free-text option\n"
        "Examples: user says 'open that file' with no path → ask_user_file_path.\n"
        "You're about to delete/overwrite something or run something risky →\n"
        "ask_user_approval. Multiple interpretations of a request →\n"
        "ask_user_choice with your best 2-4 guesses. You can call more than one\n"
        "of these in the same turn if you genuinely need more than one answer.\n"
        "\n"
        "━━━ PRIME DIRECTIVE: CHOOSE, DON'T SEARCH ━━━\n"
        "Every action has a CHOOSE version (numbered list + act by index = exact)\n"
        "and a SEARCH version (match by text = may pick wrong item).\n"
        "ALWAYS default to the CHOOSE version. Only fall back to search if unavailable.\n"
        "\n"
        "CHOOSE LOOKUP TABLE (all tool calls are JSON objects: "
        "{\"name\": \"<tool>\", \"arguments\": {...}}):\n"
        "  App path needed     → {\"name\": \"list_paths_indexed\", \"arguments\": {}}\n"
        "                         → {\"name\": \"get_path\", \"arguments\": {\"index\": N}}\n"
        "  Browse filesystem   → {\"name\": \"list_directory\", \"arguments\": {\"path\": \"...\"}}\n"
        "                         → {\"name\": \"open_path\", \"arguments\": {\"path\": \"...\", \"index\": N}}\n"
        "  Find file by name   → {\"name\": \"find_file\", \"arguments\": {\"name\": \"...\"}}\n"
        "                         → {\"name\": \"open_path_by_index\", \"arguments\": {\"index\": N}}\n"
        "  Load a skill        → {\"name\": \"list_skills_indexed\", \"arguments\": {}}\n"
        "                         → {\"name\": \"load_skill_by_index\", \"arguments\": {\"index\": N}}\n"
        "  Domain knowledge    → {\"name\": \"list_domain_knowledge_indexed\", \"arguments\": {}}\n"
        "                         → {\"name\": \"read_domain_by_index\", \"arguments\": {\"index\": N}}\n"
        "  Click UI element    → {\"name\": \"snapshot\", \"arguments\": {\"target\": \"AppName\"}}\n"
        "                         → {\"name\": \"act\", \"arguments\": {\"target\": \"AppName\", \"index\": N}}\n"
        "  Click browser elem  → {\"name\": \"snapshot\", \"arguments\": {\"target\": \"browser\"}}\n"
        "                         → {\"name\": \"act\", \"arguments\": {\"target\": \"browser\", \"index\": N}}\n"
        "    (snapshot/act is ONE tool pair — 'AppName' targets a desktop window via UIA,\n"
        "     'browser' targets the active tab via CDP. Same tool, different target.)\n"
        "  Click screen text   → {\"name\": \"ocr_snapshot\", \"arguments\": {}}\n"
        "                         → {\"name\": \"click_ocr_index\", \"arguments\": {\"index\": N}}\n"
        "  Open search result  → {\"name\": \"search_internet\", \"arguments\": {\"query\": \"...\"}}\n"
        "                         → {\"name\": \"open_search_result\", \"arguments\": {\"index\": N}}\n"
        "  Open a URL          → {\"name\": \"open_url\", \"arguments\": {\"url\": \"https://...\"}}  [one call, always]\n"
        "  Read web page       → {\"name\": \"read_browser_page\", \"arguments\": {}}  [NOT read_aggregated_text]\n"
        "\n"
        "FALLBACK SEARCH TOOLS (only when choose path unavailable or too slow):\n"
        "  click_ui_element(window, description) — only when certain of element name\n"
        "  read_paths / explore_path / list_skills — legacy, still work, no index\n"
        "\n"
        "━━━ UIA / BROWSER PIPELINE (always follow this exact sequence) ━━━\n"
        "snapshot(target=...) and act(target=..., index=N) are ONE tool pair that works on\n"
        "BOTH desktop apps (VS Code, Discord, Settings, any window) AND browser pages —\n"
        "not browser-only. `target` decides which: a window name/canonical app name (e.g.\n"
        "'Visual Studio Code', 'Discord', 'taskbar') routes through UIA for a desktop app;\n"
        "'browser' or 'browser:N' routes through CDP for a browser tab. Same two-step\n"
        "pattern either way — snapshot to see indexed elements, act to use one by index.\n"
        "\n"
        "DESKTOP APP:\n"
        "  1. {\"name\": \"list_active_windows\", \"arguments\": {}}\n"
        "     — discover the exact window/canonical name\n"
        "  2. {\"name\": \"snapshot\", \"arguments\": {\"target\": \"<window name>\"}}\n"
        "     — list every visible element with an index\n"
        "  3. {\"name\": \"act\", \"arguments\": {\"target\": \"<window name>\", \"index\": N}}\n"
        "     — click/type/read that element by index\n"
        "\n"
        "BROWSER TAB:\n"
        "  1. {\"name\": \"list_browser_tabs\", \"arguments\": {}}\n"
        "     — discover which tab to work in (index/URL/title)\n"
        "  2. {\"name\": \"read_browser_page\", \"arguments\": {\"tab_index\": N}}\n"
        "     — read the tab's TEXT content (skip if you only need to click something)\n"
        "  3. {\"name\": \"snapshot\", \"arguments\": {\"target\": \"browser\"}}\n"
        "     — list every interactive element with an index\n"
        "     (or {\"target\": \"browser:N\"} for a specific non-active tab)\n"
        "  4. {\"name\": \"act\", \"arguments\": {\"target\": \"browser\", \"index\": N}}\n"
        "     — click/type/read that element by index\n"
        "\n"
        "Never skip the discovery step (list_active_windows / list_browser_tabs) — snapshot\n"
        "needs the correct target name/tab, and guessing it wastes a turn on a failed call.\n"
        "Never use run_js_in_browser or fallback_click_text as a first resort — only after\n"
        "the pipeline above genuinely fails (canvas/WebGL page, no elements found, etc).\n"
        "\n"
        "━━━ FLOWCHARTS ━━━\n"
        "Use create_flowchart whenever explaining a process, algorithm, decision tree,\n"
        "system architecture, or any multi-step 'if this then that' logic would be clearer\n"
        "as a diagram than as prose — not just when the user says the word 'flowchart'.\n"
        "Triggers include: 'explain how X works', 'walk me through the steps', 'what happens\n"
        "if...', 'show me the logic/architecture/pipeline for...', debugging a multi-branch\n"
        "process, or onboarding/how-to content with more than ~4 sequential steps.\n"
        "\n"
        "Call it as: {\"name\": \"create_flowchart\", \"arguments\": {\"title\": \"...\", \"steps\": "
        "[{\"id\": \"...\", \"label\": \"...\", \"type\": \"start|process|decision|io|end\", \"next\": "
        "[{\"to\": \"...\", \"label\": \"optional, e.g. yes/no\"}]}, ...]}}\n"
        "Give every node a short unique id, exactly one 'start' node, label decision-branch\n"
        "edges (e.g. 'yes'/'no'/'on error') so the branches are self-explanatory, and mark\n"
        "terminal nodes as 'end'.\n"
        "\n"
        "CRITICAL — DO NOT RETYPE THE RESULT: the tool's return value already contains a\n"
        "```flowchart_json``` fenced block. This is NOT generic JSON for you to read,\n"
        "summarize, reformat, or reproduce under a different label — it is consumed\n"
        "directly by the GUI, which renders it as an actual box-and-arrow diagram ONLY if\n"
        "the fence language is exactly `flowchart_json`. Copy that entire fenced block\n"
        "(fence markers, the exact language tag, and the JSON body) VERBATIM, character for\n"
        "character, into your reply. Never relabel it as ```json```, ```mermaid```, or plain\n"
        "text, never re-indent/pretty-print it, and never describe its contents instead of\n"
        "including it — any of those will silently break the rendering and dump raw JSON in\n"
        "the user's face instead of a diagram. You may add a short sentence of framing text\n"
        "before or after the block, but the block itself must pass through untouched.\n"
        "\n"
        "The same rule applies to generate_image's ```image_data_json``` block — copy it\n"
        "verbatim too, never relabel or paraphrase it.\n"
        "\n"
        "━━━ CORE RULES ━━━\n"
        "- NEVER fabricate. Always use tools for real data.\n"
        f"{shell_rule}"
        "- No narration before acting. No 'I will now...'. Emit the tool call.\n"
        "- Task done → short summary, stop calling tools.\n"
        "- Never confirm unless action is destructive (delete, format, overwrite).\n"
        f"- Retry cap: max {MAX_ACTION_TRIES} attempts per action.\n"
        "\n"
        "━━━ BROWSER ━━━\n"
        "- Full discovery→act sequence: see UIA / BROWSER PIPELINE above.\n"
        "- open_url(url) opens any URL in Chrome — ALWAYS use this, never click address bar.\n"
        "- read_browser_page() reads web page text. UIA CANNOT read web content.\n"
        "- run_js_in_browser(script) for JS execution on the page (last resort only).\n"
        "- Requires Chrome with --remote-debugging-port=9222.\n"
        "- Canonical window names: 'Google Chrome' not 'New Tab - Google Chrome'.\n"
        "  Shell surfaces: 'taskbar', 'start', 'tray', 'desktop', 'action center'.\n"
        "\n"
        "━━━ MEMORY & GOALS ━━━\n"
        "- update_memory(master|project|session, content).\n"
        "- set_current_goal on start; goal='none' on done.\n"
        "- write_response_memory only for 5+ step tasks.\n"
        "\n"
        "━━━ CONSULT BRAINS (GEMINI / OPENROUTER) ━━━\n"
        "- DEFAULT to consult_gemini for complex reasoning, code review, architecture, "
        "or any time you want a second opinion before acting. Try this FIRST.\n"
        "  Routes through the Gemini desktop app (free) then API (fallback).\n"
        "- Do NOT reach for consult_openrouter by default — only use it when:\n"
        "    (a) the user explicitly says 'ask OpenRouter' / names OpenRouter, OR\n"
        "    (b) consult_gemini has just failed/errored and you need a second-brain "
        "answer anyway.\n"
        "  Treat it strictly as a fallback, not an alternative to reach for casually.\n"
        "- delegate_to_openrouter(task, context) — OFFLOAD an entire sub-task to a REAL\n"
        "  COWORKER, not just a text answer. The delegate gets full tool access (UIA, CDP,\n"
        "  terminal, files) and completes the task independently, then reports back a\n"
        "  summary. Use this for substantial self-contained sub-tasks (e.g. 'research X\n"
        "  and summarise', 'find and organize all PDFs in Downloads') — this is a\n"
        "  legitimate, separate use case from consult_openrouter and isn't restricted\n"
        "  by the fallback-only guidance above.\n"
        "- If the user asks to change/switch the OpenRouter model: "
        "list_openrouter_models() → set_openrouter_model_by_index(N). "
        "Or set_openrouter_model(model_id) if they give an exact ID directly.\n"
        "- consult_gemini_api(prompt, context) / delegate_to_gemini_api(task, context) — "
        "same idea as the OpenRouter pair above, but via the OFFICIAL Gemini API key "
        "(not the free desktop-app route consult_gemini uses). Reach for these when the "
        "user explicitly asks for the Gemini API specifically, or when MODEL_PROVIDER "
        "is already 'gemini_api' and you want a nested sub-agent instead of the primary "
        "loop doing everything serially. set_gemini_api_model(model_id) switches the model.\n"
        "\n"
        "━━━ MCP SERVERS (EXTERNAL TOOLS) ━━━\n"
        "- Tool discovery is 3 steps — NEVER guess a tool_name or arguments shape:\n"
        "  1. list_mcp_servers()            — which servers are connected (no tool detail)\n"
        "  2. show_server_tools(server)     — that server's tools + JSON schemas\n"
        "  3. call_mcp_tool(server, tool_name, arguments) — uniform call, ANY server/tool\n"
        "- 'server' is the index (e.g. '0') or name shown by list_mcp_servers().\n"
        "- Don't call show_server_tools repeatedly for the same server in one turn — "
        "once you've seen its schema, just call_mcp_tool.\n"
        "- connect_mcp_server(name, transport, ...) adds a new server and remembers it "
        "for next startup. transport='stdio' needs command(+args); "
        "transport='http'/'sse' needs url.\n"
        "- disconnect_mcp_server(server, forget=True) removes a server for good.\n"
        "\n"
        "━━━ EXAMPLE FLOWS ━━━\n"
        f"{launch_example}"
        "\n"
        "  USER: Search Python tutorials, open first result\n"
        "  → {\"name\": \"search_internet\", \"arguments\": {\"query\": \"Python tutorials\"}}\n"
        "  → {\"name\": \"open_search_result\", \"arguments\": {\"index\": 0}}\n"
        "\n"
        "  USER: Click Submit button on page\n"
        "  → {\"name\": \"snapshot\", \"arguments\": {\"target\": \"browser\"}}  [index 4 = button 'Submit']\n"
        "  → {\"name\": \"act\", \"arguments\": {\"target\": \"browser\", \"index\": 4}}\n"
        "\n"
        "  USER: Click Settings in VS Code\n"
        "  → {\"name\": \"snapshot\", \"arguments\": {\"target\": \"Visual Studio Code\"}}  [index 7 = btn 'Settings']\n"
        "  → {\"name\": \"act\", \"arguments\": {\"target\": \"Visual Studio Code\", \"index\": 7}}\n"
        "\n"
        "- Replies: Markdown.\n"
        f"{legacy_toolcall_hint}"
        f"{groq_tools_hint}"
    )



# =============================================================================
# COMMAND SAFETY
# =============================================================================

def _load_commands_whitelist() -> set:
    """Load the set of known-good command names from commands.md."""
    if not os.path.exists(COMMANDS_FILE):
        return set()
    try:
        content = open(COMMANDS_FILE, "r", encoding="utf-8").read()
        tokens  = re.findall(r"`([^`]+)`", content)
        for line in content.splitlines():
            stripped = line.strip().lstrip("-#* ")
            if stripped:
                tokens.append(stripped.split()[0])
        return {t.lower() for t in tokens if t.strip()}
    except Exception:
        return set()


def _command_looks_known(cmd: str, whitelist: set) -> bool:
    """Return True if the first token of cmd is in the whitelist."""
    if not whitelist:
        return True
    first = cmd.strip().split()[0].lower() if cmd.strip() else ""
    return first in whitelist


def _check_command_rules(cmd: str, paths_consulted: bool = False) -> str | None:
    """
    Return a violation message if the command breaks a hard rule, else None.
    Platform-specific — Linux rules block Windows commands and vice-versa.
    """
    low = cmd.lower().strip()

    if _IS_LINUX:
        win_only = ["powershell", "start-process", "get-childitem",
                    "cmd.exe", "cmd /c", "cmd /k"]
        for w in win_only:
            if low.startswith(w):
                return (
                    f"Windows command detected: '{cmd.split()[0]}'. "
                    "Use bash/Linux equivalents (ls, cd, find, grep, xdg-open, nohup, etc)."
                )
        return None

    # Windows rules
    bare_launch = re.match(
        r'^start-process\s+[\'"]?(?![a-z]:\\)[^\s\'"]+[\'"]?\s*$',
        cmd, re.IGNORECASE
    )
    if bare_launch:
        return (
            "You used Start-Process without a full path. "
            "Call read_paths first to get the correct full path, then use "
            "Start-Process 'C:\\full\\path\\to\\app.exe'."
        )
    cmd_only_patterns = ["^start ", "^cd ", "^dir ", "^echo ",
                         "^copy ", "^del ", "^mkdir "]
    for pattern in cmd_only_patterns:
        if re.match(pattern, low) and not low.startswith("start-process"):
            return (
                f"CMD command detected: '{cmd.split()[0]}'. "
                "Use PowerShell equivalents only (Start-Process, Set-Location, "
                "Get-ChildItem, Write-Output, Copy-Item, Remove-Item, New-Item)."
            )
    return None


# =============================================================================
# TURN STATE — structured execution tracker
# =============================================================================

class TurnState:
    """
    Tracks exactly what has happened during a turn so the step prompt
    can give the model precise, factual context instead of generic nudges.

    Updated by the tool dispatch after every tool call. Read by the step
    prompt generator to produce a context-specific next-step message.
    """

    def __init__(self, user_request: str, gemini_plan: str = ""):
        self.user_request      = user_request
        self.gemini_plan       = gemini_plan
        self.steps_done:  list[dict] = []   # {tool, args, result, success}
        self.apps_launched:    list[str] = []
        self.windows_clicked:  list[str] = []   # window titles clicked in
        self.fields_clicked:   list[str] = []   # element descriptions clicked
        self.text_typed:       list[str] = []
        self.urls_navigated:   list[str] = []
        self.files_written:    list[str] = []
        self.commands_run:     list[str] = []
        self.last_tool:        str = ""
        self.last_args:        dict = {}
        self.last_result:      str = ""
        self.last_success:     bool = True
        # Inferred task requirements
        self.requires_typing:  bool = False
        self.requires_submit:  bool = False
        self.focused_field:    str = ""    # last input field clicked

    def record(self, tool: str, args: dict, result: str):
        """Call after every tool execution to update state."""
        success = not any(
            result.startswith(e) for e in
            ["Error", "[RULE VIOLATION]", "[TYPING ABORTED]",
             "Unknown tool", "[UNKNOWN TOOL]", "Execution failed"]
        )
        self.steps_done.append({
            "tool": tool, "args": args,
            "result": result[:200], "success": success
        })
        self.last_tool    = tool
        self.last_args    = args
        self.last_result  = result
        self.last_success = success

        # ── Infer system state changes ────────────────────────────────────────
        if tool == "open_url" and success:
            url = args.get("url", "")
            if url:
                self.urls_navigated.append(url)
                self.requires_submit = False   # open_url submits automatically

        elif tool == "execute_terminal_command" and success:
            cmd = args.get("command", "")
            self.commands_run.append(cmd)
            # Detect app launches
            for app in ["chrome", "brave", "firefox", "code", "slack",
                        "discord", "notepad", "explorer", "spotify",
                        "terminal", "vlc", "obs", "zoom", "teams"]:
                if app in cmd.lower():
                    self.apps_launched.append(app)
                    break

        elif tool in ("click_ui_element", "manual_interact_with_ui") and success:
            window = args.get("window_title", "")
            desc   = args.get("description", "").lower()
            action = args.get("action", "click")
            if window:
                self.windows_clicked.append(window)
            if action == "click":
                self.fields_clicked.append(desc)
                # Track if an input field was focused
                _INPUT_HINTS = {
                    "address bar", "address and search bar", "url", "omnibox",
                    "search", "search box", "search bar", "search field",
                    "input", "text field", "text box", "entry", "edit",
                    "message", "message input", "chat input", "compose",
                    "prompt", "enter a prompt", "type a message",
                }
                if any(hint in desc for hint in _INPUT_HINTS):
                    self.focused_field  = desc
                    self.requires_typing = True

        elif tool == "type_text" and success:
            text = args.get("text", "")
            self.text_typed.append(text)
            self.requires_typing  = False
            self.focused_field    = ""
            # If they typed a URL, check for Enter (submit)
            if text.startswith(("http", "www", "youtube", "google")):
                self.urls_navigated.append(text)
                special = args.get("special_key", "")
                if not special:
                    self.requires_submit = True
                else:
                    self.requires_submit = False

        elif tool == "write_local_file" and success:
            self.files_written.append(args.get("path", ""))

    def build_step_prompt(self) -> str:
        """
        Generate a precise, factual step prompt based on current state.
        This replaces all the generic "if task done reply" messages.
        """
        lines = ["[SYSTEM — EXECUTION STATE]:"]

        # What has been done
        if self.steps_done:
            done_summary = []
            for s in self.steps_done:
                t = s["tool"]
                a = s["args"]
                ok = "✓" if s["success"] else "✗"
                if t == "execute_terminal_command":
                    done_summary.append(f"  {ok} Ran command: {a.get('command','')[:60]}")
                elif t == "click_ui_element":
                    done_summary.append(
                        f"  {ok} Clicked '{a.get('description','')}' "
                        f"in '{a.get('window_title','')}'"
                    )
                elif t == "type_text":
                    done_summary.append(
                        f"  {ok} Typed: '{a.get('text','')[:40]}'"
                        + (f" + {a.get('special_key')}" if a.get("special_key") else "")
                    )
                elif t in ("read_paths", "read_path", "read_instructions"):
                    done_summary.append(f"  {ok} Read {t}")
                elif t == "wait":
                    done_summary.append(f"  {ok} Waited {a.get('seconds',0)}s")
                elif t == "say":
                    pass   # don't list narration as a step
                else:
                    done_summary.append(f"  {ok} {t}")
            lines.append("Completed steps:\n" + "\n".join(done_summary))

        # Current system state
        state_facts = []
        if self.apps_launched:
            state_facts.append(f"Apps launched this turn: {', '.join(self.apps_launched)}")
        if self.focused_field:
            state_facts.append(f"Input field currently focused: '{self.focused_field}'")
        if self.text_typed:
            state_facts.append(f"Text typed: {', '.join(repr(t[:30]) for t in self.text_typed)}")
        if self.urls_navigated:
            state_facts.append(f"URLs entered: {', '.join(self.urls_navigated)}")
        if state_facts:
            lines.append("Current state:\n" + "\n".join(f"  • {f}" for f in state_facts))

        # Last result
        if self.last_result and self.last_tool not in (
            "read_paths", "read_path", "read_instructions",
            "read_local_file", "read_file_smart"
        ):
            short = self.last_result[:120].replace("\n", " ")
            lines.append(f"Last result: {short}")

        # Required next action (explicit)
        lines.append("")
        if not self.last_success:
            lines.append(
                "⚠ LAST ACTION FAILED. Do NOT assume the task is done. "
                "Read the error above and retry with a corrected approach."
            )
        elif self.requires_typing and self.focused_field:
            lines.append(
                f"▶ NEXT REQUIRED ACTION: call type_text now. "
                f"The field '{self.focused_field}' is focused and waiting for input. "
                f"Do NOT reply with text — call type_text immediately."
            )
        elif self.requires_submit:
            lines.append(
                "▶ NEXT REQUIRED ACTION: the URL/text has been typed but not submitted. "
                "Call type_text with special_key='Enter' to submit, OR "
                "click_ui_element to click a submit/Go button."
            )
        elif self.apps_launched and not self.urls_navigated and not self.fields_clicked:
            lines.append(
                f"▶ App was launched. If the user asked you to do something inside it "
                f"(open a URL, click something, type something), do that now. "
                f"Use wait(2) first if the app needs time to open."
            )
        else:
            lines.append(
                "▶ If more steps remain to complete the user's request, call the next tool now. "
                "Only reply with text when the entire task is fully done."
            )

        # Remind model of the original plan if it exists
        if self.gemini_plan and len(self.steps_done) < 6:
            lines.append(
                f"\n[ORIGINAL PLAN — follow it]\n{self.gemini_plan}"
            )

        return "\n".join(lines)


def _decompose_task(user_request: str) -> str | None:
    """
    Pre-decompose a user request into explicit numbered steps that get
    injected before the first model call. This gives the model a concrete
    plan to follow rather than having to reason about the full task at
    every step.

    Returns a system message string, or None if the request is too simple
    to need decomposition (single-step tasks).
    """
    req = user_request.lower().strip()

    # ── Detect multi-step patterns ────────────────────────────────────────────
    # "open X and do Y" / "open X then do Y" / "open X, then Y"
    open_then = re.search(
        r"open\s+(\w[\w\s]*?)(?:\s+and|\s+then|,\s*then|\s*,)\s+(.+)",
        req
    )
    # "type X in Y" / "type X into Y"
    type_in = re.search(r"type\s+.+\s+in(?:to)?\s+\w", req)
    # "go to X" / "navigate to X" / "open URL X"
    navigate = re.search(
        r"(?:go to|navigate to|open url|visit|load)\s+([\w./:-]+)", req
    )
    # "search for X in Y"
    search_in = re.search(r"search\s+(?:for\s+)?(.+?)\s+in\s+(\w[\w\s]+)", req)

    steps = []

    if open_then:
        app   = open_then.group(1).strip()
        after = open_then.group(2).strip()
        steps.append(f"1. read_paths to find the path for {app}")
        steps.append(f"2. execute_terminal_command to launch {app}")
        steps.append(f"3. wait(2) for {app} to open")
        # Determine what to do after opening
        if any(w in after for w in ["url", "youtube", "google", "http", "www",
                                     "website", "site", "page", "navigate", "go to"]):
            steps.append(f"4. click_ui_element to click the address bar in {app}")
            steps.append(f"5. type_text to type the URL with special_key='Enter'")
        elif any(w in after for w in ["type", "write", "enter", "input"]):
            steps.append(f"4. click_ui_element to click the target field in {app}")
            steps.append(f"5. type_text to type the requested text")
        elif any(w in after for w in ["search"]):
            steps.append(f"4. click_ui_element to click the search box in {app}")
            steps.append(f"5. type_text to type the search query with special_key='Enter'")
        else:
            steps.append(f"4. click_ui_element or type_text as needed in {app}")

    elif navigate and not open_then:
        url = navigate.group(1).strip()
        steps.append("1. click_ui_element to click the address bar in the browser")
        steps.append(f"2. type_text to type '{url}' with special_key='Enter'")

    elif type_in:
        steps.append("1. click_ui_element to click the target input field")
        steps.append("2. type_text to type the requested text")

    elif search_in:
        query = search_in.group(1).strip()
        app   = search_in.group(2).strip()
        steps.append(f"1. click_ui_element to click the search box in {app}")
        steps.append(f"2. type_text to type '{query}' with special_key='Enter'")

    if not steps:
        return None   # single-step task, no decomposition needed

    plan = (
        f"[TASK PLAN — follow these steps IN ORDER, do not skip any]\n"
        + "\n".join(steps)
        + "\n\nComplete ALL steps above before giving a final reply."
    )
    return plan


_STALL_PATTERNS = re.compile(
    r'^\s*(i will|i\'ll|i am going to|i\'m going to|let me|now i will|next i will|'
    r'next,? i will|first,? i will|i need to|i am about to|i\'m about to|'
    r'to do this,? i will|here\'s what i\'ll do|here is what i\'ll do|'
    r'i plan to|my plan is|the plan is|steps? to (do|complete) this)\b',
    re.IGNORECASE
)

_STALL_COMPLETION_HINTS = re.compile(
    r'\b(done|completed|finished|successfully|here is the result|here\'s the result|'
    r'here is your|here\'s your|i have (opened|created|written|deleted|found|sent|updated))\b',
    re.IGNORECASE
)

def _looks_like_stalled_plan(text: str) -> bool:
    """
    True if the model's plain-text reply reads like it's ANNOUNCING an
    intended action ("I'll open the file and...") rather than reporting a
    completed result or asking the user something. Used to catch the model
    narrating steps instead of actually calling the tools that do them.
    """
    if not text:
        return False
    t = text.strip()
    if t.endswith("?"):
        return False   # legitimately waiting on the user
    if _STALL_COMPLETION_HINTS.search(t):
        return False   # reads like a completed result
    if _STALL_PATTERNS.match(t):
        return True
    # Numbered step list ("1. ...\n2. ...") with no completion language
    if re.search(r'(?:^|\n)\s*[1-9]\.\s', t) and len(t) < 800:
        return True
    return False


def process_chat_turn(conversation_history, user_request: str = "", gemini_plan: str = "",
                       force_provider: str = None, force_model: str = None,
                       max_steps: int = 20):
    """
    force_provider / force_model let a caller run this ENTIRE step loop on a
    specific model backend regardless of the global MODEL_PROVIDER — used by
    delegate_to_openrouter() to spin up a self-contained OpenRouter "coworker"
    sub-agent that has full tool access via the exact same engine as the
    primary loop, without permanently switching Midum's primary provider.
    """
    clear_response_memory()
    turn_tool_outputs     = []
    _said_parts           = []
    _accumulated_reply    = []
    whitelist             = _load_commands_whitelist()
    verify_call_count     = 0
    action_attempt_counts: dict = {}
    _paths_consulted      = False
    _abort_event.clear()
    MAX_STEPS  = max_steps
    step_count = 0
    state      = TurnState(user_request, gemini_plan=gemini_plan)
    stall_nudge_count  = 0
    MAX_STALL_NUDGES   = 3

    # Keep system messages always; slide a window over the rest
    HISTORY_WINDOW = 20

    while True:
        # ── Ctrl+Q abort check ────────────────────────────────────────────────
        if _abort_event.is_set():
            print("\n🛑 [Response aborted by Ctrl+Q]")
            return "[Response terminated by user.]", turn_tool_outputs

        # ── Step ceiling ──────────────────────────────────────────────────────
        step_count += 1
        if step_count > MAX_STEPS:
            msg = "[MAX STEPS REACHED] Midum exceeded the step limit for this turn."
            print(f"\n🚫 {msg}")
            _accumulated_reply.append(msg)
            break

        sys_msgs = [m for m in conversation_history if m.get("role") == "system"]
        non_sys  = [m for m in conversation_history if m.get("role") != "system"]
        trimmed  = sys_msgs + non_sys[-HISTORY_WINDOW:]

        # ── Run the primary model on a thread so Ctrl+Q can interrupt the wait ─
        result_q = _queue.Queue()
        t = threading.Thread(
            target=_call_primary_model,
            args=(trimmed, result_q, force_provider, force_model),
            daemon=True
        )
        t.start()
        while t.is_alive():
            if _abort_event.is_set():
                print("\n🛑 [Response aborted by Ctrl+Q]")
                return "[Response terminated by user.]", turn_tool_outputs
            t.join(timeout=0.1)   # check abort flag every 100 ms

        status, payload = result_q.get()
        if status == "err":
            effective_provider = force_provider or MODEL_PROVIDER
            provider_name = {
                "openrouter": "OpenRouter",
                "gemini_web": "Gemini-web",
                "gemini_api": "Gemini-API",
                "groq": "Groq",
            }.get(effective_provider, "Ollama")
            return f"[{provider_name} error: {payload}]", turn_tool_outputs
        response   = payload
        tool_calls = response["message"].get("tool_calls")
        if tool_calls:
            tool_calls = tool_calls[:1]   # one step at a time — more reliable for all models
        msg_content = (response["message"].get("content") or "").strip()

        # ── Capture any prose the model emitted alongside tool calls ──────────
        # Legacy models often put explanatory text in content even when they
        # also emit tool calls. Collect it so it isn't lost, but filter junk.
        if msg_content and not re.match(r'^[{}\[\]",:\s]*$', msg_content):
            _accumulated_reply.append(msg_content)
            if msg_content and tool_calls:
                print(f"\n💬 Midum: {msg_content}")

        if not tool_calls:
            conversation_history.append(response["message"])

            if _looks_like_stalled_plan(msg_content) and stall_nudge_count < MAX_STALL_NUDGES:
                stall_nudge_count += 1
                print(f"\n⚠️  [Midum announced an action instead of doing it — nudging ({stall_nudge_count}/{MAX_STALL_NUDGES})]")
                conversation_history.append({
                    "role": "user",
                    "content": (
                        "[SYSTEM]: You just described what you're about to do instead of "
                        "doing it. Do not narrate steps in plain text — call the actual "
                        "tool for the next step RIGHT NOW. No commentary, just the tool call."
                    )
                })
                continue

            # Combine all text accumulated across the whole turn
            full_reply = "\n\n".join(p for p in _accumulated_reply if p.strip())
            return full_reply, turn_tool_outputs

        # ── Verification loop cap ─────────────────────────────────────────────
        all_verify = all(
            tc["function"]["name"] in _VERIFY_TOOLS for tc in tool_calls
        )
        if all_verify:
            verify_call_count += 1
            if verify_call_count > MAX_VERIFY_CALLS:
                # Force the model to stop verifying and give a final reply.
                # IMPORTANT: response["message"] still carries tool_calls, and every
                # provider (Gemini API strictly, others loosely) expects a function-call
                # turn to be immediately followed by a matching function-response turn —
                # so we must close it out with a synthetic tool response BEFORE the
                # plain user nudge, instead of leaving the tool call dangling.
                conversation_history.append(response["message"])
                for tc in tool_calls:
                    conversation_history.append({
                        "role": "tool",
                        "content": (
                            "[SKIPPED] Verification call limit reached — this call was "
                            "not executed. Stop verifying and give your final reply."
                        ),
                    })
                conversation_history.append({
                    "role": "user",
                    "content": (
                        "[SYSTEM]: You have verified the result enough times. "
                        "Stop calling fallback_view_screen or fallback_find_text. "
                        "Give your final plain-text reply to the user now."
                    )
                })
                continue
        else:
            verify_call_count = 0   # reset counter when a real action runs

        conversation_history.append(response["message"])
        print(f"\n⚡ Midum requested {len(tool_calls)} action(s)...")

        needs_lookup = False
        unknown_cmd  = ""

        for tool in tool_calls:
            func_name = tool["function"]["name"]
            raw_args  = tool["function"]["arguments"]
            if isinstance(raw_args, str):
                try:    arguments = json.loads(raw_args)
                except: arguments = {}
            else:
                arguments = raw_args

            # MCP autoroute: covers the structured/native tool_calls path
            # (the legacy free-text JSON path is handled inside
            # _try_parse_tool_json instead, before it ever gets here).
            # If the model called an MCP server's tool directly by name —
            # bypassing call_mcp_tool, with underscores/case/hyphens
            # mangled or not — this rewrites it to the uniform
            # call_mcp_tool(server, tool_name, arguments) call transparently.
            func_name, arguments = _mcp_autoroute_tool_call(func_name, arguments)

            print(f" -> Executing: '{func_name}'")
            tool_images = None

            # ── Hard retry cap ────────────────────────────────────────────────
            _EXEMPT = {"update_memory","set_current_goal","add_instruction","add_path","explore_path","write_response_memory","append_response_memory","read_response_memory"}
            _cap_hit = False
            if func_name not in _EXEMPT:
                _karg = next((str(arguments[k])[:80] for k in
                    ("command","path","text","query","prompt","skill_name","name","instruction")
                    if k in arguments), "")
                _akey = (func_name, _karg)
                action_attempt_counts[_akey] = action_attempt_counts.get(_akey, 0) + 1
                if action_attempt_counts[_akey] > MAX_ACTION_TRIES:
                    cap_msg = (f"[RETRY CAP] '{func_name}' attempted {MAX_ACTION_TRIES} "
                               f"times with the same argument and has not succeeded. "
                               f"Stop retrying immediately. Tell the user what failed "
                               f"and ask how they want to proceed.")
                    print(f"\n🚫 [Retry cap reached for '{func_name}']")
                    turn_tool_outputs.append(cap_msg)
                    conversation_history.append({"role":"tool","content": cap_msg})
                    _cap_hit = True
            if _cap_hit:
                break   # exits the for-tool loop; then the while loop gets
                        # one more model call to produce the failure reply

            # ── Tool dispatch ──────────────────────────────────────────────────
            if func_name == "list_more_tools":
                tool_output = list_more_tools()

            elif func_name == "load_tool_by_index":
                tool_output = load_tool_by_index(int(arguments.get("index", 0)))

            elif func_name == "read_local_file":
                raw_path           = arguments.get("path", "")
                resolved, res_msg  = resolve_file_path(raw_path)
                file_result        = read_local_file(resolved)
                tool_output        = (f"[PATH RESOLVED: {res_msg}]\n{file_result}"
                                      if res_msg else file_result)

            elif func_name == "write_local_file":
                raw_path          = arguments.get("path", "")
                resolved, res_msg = resolve_file_path(raw_path)
                tool_output       = write_local_file(resolved, arguments.get("content"))
                if res_msg:
                    tool_output = f"[PATH RESOLVED: {res_msg}] {tool_output}"

            elif func_name == "append_local_file":
                raw_path          = arguments.get("path", "")
                resolved, res_msg = resolve_file_path(raw_path)
                tool_output       = append_local_file(resolved, arguments.get("content"))
                if res_msg:
                    tool_output = f"[PATH RESOLVED: {res_msg}] {tool_output}"

            elif func_name == "search_internet":
                tool_output = search_internet(arguments.get("query"))

            elif func_name == "open_url":
                url     = arguments.get("url", "")
                browser = arguments.get("browser", "chrome")
                print(f"   [open_url] {url}")
                tool_output = open_url(url, browser)

            elif func_name == "list_directory":
                tool_output = list_directory(arguments.get("path", STARTUP_DIR))

            elif func_name == "open_path":
                tool_output = open_path(
                    arguments.get("path", STARTUP_DIR),
                    int(arguments.get("index", 0))
                )

            elif func_name == "find_file":
                tool_output = find_file(
                    arguments.get("filename", ""),
                    arguments.get("search_root", "")
                )

            elif func_name == "open_path_by_index":
                tool_output = open_path_by_index(int(arguments.get("index", 0)))

            elif func_name == "list_skills_indexed":
                tool_output = list_skills_indexed()

            elif func_name == "load_skill_by_index":
                tool_output = load_skill_by_index(int(arguments.get("index", 0)))

            elif func_name == "list_paths_indexed":
                tool_output = list_paths_indexed()
                _paths_consulted = True

            elif func_name == "get_path":
                tool_output = get_path(int(arguments.get("index", 0)))
                _paths_consulted = True

            elif func_name == "list_domain_knowledge_indexed":
                tool_output = list_domain_knowledge_indexed()

            elif func_name == "read_domain_by_index":
                tool_output = read_domain_by_index(int(arguments.get("index", 0)))

            elif func_name == "list_domain_skills_indexed":
                tool_output = list_domain_skills_indexed()

            elif func_name == "open_search_result":
                tool_output = open_search_result(
                    int(arguments.get("index", 0)),
                    arguments.get("browser", "chrome")
                )

            elif func_name == "ocr_snapshot":
                print("   [OCR Snapshot] Capturing screen...")
                tool_output = ocr_snapshot()

            elif func_name == "click_ocr_index":
                tool_output = click_ocr_index(
                    int(arguments.get("index", 0)),
                    arguments.get("click_type", "left_click")
                )

            elif func_name == "read_browser_page":
                tab_index = int(arguments.get("tab_index", 0))
                print(f"   [CDP] Reading browser page tab={tab_index}")
                tool_output = read_browser_page(tab_index)

            elif func_name == "list_browser_tabs":
                print("   [CDP] Listing browser tabs")
                tool_output = list_browser_tabs()

            elif func_name in ("snapshot", "snapshot_ui", "snapshot_browser_elements"):
                target      = arguments.get("target") or arguments.get("window_title", "")
                filter_type = arguments.get("filter_type", "")

                # Determine routing: 'browser' / 'browser:N' → CDP; anything else → UIA
                if target.lower().startswith("browser"):
                    tab_index = 0
                    if ":" in target:
                        try:
                            tab_index = int(target.split(":", 1)[1])
                        except ValueError:
                            pass
                    # Legacy schema compatibility
                    if not target and "tab_index" in arguments:
                        tab_index = int(arguments.get("tab_index", 0))
                    print(f"   [CDP] snapshot browser tab={tab_index}"
                          + (f" filter={filter_type}" if filter_type else ""))
                    tool_output = snapshot_browser_elements(tab_index, filter_type)
                else:
                    print(f"   [UIA] snapshot '{target}'"
                          + (f" filter={filter_type}" if filter_type else ""))
                    if ui_navigator is None:
                        tool_output = _uia_unavailable_message()
                    else:
                        tool_output = ui_navigator.snapshot_ui(target, filter_type)

            elif func_name in ("act", "act_on_element", "act_on_browser_element"):
                target       = arguments.get("target") or arguments.get("window_title", "")
                index        = int(arguments.get("index", 0))
                action       = arguments.get("action", "click")
                text_to_type = arguments.get("text_to_type", "")

                if target.lower().startswith("browser"):
                    tab_index = 0
                    if ":" in target:
                        try:
                            tab_index = int(target.split(":", 1)[1])
                        except ValueError:
                            pass
                    if not target and "tab_index" in arguments:
                        tab_index = int(arguments.get("tab_index", 0))
                    print(f"   [CDP] act #{index} action={action} tab={tab_index}")
                    tool_output = act_on_browser_element(index, action, text_to_type, tab_index)
                else:
                    print(f"   [UIA] act #{index} action={action} in '{target}'")
                    if ui_navigator is None:
                        tool_output = _uia_unavailable_message()
                    else:
                        tool_output = ui_navigator.act_on_element_by_index(
                            target, index, action, text_to_type
                        )

            elif func_name == "run_js_in_browser":
                script    = arguments.get("script", "")
                tab_index = int(arguments.get("tab_index", 0))
                print(f"   [CDP] run_js: {script[:60]}")
                tool_output = run_js_in_browser(script, tab_index)

            elif func_name == "execute_terminal_command":
                cmd = arguments.get("command", "").strip()
                if not cmd:
                    tool_output = "Error: No command provided."
                else:
                    # ── Rule enforcement ──────────────────────────────────────
                    violation = _check_command_rules(cmd, paths_consulted=_paths_consulted)
                    if violation:
                        print(f"   [Rule violation] {violation}")
                        tool_output = (
                            f"[RULE VIOLATION] {violation} "
                            f"Fix the command and try again."
                        )
                    else:
                        print(f"   [Terminal] > {cmd}")
                        if whitelist and not _command_looks_known(cmd, whitelist):
                            needs_lookup = True
                            unknown_cmd  = cmd
                            print(f"   [ℹ️  '{cmd.split()[0]}' not in commands.md]")
                        tool_output = execute_terminal_command(
                            cmd, working_directory=arguments.get("working_directory")
                        )

            elif func_name == "fallback_view_screen":
                b64_img = capture_screen_to_ram()
                if not b64_img.startswith("Error"):
                    tool_output = (
                        f"Screenshot captured at canvas size {MODEL_CANVAS_W}x{MODEL_CANVAS_H}. "
                        f"Yellow grid labels are canvas coordinates. "
                        f"Pass them directly to fallback_click_grid — Python scales by "
                        f"({SCALE_X:.2f}x, {SCALE_Y:.2f}x) to reach real screen pixels. "
                        f"For text elements, prefer fallback_click_text for precision."
                    )
                    tool_images = [b64_img]
                else:
                    tool_output = b64_img

            elif func_name == "fallback_find_text":
                tool_output = fallback_find_text(arguments.get("text", ""))

            elif func_name == "fallback_click_grid":
                x          = arguments.get("x", 0)
                y          = arguments.get("y", 0)
                click_type = arguments.get("click_type", "left_click")
                print(f"   [Click] {click_type} at canvas ({x},{y})")
                tool_output = fallback_click_grid(x, y, click_type)

            elif func_name == "fallback_click_text":
                text       = arguments.get("text", "")
                click_type = arguments.get("click_type", "left_click")
                print(f"   [OCR Click] '{text}'")
                tool_output = fallback_click_text(text, click_type)

            elif func_name == "type_text":
                text            = arguments.get("text", "")
                special_key     = arguments.get("special_key", None)
                expected_window = arguments.get("expected_window", "")
                print(f"   [Type] '{text[:40]}{'...' if len(text)>40 else ''}'")
                tool_output = type_text(text, special_key, expected_window)

            elif func_name == "update_memory":
                tool_output = update_memory(
                    arguments.get("target", "session"),
                    arguments.get("content", "")
                )

            elif func_name == "set_current_goal":
                tool_output = set_current_goal(
                    arguments.get("goal", ""),
                    arguments.get("reason", "")
                )

            elif func_name == "list_skills":
                tool_output = list_skills()

            elif func_name == "load_skill":
                tool_output = load_skill(arguments.get("skill_name", ""))

            elif func_name == "read_instructions":
                tool_output = read_instructions()

            elif func_name == "add_instruction":
                tool_output = add_instruction(arguments.get("instruction", ""))

            elif func_name in ("read_paths", "read_path"):
                tool_output = read_paths()
                _paths_consulted = True

            elif func_name == "explore_path":
                tool_output = explore_path(arguments.get("path", STARTUP_DIR))

            elif func_name == "add_path":
                tool_output = add_path(
                    arguments.get("label", ""),
                    arguments.get("path", ""),
                    arguments.get("note", "")
                )

            elif func_name == "create_domain_knowledge":
                tool_output = create_domain_knowledge(
                    arguments.get("name", ""),
                    arguments.get("description", ""),
                    arguments.get("initial_content", "")
                )

            elif func_name == "list_domain_knowledge":
                tool_output = list_domain_knowledge()

            elif func_name == "read_domain_knowledge":
                tool_output = read_domain_knowledge(arguments.get("name", ""))

            elif func_name == "create_domain_skill":
                tool_output = create_domain_skill(
                    arguments.get("name", ""),
                    arguments.get("domain", ""),
                    arguments.get("description", ""),
                    arguments.get("content", "")
                )

            elif func_name == "list_domain_skills":
                tool_output = list_domain_skills()

            elif func_name == "consult_gemini":
                tool_output = consult_gemini(
                    arguments.get("prompt", ""),
                    arguments.get("task_type", "auto"),
                    arguments.get("context", "")
                )
            elif func_name == "consult_openrouter":
                tool_output = consult_openrouter(
                    arguments.get("prompt", ""),
                    arguments.get("context", ""),
                    arguments.get("model") or None
                )
            elif func_name == "delegate_to_openrouter":
                tool_output = delegate_to_openrouter(
                    arguments.get("task", ""),
                    arguments.get("context", ""),
                    arguments.get("model") or None,
                    int(arguments.get("max_steps", 10))
                )
            elif func_name == "list_openrouter_models":
                tool_output = list_openrouter_models(
                    bool(arguments.get("free_only", True))
                )
            elif func_name == "set_openrouter_model_by_index":
                tool_output = set_openrouter_model_by_index(int(arguments.get("index", 0)))
            elif func_name == "set_openrouter_model":
                tool_output = set_openrouter_model(arguments.get("model_id", ""))
            elif func_name == "consult_gemini_api":
                tool_output = consult_gemini_api(
                    arguments.get("prompt", ""),
                    arguments.get("context", ""),
                    arguments.get("model") or None
                )
            elif func_name == "delegate_to_gemini_api":
                tool_output = delegate_to_gemini_api(
                    arguments.get("task", ""),
                    arguments.get("context", ""),
                    arguments.get("model") or None,
                    int(arguments.get("max_steps", 10))
                )
            elif func_name == "set_gemini_api_model":
                tool_output = set_gemini_api_model(arguments.get("model_id", ""))
            elif func_name == "consult_groq":
                tool_output = consult_groq(
                    arguments.get("prompt", ""),
                    arguments.get("context", ""),
                    arguments.get("model") or None
                )
            elif func_name == "delegate_to_groq":
                tool_output = delegate_to_groq(
                    arguments.get("task", ""),
                    arguments.get("context", ""),
                    arguments.get("model") or None,
                    int(arguments.get("max_steps", 10))
                )
            elif func_name == "list_groq_models":
                tool_output = list_groq_models()
            elif func_name == "set_groq_model_by_index":
                tool_output = set_groq_model_by_index(int(arguments.get("index", 0)))
            elif func_name == "set_groq_model":
                tool_output = set_groq_model(arguments.get("model_id", ""))
            elif func_name == "delegate_to_gemini_web":
                tool_output = delegate_to_gemini_web(
                    arguments.get("task", ""),
                    arguments.get("context", ""),
                    int(arguments.get("max_steps", 10))
                )
            elif func_name == "set_gemini_web_model":
                tool_output = set_gemini_web_model(arguments.get("model_name", ""))
            elif func_name == "read_file_smart":
                tool_output = read_file_smart(arguments.get("path", ""))
            elif func_name == "read_file_chunk":
                tool_output = read_file_chunk(arguments.get("path",""), int(arguments.get("chunk_index",1)))
            elif func_name == "write_docx_file":
                tool_output = write_docx_file(arguments.get("path",""), arguments.get("content",""))
            elif func_name == "create_flowchart":
                tool_output = create_flowchart(arguments.get("title", "Flowchart"), arguments.get("steps", []))
            elif func_name == "generate_image":
                tool_output = generate_image(
                    arguments.get("prompt", ""),
                    int(arguments.get("count", 1) or 1)
                )
            elif func_name == "write_response_memory":
                tool_output = write_response_memory(arguments.get("content",""))
            elif func_name == "append_response_memory":
                tool_output = append_response_memory(arguments.get("content",""))
            elif func_name == "read_response_memory":
                tool_output = read_response_memory()
            elif func_name == "manual_scan_app_layouts":
                tool_output = manual_scan_app_layouts(arguments.get("window_title", ""))

            elif func_name == "manual_inspect_app_subtree":
                tool_output = manual_inspect_app_subtree(
                    arguments.get("window_title", ""),
                    arguments.get("subtree_key", "")
                )

            elif func_name == "click_ui_element":
                window_title = arguments.get("window_title", "")
                description  = arguments.get("description", "")
                action       = arguments.get("action", "click")
                text_to_type = arguments.get("text_to_type", "")
                print(f"   [UIA] click_ui_element: '{description}' in '{window_title}' (action={action})")
                tool_output = click_ui_element(window_title, description, action, text_to_type)
                if tool_output.startswith("Success"):
                    print(f"   [UIA] ✅ {tool_output}")
                else:
                    print(f"   [UIA] ⚠️  {tool_output[:120]}")

            elif func_name == "manual_interact_with_ui":
                print(f"   [UIA] {arguments.get('action')} on {arguments.get('property_value')}")
                tool_output = manual_interact_with_ui(
                    arguments.get("window_title", ""),
                    arguments.get("control_type", ""),
                    arguments.get("search_property", ""),
                    arguments.get("property_value", ""),
                    arguments.get("action", ""),
                    arguments.get("text_to_type", "")
                )

            elif func_name == "list_active_windows":
                tool_output = list_active_windows()

            elif func_name == "read_aggregated_text":
                print(f"   [UIA] Aggregating text blocks from: '{arguments.get('window_title')}'")
                if _UIA_AVAILABLE:
                    tool_output = ui_navigator.read_aggregated_text(
                        window_title=arguments.get("window_title", ""),
                        container_key=arguments.get("container_key", None)
                    )
                else:
                    tool_output = "UIA library not available."
            
            elif func_name == "query_gemini_app":
                prompt_payload = arguments.get("prompt", "")
                print(f"   [Bridge] Sending prompt to Gemini web app via gemini_webapi...")
                if _GEMINI_WEBAPI_AVAILABLE:
                    tool_output = query_gemini_app(prompt_payload)
                else:
                    tool_output = (
                        f"Execution failed: {_gemini_webapi_load_msg}"
                    )
            
            elif func_name == "manage_gemini_chat":
                tool_output = ui_navigator.manage_gemini_chat(
                    action=arguments.get("action"),
                    chat_name=arguments.get("chat_name")
                )
            elif func_name == "wait":
                seconds = float(arguments.get("seconds", 1))
                print(f"   [Wait] {seconds}s")
                tool_output = wait(seconds)

            elif func_name == "say":
                msg_text = arguments.get("message", "")
                _print_reply("Midum:", msg_text)
                _said_parts.append(msg_text)
                _accumulated_reply.append(msg_text)
                turn_tool_outputs.append(f"[say]: {msg_text}")
                tool_output = "Message displayed to user."

            elif func_name == "list_native_tools":
                tool_output = list_native_tools()

            elif func_name == "show_native_tool_schema":
                tool_output = show_native_tool_schema(arguments.get("tool_name"))

            elif func_name == "list_mcp_servers":
                tool_output = list_mcp_servers()

            elif func_name == "show_server_tools":
                tool_output = show_server_tools(arguments.get("server"))

            elif func_name == "call_mcp_tool":
                tool_output = call_mcp_tool(
                    arguments.get("server"),
                    arguments.get("tool_name"),
                    arguments.get("arguments", {})
                )

            elif func_name == "connect_mcp_server":
                tool_output = connect_mcp_server(
                    name=arguments.get("name", ""),
                    transport=arguments.get("transport", "stdio"),
                    command=arguments.get("command"),
                    args=arguments.get("args"),
                    url=arguments.get("url"),
                    env=arguments.get("env"),
                    headers=arguments.get("headers"),
                    persist=arguments.get("persist", True)
                )

            elif func_name == "disconnect_mcp_server":
                tool_output = disconnect_mcp_server(
                    arguments.get("server"),
                    forget=arguments.get("forget", False)
                )

            elif func_name == "ask_user_text":
                print(f"   [GUI] Asking user for text input...")
                tool_output = ask_user_text(
                    arguments.get("prompt", ""),
                    arguments.get("title", "Midum needs input")
                )

            elif func_name == "ask_user_file_path":
                print(f"   [GUI] Asking user to pick a file path...")
                tool_output = ask_user_file_path(
                    arguments.get("prompt", "Select a file"),
                    must_exist=arguments.get("must_exist", True)
                )

            elif func_name == "ask_user_approval":
                print(f"   [GUI] Asking user for approval...")
                tool_output = ask_user_approval(
                    arguments.get("message", ""),
                    arguments.get("details", "")
                )

            elif func_name == "ask_user_choice":
                print(f"   [GUI] Asking user to choose...")
                tool_output = ask_user_choice(
                    arguments.get("question", ""),
                    arguments.get("choice_1", ""),
                    arguments.get("choice_2", ""),
                    arguments.get("choice_3", ""),
                    arguments.get("choice_4", ""),
                    allow_custom=arguments.get("allow_custom", True)
                )

            else:
                # ── Fuzzy tool name resolver ──────────────────────────────────
                # The model sometimes drops or adds an 's', swaps underscores for
                # spaces, or abbreviates tool names. Rather than returning a hard
                # "Unknown tool" failure that breaks the turn, we find the closest
                # real tool name and tell the model to retry with the correct one.
                _KNOWN_TOOLS = {t["function"]["name"] for t in tools}

                # Build common aliases explicitly for the most-misspelled tools
                _TOOL_ALIASES: dict[str, str] = {
                    # Missing/extra 's'
                    "read_path":            "read_paths",
                    "read_instruction":     "read_instructions",
                    "list_skill":           "list_skills",
                    "load_skills":          "load_skill",
                    "list_domain_skill":    "list_domain_skills",
                    "list_domain_knowledges": "list_domain_knowledge",
                    "add_paths":            "add_path",
                    "add_instructions":     "add_instruction",
                    # Spaces instead of underscores
                    "read paths":           "read_paths",
                    "list skills":          "list_skills",
                    "load skill":           "load_skill",
                    "search internet":      "search_internet",
                    "execute terminal":     "execute_terminal_command",
                    "execute command":      "execute_terminal_command",
                    "run command":          "execute_terminal_command",
                    "terminal command":     "execute_terminal_command",
                    "view screen":          "fallback_view_screen",
                    "click text":           "fallback_click_text",
                    "find text":            "fallback_find_text",
                    "click grid":           "fallback_click_grid",
                    "update memory":        "update_memory",
                    "set goal":             "set_current_goal",
                    "write file":           "write_local_file",
                    "read file":            "read_local_file",
                    "append file":          "append_local_file",
                    # Memory
                    "read_mem":             "read_response_memory",
                    "write_mem":            "write_response_memory",
                    "append_mem":           "append_response_memory",
                    # Gemini
                    "gemini":               "consult_gemini",
                    "ask_gemini":           "consult_gemini",
                    # OpenRouter
                    "openrouter":           "consult_openrouter",
                    "ask_openrouter":       "consult_openrouter",
                    "consult_or":           "consult_openrouter",
                    "delegate":             "delegate_to_openrouter",
                    "delegate_task":        "delegate_to_openrouter",
                    "assign_task":          "delegate_to_openrouter",
                    "offload_task":         "delegate_to_openrouter",
                    "offload_to_openrouter": "delegate_to_openrouter",
                    "delegate_or":          "delegate_to_openrouter",
                    "list_models":          "list_openrouter_models",
                    "list_or_models":       "list_openrouter_models",
                    "switch_model":         "set_openrouter_model_by_index",
                    "change_model":         "set_openrouter_model_by_index",
                    "select_model":         "set_openrouter_model_by_index",
                    "set_model":            "set_openrouter_model",
                    # GroqCloud
                    "groq":                 "consult_groq",
                    "ask_groq":             "consult_groq",
                    "consult_gq":           "consult_groq",
                    "delegate_groq":        "delegate_to_groq",
                    "offload_to_groq":      "delegate_to_groq",
                    "list_groq_models":     "list_groq_models",
                    "switch_groq_model":    "set_groq_model_by_index",
                    "set_groq_model":       "set_groq_model",
                    # Gemini-web (primary-execution delegate, distinct from consult_gemini)
                    "delegate_gemini":      "delegate_to_gemini_web",
                    "delegate_to_gemini":   "delegate_to_gemini_web",
                    "offload_to_gemini":    "delegate_to_gemini_web",
                    "delegate_gw":          "delegate_to_gemini_web",
                    "set_gemini_model":     "set_gemini_web_model",
                    # UI snapshot/act — all old names → unified tools
                    "snapshot_ui":              "snapshot",
                    "snapshot_browser_elements": "snapshot",
                    "list_elements":            "snapshot",
                    "list_ui":                  "snapshot",
                    "ui_snapshot":              "snapshot",
                    "scan_ui":                  "snapshot",
                    "browser_snapshot":         "snapshot",
                    "snapshot_page":            "snapshot",
                    "act_on_element":           "act",
                    "act_on_browser_element":   "act",
                    "click_index":              "act",
                    "act_by_index":             "act",
                    "act_element":              "act",
                    "click_browser":            "act",
                    "browser_click":            "act",
                    # UI direct
                    "click_element":            "click_ui_element",
                    "click_ui":                 "click_ui_element",
                    "ui_click":                 "click_ui_element",
                    # Web / browser
                    "navigate":             "open_url",
                    "navigate_to":          "open_url",
                    "go_to_url":            "open_url",
                    "open_link":            "open_url",
                    "browse":               "open_url",
                    "open_browser":         "open_url",
                    "open_result":          "open_search_result",
                    "read_page":            "read_browser_page",
                    "read_browser":         "read_browser_page",
                    "read_browser_tab":     "read_browser_page",
                    "get_page_content":     "read_browser_page",
                    "browser_tabs":         "list_browser_tabs",
                    "list_tabs":            "list_browser_tabs",
                    "run_js":               "run_js_in_browser",
                    "execute_js":           "run_js_in_browser",
                    "js":                   "run_js_in_browser",
                    # Directory / file
                    "list_dir":             "list_directory",
                    "ls":                   "list_directory",
                    "explore_path":         "list_directory",
                    "open_file":            "open_path",
                    "search_file":          "find_file",
                    "search_files":         "find_file",
                    "locate_file":          "find_file",
                    "open_file_index":      "open_path_by_index",
                    # Skills / knowledge
                    "list_skills":          "list_skills_indexed",
                    "list_skill":           "list_skills_indexed",
                    "load_skill_index":     "load_skill_by_index",
                    "list_domain_knowledge": "list_domain_knowledge_indexed",
                    "list_domain_skills":   "list_domain_skills_indexed",
                    "read_domain_index":    "read_domain_by_index",
                    # Paths
                    "read_paths":           "list_paths_indexed",
                    "read_path":            "list_paths_indexed",
                    "get_path_index":       "get_path",
                    # OCR
                    "ocr_screen":           "ocr_snapshot",
                    "screen_snapshot":      "ocr_snapshot",
                    "click_ocr":            "click_ocr_index",
                    # Misc
                    "type":                 "type_text",
                    "screenshot":           "fallback_view_screen",
                }

                normalised = func_name.lower().replace("-", "_")
                resolved   = _TOOL_ALIASES.get(normalised)

                if not resolved:
                    # Fuzzy: find the real tool name with the most character overlap
                    from difflib import get_close_matches
                    close = get_close_matches(normalised, _KNOWN_TOOLS, n=1, cutoff=0.6)
                    resolved = close[0] if close else None

                if resolved and resolved in _KNOWN_TOOLS:
                    if resolved == func_name or resolved == normalised:
                        # Resolved to itself — tool exists but has no dispatch case
                        tool_output = (
                            f"[INTERNAL ERROR] Tool '{func_name}' is registered but has no "
                            f"dispatch handler. This is a Midum bug — report it."
                        )
                    else:
                        print(f"   [Tool resolver] '{func_name}' → '{resolved}'")
                        tool_output = (
                            f"[TOOL NAME ERROR] You called '{func_name}' which does not exist. "
                            f"The correct tool name is '{resolved}'. "
                            f"Call '{resolved}' now with the same arguments."
                        )
                else:
                    mcp_matches = _mcp_find_tool_matches(func_name)
                    if len(mcp_matches) > 1:
                        options = "; ".join(
                            f"server='{s}' tool_name='{t}'" for s, t in mcp_matches
                        )
                        tool_output = (
                            f"[AMBIGUOUS MCP TOOL] '{func_name}' matches a tool on more than "
                            f"one connected server: {options}. Call call_mcp_tool with the "
                            f"specific server you meant."
                        )
                    else:
                        known_list = ", ".join(sorted(_KNOWN_TOOLS))
                        tool_output = (
                            f"[UNKNOWN TOOL] '{func_name}' is not a valid tool. "
                            f"Available tools: {known_list}. If this was meant to be an MCP "
                            f"server tool, call list_mcp_servers() then show_server_tools(server) "
                            f"to confirm the exact name, then call_mcp_tool(server, tool_name, arguments)."
                        )

            # ── Record in TurnState ───────────────────────────────────────────
            state.record(func_name, arguments, tool_output)

            # ──────────────────────────────────────────────────────────────────

            tool_output = tool_output.replace("<", "&lt;").replace(">", "&gt;")
            turn_tool_outputs.append(tool_output)

            msg = {"role": "tool", "content": tool_output}
            if tool_images:
                msg["images"] = tool_images
            conversation_history.append(msg)

        if needs_lookup:
            conversation_history.append({
                "role": "user",
                "content": (
                    f"[SYSTEM NOTE]: '{unknown_cmd.split()[0]}' was not in commands.md. "
                    "Check commands.md or search online if you are unsure it was correct."
                )
            })

        # ── Step prompt: state-aware next-step message ────────────────────────
        conversation_history.append({
            "role": "user",
            "content": state.build_step_prompt()
        })

    # ── Reached via break (MAX_STEPS or cap_hit) ─────────────────────────────
    full_reply = "\n\n".join(p for p in _accumulated_reply if p.strip())
    return full_reply or "[Task stopped — see above for details.]", turn_tool_outputs


# =============================================================================
# 10. INTERACTIVE MAIN LOOP
# =============================================================================

if __name__ == "__main__":
    os.makedirs(TARGET_DIR, exist_ok=True)

    # ── Ctrl+Q abort hotkey ────────────────────────────────────────────────────
    if _KEYBOARD_AVAILABLE:
        _keyboard.add_hotkey("ctrl+q", lambda: _abort_event.set())
        print("⌨️  [Ctrl+Q registered — press to abort the current response]")
    else:
        print("⚠️  [keyboard package not found — Ctrl+Q abort unavailable]")
        print("    Install with: pip install keyboard")

    # ── Tesseract status ───────────────────────────────────────────────────────
    if _TESSERACT_AVAILABLE:
        print("👁️  [Tesseract OCR: available — fallback_click_text is active]")
    else:
        print("⚠️  [Tesseract OCR not found — fallback_click_text will not work]")
        if _IS_LINUX:
            print("    Install with: sudo apt install tesseract-ocr && pip install pytesseract")
        else:
            print("    Install from: https://github.com/UB-Mannheim/tesseract/wiki")
            print("    Then: pip install pytesseract")

    # ── UI automation status ───────────────────────────────────────────────────
    if _IS_LINUX:
        if _PYATSPI_AVAILABLE:
            print("🖱️  [Linux UI automation: AT-SPI2 + xdotool available]")
        else:
            print("⚠️  [pyatspi not found — AT-SPI tree inspection unavailable]")
            print("    Install: sudo apt install python3-pyatspi xdotool xclip")
    else:
        if _UIA_AVAILABLE:
            print("🖱️  [Windows UI automation: uiautomation available]")
        else:
            print("⚠️  [UI automation unavailable — click_ui_element/snapshot will not work]")
            if _UIA_INIT_ERROR:
                print(f"    Real cause: {_UIA_INIT_ERROR}")
                low = _UIA_INIT_ERROR.lower()
                if "pywin32 not installed" in low:
                    print("    Fix: pip install pywin32 uiautomation")
                elif "coinitialize" in low or "com" in low:
                    print("    Fix: this is a COM threading conflict, not a Windows version issue.")
                    print("    Try: pip install --force-reinstall pywin32, then run:")
                    print("      python <python_dir>\\Scripts\\pywin32_postinstall.py -install")
                    print("    (run that command as Administrator)")
                else:
                    print("    This is NOT actually a 'Windows needs updating' issue — ignore that advice.")
                    print("    Try: pip install --force-reinstall uiautomation pywin32")
            else:
                print("    pip install uiautomation pywin32")

    # ── Gemini status (web chat only, via gemini_webapi — no API fallback) ─────
    if _GEMINI_WEBAPI_AVAILABLE:
        print("🤖 [Gemini: gemini_webapi installed — consult_gemini is active "
              "(client/cookies initialise lazily on first use)]")
    else:
        print(f"⚠️  [Gemini web chat unavailable: {_gemini_webapi_load_msg}]")
        print( "    Install: pip install -U gemini_webapi")
        print( "    Optional (auto cookie import): pip install -U browser-cookie3")
        print(f"    Or add cookies manually to the secrets file "
              f"({os.path.abspath(SECRETS_FILE)}):")
        print( "      { \"GEMINI_SECURE_1PSID\": \"...\", \"GEMINI_SECURE_1PSIDTS\": \"...\" }")
    if _GEMINI_AVAILABLE:
        print("    (Gemini API key is also configured, but is unused — consult_gemini "
              "only uses the web chat interface now.)")

    # ── Gemini API (official) status ────────────────────────────────────────────
    if _GEMINI_API_AVAILABLE:
        print(f"🔑 [Gemini API (official): available — model={GEMINI_API_MODEL}]")
    else:
        print(f"⚠️  [Gemini API (official) not available: {_gemini_api_load_msg}]")
        print(f"    Secrets file expected at: {os.path.abspath(SECRETS_FILE)}")
        print( "    Add key: { \"GEMINI_API_KEY\": \"AIza...\" } (same file as everything else)")
        print( "    Get a key: https://aistudio.google.com/app/apikey")

    # ── OpenRouter status ──────────────────────────────────────────────────────
    if _OPENROUTER_AVAILABLE:
        print(f"🌍 [OpenRouter: available — model={OPENROUTER_MODEL}, "
              f"consult_mode={OPENROUTER_CONSULT_MODE}]")
    else:
        print(f"⚠️  [OpenRouter not available: {_openrouter_load_msg}]")
        print(f"    Secrets file expected at: {os.path.abspath(SECRETS_FILE)}")
        print( "    Add key: { \"OPENROUTER_API_KEY\": \"sk-or-v1-...\" } (same file as Gemini)")
        print( "    Get a key: https://openrouter.ai/keys")

    # ── GroqCloud status ────────────────────────────────────────────────────────
    if _GROQ_AVAILABLE:
        print(f"⚡ [GroqCloud: available — model={GROQ_MODEL}]")
    else:
        print(f"⚠️  [GroqCloud not available: {_groq_load_msg}]")
        print(f"    Secrets file expected at: {os.path.abspath(SECRETS_FILE)}")
        print( "    Add key: { \"GROQ_API_KEY\": \"gsk_...\" } (same file as everything else)")
        print( "    Get a free key: https://console.groq.com/keys")

    # ── Primary model provider summary ────────────────────────────────────────
    if MODEL_PROVIDER == "openrouter":
        if not _OPENROUTER_AVAILABLE:
            print("🛑 [MODEL_PROVIDER='openrouter' but OpenRouter is NOT configured — "
                  "Midum cannot run until OPENROUTER_API_KEY is set!]")
        else:
            print(f"🧠 [PRIMARY MODEL: OpenRouter/{OPENROUTER_MODEL} — driving Midum directly]")
    elif MODEL_PROVIDER == "gemini_web":
        if not _GEMINI_WEBAPI_AVAILABLE:
            print(f"🛑 [MODEL_PROVIDER='gemini_web' but gemini_webapi is NOT installed — "
                  f"{_gemini_webapi_load_msg}]")
        else:
            _gw_client, _gw_err = _get_gemini_web_client()
            if _gw_err:
                print(f"🛑 [MODEL_PROVIDER='gemini_web' but the client failed to initialize: {_gw_err}]")
            else:
                print(f"🧠 [PRIMARY MODEL: Gemini-web/{GEMINI_WEB_MODEL or 'auto'} — "
                      f"driving Midum directly via gemini_webapi ChatSession]")
    elif MODEL_PROVIDER == "gemini_api":
        if not _GEMINI_API_AVAILABLE:
            print("🛑 [MODEL_PROVIDER='gemini_api' but the Gemini API is NOT configured — "
                  "Midum cannot run until GEMINI_API_KEY is set!]")
        else:
            print(f"🧠 [PRIMARY MODEL: Gemini-API/{GEMINI_API_MODEL} — driving Midum directly "
                  f"via the official API]")
    elif MODEL_PROVIDER == "groq":
        if not _GROQ_AVAILABLE:
            print("🛑 [MODEL_PROVIDER='groq' but GroqCloud is NOT configured — "
                  "Midum cannot run until GROQ_API_KEY is set!]")
        else:
            print(f"🧠 [PRIMARY MODEL: Groq/{GROQ_MODEL} — driving Midum directly "
                  f"via GroqCloud's free-tier API]")
    else:
        print(f"🧠 [PRIMARY MODEL: Ollama/{MODEL_NAME} — local execution brain]")
        if OPENROUTER_CONSULT_MODE != "off" and _OPENROUTER_AVAILABLE:
            print(f"    Planning consult: OpenRouter/{OPENROUTER_MODEL} "
                  f"({OPENROUTER_CONSULT_MODE}), Gemini as {'primary consult' if OPENROUTER_CONSULT_MODE=='fallback' else 'fallback'}")

    # ── CDP browser status ─────────────────────────────────────────────────────
    if _CDP_AVAILABLE:
        tabs = _cdp_get_tabs()
        if tabs:
            print(f"🌐 [CDP: connected — {len(tabs)} tab(s) open]")
        else:
            print("🌐 [CDP: installed but Chrome not in debug mode]")
            print("    Launch Chrome with: --remote-debugging-port=9222")
    else:
        print("⚠️  [CDP not available — browser DOM tools disabled]")
        print("    Install: pip install websocket-client requests")

    # ── MCP servers ─────────────────────────────────────────────────────────
    if _MCP_SDK_AVAILABLE:
        init_mcp_servers_from_config()
        if _MCP_SERVER_ORDER:
            connected_n = sum(1 for n in _MCP_SERVER_ORDER if _MCP_SERVERS[n].connected)
            print(f"🧩 [MCP: {connected_n}/{len(_MCP_SERVER_ORDER)} server(s) connected — "
                  f"list_mcp_servers() for details]")
        else:
            print(f"🧩 [MCP: SDK installed, no servers configured yet — "
                  f"connect_mcp_server(...) to add one, or edit "
                  f"{os.path.abspath(MCP_SERVERS_FILE)}]")
    else:
        print("⚠️  [MCP not available — 'mcp' package not installed]")
        print("    Install: pip install mcp")

    print(f"🖥️  [Platform: {'Linux' if _IS_LINUX else 'Windows'} | "
          f"Screen: {SCREEN_W}x{SCREEN_H} | "
          f"Shell: {'bash' if _IS_LINUX else 'PowerShell'}]")

    memory_injections = init_memory_at_startup()

    try:
        with open(LOG_FILE, "w", encoding="utf-8") as f:
            f.write("# Midum Master Interaction Log\n")
            f.write(f"Session started: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
            f.write("=========================================\n\n")
    except Exception as e:
        print(f"⚠️ Warning: Could not initialise log file: {e}")

    print("\n====================================================")
    print("Midum local agent started. Persistent Chat Ready.")
    print(f"Tracking live session in: {LOG_FILE}")
    print("Type 'new session' to wipe session memory.")
    print("Type 'exit' or 'quit' to close.")
    print("====================================================\n")

    system_prompt = get_system_prompt()

    goal_reminder = ""
    if _current_goal:
        goal_reminder = (
            f"\n\n[GOAL REMINDER]\nCurrent goal: {_current_goal}\n"
            "Continue unless redirected."
        )

    history = [{"role": "system", "content": system_prompt + goal_reminder}]
    for inj in memory_injections:
        history.append({"role": "system", "content": inj})

    turn_counter = 1

    while True:
        try:
            user_input = input("\nYou: ").strip()

            if user_input.lower() == "new session":
                if os.path.exists(SESSION_MEMORY):
                    os.remove(SESSION_MEMORY)
                    print("🗑️ Session memory cleared.")
                _current_goal = None
                reset_groq_loaded_tools()
                ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                write_local_file(
                    SESSION_MEMORY,
                    f"# Midum Session Memory\nSession started: {ts}\n\n"
                    f"{GOAL_SECTION_HEADER}\n_No active goal._\n\n{GOAL_SECTION_END}\n"
                )
                print("🧠 New session started.\n")
                history = [{"role": "system", "content": system_prompt}]
                for inj in memory_injections[:1]:
                    history.append({"role": "system", "content": inj})
                turn_counter = 1
                continue

            if user_input.lower() in ("exit", "quit"):
                print("\nCleaning up...")
                if os.path.exists(LOG_FILE):
                    try:
                        os.remove(LOG_FILE)
                        print(f"🗑️ Deleted: {LOG_FILE}")
                    except Exception as e:
                        print(f"⚠️ Could not delete log: {e}")
                print("Goodbye!")
                break

            if not user_input:
                continue

            gemini_plan = ""   # always defined before process_chat_turn
            approval_keywords = ["yes", "grant", "approve", "run it", "go ahead", "y"]
            if any(kw in user_input.lower() for kw in approval_keywords):
                payload = f"{user_input} [USER MANUALLY GRANTED BYPASS]"
                history.append({"role": "user", "content": payload})
            else:
                # ── Gemini pre-planning ───────────────────────────────────────
                if not _is_trivial_input(user_input):
                    plan = get_gemini_reasoning(user_input, history)
                    gemini_plan = plan or ""
                else:
                    gemini_plan = ""

                # ── Task decomposition (fallback if Gemini unavailable) ────────
                task_plan = _decompose_task(user_input) if not gemini_plan else None

                plan_text = ""
                if gemini_plan:
                    plan_text = f"\n\n[EXECUTION PLAN FROM PLANNING BRAIN]\n{gemini_plan}"
                elif task_plan:
                    plan_text = f"\n\n{task_plan}"

                payload = (
                    f"{user_input}{plan_text}\n\n"
                    "[SYSTEM]: Follow the execution plan above step by step. "
                    "Execute the first tool call now. Do not explain — just act."
                )
                history.append({"role": "user", "content": payload})

            print("\n[Thinking...]")
            assistant_reply, turn_tool_outputs = process_chat_turn(
                history,
                user_request=user_input,
                gemini_plan=gemini_plan
            )
            _print_reply("Midum:", assistant_reply)

            python_trigger_memory_update(turn_tool_outputs, assistant_reply)

            try:
                with open(LOG_FILE, "a", encoding="utf-8") as f:
                    f.write(f"# Response {turn_counter}\n\n")
                    f.write(f"### **User Prompt:**\n> {user_input}\n\n")
                    f.write(f"_Goal: {_current_goal or 'none'}_\n\n")
                    f.write(f"### **Midum Reply:**\n{assistant_reply}\n\n---\n\n")
                print(f"💾 [Logged response {turn_counter}]")
                turn_counter += 1
            except Exception as e:
                print(f"⚠️ Could not append to log: {e}")

        except KeyboardInterrupt:
            print("\n\nAborted.")
            if os.path.exists(LOG_FILE):
                try: os.remove(LOG_FILE)
                except: pass
            break
        except Exception as e:
            print(f"\nUnexpected error: {e}")